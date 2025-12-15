#!/usr/bin/env python3
"""
Test Word to PDF với nhiều định dạng phức tạp
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
from pathlib import Path

def create_complex_word_file():
    """Tạo file Word với nhiều định dạng phức tạp"""
    doc = Document()
    
    # Title với formatting
    title = doc.add_heading('Test Document - Định Dạng Phức Tạp', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Paragraph với bold, italic, underline
    p1 = doc.add_paragraph()
    p1.add_run('Text bình thường, ')
    p1.add_run('text in đậm').bold = True
    p1.add_run(', ')
    p1.add_run('text in nghiêng').italic = True
    p1.add_run(', và ')
    p1.add_run('text gạch chân').underline = True
    
    # Text với màu sắc
    p2 = doc.add_paragraph()
    run = p2.add_run('Text màu đỏ')
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.font.size = Pt(14)
    
    # Heading cấp 1
    doc.add_heading('1. Section với Tiếng Việt', level=1)
    
    # Paragraph với tiếng Việt đầy đủ dấu
    doc.add_paragraph(
        'Đây là đoạn văn tiếng Việt có đầy đủ dấu: '
        'àáảãạăằắẳẵặâầấẩẫậèéẻẽẹêềếểễệìíỉĩịòóỏõọôồốổỗộơờớởỡợùúủũụưừứửữựỳýỷỹỵ'
    )
    
    # Bullet list
    doc.add_heading('2. Danh sách gạch đầu dòng', level=1)
    doc.add_paragraph('Item 1 với text bình thường', style='List Bullet')
    doc.add_paragraph('Item 2 với tiếng Việt: Nguyễn Văn A', style='List Bullet')
    doc.add_paragraph('Item 3 với số: 123,456,789 VNĐ', style='List Bullet')
    
    # Numbered list
    doc.add_heading('3. Danh sách đánh số', level=1)
    doc.add_paragraph('Bước 1: Chuẩn bị', style='List Number')
    doc.add_paragraph('Bước 2: Thực hiện', style='List Number')
    doc.add_paragraph('Bước 3: Hoàn thành', style='List Number')
    
    # Table
    doc.add_heading('4. Bảng biểu', level=1)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tên'
    hdr_cells[1].text = 'Tuổi'
    hdr_cells[2].text = 'Địa chỉ'
    
    # Data rows
    data = [
        ('Nguyễn Văn A', '25', 'Hà Nội'),
        ('Trần Thị B', '30', 'TP HCM'),
        ('Lê Văn C', '28', 'Đà Nẵng')
    ]
    
    for i, (name, age, address) in enumerate(data, start=1):
        row = table.rows[i].cells
        row[0].text = name
        row[1].text = age
        row[2].text = address
    
    # Paragraph với alignment khác nhau
    doc.add_heading('5. Căn lề khác nhau', level=1)
    
    p_left = doc.add_paragraph('Text căn trái (Left aligned)')
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    p_center = doc.add_paragraph('Text căn giữa (Center aligned)')
    p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_right = doc.add_paragraph('Text căn phải (Right aligned)')
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    p_justify = doc.add_paragraph(
        'Text căn đều hai bên (Justified). '
        'Đây là một đoạn văn dài hơn để thấy rõ hiệu quả của căn đều. '
        'Các dòng sẽ được căn chỉnh để vừa khít hai bên lề.'
    )
    p_justify.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Multiple font sizes
    doc.add_heading('6. Kích thước font khác nhau', level=1)
    for size in [10, 12, 14, 16, 18]:
        p = doc.add_paragraph()
        run = p.add_run(f'Font size {size}pt')
        run.font.size = Pt(size)
    
    # Save
    filename = 'test_complex_word.docx'
    doc.save(filename)
    print(f'✅ Created complex Word file: {filename}')
    return Path(filename)

def test_conversion():
    """Test convert và so sánh"""
    BASE_URL = "http://165.99.59.47"
    
    # Login
    print('\n🔐 Logging in...')
    login_response = requests.post(
        f"{BASE_URL}/api/auth/login",
        json={"username": "admin", "password": "admin123"}
    )
    
    if login_response.status_code != 200:
        print(f'❌ Login failed: {login_response.status_code}')
        return
    
    token = login_response.json()['token']['access_token']
    headers = {"Authorization": f"Bearer {token}"}
    print('✅ Login successful')
    
    # Create test file
    print('\n📝 Creating complex Word document...')
    word_file = create_complex_word_file()
    
    # Convert
    print(f'\n🔄 Converting {word_file.name} to PDF...')
    with open(word_file, 'rb') as f:
        files = {
            'file': (word_file.name, f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        }
        response = requests.post(
            f"{BASE_URL}/api/documents/convert/word-to-pdf",
            files=files,
            headers=headers,
            timeout=60
        )
    
    if response.status_code == 200:
        pdf_file = word_file.with_suffix('.pdf')
        with open(pdf_file, 'wb') as f:
            f.write(response.content)
        
        print(f'\n✅ Conversion SUCCESS!')
        print(f'   Input:  {word_file} ({word_file.stat().st_size:,} bytes)')
        print(f'   Output: {pdf_file} ({pdf_file.stat().st_size:,} bytes)')
        print(f'\n📂 Files created:')
        print(f'   - {word_file} (Word original)')
        print(f'   - {pdf_file} (PDF converted)')
        print(f'\n💡 Hãy mở 2 files này để so sánh định dạng:')
        print(f'   1. Mở {word_file} trong Word')
        print(f'   2. Mở {pdf_file} trong PDF viewer')
        print(f'   3. So sánh: Bold, Italic, Colors, Tables, Lists, Alignment')
        
        return True
    else:
        print(f'❌ Conversion FAILED: {response.status_code}')
        print(f'   Response: {response.text[:500]}')
        return False

if __name__ == '__main__':
    print('=' * 60)
    print('🧪 TEST WORD FORMATTING IN PDF CONVERSION')
    print('=' * 60)
    
    try:
        success = test_conversion()
        if success:
            print('\n✅ Test completed! Please check the files manually.')
        else:
            print('\n❌ Test failed!')
    except Exception as e:
        print(f'\n❌ Error: {str(e)}')
        import traceback
        traceback.print_exc()
