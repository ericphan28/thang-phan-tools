"""
Test All Document API Endpoints
Kiểm tra toàn bộ tính năng trong /tools
"""

import requests
import json
from pathlib import Path
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from docx import Document
from openpyxl import Workbook
from PIL import Image

BASE_URL = "http://localhost:8000/api/v1/documents"

# Colors for output
GREEN = "\033[92m"
RED = "\033[91m"
YELLOW = "\033[93m"
RESET = "\033[0m"

def print_test(name, status, message=""):
    """Print test result with colors"""
    color = GREEN if status == "✅" else RED if status == "❌" else YELLOW
    print(f"{color}{status} {name}{RESET} {message}")

def create_test_pdf(filename="test.pdf"):
    """Tạo file PDF test"""
    pdf_path = Path("uploads") / filename
    pdf_path.parent.mkdir(exist_ok=True)
    
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    c.drawString(100, 750, "Test PDF Document")
    c.drawString(100, 730, "This is a test file for API testing")
    c.showPage()
    c.drawString(100, 750, "Page 2")
    c.showPage()
    c.drawString(100, 750, "Page 3")
    c.save()
    
    with open(pdf_path, 'wb') as f:
        f.write(buffer.getvalue())
    
    return pdf_path

def create_test_word(filename="test.docx"):
    """Tạo file Word test"""
    doc_path = Path("uploads") / filename
    doc_path.parent.mkdir(exist_ok=True)
    
    doc = Document()
    doc.add_heading("Test Word Document", 0)
    doc.add_paragraph("This is a test file for API testing")
    doc.add_paragraph("Tiếng Việt có dấu: Đây là văn bản tiếng Việt")
    doc.save(doc_path)
    
    return doc_path

def create_test_excel(filename="test.xlsx"):
    """Tạo file Excel test"""
    excel_path = Path("uploads") / filename
    excel_path.parent.mkdir(exist_ok=True)
    
    wb = Workbook()
    ws = wb.active
    ws['A1'] = "Test Excel"
    ws['A2'] = "Data 1"
    ws['B2'] = "Data 2"
    wb.save(excel_path)
    
    return excel_path

def create_test_image(filename="test.png"):
    """Tạo file ảnh test"""
    img_path = Path("uploads") / filename
    img_path.parent.mkdir(exist_ok=True)
    
    img = Image.new('RGB', (800, 600), color='white')
    img.save(img_path)
    
    return img_path

def test_pdf_to_word():
    """Test PDF → Word"""
    try:
        pdf_path = create_test_pdf("test_pdf_to_word.pdf")
        
        with open(pdf_path, 'rb') as f:
            response = requests.post(
                f"{BASE_URL}/convert/pdf-to-word",
                files={"file": ("test.pdf", f, "application/pdf")}
            )
        
        if response.status_code == 200:
            print_test("PDF → Word", "✅", f"({len(response.content)} bytes)")
            return True
        else:
            print_test("PDF → Word", "❌", f"Status {response.status_code}: {response.text[:100]}")
            return False
    except Exception as e:
        print_test("PDF → Word", "❌", str(e))
        return False

def test_word_to_pdf():
    """Test Word → PDF"""
    try:
        doc_path = create_test_word("test_word_to_pdf.docx")
        
        with open(doc_path, 'rb') as f:
            response = requests.post(
                f"{BASE_URL}/convert/word-to-pdf",
                files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            )
        
        if response.status_code == 200:
            print_test("Word → PDF", "✅", f"({len(response.content)} bytes)")
            return True
        else:
            print_test("Word → PDF", "❌", f"Status {response.status_code}: {response.text[:100]}")
            return False
    except Exception as e:
        print_test("Word → PDF", "❌", str(e))
        return False

def test_excel_to_pdf():
    """Test Excel → PDF"""
    try:
        excel_path = create_test_excel("test_excel_to_pdf.xlsx")
        
        with open(excel_path, 'rb') as f:
            response = requests.post(
                f"{BASE_URL}/convert/excel-to-pdf",
                files={"file": ("test.xlsx", f, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")}
            )
        
        if response.status_code == 200:
            print_test("Excel → PDF", "✅", f"({len(response.content)} bytes)")
            return True
        else:
            print_test("Excel → PDF", "❌", f"Status {response.status_code}: {response.text[:100]}")
            return False
    except Exception as e:
        print_test("Excel → PDF", "❌", str(e))
        return False

def test_merge_pdfs():
    """Test Gộp PDF"""
    try:
        pdf1 = create_test_pdf("merge1.pdf")
        pdf2 = create_test_pdf("merge2.pdf")
        
        with open(pdf1, 'rb') as f1, open(pdf2, 'rb') as f2:
            response = requests.post(
                f"{BASE_URL}/pdf/merge",
                files=[
                    ("files", ("merge1.pdf", f1, "application/pdf")),
                    ("files", ("merge2.pdf", f2, "application/pdf"))
                ]
            )
        
        if response.status_code == 200:
            print_test("Gộp PDF", "✅", f"({len(response.content)} bytes)")
            return True
        else:
            print_test("Gộp PDF", "❌", f"Status {response.status_code}: {response.text[:100]}")
            return False
    except Exception as e:
        print_test("Gộp PDF", "❌", str(e))
        return False

def test_split_pdf():
    """Test Tách PDF"""
    try:
        pdf_path = create_test_pdf("test_split.pdf")
        
        with open(pdf_path, 'rb') as f:
            response = requests.post(
                f"{BASE_URL}/pdf/split",
                files={"file": ("test.pdf", f, "application/pdf")},
                data={"page_ranges": "1,2,3"}
            )
        
        if response.status_code == 200:
            print_test("Tách PDF", "✅", f"({len(response.content)} bytes)")
            return True
        else:
            print_test("Tách PDF", "❌", f"Status {response.status_code}: {response.text[:100]}")
            return False
    except Exception as e:
        print_test("Tách PDF", "❌", str(e))
        return False

def test_compress_pdf():
    """Test Nén PDF"""
    try:
        pdf_path = create_test_pdf("test_compress.pdf")
        
        with open(pdf_path, 'rb') as f:
            response = requests.post(
                f"{BASE_URL}/pdf/compress",
                files={"file": ("test.pdf", f, "application/pdf")},
                data={"quality": "medium"}
            )
        
        if response.status_code == 200:
            print_test("Nén PDF", "✅", f"({len(response.content)} bytes)")
            return True
        else:
            print_test("Nén PDF", "❌", f"Status {response.status_code}: {response.text[:100]}")
            return False
    except Exception as e:
        print_test("Nén PDF", "❌", str(e))
        return False

def test_extract_text():
    """Test Trích xuất text từ PDF"""
    try:
        pdf_path = create_test_pdf("test_extract.pdf")
        
        with open(pdf_path, 'rb') as f:
            response = requests.post(
                f"{BASE_URL}/pdf/extract-text",
                files={"file": ("test.pdf", f, "application/pdf")}
            )
        
        if response.status_code == 200:
            data = response.json()
            print_test("Trích xuất text", "✅", f"({len(data.get('text', ''))} chars)")
            return True
        else:
            print_test("Trích xuất text", "❌", f"Status {response.status_code}: {response.text[:100]}")
            return False
    except Exception as e:
        print_test("Trích xuất text", "❌", str(e))
        return False

def test_pdf_to_images():
    """Test PDF → Images"""
    try:
        pdf_path = create_test_pdf("test_to_images.pdf")
        
        with open(pdf_path, 'rb') as f:
            response = requests.post(
                f"{BASE_URL}/pdf/to-images",
                files={"file": ("test.pdf", f, "application/pdf")},
                data={"format": "png", "dpi": "150"}
            )
        
        if response.status_code == 200:
            print_test("PDF → Images", "✅", f"({len(response.content)} bytes)")
            return True
        else:
            print_test("PDF → Images", "❌", f"Status {response.status_code}: {response.text[:100]}")
            return False
    except Exception as e:
        print_test("PDF → Images", "❌", str(e))
        return False

def test_rotate_pdf():
    """Test Xoay PDF"""
    try:
        pdf_path = create_test_pdf("test_rotate.pdf")
        
        with open(pdf_path, 'rb') as f:
            response = requests.post(
                f"{BASE_URL}/pdf/rotate",
                files={"file": ("test.pdf", f, "application/pdf")},
                data={"angle": "90"}
            )
        
        if response.status_code == 200:
            print_test("Xoay PDF", "✅", f"({len(response.content)} bytes)")
            return True
        else:
            print_test("Xoay PDF", "❌", f"Status {response.status_code}: {response.text[:100]}")
            return False
    except Exception as e:
        print_test("Xoay PDF", "❌", str(e))
        return False

def test_images_to_pdf():
    """Test Images → PDF"""
    try:
        img1 = create_test_image("img1.png")
        
        with open(img1, 'rb') as f1:
            response = requests.post(
                f"{BASE_URL}/convert/image-to-pdf",
                files={"file": ("img1.png", f1, "image/png")}
            )
        
        if response.status_code == 200:
            print_test("Images → PDF", "✅", f"({len(response.content)} bytes)")
            return True
        else:
            print_test("Images → PDF", "❌", f"Status {response.status_code}: {response.text[:100]}")
            return False
    except Exception as e:
        print_test("Images → PDF", "❌", str(e))
        return False

def main():
    """Chạy tất cả tests"""
    print("\n" + "="*60)
    print("🧪 TESTING ALL DOCUMENT API ENDPOINTS")
    print("="*60 + "\n")
    
    tests = [
        ("PDF → Word", test_pdf_to_word),
        ("Word → PDF", test_word_to_pdf),
        ("Excel → PDF", test_excel_to_pdf),
        ("Gộp PDF", test_merge_pdfs),
        ("Tách PDF", test_split_pdf),
        ("Nén PDF", test_compress_pdf),
        ("Trích xuất text", test_extract_text),
        ("PDF → Images", test_pdf_to_images),
        ("Xoay PDF", test_rotate_pdf),
        ("Images → PDF", test_images_to_pdf),
    ]
    
    results = []
    for name, test_func in tests:
        try:
            result = test_func()
            results.append((name, result))
        except Exception as e:
            print_test(name, "❌", f"Exception: {e}")
            results.append((name, False))
    
    # Summary
    print("\n" + "="*60)
    print("📊 TEST SUMMARY")
    print("="*60)
    
    passed = sum(1 for _, result in results if result)
    failed = len(results) - passed
    
    print(f"\n{GREEN}✅ Passed: {passed}{RESET}")
    print(f"{RED}❌ Failed: {failed}{RESET}")
    print(f"Total: {len(results)}")
    
    if failed > 0:
        print(f"\n{RED}Failed tests:{RESET}")
        for name, result in results:
            if not result:
                print(f"  - {name}")
    
    print("\n" + "="*60 + "\n")

if __name__ == "__main__":
    main()
