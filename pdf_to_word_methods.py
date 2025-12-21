"""
PDF to Word Smart Conversion Methods - Add to DocumentService
"""

# Add this after cleanup_old_files method in document_service.py

    # ==================== Smart PDF to Word ====================
    
    async def pdf_to_word_smart(
        self,
        input_file: Path,
        language: str = "vi"
    ) -> Path:
        """
        Convert scanned PDF to Word using Gemini AI
        
        Strategy:
        1. Upload PDF to Gemini Files API
        2. Generate Markdown with structure (headings, tables, formatting)
        3. Parse Markdown and create Word document with python-docx
        
        Args:
            input_file: Path to PDF file
            language: Language (vi or en)
            
        Returns:
            Path to generated .docx file
        """
        if input_file.suffix.lower() != '.pdf':
            raise HTTPException(400, "File must be .pdf")
        
        start_time = time.time()
        output_file = self.output_dir / f"{input_file.stem}_converted.docx"
        
        try:
            # Step 1: Upload PDF to Gemini and generate Markdown
            logger.info(f"📄 Converting PDF to Markdown with Gemini...")
            markdown_result = await self._pdf_to_markdown_gemini(input_file, language)
            
            # Step 2: Convert Markdown to Word
            logger.info(f"📝 Converting Markdown to Word...")
            await self._markdown_to_word(markdown_result["markdown"], output_file)
            
            processing_time = time.time() - start_time
            logger.info(f"✅ PDF → Word completed in {processing_time:.2f}s")
            
            # Track AI usage
            await self._track_pdf_to_word_usage(
                markdown_result.get("ai_usage", {}),
                processing_time
            )
            
            return output_file
            
        except Exception as e:
            logger.error(f"PDF to Word conversion failed: {e}")
            raise HTTPException(500, f"Conversion failed: {str(e)}")
    
    async def _pdf_to_markdown_gemini(self, pdf_path: Path, language: str) -> dict:
        """Generate Markdown from PDF using Gemini native PDF support"""
        try:
            import google.generativeai as genai
            
            # Get API key
            from app.services.ai_usage_service import get_api_key
            api_key = get_api_key("gemini")
            if not api_key:
                raise HTTPException(500, "Gemini API key not configured")
                
            genai.configure(api_key=api_key)
            
            # Upload PDF
            logger.info(f"📤 Uploading PDF to Gemini...")
            uploaded_file = genai.upload_file(str(pdf_path))
            
            # Create model
            model = genai.GenerativeModel("gemini-2.0-flash-exp")
            
            # Prompt for Markdown generation
            lang_desc = "Tiếng Việt" if language == "vi" else "English"
            prompt = f"""Chuyển đổi toàn bộ nội dung tài liệu PDF này sang định dạng Markdown.

YÊU CẦU QUAN TRỌNG:
1. **Tiêu đề (Headings):**
   - Heading 1: # Text
   - Heading 2: ## Text
   - Heading 3: ### Text

2. **Bảng (Tables):**
   - Dùng Markdown table syntax:
   ```
   | Cột 1 | Cột 2 | Cột 3 |
   |-------|-------|-------|
   | Dữ liệu | Dữ liệu | Dữ liệu |
   ```
   - Giữ nguyên cấu trúc bảng gốc
   - Merge cells: dùng colspan notation nếu có

3. **Format text:**
   - Bold: **text**
   - Italic: *text*
   - Bold + Italic: ***text***

4. **Lists:**
   - Unordered: - Item hoặc * Item
   - Ordered: 1. Item, 2. Item

5. **Paragraphs:**
   - Tách đoạn văn bằng dòng trống
   - Giữ xuống dòng trong đoạn

6. **Độ chính xác:**
   - Giữ chính xác 100% ký tự {lang_desc}
   - Giữ dấu thanh, dấu câu
   - Giữ số liệu, ngày tháng chính xác

7. **Cấu trúc:**
   - Giữ nguyên thứ tự nội dung từ trên xuống dưới
   - Phân biệt rõ header, body, footer

CHỈ trả về Markdown text, KHÔNG thêm giải thích hay comment.

Markdown output:"""

            response = model.generate_content([uploaded_file, prompt])
            markdown = response.text.strip()
            
            # Calculate usage
            usage = response.usage_metadata
            total_tokens = usage.prompt_token_count + usage.candidates_token_count
            cost = (usage.prompt_token_count / 1_000_000 * 0.075) + (usage.candidates_token_count / 1_000_000 * 0.30)
            
            # Cleanup
            try:
                genai.delete_file(uploaded_file.name)
            except:
                pass
            
            return {
                "markdown": markdown,
                "char_count": len(markdown),
                "ai_usage": {
                    "engine": "gemini",
                    "model": "gemini-2.0-flash-exp",
                    "total_tokens": total_tokens,
                    "total_cost_usd": round(cost, 6)
                }
            }
            
        except Exception as e:
            logger.error(f"Gemini Markdown generation failed: {e}")
            raise HTTPException(500, f"Markdown generation failed: {str(e)}")
    
    async def _markdown_to_word(self, markdown: str, output_path: Path):
        """Convert Markdown to Word document using python-docx"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            import re
            
            doc = Document()
            
            # Parse Markdown line by line
            lines = markdown.split('\n')
            i = 0
            
            while i < len(lines):
                line = lines[i].rstrip()
                
                # Skip empty lines
                if not line:
                    i += 1
                    continue
                
                # Heading 1
                if line.startswith('# ') and not line.startswith('##'):
                    text = line[2:].strip()
                    heading = doc.add_heading(text, level=1)
                    i += 1
                    
                # Heading 2
                elif line.startswith('## ') and not line.startswith('###'):
                    text = line[3:].strip()
                    heading = doc.add_heading(text, level=2)
                    i += 1
                    
                # Heading 3
                elif line.startswith('### '):
                    text = line[4:].strip()
                    heading = doc.add_heading(text, level=3)
                    i += 1
                    
                # Table
                elif line.startswith('|'):
                    table_lines = []
                    while i < len(lines) and lines[i].strip().startswith('|'):
                        table_lines.append(lines[i].strip())
                        i += 1
                    
                    # Parse table
                    if len(table_lines) >= 2:
                        # Extract header
                        header_cells = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
                        
                        # Skip separator line
                        data_lines = table_lines[2:] if len(table_lines) > 2 else []
                        
                        # Create table
                        if data_lines:
                            table = doc.add_table(rows=1 + len(data_lines), cols=len(header_cells))
                            table.style = 'Light Grid Accent 1'
                            
                            # Add header
                            for col_idx, header_text in enumerate(header_cells):
                                cell = table.rows[0].cells[col_idx]
                                cell.text = header_text
                                # Bold header
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                            
                            # Add data rows
                            for row_idx, data_line in enumerate(data_lines, start=1):
                                cells_data = [cell.strip() for cell in data_line.split('|')[1:-1]]
                                for col_idx, cell_data in enumerate(cells_data):
                                    if col_idx < len(header_cells):
                                        table.rows[row_idx].cells[col_idx].text = cell_data
                    
                # List (unordered)
                elif line.startswith('- ') or line.startswith('* '):
                    text = line[2:].strip()
                    p = doc.add_paragraph(style='List Bullet')
                    self._add_formatted_text(p, text)
                    i += 1
                    
                # List (ordered)
                elif re.match(r'^\d+\.\s', line):
                    text = re.sub(r'^\d+\.\s', '', line).strip()
                    p = doc.add_paragraph(style='List Number')
                    self._add_formatted_text(p, text)
                    i += 1
                    
                # Regular paragraph
                else:
                    p = doc.add_paragraph()
                    self._add_formatted_text(p, line)
                    i += 1
            
            # Save document
            doc.save(str(output_path))
            logger.info(f"✅ Word document saved: {output_path}")
            
        except Exception as e:
            logger.error(f"Markdown to Word conversion failed: {e}")
            raise HTTPException(500, f"Word generation failed: {str(e)}")
    
    def _add_formatted_text(self, paragraph, text: str):
        """Add text to paragraph with Markdown formatting (bold, italic)"""
        import re
        
        # Pattern to match: ***text*** (bold+italic), **text** (bold), *text* (italic)
        pattern = r'(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*)'
        
        parts = re.split(pattern, text)
        
        for part in parts:
            if not part:
                continue
                
            # Bold + Italic
            if part.startswith('***') and part.endswith('***'):
                run = paragraph.add_run(part[3:-3])
                run.bold = True
                run.italic = True
                
            # Bold
            elif part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
                
            # Italic
            elif part.startswith('*') and part.endswith('*'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
                
            # Normal text
            else:
                paragraph.add_run(part)
    
    async def _track_pdf_to_word_usage(self, ai_usage: dict, processing_time: float):
        """Track PDF to Word AI usage"""
        try:
            from app.services.ai_usage_service import log_usage
            from app.core.database import SessionLocal
            
            if not ai_usage:
                return
            
            total_tokens = ai_usage.get("total_tokens", 0)
            input_tokens = int(total_tokens * 0.7)
            output_tokens = int(total_tokens * 0.3)
            
            db = SessionLocal()
            try:
                log_usage(
                    db=db,
                    provider="gemini",
                    model=ai_usage.get("model", "gemini-2.0-flash-exp"),
                    endpoint="PDF to Word",
                    input_tokens=input_tokens,
                    output_tokens=output_tokens,
                    processing_time=processing_time,
                    status="success",
                    request_metadata={
                        "cost_usd": ai_usage.get("total_cost_usd", 0.0)
                    }
                )
                db.commit()
            finally:
                db.close()
                
        except Exception as e:
            logger.error(f"Failed to track PDF to Word usage: {e}")
