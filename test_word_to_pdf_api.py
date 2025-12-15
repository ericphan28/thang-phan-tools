"""
Test Word to PDF conversion API
"""

import asyncio
import requests
from pathlib import Path
from docx import Document

async def test_word_to_pdf_api():
    """Test Word to PDF conversion via API"""
    
    # Create test Word document
    doc = Document()
    doc.add_heading("Test Document", 0)
    doc.add_paragraph("Đây là tài liệu tiếng Việt để test chuyển đổi Word sang PDF.")
    doc.add_paragraph("This is a test document for Word to PDF conversion.")
    
    test_docx = Path("test_word.docx")
    doc.save(str(test_docx))
    print(f"✅ Created test Word document: {test_docx}")
    
    try:
        # Test API endpoint
        with open(test_docx, "rb") as f:
            files = {"file": (test_docx.name, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            
            response = requests.post(
                "http://localhost:8000/api/v1/documents/convert/word-to-pdf",
                files=files,
                timeout=30
            )
        
        print(f"📡 API Response Status: {response.status_code}")
        
        if response.status_code == 200:
            # Save PDF response
            output_pdf = Path("test_output.pdf")
            with open(output_pdf, "wb") as f:
                f.write(response.content)
            
            file_size = output_pdf.stat().st_size
            print(f"✅ PDF created successfully: {output_pdf} ({file_size} bytes)")
            
            # Clean up
            output_pdf.unlink()
        else:
            print(f"❌ API Error: {response.status_code}")
            print(f"Response: {response.text}")
            
    except requests.exceptions.ConnectionError:
        print("❌ Connection Error - Backend server not running?")
    except Exception as e:
        print(f"❌ Error: {e}")
    finally:
        # Clean up test file
        if test_docx.exists():
            test_docx.unlink()
            print(f"🧹 Cleaned up: {test_docx}")

if __name__ == "__main__":
    asyncio.run(test_word_to_pdf_api())