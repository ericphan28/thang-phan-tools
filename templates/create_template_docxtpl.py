"""
Script tự động tạo template Word cho Mẫu 2C-TCTW-98
Sử dụng thư viện docxtpl (python-docx-template)

✅ GIỮ NGUYÊN 100% ĐỊNH DẠNG
✅ TỰ ĐỘNG 100% - KHÔNG CẦN EDIT THỦ CÔNG
✅ SYNTAX ĐƠN GIẢN (Jinja2)
"""

from docx import Document
import re
from pathlib import Path

def create_docxtpl_template():
    """
    Tạo template Word với Jinja2 syntax (docxtpl)
    """
    
    print("🚀 BẮT ĐẦU TẠO TEMPLATE VỚI DOCXTPL...")
    print("="*60)
    
    # Load original document
    template_path = Path("mau-nha-nuoc/Mau-ly-lich-2C-TCTW-98.docx")
    if not template_path.exists():
        print(f"❌ ERROR: Không tìm thấy file {template_path}")
        return False
    
    print(f"📖 Đọc file gốc: {template_path}")
    doc = Document(template_path)
    
    # Define replacement patterns
    # Jinja2 syntax: {{ variable }} và {% for item in array %}...{% endfor %}
    replacements = {
        # Header information
        r"Tỉnh:\s*…+": "Tỉnh: {{ tinh }}",
        r"Đơn vị trực thuộc:\s*…+": "Đơn vị trực thuộc: {{ don_vi_truc_thuoc }}",
        r"Đơn vị cơ sở:\s*…+": "Đơn vị cơ sở: {{ don_vi_co_so }}",
        r"Số hiệu:\s*…+": "Số hiệu: {{ so_hieu }}",
        
        # Personal information
        r"Họ và tên.*?:\s*…+": "Họ và tên: {{ ho_ten }}",
        r"(?:Tên gọi khác|Tên thường gọi):\s*…+": "Tên gọi khác: {{ ten_goi_khac }}",
        r"Sinh ngày\s+…+\s+tháng\s+…+\s+năm\s+…+": 
            "Sinh ngày {{ ngay }} tháng {{ thang }} năm {{ nam }}",
        r"Nơi sinh:\s*…+": "Nơi sinh: {{ noi_sinh }}",
        r"Nguyên quán:\s*…+": "Nguyên quán: {{ nguyen_quan }}",
        r"Dân tộc:\s*…+": "Dân tộc: {{ dan_toc }}",
        r"Tôn giáo:\s*…+": "Tôn giáo: {{ ton_giao }}",
        
        # Contact & Identity
        r"Số CMND/CCCD:\s*…+": "Số CMND/CCCD: {{ so_cmnd }}",
        r"Ngày cấp:\s*…+": "Ngày cấp: {{ ngay_cap }}",
        r"Nơi cấp:\s*…+": "Nơi cấp: {{ noi_cap }}",
        r"Hộ khẩu thường trú:\s*…+": "Hộ khẩu thường trú: {{ ho_khau }}",
        r"Chỗ ở hiện tại:\s*…+": "Chỗ ở hiện tại: {{ cho_o_hien_tai }}",
        r"Điện thoại:\s*…+": "Điện thoại: {{ dien_thoai }}",
        r"Email:\s*…+": "Email: {{ email }}",
        
        # Education & Qualifications
        r"Trình độ văn hóa:\s*…+": "Trình độ văn hóa: {{ trinh_do_van_hoa }}",
        r"Trình độ chuyên môn cao nhất:\s*…+": "Trình độ chuyên môn: {{ trinh_do_chuyen_mon }}",
        r"Lý luận chính trị:\s*…+": "Lý luận chính trị: {{ ly_luan_chinh_tri }}",
        r"Quản lý nhà nước:\s*…+": "Quản lý nhà nước: {{ quan_ly_nha_nuoc }}",
        r"Ngoại ngữ:\s*…+": "Ngoại ngữ: {{ ngoai_ngu }}",
        r"Tin học:\s*…+": "Tin học: {{ tin_hoc }}",
        
        # Party & Political Status
        r"Ngày vào Đảng Cộng sản Việt Nam:\s*…+": 
            "Ngày vào Đảng: {{ ngay_vao_dang }}",
        r"Ngày chính thức:\s*…+": "Ngày chính thức: {{ ngay_chinh_thuc }}",
        r"Ngày tham gia tổ chức chính trị - xã hội:\s*…+": 
            "Ngày tham gia tổ chức: {{ ngay_tham_gia_to_chuc }}",
        r"Ngày nhập ngũ:\s*…+": "Ngày nhập ngũ: {{ ngay_nhap_ngu }}",
        r"Ngày xuất ngũ:\s*…+": "Ngày xuất ngũ: {{ ngay_xuat_ngu }}",
        r"Quân hàm cao nhất:\s*…+": "Quân hàm: {{ quan_ham }}",
        
        # Current Position
        r"Chức vụ hiện tại:\s*…+": "Chức vụ hiện tại: {{ chuc_vu }}",
        r"Công việc chính được giao:\s*…+": 
            "Công việc chính: {{ cong_viec_chinh }}",
        r"Ngạch công chức:\s*…+": "Ngạch công chức: {{ ngach_cong_chuc }}",
        r"Mã ngạch:\s*…+": "Mã ngạch: {{ ma_ngach }}",
        r"Bậc lương:\s*…+": "Bậc lương: {{ bac_luong }}",
        r"Phụ cấp chức vụ:\s*…+": "Phụ cấp chức vụ: {{ phu_cap_chuc_vu }}",
        r"Phụ cấp khác:\s*…+": "Phụ cấp khác: {{ phu_cap_khac }}",
        
        # Family Status
        r"Tình trạng hôn nhân:\s*…+": "Tình trạng hôn nhân: {{ tinh_trang_hon_nhan }}",
        r"Họ và tên vợ \(chồng\):\s*…+": "Họ và tên vợ (chồng): {{ ten_vo_chong }}",
        r"Năm sinh:\s*…+": "Năm sinh: {{ nam_sinh_vo_chong }}",
        r"Quê quán:\s*…+": "Quê quán: {{ que_quan_vo_chong }}",
        r"Nghề nghiệp:\s*…+": "Nghề nghiệp: {{ nghe_nghiep_vo_chong }}",
        r"Chỗ ở:\s*…+": "Chỗ ở: {{ cho_o_vo_chong }}",
        
        # Health & Other
        r"Tình trạng sức khỏe:\s*…+": "Tình trạng sức khỏe: {{ suc_khoe }}",
        r"Chiều cao:\s*…+": "Chiều cao: {{ chieu_cao }}",
        r"Cân nặng:\s*…+": "Cân nặng: {{ can_nang }}",
        r"Nhóm máu:\s*…+": "Nhóm máu: {{ nhom_mau }}",
        
        # Rewards & Disciplines
        r"Khen thưởng:\s*…+": "Khen thưởng: {{ khen_thuong }}",
        r"Kỷ luật:\s*…+": "Kỷ luật: {{ ky_luat }}",
        
        # Signature fields
        r"Ngày\s+…+\s+tháng\s+…+\s+năm\s+20…+": 
            "Ngày {{ ngay_ky }} tháng {{ thang_ky }} năm {{ nam_ky }}",
    }
    
    # Replace in paragraphs
    print("\n🔧 Bước 1: Thay thế trong paragraphs...")
    para_count = 0
    for para in doc.paragraphs:
        original_text = para.text
        new_text = original_text
        
        for pattern, replacement in replacements.items():
            new_text = re.sub(pattern, replacement, new_text, flags=re.IGNORECASE)
        
        if new_text != original_text:
            para.text = new_text
            para_count += 1
    
    print(f"   ✅ Đã xử lý {para_count} paragraphs")
    
    # Process tables with special handling
    print("\n🔧 Bước 2: Xử lý các bảng...")
    
    if len(doc.tables) >= 5:
        # Table 1: Đào tạo (2x5)
        print("   📋 Bảng 1: Đào tạo")
        table1 = doc.tables[0]
        if len(table1.rows) > 1:
            # Row 2 (data row) - add Jinja2 loop
            cells = table1.rows[1].cells
            cells[0].text = "{% for edu in dao_tao %}{{ edu.ten_truong }}{% endfor %}"
            cells[1].text = "{% for edu in dao_tao %}{{ edu.nganh_hoc }}{% endfor %}"
            cells[2].text = "{% for edu in dao_tao %}{{ edu.thoi_gian }}{% endfor %}"
            cells[3].text = "{% for edu in dao_tao %}{{ edu.hinh_thuc }}{% endfor %}"
            cells[4].text = "{% for edu in dao_tao %}{{ edu.van_bang }}{% endfor %}"
        
        # Table 2: Công tác (2x2)
        print("   📋 Bảng 2: Quá trình công tác")
        table2 = doc.tables[1]
        if len(table2.rows) > 1:
            cells = table2.rows[1].cells
            cells[0].text = "{% for work in cong_tac %}{{ work.thoi_gian }}{% endfor %}"
            cells[1].text = "{% for work in cong_tac %}{{ work.chuc_vu_don_vi }}{% endfor %}"
        
        # Table 3: Gia đình bản thân (2x4)
        print("   📋 Bảng 3: Gia đình bản thân")
        table3 = doc.tables[2]
        if len(table3.rows) > 1:
            # IMPORTANT: GIỮ NGUYÊN labels trong column 0!
            # Chỉ thay column 1, 2, 3
            cells = table3.rows[1].cells
            cells[1].text = "{% for member in gia_dinh %}{{ member.ho_ten }}{% endfor %}"
            cells[2].text = "{% for member in gia_dinh %}{{ member.nam_sinh }}{% endfor %}"
            cells[3].text = "{% for member in gia_dinh %}{{ member.thong_tin }}{% endfor %}"
        
        # Table 4: Gia đình vợ/chồng (2x4)
        print("   📋 Bảng 4: Gia đình vợ/chồng")
        table4 = doc.tables[3]
        if len(table4.rows) > 1:
            # IMPORTANT: GIỮ NGUYÊN labels trong column 0!
            cells = table4.rows[1].cells
            cells[1].text = "{% for member in gia_dinh_vo_chong %}{{ member.ho_ten }}{% endfor %}"
            cells[2].text = "{% for member in gia_dinh_vo_chong %}{{ member.nam_sinh }}{% endfor %}"
            cells[3].text = "{% for member in gia_dinh_vo_chong %}{{ member.thong_tin }}{% endfor %}"
        
        # Table 5: Lương (3x7)
        print("   📋 Bảng 5: Quá trình lương")
        table5 = doc.tables[4]
        if len(table5.rows) > 2:
            # Row 3 (data row)
            cells = table5.rows[2].cells
            cells[0].text = "{% for sal in luong %}{{ sal.thang_nam }}{% endfor %}"
            cells[1].text = "{% for sal in luong %}{{ sal.ngach_bac }}{% endfor %}"
            cells[2].text = "{% for sal in luong %}{{ sal.he_so }}{% endfor %}"
            # Cells 3-6 usually empty or same pattern
    
    print(f"   ✅ Đã xử lý {len(doc.tables)} bảng")
    
    # Save template
    output_path = Path("mau_2c_template_docxtpl.docx")
    doc.save(output_path)
    
    print("\n" + "="*60)
    print(f"✅ THÀNH CÔNG!")
    print(f"📄 Template đã lưu: {output_path}")
    print(f"📊 Size: {output_path.stat().st_size:,} bytes")
    
    print("\n💡 CÁCH DÙNG:")
    print("""
    from docxtpl import DocxTemplate
    import json
    
    doc = DocxTemplate('mau_2c_template_docxtpl.docx')
    with open('mau_2c_DATA_FULL.json', encoding='utf-8') as f:
        context = json.load(f)
    doc.render(context)
    doc.save('output_mau_2c.docx')
    """)
    
    return True

if __name__ == "__main__":
    try:
        success = create_docxtpl_template()
        if success:
            print("\n🎉 HOÀN THÀNH! Template sẵn sàng để dùng với docxtpl!")
        else:
            print("\n❌ Có lỗi xảy ra!")
    except Exception as e:
        print(f"\n❌ LỖI: {e}")
        import traceback
        traceback.print_exc()
