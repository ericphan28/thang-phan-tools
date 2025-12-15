"""
TỰ ĐỘNG TẠO MAILMERGE TEMPLATE
================================
Thêm MergeField vào Word document bằng cách thay thế text patterns
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

def add_merge_field(paragraph, field_name):
    """Thêm MergeField vào paragraph"""
    run = paragraph.add_run()
    
    # Create field character elements
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' MERGEFIELD  {field_name}  \\* MERGEFORMAT '
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    # Placeholder text
    run_text = OxmlElement('w:t')
    run_text.text = f'«{field_name}»'
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    # Add all elements to run
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(run_text)
    run._r.append(fldChar3)
    
    return run

def smart_replace_with_mergefield(paragraph, pattern, field_name):
    """
    Tìm pattern trong paragraph và thay bằng MergeField
    Giữ format của text xung quanh
    """
    full_text = paragraph.text
    
    if not re.search(pattern, full_text):
        return False
    
    # Find match position
    match = re.search(pattern, full_text)
    if not match:
        return False
    
    start_pos = match.start()
    end_pos = match.end()
    
    # Clear paragraph but keep format
    runs_before = []
    runs_after = []
    
    current_pos = 0
    for run in paragraph.runs:
        run_end = current_pos + len(run.text)
        
        if run_end <= start_pos:
            # Run is completely before match
            runs_before.append((run.text, run))
        elif current_pos >= end_pos:
            # Run is completely after match
            runs_after.append((run.text, run))
        elif current_pos < start_pos and run_end > start_pos:
            # Run contains start of match
            runs_before.append((run.text[:start_pos - current_pos], run))
            if run_end > end_pos:
                runs_after.append((run.text[end_pos - current_pos:], run))
        elif current_pos < end_pos and run_end > end_pos:
            # Run contains end of match
            runs_after.append((run.text[end_pos - current_pos:], run))
        
        current_pos = run_end
    
    # Clear all runs
    for run in paragraph.runs:
        run.text = ''
    
    # Re-add before text
    for text, original_run in runs_before:
        paragraph.runs[0].text += text
    
    # Add merge field
    add_merge_field(paragraph, field_name)
    
    # Add after text
    for text, original_run in runs_after:
        paragraph.add_run(text)
    
    return True

# Field patterns - comprehensive list
FIELD_PATTERNS = [
    # Header info
    (r"Tỉnh:\s*[\.…\s]+", "tinh"),
    (r"Đơn vị trực thuộc:\s*[\.…\s]+", "don_vi_truc_thuoc"),
    (r"Đơn vị cơ sở:\s*[\.…\s]+", "don_vi_co_so"),
    
    # Personal info
    (r"Họ và tên khai sinh:\s*[\.…\s]+", "ho_ten"),
    (r"Nam, nữ:\s*[\.…\s]+", "gioi_tinh"),
    (r"Sinh ngày:\s*[\.…]{2,}\s*tháng:\s*[\.…]{2,}\s*năm:\s*[\.…]{2,}", 
     "ngay_thang_nam_sinh"),
    (r"Các tên gọi khác:\s*[\.…\s]+", "ten_goi_khac"),
    (r"Nơi sinh:\s*[\.…\s]+", "noi_sinh"),
    (r"Quê quán:\s*[\.…\s]+", "que_quan"),
    (r"Dân tộc:\s*[\.…\s]+", "dan_toc"),
    (r"Tôn giáo:\s*[\.…\s]+", "ton_giao"),
    
    # Contact
    (r"Nơi đăng ký hộ khẩu thường trú:\s*[\.…\s]+", "ho_khau_thuong_tru"),
    (r"Nơi ở hiện nay:\s*[\.…\s]+", "noi_o_hien_nay"),
    (r"Điện thoại:\s*[\.…\s]+", "dien_thoai"),
    (r"Email:\s*[\.…\s]+", "email"),
    
    # Party info
    (r"Ngày vào Đảng Cộng sản Việt Nam:\s*[\.…\s]+", "ngay_vao_dang"),
    (r"Ngày chính thức:\s*[\.…\s]+", "ngay_chinh_thuc"),
    (r"Ngày vào Đảng cộng sản Việt Nam:\s*[\.…\s]+", "ngay_vao_dang_cu"),
    (r"Ngày tham gia tổ chức chính trị - xã hội:\s*[\.…\s]+", "ngay_tham_gia_to_chuc"),
    (r"Ngày nhập ngũ:\s*[\.…\s]+", "ngay_nhap_ngu"),
    (r"Ngày xuất ngũ:\s*[\.…\s]+", "ngay_xuat_ngu"),
    (r"Quân hàm cao nhất:\s*[\.…\s]+", "quan_ham"),
    
    # Education
    (r"Trình độ giáo dục phổ thông:\s*[\.…\s]+", "trinh_do_giao_duc"),
    (r"Trình độ chuyên môn cao nhất:\s*[\.…\s]+", "trinh_do_chuyen_mon"),
    (r"Học hàm, học vị cao nhất:\s*[\.…\s]+", "hoc_ham_hoc_vi"),
    (r"Lý luận chính trị:\s*[\.…\s]+", "ly_luan_chinh_tri"),
    (r"Ngoại ngữ:\s*[\.…\s]+", "ngoai_ngu"),
    (r"Dân tộc thiểu số:\s*[\.…\s]+", "dan_toc_thieu_so"),
    (r"Trình độ tin học:\s*[\.…\s]+", "trinh_do_tin_hoc"),
    
    # Current position
    (r"Cấp ủy hiện tại:\s*[\.…\s]+", "cap_uy_hien_tai"),
    (r"Cấp ủy kiêm:\s*[\.…\s]+", "cap_uy_kiem"),
    (r"Chức vụ \(Đảng, đoàn thể, Chính quyền.*?\):\s*[\.…\s]+", "chuc_vu"),
    (r"Phụ cấp chức vụ:\s*[\.…\s]+", "phu_cap_chuc_vu"),
    (r"Phụ cấp khác:\s*[\.…\s]+", "phu_cap_khac"),
    (r"Ngạch, bậc, lương:\s*[\.…\s]+", "ngach_bac_luong"),
    (r"Ngày bổ nhiệm, bầu cử, phê chuẩn:\s*[\.…\s]+", "ngay_bo_nhiem"),
    
    # Career history
    (r"Khen thưởng:\s*[\.…\s]+", "khen_thuong"),
    (r"Kỷ luật:\s*[\.…\s]+", "ky_luat"),
    
    # Summary
    (r"Đặc điểm lịch sử bản thân:\s*[\.…\s]+", "dac_diem_lich_su"),
]

print("🚀 TẠO MAILMERGE TEMPLATE TỰ ĐỘNG")
print("=" * 60)

# Load original document
print("\n📖 Đang load file gốc...")
doc = Document('mau-nha-nuoc/Mau-ly-lich-2C-TCTW-98.docx')
print(f"   ✅ Loaded: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

# Process paragraphs
print("\n🔧 Đang xử lý paragraphs...")
replaced_count = 0
for i, para in enumerate(doc.paragraphs):
    for pattern, field_name in FIELD_PATTERNS:
        if smart_replace_with_mergefield(para, pattern, field_name):
            print(f"   ✅ P{i}: Replaced with <<{field_name}>>")
            replaced_count += 1

# Process tables
print("\n🔧 Đang xử lý tables...")
for table_idx, table in enumerate(doc.tables):
    print(f"   📋 Table {table_idx + 1}:")
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                for pattern, field_name in FIELD_PATTERNS:
                    if smart_replace_with_mergefield(para, pattern, field_name):
                        print(f"      ✅ Replaced with <<{field_name}>>")
                        replaced_count += 1

# Save
output_file = 'mau_2c_MAILMERGE_TEMPLATE.docx'
print(f"\n💾 Đang save...")
doc.save(output_file)

print("\n" + "=" * 60)
print("✅ HOÀN THÀNH!")
print(f"📄 File: {output_file}")
print(f"🔄 Số field được thay thế: {replaced_count}")
print("\n💡 BƯỚC KẾ TIẾP:")
print("   1. Mở file template trong Word")
print("   2. Kiểm tra các MergeField (bấm Alt+F9 để toggle field codes)")
print("   3. Chạy test_mailmerge.py để render với data thật")
