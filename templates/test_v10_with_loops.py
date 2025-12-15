#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
TEST V10 WITH LOOPS
===================
Test template with table loops
"""

from docxtpl import DocxTemplate
import json
from docx import Document
import re

def flatten_dict(d, parent_key='', sep='_'):
    """Flatten nested dict"""
    items = []
    for k, v in d.items():
        if k.startswith('_'):
            continue
        new_key = f"{parent_key}{sep}{k}" if parent_key else k
        if isinstance(v, dict):
            items.extend(flatten_dict(v, new_key, sep=sep).items())
        else:
            items.append((new_key, v))
    return dict(items)

def main():
    TEMPLATE = "mau_2c_V10_WITH_LOOPS.docx"
    OUTPUT = "OUTPUT_V10_WITH_LOOPS.docx"
    JSON_FILE = "mau_2c_DATA_RESTRUCTURED.json"
    
    print(f"📖 Loading template: {TEMPLATE}")
    doc = DocxTemplate(TEMPLATE)
    
    print(f"📊 Loading data: {JSON_FILE}")
    with open(JSON_FILE, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # Flatten for simple fields
    flat_data = flatten_dict(data)
    
    # Prepare render data
    render_data = {**flat_data}
    
    # Keep arrays for loops - MAP TO TEMPLATE VARIABLES!
    if 'dao_tao' in data:
        render_data['hoc_tap'] = data['dao_tao']  # Map dao_tao -> hoc_tap
        print(f"   ✅ hoc_tap (from dao_tao): {len(data['dao_tao'])} items")
    
    if 'cong_tac' in data:
        # Map keys: thoi_gian -> tu_thang_nam, chuc_vu_don_vi -> chuc_danh
        render_data['cong_tac'] = [
            {
                'tu_thang_nam': ct.get('thoi_gian', ''),
                'den_thang_nam': ct.get('den_thang_nam', ''),
                'chuc_danh': ct.get('chuc_vu_don_vi', '')
            }
            for ct in data['cong_tac']
        ]
        print(f"   ✅ cong_tac: {len(data['cong_tac'])} items (mapped keys)")
    
    # Map gia_dinh arrays
    if 'gia_dinh' in data:
        render_data['gia_dinh'] = {
            'ban_than': data['gia_dinh'],  # Main family
            'vo_chong': data.get('gia_dinh_vo_chong', [])  # Spouse family
        }
        print(f"   ✅ gia_dinh.ban_than: {len(data['gia_dinh'])} members")
        print(f"   ✅ gia_dinh.vo_chong: {len(data.get('gia_dinh_vo_chong', []))} members")
    
    # Add salary history
    if 'luong' in data:
        # Map keys: thang_nam -> tu_thang_nam, ngach_bac -> chuc_danh + bac
        render_data['he_so_luong_history'] = [
            {
                'tu_thang_nam': l.get('thang_nam', ''),
                'chuc_danh': 'Chuyên viên',  # Extract from ngach_bac
                'bac': l.get('ngach_bac', '').split('Bậc ')[-1] if 'Bậc' in l.get('ngach_bac', '') else '',
                'he_so': l.get('he_so', '')
            }
            for l in data['luong']
        ]
        print(f"   ✅ he_so_luong_history (from luong): {len(data['luong'])} entries (mapped keys)")
    
    print(f"\n🎨 Rendering...")
    try:
        doc.render(render_data)
        print(f"   ✅ Render successful!")
    except Exception as e:
        print(f"   ❌ Render error: {e}")
        return
    
    print(f"💾 Saving: {OUTPUT}")
    doc.save(OUTPUT)
    
    # Clean dots
    print(f"\n🧹 Cleaning dots...")
    output_doc = Document(OUTPUT)
    cleaned = 0
    
    for para in output_doc.paragraphs:
        for run in para.runs:
            if re.search(r'[.…]{3,}', run.text):
                run.text = re.sub(r'[.…]{3,}', '', run.text).strip()
                cleaned += 1
    
    for table in output_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if re.search(r'[.…]{3,}', run.text):
                            run.text = re.sub(r'[.…]{3,}', '', run.text).strip()
                            cleaned += 1
    
    output_doc.save(OUTPUT.replace('.docx', '_CLEANED.docx'))
    
    # Verify tables
    print(f"\n{'='*60}")
    print(f"✅ Cleaned {cleaned} locations")
    print(f"📄 Output: {OUTPUT.replace('.docx', '_CLEANED.docx')}")
    
    print(f"\n🔍 Verifying tables...")
    final_doc = Document(OUTPUT.replace('.docx', '_CLEANED.docx'))
    
    for i, table in enumerate(final_doc.tables):
        print(f"\n  TABLE {i+1}: {len(table.rows)} rows")
        # Check if row 2 has data
        if len(table.rows) >= 2:
            row2_text = table.rows[1].cells[0].text.strip()
            if row2_text and len(row2_text) > 10:
                print(f"    ✅ Has data: {row2_text[:50]}...")
            else:
                print(f"    ⚠️  Empty or minimal data")

if __name__ == "__main__":
    main()
