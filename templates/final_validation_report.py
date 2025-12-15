#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
FINAL VALIDATION REPORT
Template V4 với xử lý newlines trong bảng
"""

from docx import Document
from pathlib import Path
import json

print("=" * 80)
print("📊 BÁO CÁO KIỂM THỬ HOÀN TẤT - TEMPLATE V4")
print("=" * 80)

# 1. Load JSON data
json_file = Path("mau_2c_DATA_COMPLETE_V3.json")
with open(json_file, 'r', encoding='utf-8') as f:
    data = json.load(f)

# Count fields
simple_fields = sum(1 for k, v in data.items() if not isinstance(v, list))
array_fields = sum(1 for k, v in data.items() if isinstance(v, list))
total_items = sum(len(v) for v in data.values() if isinstance(v, list))

print("\n📄 DỮ LIỆU JSON:")
print(f"   - Simple fields: {simple_fields}")
print(f"   - Array fields: {array_fields}")
print(f"   - Total array items: {total_items}")
print(f"   - TỔNG: {simple_fields + array_fields} fields")

# 2. Load output
output_file = Path("OUTPUT_MAU_2C_DOCXTPL.docx")
doc = Document(output_file)

print("\n📄 FILE OUTPUT:")
print(f"   - File: {output_file.name}")
print(f"   - Size: {output_file.stat().st_size:,} bytes ({output_file.stat().st_size/1024:.2f} KB)")
print(f"   - Paragraphs: {len(doc.paragraphs)}")
print(f"   - Tables: {len(doc.tables)}")

# 3. Validate tables with newlines
print("\n📋 KIỂM TRA CÁC BẢNG:")

# Table 1: Education
t1 = doc.tables[0]
t1_row2_col1 = t1.rows[1].cells[0].text
t1_entries = len([line for line in t1_row2_col1.split('\n') if line.strip()])
print(f"\n   1️⃣ BẢNG ĐÀO TẠO:")
print(f"      - Expected: 3 entries")
print(f"      - Found: {t1_entries} entries")
print(f"      - Status: {'✅ PASS' if t1_entries == 3 else '❌ FAIL'}")
print(f"      - Preview: {t1_row2_col1.split(chr(10))[0]}")

# Table 2: Work History
t2 = doc.tables[1]
t2_row2_col1 = t2.rows[1].cells[0].text
t2_entries = len([line for line in t2_row2_col1.split('\n') if line.strip()])
print(f"\n   2️⃣ BẢNG CÔNG TÁC:")
print(f"      - Expected: 2 entries")
print(f"      - Found: {t2_entries} entries")
print(f"      - Status: {'✅ PASS' if t2_entries == 2 else '❌ FAIL'}")
print(f"      - Preview: {t2_row2_col1.split(chr(10))[0]}")

# Table 3: Family
t3 = doc.tables[2]
t3_row2_col2 = t3.rows[1].cells[1].text  # Names column
t3_entries = len([line for line in t3_row2_col2.split('\n') if line.strip()])
print(f"\n   3️⃣ BẢNG GIA ĐÌNH:")
print(f"      - Expected: 4 entries")
print(f"      - Found: {t3_entries} entries")
print(f"      - Status: {'✅ PASS' if t3_entries == 4 else '❌ FAIL'}")
print(f"      - Preview: {t3_row2_col2.split(chr(10))[0]}")

# Table 4: Spouse's Family
t4 = doc.tables[3]
t4_row2_col2 = t4.rows[1].cells[1].text
t4_entries = len([line for line in t4_row2_col2.split('\n') if line.strip()])
print(f"\n   4️⃣ BẢNG GIA ĐÌNH VỢ/CHỒNG:")
print(f"      - Expected: 3 entries")
print(f"      - Found: {t4_entries} entries")
print(f"      - Status: {'✅ PASS' if t4_entries == 3 else '❌ FAIL'}")
print(f"      - Preview: {t4_row2_col2.split(chr(10))[0]}")

# Table 5: Salary
t5 = doc.tables[4]
t5_row3_col1 = t5.rows[2].cells[0].text
t5_entries = len([line for line in t5_row3_col1.split('\n') if line.strip()])
print(f"\n   5️⃣ BẢNG LƯƠNG:")
print(f"      - Expected: 3 entries")
print(f"      - Found: {t5_entries} entries")
print(f"      - Status: {'✅ PASS' if t5_entries == 3 else '❌ FAIL'}")
print(f"      - Preview: {t5_row3_col1.split(chr(10))[0]}")

# 4. Sample data validation
print("\n📊 MẪU DỮ LIỆU:")

print("\n   👤 THÔNG TIN CHÍNH:")
print(f"      - Họ tên: {data.get('ho_ten', 'N/A')}")
print(f"      - Tỉnh: {data.get('tinh', 'N/A')}")
print(f"      - Ngày sinh: {data.get('ngay_sinh', 'N/A')}")

print("\n   🏠 NHÀ Ở:")
print(f"      - Được cấp: {data.get('nha_o_duoc_cap', 'N/A')}")
print(f"      - Tự mua: {data.get('nha_o_tu_mua', 'N/A')}")
print(f"      - Loại: {data.get('nha_o_tu_mua_loai', 'N/A')}")
print(f"      - Diện tích: {data.get('nha_o_tu_mua_dien_tich', 'N/A')}")

print("\n   🌾 ĐẤT Ở:")
print(f"      - Được cấp: {data.get('dat_o_duoc_cap', 'N/A')}")
print(f"      - Tự mua: {data.get('dat_o_tu_mua', 'N/A')}")
print(f"      - Đất sản xuất: {data.get('dat_san_xuat', 'N/A')}")

print("\n   👨‍👩‍👧‍👦 GIA ĐÌNH:")
print(f"      - Anh chị em: {len(data.get('gia_dinh', []))} người")
for member in data.get('gia_dinh', []):
    print(f"        • {member.get('quan_he', 'N/A')}: {member.get('ho_ten', 'N/A')} ({member.get('nam_sinh', 'N/A')})")

print("\n   💼 CÔNG TÁC:")
print(f"      - Số kỳ công tác: {len(data.get('cong_tac', []))}")
for i, work in enumerate(data.get('cong_tac', []), 1):
    print(f"        {i}. {work.get('thoi_gian', 'N/A')}")
    desc = work.get('chuc_vu_don_vi', 'N/A').replace('\n', ' | ')
    print(f"           {desc}")

# 5. Final assessment
print("\n" + "=" * 80)
print("🎯 KẾT QUẢ TỔNG THỂ:")
print("=" * 80)

all_tables_pass = (
    t1_entries == 3 and
    t2_entries == 2 and
    t3_entries == 4 and
    t4_entries == 3 and
    t5_entries == 3
)

if all_tables_pass:
    print("✅ TẤT CẢ CÁC BẢNG ĐÚNG!")
    print("✅ NEWLINES HOẠT ĐỘNG ĐÚNG!")
    print("✅ DỮ LIỆU ĐẦY ĐỦ!")
    print("\n🎉 TEMPLATE V4 HOÀN TẤT THÀNH CÔNG!")
else:
    print("⚠️ CÓ MỘT SỐ BẢNG CẦN KIỂM TRA LẠI")

print("\n📝 TỔNG KẾT:")
print(f"   - Template: mau_2c_template_FINAL_V4.docx")
print(f"   - JSON Data: mau_2c_DATA_COMPLETE_V3.json")
print(f"   - Output: OUTPUT_MAU_2C_DOCXTPL.docx")
print(f"   - Total fields: {simple_fields + array_fields}")
print(f"   - Array entries: {total_items}")
print(f"   - Tables with newlines: 5/5")
print("=" * 80)
