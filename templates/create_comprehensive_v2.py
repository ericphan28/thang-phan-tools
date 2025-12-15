"""
VERSION 2 - COMPREHENSIVE PATTERNS
===================================
Thêm TẤT CẢ patterns còn thiếu
"""

from docx import Document
import re

def simple_replace_in_run(paragraph, pattern, jinja_var):
    """Replace trong run, giữ format"""
    replaced = False
    for run in paragraph.runs:
        if re.search(pattern, run.text):
            run.text = re.sub(pattern, f'{{{{ {jinja_var} }}}}', run.text)
            replaced = True
            break
    return replaced

# COMPREHENSIVE PATTERNS - ALL 100+ FIELDS
FIELD_PATTERNS = [
    # === HEADER (3) ===
    (r"Tỉnh:\s*[\.…]+", "tinh"),
    (r"Đơn vị trực thuộc:\s*[\.…]+", "don_vi_truc_thuoc"),
    (r"Đơn vị cơ sở:\s*[\.…]+", "don_vi_co_so"),
    (r"Số hiệu cán bộ, công chức", "Số hiệu: {{ so_hieu_can_bo }}"),
    
    # === PERSONAL INFO (25) ===
    (r"1\)\s+Họ và tên khai sinh:\s*[\.…]+", "1) Họ và tên khai sinh: {{ ho_ten }}"),
    (r"Nam, nữ:\s*[\.…]+", "Nam, nữ: {{ gioi_tinh }}"),
    (r"Sinh ngày:\s*[\.…]+\s*tháng\s*[\.…]+\s*năm\s*[\.…]+", "Sinh ngày: {{ ngay }} tháng {{ thang }} năm {{ nam }}"),
    (r"2\)\s*Các tên gọi khác:\s*[\.…]+", "2) Các tên gọi khác: {{ ten_goi_khac }}"),
    (r"3\)\s*Cấp ủy hiện tại:\s*[\.…]+", "3) Cấp ủy hiện tại: {{ cap_uy_hien_tai }}"),
    (r"Cấp ủy kiêm:\s*[\.…]+", "Cấp ủy kiêm: {{ cap_uy_kiem }}"),
    (r"Chức vụ \(Đảng, đoàn thể[^:]+:\s*[\.…]+", "Chức vụ (Đảng, đoàn thể, Chính quyền, kể cả chức vụ kiêm nhiệm): {{ chuc_vu }}"),
    (r"Phụ cấp chức vụ:\s*[\.…]+", "Phụ cấp chức vụ: {{ phu_cap_chuc_vu }}"),
    (r"Phụ cấp khác:\s*[\.…]+", "Phụ cấp khác: {{ phu_cap_khac }}"),
    (r"4\)\s*Sinh ngày", ""),  # Skip, already handled
    (r"5\)\s+([^\.]+?)\s*$", r"5) {{ noi_sinh }}"),  # Noi sinh at line end
    (r"6\)\s*Quê quán \(xã, phường\):\s*[\.…]+", "6) Quê quán (xã, phường): {{ que_quan_xa }}"),
    (r"\(huyện, quận\):\s*[\.…]+", "(huyện, quận): {{ que_quan_huyen }}"),
    (r"\(tỉnh, TP\):\s*[\.…]+", "(tỉnh, TP): {{ que_quan_tinh }}"),
    (r"7\)\s*Nơi ở hiện nay[^:]*:\s*[\.…]+", "7) Nơi ở hiện nay (Xã, huyện, tỉnh hoặc số nhà, đường phố, TP): {{ noi_o_hien_nay }}"),
    (r"đ/thoại:\s*[\.…]+", "đ/thoại: {{ dien_thoai }}"),
    (r"Email:\s*[\.…]+", "Email: {{ email }}"),
    (r"8\)\s*Dân tộc:\s*\([^)]+\):\s*[\.…]+", "8) Dân tộc: (Kinh, Tày, Mông, Ê đê...): {{ dan_toc }}"),
    (r"9\)\s*Tôn giáo:\s*[\.…]+", "9) Tôn giáo: {{ ton_giao }}"),
    (r"10\)\s*Thành phần gia đình xuất thân:\s*[\.…]+", "10) Thành phần gia đình xuất thân: {{ thanh_phan_gia_dinh }}"),
    (r"\(Ghi là công nhân[^)]+\)", "({{ ghi_chu_thanh_phan }})"),
    (r"11\)\s*Nghề nghiệp bản thân[^:]*:\s*[\.…]+", "11) Nghề nghiệp bản thân trước khi được tuyển dụng: {{ nghe_nghiep_ban_than }}"),
    (r"\(Ghi nghề được đào tạo[^)]+\)", "({{ ghi_chu_nghe_nghiep }})"),
    (r"12\)\s*Ngày được tuyển dụng:\s*[\.…/]+", "12) Ngày được tuyển dụng: {{ ngay_tuyen_dung }}"),
    (r"Vào cơ quan nào, ở đâu:\s*[\.…]+", "Vào cơ quan nào, ở đâu: {{ vao_co_quan }}"),
    
    # === PARTY & MILITARY (15) ===
    (r"13\)\s*Ngày vào cơ quan hiện đang công tác:\s*[\.…/]+", "13) Ngày vào cơ quan hiện đang công tác: {{ ngay_vao_co_quan_hien_tai }}"),
    (r"Ngày tham gia cách mạng:\s*[\.…/]+", "Ngày tham gia cách mạng: {{ ngay_tham_gia_cach_mang }}"),
    (r"14\)\s*Ngày vào Đảng Cộng sản Việt Nam:\s*[\.…/]+", "14) Ngày vào Đảng Cộng sản Việt Nam: {{ ngay_vao_dang }}"),
    (r"Ngày chính thức:\s*[\.…/]+", "Ngày chính thức: {{ ngay_chinh_thuc }}"),
    (r"15\)\s*Ngày vào Đảng cộng sản Việt Nam \(nếu có\):\s*[\.…/]+", "15) Ngày vào Đảng cộng sản Việt Nam (nếu có): {{ ngay_vao_dang_cu }}"),
    (r"Ngày tham gia tổ chức chính trị - xã hội:\s*[\.…/]+", "Ngày tham gia tổ chức chính trị - xã hội: {{ ngay_tham_gia_to_chuc }}"),
    (r"16\)\s*«ngay_nhap_ngu»\s*/\s*[\.…/]+", "16) {{ ngay_nhap_ngu }}"),
    (r"«ngay_xuat_ngu»\s*/\s*[\.…/]+", "{{ ngay_xuat_ngu }}"),
    (r"Quân hàm, chức vụ cao nhất[^:]*:\s*[\.…]+", "Quân hàm, chức vụ cao nhất (sĩ quan, quân nhân...): {{ quan_ham }}"),
    (r"17\)\s*Trình độ học vấn: Giáo dục phổ thông:\s*[\.…]+", "17) Trình độ học vấn: Giáo dục phổ thông: {{ trinh_do_giao_duc }}"),
    (r"Trình độ chuyên môn cao nhất:\s*[\.…]+", "Trình độ chuyên môn cao nhất: {{ trinh_do_chuyen_mon }}"),
    (r"«hoc_ham_hoc_vi»", "{{ hoc_ham_hoc_vi }}"),
    (r"«ly_luan_chinh_tri»", "{{ ly_luan_chinh_tri }}"),
    (r"«ngoai_ngu»", "{{ ngoai_ngu }}"),
    (r"Dân tộc thiểu số:\s*[\.…]+", "Dân tộc thiểu số: {{ dan_toc_thieu_so }}"),
    (r"Tin học:\s*[\.…]+", "Tin học: {{ tin_hoc }}"),
    
    # === CURRENT POSITION (10) ===
    (r"Ngạch, bậc, lương:\s*[\.…]+", "Ngạch, bậc, lương: {{ ngach_bac_luong }}"),
    (r"Ngày bổ nhiệm, bầu cử, phê chuẩn:\s*[\.…/]+", "Ngày bổ nhiệm, bầu cử, phê chuẩn: {{ ngay_bo_nhiem }}"),
    (r"18\)\s*Công tác chính đang làm:\s*[\.…]+", "18) Công tác chính đang làm: {{ cong_tac_chinh }}"),
    (r"19\)\s*Ngạch công chức:\s*[\.…]+", "19) Ngạch công chức: {{ ngach_cong_chuc }}"),
    (r"Bậc lương:\s*[\.…]+", "Bậc lương: {{ bac_luong }}"),
    (r"Hệ số:\s*[\.…]+", "Hệ số: {{ he_so }}"),
    (r"Từ tháng[^:]*:\s*[\.…]+", "Từ tháng ..., năm ...: {{ tu_thang }}"),
    (r"20\)\s*Danh hiệu được phong[^:]*:\s*[\.…]+", "20) Danh hiệu được phong (tặng): {{ danh_hieu }}"),
    (r"21\)\s*Sở trường công tác:\s*[\.…]+", "21) Sở trường công tác: {{ so_truong }}"),
    (r"22\)\s*«khen_thuong»", "22) {{ khen_thuong }}"),
    (r"23\)\s*Kỷ luật: Cảnh cáo[^:]*:\s*[\.…]+", "23) Kỷ luật: Cảnh cáo, khiển trách...: {{ ky_luat }}"),
    (r"24\)\s*Trình trạng sức khoẻ:\s*[\.…]+", "24) Trình trạng sức khoẻ: {{ trinh_trang_suc_khoe }}"),
    (r"Chiều cao:\s*[\.…]+", "Chiều cao: {{ chieu_cao }}"),
    (r"Cân nặng:\s*[\.…]+", "Cân nặng: {{ can_nang }}"),
    (r"Nhóm máu:\s*[\.…]+", "Nhóm máu: {{ nhom_mau }}"),
]

print("🚀 CREATE COMPREHENSIVE AUTO TEMPLATE V2")
print("=" * 70)

# Load
doc = Document('mau-nha-nuoc/Mau-ly-lich-2C-TCTW-98.docx')
print(f"✅ Loaded: {len(doc.paragraphs)} paragraphs")

# Process
replaced = 0
for i, para in enumerate(doc.paragraphs):
    for pattern, replacement in FIELD_PATTERNS:
        if simple_replace_in_run(para, pattern, replacement):
            # Extract var name from replacement
            var_match = re.search(r'\{\{\s*(\w+)\s*\}\}', replacement)
            var_name = var_match.group(1) if var_match else "unknown"
            print(f"✅ P{i}: {var_name}")
            replaced += 1
            break

# Save
output = 'mau_2c_COMPREHENSIVE_TEMPLATE.docx'
doc.save(output)

print(f"\n✅ DONE!")
print(f"📄 File: {output}")
print(f"🔄 Replaced: {replaced} fields")
