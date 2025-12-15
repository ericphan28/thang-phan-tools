#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Check table 3 original structure"""

from docx import Document
from pathlib import Path

# Check original file first
orig_file = Path("Mau-ly-lich-2C-TCTW-98.docx")
if orig_file.exists():
    print("=" * 80)
    print("📋 BẢNG 3 GIA ĐÌNH - FILE GỐC")
    print("=" * 80)
    
    doc = Document(orig_file)
    if len(doc.tables) > 2:
        t3 = doc.tables[2]
        print(f"\nRows: {len(t3.rows)}")
        print(f"Cols: {len(t3.rows[0].cells)}")
        
        for i in range(len(t3.rows)):
            print(f"\n{'='*70}")
            print(f"ROW {i}:")
            for j in range(min(4, len(t3.rows[i].cells))):
                text = t3.rows[i].cells[j].text.strip()
                if text:
                    print(f"  Col {j}: {text[:100]}")

print("\n" + "=" * 80)
print("📋 BẢNG 3 GIA ĐÌNH - TEMPLATE V4")
print("=" * 80)

template_file = Path("mau_2c_template_FINAL_V4.docx")
doc = Document(template_file)
t3 = doc.tables[2]

print(f"\nRows: {len(t3.rows)}")
print(f"Cols: {len(t3.rows[0].cells)}")

for i in range(len(t3.rows)):
    print(f"\n{'='*70}")
    print(f"ROW {i}:")
    for j in range(min(4, len(t3.rows[i].cells))):
        text = t3.rows[i].cells[j].text.strip()
        if text:
            print(f"  Col {j}: {text[:150]}")
