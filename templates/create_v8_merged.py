#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
V8 - MERGE RUNS + EXACT REPLACE
================================
Strategy: Merge all runs in paragraph để text liền, rồi mới replace
"""

from docx import Document
from copy import deepcopy

def merge_runs_in_paragraph(paragraph):
    """
    Merge all runs thành 1 run duy nhất
    Giữ format của run đầu tiên
    """
    if len(paragraph.runs) <= 1:
        return False
    
    # Get text from all runs
    full_text = ''.join([run.text for run in paragraph.runs])
    
    # Keep first run, copy its format
    first_run = paragraph.runs[0]
    
    # Delete all other runs
    for i in range(len(paragraph.runs) - 1, 0, -1):
        paragraph._element.remove(paragraph.runs[i]._element)
    
    # Set merged text to first run
    first_run.text = full_text
    
    return True

def replace_exact_in_paragraph(paragraph, old_text, new_text):
    """
    Replace EXACT text in paragraph (sau khi đã merge runs)
    """
    if old_text in paragraph.text:
        # Text exists - replace in runs
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                return True
    return False

def process_document():
    INPUT = "mau-nha-nuoc/Mau-ly-lich-2C-TCTW-98.docx"
    OUTPUT = "mau_2c_V8_MERGED_RUNS.docx"
    
    print(f"📖 Loading: {INPUT}")
    doc = Document(INPUT)
    
    print(f"🔄 Step 1: Merge runs...")
    merged_count = 0
    for para in doc.paragraphs:
        if merge_runs_in_paragraph(para):
            merged_count += 1
    print(f"   ✅ Merged {merged_count} paragraphs")
    
    # Also merge in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if merge_runs_in_paragraph(para):
                        merged_count += 1
    print(f"   ✅ Total merged: {merged_count}")
    
    print(f"\n🔄 Step 2: Exact replacements...")
    
    # EXACT patterns - sau khi merge thì text liền
    replacements = [
        ("Tỉnh: ………………….", "Tỉnh: {{ tinh }}"),
        ("Đơn vị trực thuộc: ..........................", "Đơn vị trực thuộc: {{ don_vi_truc_thuoc }}"),
        ("Đơn vị cơ sở: ................................", "Đơn vị cơ sở: {{ don_vi_co_so }}"),
        ("Số hiệu cán bộ, công chức", "Số hiệu: {{ so_hieu }}"),
        
        # Line 5
        ("1) Họ và tên khai sinh: ……………………………………..", "1) Họ và tên khai sinh: {{ ho_ten }}"),
        ("Nam, nữ: .....................", "Nam, nữ: {{ gioi_tinh }}"),
        
        # Line 6
        ("2) Các tên gọi khác:", "2) Các tên gọi khác: {{ ten_goi_khac }}"),
        
        # Line 7
        ("3) Cấp ủy hiện tại: .......................................", "3) Cấp ủy hiện tại: {{ cap_uy_hien_tai }}"),
        ("Cấp ủy kiêm: .........................................", "Cấp ủy kiêm: {{ cap_uy_kiem }}"),
        
        # Line 8
        ("Chức vụ (Đảng, đoàn thể, Chính quyền, kể cả chức vụ kiêm nhiệm):", 
         "Chức vụ: {{ chuc_vu_full }}"),
        
        # Line 9 - phụ cấp
        ("Phụ cấp chức vụ: ...........................", "Phụ cấp chức vụ: {{ phu_cap_chuc_vu }}"),
        
        # Line 10 - sinh ngày
        ("4) Sinh ngày: ..........", "4) Sinh ngày: {{ ngay }}"),
        ("tháng ..........", "tháng {{ thang }}"),
        ("năm ...............", "năm {{ nam }}"),
        ("5) Nơi sinh: ..................................................", "5) Nơi sinh: {{ noi_sinh }}"),
        
        # Line 11 - quê quán
        ("6) Quê quán (xã, phường): .......................................", "6) Quê quán (xã, phường): {{ que_quan_xa }}"),
        ("(huyện, quận): ........................", "(huyện, quận): {{ que_quan_huyen }}"),
        ("(tỉnh, TP): ...............................", "(tỉnh, TP): {{ que_quan_tinh }}"),
        
        # Line 12
        ("7) Nơi ở hiện nay (Xã, huyện, tỉnh hoặc số nhà, đường phố, TP): ..............................................",
         "7) Nơi ở hiện nay: {{ noi_o_hien_nay }}"),
        ("đ/thoại: ....................", "đ/thoại: {{ dien_thoai }}"),
        
        # Line 13
        ("8) Dân tộc: (Kinh, Tày, Mông, Ê đê...): ..............................", "8) Dân tộc: {{ dan_toc }}"),
        ("9) Tôn giáo: ......................................................", "9) Tôn giáo: {{ ton_giao }}"),
        
        # More patterns...
        ("10) Thành phần gia đình xuất thân:", "10) Thành phần gia đình xuất thân: {{ thanh_phan_xuat_than }}"),
        ("11) Nghề nghiệp bản thân trước khi được tuyển dụng:", "11) Nghề nghiệp bản thân: {{ nghe_nghiep_ban_than }}"),
        
        # Dates
        ("12) Ngày được tuyển dụng:", "12) Ngày được tuyển dụng: {{ ngay_tuyen_dung }}"),
        ("Vào cơ quan nào, ở đâu:", "Vào cơ quan: {{ co_quan_tuyen_dung }}"),
        ("13) Ngày vào cơ quan hiện đang công tác:", "13) Ngày vào cơ quan: {{ ngay_vao_co_quan }}"),
        ("Ngày tham gia cách mạng:", "Ngày tham gia cách mạng: {{ ngay_tham_gia_cach_mang }}"),
        ("14) Ngày vào Đảng Cộng sản Việt Nam:", "14) Ngày vào Đảng: {{ ngay_vao_dang }}"),
        ("Ngày chính thức:", "Ngày chính thức: {{ ngay_chinh_thuc_dang }}"),
        ("15) Ngày tham gia các tổ chức chính trị, xã hội:", "15) Tổ chức: {{ ngay_tham_gia_to_chuc }}"),
        ("16) Ngày nhập ngũ:", "16) Ngày nhập ngũ: {{ ngay_nhap_ngu }}"),
        ("Ngày xuất ngũ:", "Ngày xuất ngũ: {{ ngay_xuat_ngu }}"),
        ("Quân hàm, chức vụ cao nhất (năm):", "Quân hàm: {{ quan_ham }}"),
        
        # Education
        ("17)Trình độ học vấn: Giáo dục phổ thông:", "17) Giáo dục phổ thông: {{ trinh_do_giao_duc_pho_thong }}"),
        ("Học hàm, học vị cao nhất:", "Học hàm, học vị: {{ hoc_ham_hoc_vi }}"),
        ("- Lý luận chính trị:", "- Lý luận chính trị: {{ ly_luan_chinh_tri }}"),
        ("- Ngoại ngữ:", "- Ngoại ngữ: {{ ngoai_ngu }}"),
        
        # Work
        ("18) Công tác chính đảng làm:", "18) Công tác: {{ cong_tac_chinh_dang }}"),
        ("19) Ngạch công chức:", "19) Ngạch: {{ nguoi_cong_chuc_vien_chuc }}"),
        ("Bậc lương:", "Bậc lương: {{ bac_luong }}"),
        ("hệ số:", "hệ số: {{ he_so }}"),
        
        # Health
        ("24) Tình trạng sức khỏe:", "24) Tình trạng sức khỏe: {{ tinh_trang_suc_khoe }}"),
        ("Cao:", "Cao: {{ chieu_cao }}"),
        ("Cân nặng:", "Cân nặng: {{ can_nang }}"),
        ("Nhóm máu:", "Nhóm máu: {{ nhom_mau }}"),
    ]
    
    replace_count = 0
    for i, para in enumerate(doc.paragraphs):
        for old, new in replacements:
            if replace_exact_in_paragraph(para, old, new):
                print(f"✅ P{i}: {new[:60]}")
                replace_count += 1
    
    # Tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    for old, new in replacements:
                        if replace_exact_in_paragraph(para, old, new):
                            print(f"✅ T{table_idx}R{row_idx}C{cell_idx}: {new[:40]}")
                            replace_count += 1
    
    # Save
    doc.save(OUTPUT)
    
    print(f"\n{'='*60}")
    print(f"📄 File: {OUTPUT}")
    print(f"🔄 Total replacements: {replace_count}")
    print(f"✅ V8 COMPLETE - With merged runs!")

if __name__ == "__main__":
    process_document()
