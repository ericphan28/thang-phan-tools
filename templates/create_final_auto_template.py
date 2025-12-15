"""
GIẢI PHÁP TỰ ĐỘNG 100% - FORMAT PRESERVED
==========================================
Clone runs với format gốc thay vì tạo mới
"""

from docx import Document
from docx.shared import RGBColor
import re
import json

def copy_run_format(source_run, target_run):
    """
    Copy TẤT CẢ format properties từ source sang target run
    """
    # Font properties
    if source_run.font.name:
        target_run.font.name = source_run.font.name
    if source_run.font.size:
        target_run.font.size = source_run.font.size
    
    # Style properties
    target_run.bold = source_run.bold
    target_run.italic = source_run.italic  
    target_run.underline = source_run.underline
    
    # Color
    if source_run.font.color and source_run.font.color.rgb:
        target_run.font.color.rgb = source_run.font.color.rgb
    
    # More properties
    if source_run.font.all_caps is not None:
        target_run.font.all_caps = source_run.font.all_caps
    if source_run.font.small_caps is not None:
        target_run.font.small_caps = source_run.font.small_caps
    if source_run.font.strike is not None:
        target_run.font.strike = source_run.font.strike

def smart_replace_in_paragraph(paragraph, pattern, jinja_var):
    """
    Replace text trong paragraph nhưng GIỮ NGUYÊN format
    Bằng cách clone run thay vì tạo mới
    """
    full_text = paragraph.text
    
    # Check if pattern exists
    match = re.search(pattern, full_text)
    if not match:
        return False
    
    match_start = match.start()
    match_end = match.end()
    
    # Find which run(s) contain the match
    current_pos = 0
    runs_to_process = []
    
    for idx, run in enumerate(paragraph.runs):
        run_start = current_pos
        run_end = current_pos + len(run.text)
        
        # Check if this run overlaps with match
        if run_start < match_end and run_end > match_start:
            runs_to_process.append({
                'idx': idx,
                'run': run,
                'run_start': run_start,
                'run_end': run_end,
                'overlap_start': max(run_start, match_start),
                'overlap_end': min(run_end, match_end)
            })
        
        current_pos = run_end
    
    if not runs_to_process:
        return False
    
    # Simple case: match is entirely within one run
    if len(runs_to_process) == 1:
        run_info = runs_to_process[0]
        run = run_info['run']
        run_start = run_info['run_start']
        
        # Split into: before | jinja_var | after
        before_text = run.text[:match_start - run_start]
        after_text = run.text[match_end - run_start:]
        jinja_text = f'{{{{ {jinja_var} }}}}'
        
        # Clear original run but keep it for format reference
        original_format_run = run
        run.text = ''
        
        # Get parent paragraph and insert position
        para = run._element.getparent()
        run_elem = run._element
        run_idx = list(para).index(run_elem)
        
        # Remove original run element
        para.remove(run_elem)
        
        # Add new runs with preserved format
        if before_text:
            new_run = paragraph.add_run(before_text)
            copy_run_format(original_format_run, new_run)
            # Move to correct position
            para.insert(run_idx, new_run._element)
            run_idx += 1
        
        # Add Jinja variable run
        jinja_run = paragraph.add_run(jinja_text)
        copy_run_format(original_format_run, jinja_run)
        para.insert(run_idx, jinja_run._element)
        run_idx += 1
        
        if after_text:
            new_run = paragraph.add_run(after_text)
            copy_run_format(original_format_run, new_run)
            para.insert(run_idx, new_run._element)
        
        return True
    
    # Complex case: match spans multiple runs
    # For now, handle by replacing entire paragraph
    # (Can improve later if needed)
    
    return False

def simple_replace_in_run(paragraph, pattern, jinja_var):
    """
    Phương pháp đơn giản hơn: Replace trong từng run
    Giữ format của run đó
    """
    replaced = False
    
    for run in paragraph.runs:
        if re.search(pattern, run.text):
            # Replace text trong run, giữ format
            run.text = re.sub(pattern, f'{{{{ {jinja_var} }}}}', run.text)
            replaced = True
            # Chỉ replace lần đầu tìm thấy
            break
    
    return replaced

# Comprehensive field patterns - 50 most common fields
FIELD_PATTERNS = [
    # Header (3)
    (r"Tỉnh:\s*[\.…]+", "tinh"),
    (r"Đơn vị trực thuộc:\s*[\.…]+", "don_vi_truc_thuoc"),
    (r"Đơn vị cơ sở:\s*[\.…]+", "don_vi_co_so"),
    
    # Personal info (20)
    (r"Họ và tên khai sinh:\s*[\.…]+", "ho_ten"),
    (r"Nam, nữ:\s*[\.…]+", "gioi_tinh"),
    (r"Sinh ngày:\s*[\.…]+\s*tháng:\s*[\.…]+\s*năm:\s*[\.…]+", "sinh_ngay_thang_nam"),
    (r"Các tên gọi khác:\s*[\.…]+", "ten_goi_khac"),
    (r"Nơi sinh:\s*[\.…]+", "noi_sinh"),
    (r"Quê quán \(xã, phường\):\s*[\.…]+", "que_quan_xa"),
    (r"\(huyện, quận\):\s*[\.…]+", "que_quan_huyen"),
    (r"\(tỉnh, TP\):\s*[\.…]+", "que_quan_tinh"),
    (r"Nơi ở hiện nay[^:]*:\s*[\.…]+", "noi_o_hien_nay"),
    (r"đ/thoại:\s*[\.…]+", "dien_thoai"),
    (r"Email:\s*[\.…]+", "email"),
    (r"\(Kinh, Tày, Mông[^)]*\):\s*[\.…]+", "dan_toc"),
    (r"9\)\s*[\.…]+", "ton_giao"),  # After dan_toc
    (r"Thành phần gia đình xuất thân:\s*[\.…]+", "thanh_phan_gia_dinh"),
    (r"Nghề nghiệp bản thân[^:]*:\s*[\.…]+", "nghe_nghiep_ban_than"),
    (r"Ngày được tuyển dụng:\s*[\.…]+", "ngay_tuyen_dung"),
    (r"Vào cơ quan nào, ở đâu:\s*[\.…]+", "vao_co_quan"),
    
    # Party & Military (10)
    (r"Ngày vào cơ quan hiện đang công tác:\s*[\.…/]+", "ngay_vao_co_quan_hien_tai"),
    (r"Ngày tham gia cách mạng:\s*[\.…/]+", "ngay_tham_gia_cach_mang"),
    (r"14\)\s*[\.…/]+", "ngay_vao_dang"),  # After item 14
    (r"Ngày chính thức:\s*[\.…/]+", "ngay_chinh_thuc"),
    (r"15\)\s*[\.…]+", "ngay_vao_dang_cu"),  # After item 15
    (r"16\)\s*«ngay_nhap_ngu»\s*/\s*[\.…]+", "ngay_nhap_ngu"),
    (r"«ngay_xuat_ngu»\s*/\s*[\.…]+", "ngay_xuat_ngu"),
    (r"Quân hàm, chức vụ cao nhất[^:]*:\s*[\.…]+", "quan_ham"),
    (r"17\)Trình độ học vấn: Giáo dục phổ thông:\s*[\.…]+", "trinh_do_giao_duc"),
    
    # Education & Skills (10)
    (r"«hoc_ham_hoc_vi»", "hoc_ham_hoc_vi"),
    (r"«ly_luan_chinh_tri»", "ly_luan_chinh_tri"),
    (r"«ngoai_ngu»", "ngoai_ngu"),
    (r"Dân tộc thiểu số:\s*[\.…]+", "dan_toc_thieu_so"),
    (r"Tin học:\s*[\.…]+", "tin_hoc"),
    
    # Current position (7)
    (r"3\)\s*«cap_uy_hien_tai»", "cap_uy_hien_tai"),
    (r"«cap_uy_kiem»", "cap_uy_kiem"),
    (r"«chuc_vu»", "chuc_vu"),
    (r"«phu_cap_chuc_vu»", "phu_cap_chuc_vu"),
    (r"Phụ cấp khác:\s*[\.…]+", "phu_cap_khac"),
    (r"Ngạch, bậc, lương:\s*[\.…]+", "ngach_bac_luong"),
    (r"Ngày bổ nhiệm[^:]*:\s*[\.…]+", "ngay_bo_nhiem"),
]

print("🚀 TẠO TEMPLATE TỰ ĐỘNG - FORMAT 100% PRESERVED")
print("=" * 70)

# Load original document
print("\n📖 Loading original document...")
doc = Document('mau-nha-nuoc/Mau-ly-lich-2C-TCTW-98.docx')
print(f"   ✅ Loaded: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

# Process paragraphs
print("\n🔧 Processing paragraphs...")
replaced_count = 0

for i, para in enumerate(doc.paragraphs):
    for pattern, var_name in FIELD_PATTERNS:
        # Try simple replace first (faster, preserves format within run)
        if simple_replace_in_run(para, pattern, var_name):
            print(f"   ✅ P{i}: {var_name}")
            replaced_count += 1
            break  # Only replace once per paragraph

# Process tables
print("\n🔧 Processing tables...")
for table_idx, table in enumerate(doc.tables):
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                for pattern, var_name in FIELD_PATTERNS:
                    if simple_replace_in_run(para, pattern, var_name):
                        print(f"   ✅ Table {table_idx+1}, Row {row_idx+1}, Cell {cell_idx+1}: {var_name}")
                        replaced_count += 1

# Save template
output_file = 'mau_2c_FINAL_AUTO_TEMPLATE.docx'
print(f"\n💾 Saving template...")
doc.save(output_file)

print("\n" + "=" * 70)
print("✅ TEMPLATE CREATED SUCCESSFULLY!")
print(f"📄 File: {output_file}")
print(f"🔄 Replacements: {replaced_count}")
print(f"\n💡 Next step: Test render with docxtpl")
print(f"   python test_final_template.py")
