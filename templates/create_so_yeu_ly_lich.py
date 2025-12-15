"""
Tạo Template Sơ Yếu Lý Lịch - Mẫu Nhà Nước
Theo format chuẩn của Việt Nam
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_table_border(table):
    """Thêm viền cho bảng"""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)

# Tạo document
doc = Document()

# Setup margins
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(1.5)
section.page_width = Cm(21)
section.page_height = Cm(29.7)

# ================ HEADER ================
header1 = doc.add_paragraph()
header1.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = header1.add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM')
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Times New Roman'

header2 = doc.add_paragraph()
header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = header2.add_run('Độc lập - Tự do - Hạnh phúc')
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Times New Roman'

underline = doc.add_paragraph()
underline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = underline.add_run('________________')
run.font.size = Pt(13)

doc.add_paragraph()  # Space

# ================ TITLE ================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('SƠ YẾU LÝ LỊCH')
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0, 0, 139)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('(Dùng cho cán bộ, công chức, viên chức)')
run.italic = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()  # Space

# ================ THÔNG TIN CÁ NHÂN ================
# 1. Họ và tên
p1 = doc.add_paragraph()
p1.add_run('1. Họ và tên khai sinh (viết chữ IN HOA): ').bold = True
p1.add_run('{{ho_ten}}').font.color.rgb = RGBColor(0, 0, 255)
p1.paragraph_format.left_indent = Cm(1)

# 2. Tên gọi khác
p2 = doc.add_paragraph()
p2.add_run('2. Tên gọi khác: ').bold = True
p2.add_run('{{ten_goi_khac}}').font.color.rgb = RGBColor(0, 0, 255)
p2.paragraph_format.left_indent = Cm(1)

# 3. Sinh ngày
p3 = doc.add_paragraph()
p3.add_run('3. Sinh ngày: ').bold = True
p3.add_run('{{ngay_sinh}}').font.color.rgb = RGBColor(0, 0, 255)
p3.add_run('    Nơi sinh: ').bold = True
p3.add_run('{{noi_sinh}}').font.color.rgb = RGBColor(0, 0, 255)
p3.paragraph_format.left_indent = Cm(1)

# 4. Quê quán
p4 = doc.add_paragraph()
p4.add_run('4. Quê quán: ').bold = True
p4.add_run('{{que_quan}}').font.color.rgb = RGBColor(0, 0, 255)
p4.paragraph_format.left_indent = Cm(1)

# 5. Hộ khẩu thường trú
p5 = doc.add_paragraph()
p5.add_run('5. Hộ khẩu thường trú: ').bold = True
p5.add_run('{{ho_khau}}').font.color.rgb = RGBColor(0, 0, 255)
p5.paragraph_format.left_indent = Cm(1)

# 6. Nơi ở hiện nay
p6 = doc.add_paragraph()
p6.add_run('6. Nơi ở hiện nay: ').bold = True
p6.add_run('{{noi_o_hien_nay}}').font.color.rgb = RGBColor(0, 0, 255)
p6.paragraph_format.left_indent = Cm(1)

# 7. Dân tộc
p7 = doc.add_paragraph()
p7.add_run('7. Dân tộc: ').bold = True
p7.add_run('{{dan_toc}}').font.color.rgb = RGBColor(0, 0, 255)
p7.add_run('    Tôn giáo: ').bold = True
p7.add_run('{{ton_giao}}').font.color.rgb = RGBColor(0, 0, 255)
p7.paragraph_format.left_indent = Cm(1)

# 8. Số CMND/CCCD
p8 = doc.add_paragraph()
p8.add_run('8. Số CMND/CCCD: ').bold = True
p8.add_run('{{cmnd}}').font.color.rgb = RGBColor(0, 0, 255)
p8.add_run('    Ngày cấp: ').bold = True
p8.add_run('{{cmnd_ngay_cap}}').font.color.rgb = RGBColor(0, 0, 255)
p8.add_run('    Nơi cấp: ').bold = True
p8.add_run('{{cmnd_noi_cap}}').font.color.rgb = RGBColor(0, 0, 255)
p8.paragraph_format.left_indent = Cm(1)

# 9. Nghề nghiệp
p9 = doc.add_paragraph()
p9.add_run('9. Nghề nghiệp, công việc hiện nay: ').bold = True
p9.add_run('{{nghe_nghiep}}').font.color.rgb = RGBColor(0, 0, 255)
p9.paragraph_format.left_indent = Cm(1)

# 10. Đơn vị công tác
p10 = doc.add_paragraph()
p10.add_run('10. Đơn vị công tác: ').bold = True
p10.add_run('{{don_vi_cong_tac}}').font.color.rgb = RGBColor(0, 0, 255)
p10.paragraph_format.left_indent = Cm(1)

# 11. Điện thoại
p11 = doc.add_paragraph()
p11.add_run('11. Điện thoại: ').bold = True
p11.add_run('{{dien_thoai}}').font.color.rgb = RGBColor(0, 0, 255)
p11.add_run('    Email: ').bold = True
p11.add_run('{{email}}').font.color.rgb = RGBColor(0, 0, 255)
p11.paragraph_format.left_indent = Cm(1)

doc.add_paragraph()  # Space

# ================ TRÌNH ĐỘ HỌC VẤN ================
section_title = doc.add_paragraph()
section_title.add_run('12. Trình độ học vấn, chuyên môn:').bold = True
section_title.paragraph_format.left_indent = Cm(1)

# Table cho học vấn
edu_table = doc.add_table(rows=1, cols=4)
edu_table.style = 'Table Grid'
add_table_border(edu_table)

# Header row
header_cells = edu_table.rows[0].cells
headers = ['Trình độ', 'Chuyên ngành', 'Trường đào tạo', 'Năm tốt nghiệp']
for i, header in enumerate(headers):
    cell = header_cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(11)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add loop for education
edu_row = edu_table.add_row()
edu_row.cells[0].text = '{{#hoc_van}}{{trinh_do}}'
edu_row.cells[1].text = '{{chuyen_nganh}}'
edu_row.cells[2].text = '{{truong}}'
edu_row.cells[3].text = '{{nam_tot_nghiep}}{{/hoc_van}}'

doc.add_paragraph()  # Space

# ================ QUÁ TRÌNH CÔNG TÁC ================
work_title = doc.add_paragraph()
work_title.add_run('13. Quá trình công tác:').bold = True
work_title.paragraph_format.left_indent = Cm(1)

# Table cho công tác
work_table = doc.add_table(rows=1, cols=3)
work_table.style = 'Table Grid'
add_table_border(work_table)

# Header row
header_cells = work_table.rows[0].cells
headers = ['Thời gian', 'Đơn vị công tác', 'Chức vụ']
for i, header in enumerate(headers):
    cell = header_cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(11)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add loop for work history
work_row = work_table.add_row()
work_row.cells[0].text = '{{#qua_trinh_cong_tac}}{{thoi_gian}}'
work_row.cells[1].text = '{{don_vi}}'
work_row.cells[2].text = '{{chuc_vu}}{{/qua_trinh_cong_tac}}'

doc.add_paragraph()  # Space

# ================ GIA ĐÌNH ================
family_title = doc.add_paragraph()
family_title.add_run('14. Quan hệ gia đình:').bold = True
family_title.paragraph_format.left_indent = Cm(1)

# a. Bố
father = doc.add_paragraph()
father.add_run('    a. Bố: ').bold = True
father.add_run('{{bo.ho_ten}}').font.color.rgb = RGBColor(0, 0, 255)
father.paragraph_format.left_indent = Cm(2)

father2 = doc.add_paragraph()
father2.add_run('        Năm sinh: ').font.size = Pt(11)
father2.add_run('{{bo.nam_sinh}}').font.color.rgb = RGBColor(0, 0, 255)
father2.add_run('    Nghề nghiệp: ').font.size = Pt(11)
father2.add_run('{{bo.nghe_nghiep}}').font.color.rgb = RGBColor(0, 0, 255)
father2.paragraph_format.left_indent = Cm(2)

# b. Mẹ
mother = doc.add_paragraph()
mother.add_run('    b. Mẹ: ').bold = True
mother.add_run('{{me.ho_ten}}').font.color.rgb = RGBColor(0, 0, 255)
mother.paragraph_format.left_indent = Cm(2)

mother2 = doc.add_paragraph()
mother2.add_run('        Năm sinh: ').font.size = Pt(11)
mother2.add_run('{{me.nam_sinh}}').font.color.rgb = RGBColor(0, 0, 255)
mother2.add_run('    Nghề nghiệp: ').font.size = Pt(11)
mother2.add_run('{{me.nghe_nghiep}}').font.color.rgb = RGBColor(0, 0, 255)
mother2.paragraph_format.left_indent = Cm(2)

# c. Vợ/Chồng
spouse = doc.add_paragraph()
spouse.add_run('    c. Vợ/Chồng: ').bold = True
spouse.add_run('{{vo_chong.ho_ten}}').font.color.rgb = RGBColor(0, 0, 255)
spouse.paragraph_format.left_indent = Cm(2)

spouse2 = doc.add_paragraph()
spouse2.add_run('        Năm sinh: ').font.size = Pt(11)
spouse2.add_run('{{vo_chong.nam_sinh}}').font.color.rgb = RGBColor(0, 0, 255)
spouse2.add_run('    Nghề nghiệp: ').font.size = Pt(11)
spouse2.add_run('{{vo_chong.nghe_nghiep}}').font.color.rgb = RGBColor(0, 0, 255)
spouse2.paragraph_format.left_indent = Cm(2)

# d. Con
children_title = doc.add_paragraph()
children_title.add_run('    d. Con:').bold = True
children_title.paragraph_format.left_indent = Cm(2)

# Table cho con cái
children_table = doc.add_table(rows=1, cols=3)
children_table.style = 'Table Grid'
add_table_border(children_table)
children_table.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Header row
header_cells = children_table.rows[0].cells
headers = ['Họ và tên', 'Năm sinh', 'Nghề nghiệp']
for i, header in enumerate(headers):
    cell = header_cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(11)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add loop for children
child_row = children_table.add_row()
child_row.cells[0].text = '{{#con}}{{ho_ten}}'
child_row.cells[1].text = '{{nam_sinh}}'
child_row.cells[2].text = '{{nghe_nghiep}}{{/con}}'

doc.add_paragraph()  # Space

# ================ KỶ LUẬT/KHEN THƯỞNG ================
p15 = doc.add_paragraph()
p15.add_run('15. Khen thưởng: ').bold = True
p15.add_run('{{khen_thuong}}').font.color.rgb = RGBColor(0, 0, 255)
p15.paragraph_format.left_indent = Cm(1)

p16 = doc.add_paragraph()
p16.add_run('16. Kỷ luật: ').bold = True
p16.add_run('{{ky_luat}}').font.color.rgb = RGBColor(0, 0, 255)
p16.paragraph_format.left_indent = Cm(1)

doc.add_paragraph()
doc.add_paragraph()

# ================ FOOTER ================
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = footer.add_run('{{noi_lam}}, ngày {{ngay}} tháng {{thang}} năm {{nam}}')
run.italic = True
run.font.size = Pt(12)

declare = doc.add_paragraph()
declare.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = declare.add_run('Người khai')
run.bold = True
run.font.size = Pt(12)

sign_note = doc.add_paragraph()
sign_note.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = sign_note.add_run('(Ký và ghi rõ họ tên)')
run.italic = True
run.font.size = Pt(11)

# Save
output_path = r'd:\thang\utility-server\templates\so_yeu_ly_lich.docx'
doc.save(output_path)
print(f'✅ Đã tạo template: {output_path}')
print('📄 Cấu trúc:')
print('   - Thông tin cá nhân (16 fields)')
print('   - Trình độ học vấn (table with loop)')
print('   - Quá trình công tác (table with loop)')
print('   - Quan hệ gia đình (bố, mẹ, vợ/chồng, con)')
print('   - Khen thưởng, kỷ luật')
