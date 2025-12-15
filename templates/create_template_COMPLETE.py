"""
Script tạo template HOÀN CHỈNH cho Mẫu 2C với docxtpl
Version 2.0 - Improved with better pattern matching
"""

from docx import Document
import re
from pathlib import Path

def create_complete_template():
    """
    Tạo template HOÀN CHỈNH với tất cả 60+ fields được map đúng
    """
    
    print("🚀 TẠO TEMPLATE HOÀN CHỈNH - VERSION 2.0")
    print("="*70)
    
    # Load original
    template_path = Path("mau-nha-nuoc/Mau-ly-lich-2C-TCTW-98.docx")
    if not template_path.exists():
        print(f"❌ ERROR: Không tìm thấy {template_path}")
        return False
    
    print(f"📖 Đọc file gốc: {template_path}")
    doc = Document(template_path)
    
    print(f"📊 Tổng: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
    
    # Enhanced replacement patterns - MORE SPECIFIC
    replacements = {
        # === HEADER ===
        r"Tỉnh:\s*\.{3,}|Tỉnh:\s*…+": "Tỉnh: {{ tinh }}",
        r"Đơn vị trực thuộc:\s*\.{3,}|Đơn vị trực thuộc:\s*…+": "Đơn vị trực thuộc: {{ don_vi_truc_thuoc }}",
        r"Đơn vị cơ sở:\s*\.{3,}|Đơn vị cơ sở:\s*…+": "Đơn vị cơ sở: {{ don_vi_co_so }}",
        r"Số hiệu:\s*\.{3,}|Số hiệu:\s*…+": "Số hiệu: {{ so_hieu }}",
        
        # === MỤC 1: HỌ TÊN ===
        r"Họ và tên.*?:\s*\.{3,}|Họ và tên.*?:\s*…+": "Họ và tên: {{ ho_ten }}",
        r"(?:Tên gọi khác|Tên thường gọi):\s*\.{3,}|(?:Tên gọi khác|Tên thường gọi):\s*…+": "Tên gọi khác: {{ ten_goi_khac }}",
        
        # === MỤC 2: NGÀY SINH (SPECIAL PATTERN) ===
        r"Sinh ngày\s+\.{3,}\s+tháng\s+\.{3,}\s+năm\s+\.{3,}": "Sinh ngày {{ ngay }} tháng {{ thang }} năm {{ nam }}",
        r"Sinh ngày\s+…+\s+tháng\s+…+\s+năm\s+…+": "Sinh ngày {{ ngay }} tháng {{ thang }} năm {{ nam }}",
        
        r"Nơi sinh:\s*\.{3,}|Nơi sinh:\s*…+": "Nơi sinh: {{ noi_sinh }}",
        r"Nguyên quán:\s*\.{3,}|Nguyên quán:\s*…+": "Nguyên quán: {{ nguyen_quan }}",
        
        # === MỤC 3: DÂN TỘC, TÔN GIÁO ===
        r"Dân tộc:\s*\.{3,}|Dân tộc:\s*…+": "Dân tộc: {{ dan_toc }}",
        r"Tôn giáo:\s*\.{3,}|Tôn giáo:\s*…+": "Tôn giáo: {{ ton_giao }}",
        
        # === MỤC 4: CMND/CCCD ===
        r"Số CMND/CCCD:\s*\.{3,}|Số CMND/CCCD:\s*…+": "Số CMND/CCCD: {{ so_cmnd }}",
        r"Ngày cấp:\s*\.{3,}|Ngày cấp:\s*…+": "Ngày cấp: {{ ngay_cap }}",
        r"Nơi cấp:\s*\.{3,}|Nơi cấp:\s*…+": "Nơi cấp: {{ noi_cap }}",
        
        # === MỤC 5: HỘ KHẨU ===
        r"Hộ khẩu thường trú:\s*\.{3,}|Hộ khẩu thường trú:\s*…+": "Hộ khẩu thường trú: {{ ho_khau }}",
        r"Chỗ ở hiện tại:\s*\.{3,}|Chỗ ở hiện tại:\s*…+": "Chỗ ở hiện tại: {{ cho_o_hien_tai }}",
        
        # === MỤC 6: LIÊN HỆ ===
        r"Điện thoại:\s*\.{3,}|Điện thoại:\s*…+": "Điện thoại: {{ dien_thoai }}",
        r"Email:\s*\.{3,}|Email:\s*…+": "Email: {{ email }}",
        
        # === MỤC 7-12: TRÌNH ĐỘ ===
        r"Trình độ văn hóa:\s*\.{3,}|Trình độ văn hóa:\s*…+": "Trình độ văn hóa: {{ trinh_do_van_hoa }}",
        r"Trình độ chuyên môn.*?:\s*\.{3,}|Trình độ chuyên môn.*?:\s*…+": "Trình độ chuyên môn cao nhất: {{ trinh_do_chuyen_mon }}",
        r"Lý luận chính trị:\s*\.{3,}|Lý luận chính trị:\s*…+": "Lý luận chính trị: {{ ly_luan_chinh_tri }}",
        r"Quản lý nhà nước:\s*\.{3,}|Quản lý nhà nước:\s*…+": "Quản lý nhà nước: {{ quan_ly_nha_nuoc }}",
        r"Ngoại ngữ:\s*\.{3,}|Ngoại ngữ:\s*…+": "Ngoại ngữ: {{ ngoai_ngu }}",
        r"Tin học:\s*\.{3,}|Tin học:\s*…+": "Tin học: {{ tin_hoc }}",
        
        # === MỤC 13-15: ĐẢNG, ĐOÀN THỂ ===
        r"Ngày vào Đảng.*?:\s*\.{3,}|Ngày vào Đảng.*?:\s*…+": "Ngày vào Đảng Cộng sản Việt Nam: {{ ngay_vao_dang }}",
        r"Ngày chính thức:\s*\.{3,}|Ngày chính thức:\s*…+": "Ngày chính thức: {{ ngay_chinh_thuc }}",
        r"Ngày tham gia.*?chính trị.*?:\s*\.{3,}|Ngày tham gia.*?chính trị.*?:\s*…+": "Ngày tham gia tổ chức chính trị - xã hội: {{ ngay_tham_gia_to_chuc }}",
        r"Ngày nhập ngũ:\s*\.{3,}|Ngày nhập ngũ:\s*…+": "Ngày nhập ngũ: {{ ngay_nhap_ngu }}",
        r"Ngày xuất ngũ:\s*\.{3,}|Ngày xuất ngũ:\s*…+": "Ngày xuất ngũ: {{ ngay_xuat_ngu }}",
        r"Quân hàm.*?:\s*\.{3,}|Quân hàm.*?:\s*…+": "Quân hàm cao nhất: {{ quan_ham }}",
        
        # === MỤC 16-19: CÔNG VIỆC ===
        r"Chức vụ hiện tại:\s*\.{3,}|Chức vụ hiện tại:\s*…+": "Chức vụ hiện tại: {{ chuc_vu }}",
        r"Công việc chính.*?:\s*\.{3,}|Công việc chính.*?:\s*…+": "Công việc chính được giao: {{ cong_viec_chinh }}",
        r"Ngạch công chức:\s*\.{3,}|Ngạch công chức:\s*…+": "Ngạch công chức: {{ ngach_cong_chuc }}",
        r"Mã ngạch:\s*\.{3,}|Mã ngạch:\s*…+": "Mã ngạch: {{ ma_ngach }}",
        r"Bậc lương:\s*\.{3,}|Bậc lương:\s*…+": "Bậc lương: {{ bac_luong }}",
        r"Phụ cấp chức vụ:\s*\.{3,}|Phụ cấp chức vụ:\s*…+": "Phụ cấp chức vụ: {{ phu_cap_chuc_vu }}",
        r"Phụ cấp khác:\s*\.{3,}|Phụ cấp khác:\s*…+": "Phụ cấp khác: {{ phu_cap_khac }}",
        
        # === MỤC 20-21: GIA ĐÌNH ===
        r"Tình trạng hôn nhân:\s*\.{3,}|Tình trạng hôn nhân:\s*…+": "Tình trạng hôn nhân: {{ tinh_trang_hon_nhan }}",
        r"Họ và tên vợ.*?:\s*\.{3,}|Họ và tên vợ.*?:\s*…+": "Họ và tên vợ (chồng): {{ ten_vo_chong }}",
        r"Năm sinh:\s*\.{3,}|Năm sinh:\s*…+": "Năm sinh: {{ nam_sinh_vo_chong }}",
        r"Quê quán:\s*\.{3,}|Quê quán:\s*…+": "Quê quán: {{ que_quan_vo_chong }}",
        r"Nghề nghiệp:\s*\.{3,}|Nghề nghiệp:\s*…+": "Nghề nghiệp: {{ nghe_nghiep_vo_chong }}",
        r"Chỗ ở:\s*\.{3,}|Chỗ ở:\s*…+": "Chỗ ở: {{ cho_o_vo_chong }}",
        
        # === MỤC 22-25: SỨC KHỎE ===
        r"Tình trạng sức khỏe:\s*\.{3,}|Tình trạng sức khỏe:\s*…+": "Tình trạng sức khỏe: {{ suc_khoe }}",
        r"Chiều cao:\s*\.{3,}|Chiều cao:\s*…+": "Chiều cao: {{ chieu_cao }}",
        r"Cân nặng:\s*\.{3,}|Cân nặng:\s*…+": "Cân nặng: {{ can_nang }}",
        r"Nhóm máu:\s*\.{3,}|Nhóm máu:\s*…+": "Nhóm máu: {{ nhom_mau }}",
        
        # === MỤC 26-27: KHEN THƯỞNG, KỶ LUẬT ===
        r"Khen thưởng:\s*\.{3,}|Khen thưởng:\s*…+": "Khen thưởng: {{ khen_thuong }}",
        r"Kỷ luật:\s*\.{3,}|Kỷ luật:\s*…+": "Kỷ luật: {{ ky_luat }}",
        
        # === CHỮ KÝ ===
        r"Ngày\s+\.{3,}\s+tháng\s+\.{3,}\s+năm\s+20\.{2,}": "Ngày {{ ngay_ky }} tháng {{ thang_ky }} năm {{ nam_ky }}",
        r"Ngày\s+…+\s+tháng\s+…+\s+năm\s+20…+": "Ngày {{ ngay_ky }} tháng {{ thang_ky }} năm {{ nam_ky }}",
    }
    
    # Apply replacements to paragraphs
    print("\n🔧 Bước 1: Thay thế trong paragraphs...")
    replaced_count = 0
    
    for para in doc.paragraphs:
        original = para.text
        new_text = original
        
        for pattern, replacement in replacements.items():
            new_text = re.sub(pattern, replacement, new_text, flags=re.IGNORECASE)
        
        if new_text != original:
            para.text = new_text
            replaced_count += 1
            print(f"   ✓ {original[:50]}... → {new_text[:50]}...")
    
    print(f"   ✅ Đã thay thế {replaced_count} paragraphs")
    
    # Process tables
    print("\n🔧 Bước 2: Xử lý 5 bảng...")
    
    if len(doc.tables) >= 5:
        # Table 1: Đào tạo (2x5)
        print("   📋 Bảng 1: Đào tạo, bồi dưỡng (2×5)")
        table1 = doc.tables[0]
        if len(table1.rows) > 1:
            row = table1.rows[1]
            row.cells[0].text = "{% for edu in dao_tao %}{{ edu.ten_truong }}{% endfor %}"
            row.cells[1].text = "{% for edu in dao_tao %}{{ edu.nganh_hoc }}{% endfor %}"
            row.cells[2].text = "{% for edu in dao_tao %}{{ edu.thoi_gian }}{% endfor %}"
            row.cells[3].text = "{% for edu in dao_tao %}{{ edu.hinh_thuc }}{% endfor %}"
            row.cells[4].text = "{% for edu in dao_tao %}{{ edu.van_bang }}{% endfor %}"
        
        # Table 2: Công tác (2x2)
        print("   📋 Bảng 2: Quá trình công tác (2×2)")
        table2 = doc.tables[1]
        if len(table2.rows) > 1:
            row = table2.rows[1]
            row.cells[0].text = "{% for work in cong_tac %}{{ work.thoi_gian }}{% endfor %}"
            row.cells[1].text = "{% for work in cong_tac %}{{ work.chuc_vu_don_vi }}{% endfor %}"
        
        # Table 3: Gia đình bản thân (2x4) - GIỮ NGUYÊN column 0!
        print("   📋 Bảng 3: Gia đình bản thân (2×4) - ⚠️ GIỮ labels")
        table3 = doc.tables[2]
        if len(table3.rows) > 1:
            row = table3.rows[1]
            # Column 0 = GIỮ NGUYÊN (có labels "Bố, mẹ", "Vợ", etc)
            row.cells[1].text = "{% for member in gia_dinh %}{{ member.ho_ten }}{% endfor %}"
            row.cells[2].text = "{% for member in gia_dinh %}{{ member.nam_sinh }}{% endfor %}"
            row.cells[3].text = "{% for member in gia_dinh %}{{ member.thong_tin }}{% endfor %}"
        
        # Table 4: Gia đình vợ/chồng (2x4) - GIỮ NGUYÊN column 0!
        print("   📋 Bảng 4: Gia đình vợ/chồng (2×4) - ⚠️ GIỮ labels")
        table4 = doc.tables[3]
        if len(table4.rows) > 1:
            row = table4.rows[1]
            # Column 0 = GIỮ NGUYÊN
            row.cells[1].text = "{% for member in gia_dinh_vo_chong %}{{ member.ho_ten }}{% endfor %}"
            row.cells[2].text = "{% for member in gia_dinh_vo_chong %}{{ member.nam_sinh }}{% endfor %}"
            row.cells[3].text = "{% for member in gia_dinh_vo_chong %}{{ member.thong_tin }}{% endfor %}"
        
        # Table 5: Lương (3x7)
        print("   📋 Bảng 5: Quá trình lương (3×7)")
        table5 = doc.tables[4]
        if len(table5.rows) > 2:
            row = table5.rows[2]  # Row 3 = data row
            row.cells[0].text = "{% for sal in luong %}{{ sal.thang_nam }}{% endfor %}"
            row.cells[1].text = "{% for sal in luong %}{{ sal.ngach_bac }}{% endfor %}"
            row.cells[2].text = "{% for sal in luong %}{{ sal.he_so }}{% endfor %}"
    
    print(f"   ✅ Đã xử lý {len(doc.tables)} bảng")
    
    # Save
    output_path = Path("mau_2c_template_COMPLETE.docx")
    doc.save(str(output_path))
    
    file_size = output_path.stat().st_size
    
    print("\n" + "="*70)
    print("✅ THÀNH CÔNG!")
    print(f"📄 Template hoàn chỉnh: {output_path}")
    print(f"📊 Size: {file_size:,} bytes ({file_size/1024:.2f} KB)")
    print(f"\n💡 Đã thay thế: {replaced_count} paragraphs + 5 tables")
    print("\n🎯 SẴN SÀNG để dùng với docxtpl!")
    
    return True

if __name__ == "__main__":
    try:
        success = create_complete_template()
        if success:
            print("\n" + "="*70)
            print("🎉 TEMPLATE HOÀN CHỈNH ĐÃ SẴN SÀNG!")
            print("\nTest ngay:")
            print("  python test_docxtpl.py")
        else:
            print("\n❌ Có lỗi xảy ra!")
    except Exception as e:
        print(f"\n❌ LỖI: {e}")
        import traceback
        traceback.print_exc()
