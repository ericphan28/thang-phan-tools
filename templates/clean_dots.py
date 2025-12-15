#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
POST-PROCESS: Clean up remaining dots
======================================
Sau khi render, xóa tất cả sequences of 3+ dots
"""

from docx import Document
import re

def clean_dots_in_run(run):
    """
    Remove sequences of 3+ dots from run text
    """
    # Pattern: 3+ consecutive dots or ellipsis
    pattern = r'[.…]{3,}'
    
    if re.search(pattern, run.text):
        # Replace with empty space
        old_text = run.text
        run.text = re.sub(pattern, '', run.text).strip()
        return old_text != run.text
    return False

def clean_document(input_path, output_path):
    """
    Clean all dots from rendered document
    """
    print(f"📖 Loading: {input_path}")
    doc = Document(input_path)
    
    cleaned_count = 0
    
    # Clean paragraphs
    for i, para in enumerate(doc.paragraphs):
        for run in para.runs:
            if clean_dots_in_run(run):
                cleaned_count += 1
                print(f"✅ P{i}: Cleaned dots")
    
    # Clean tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    for run in para.runs:
                        if clean_dots_in_run(run):
                            cleaned_count += 1
                            print(f"✅ T{table_idx}R{row_idx}C{cell_idx}: Cleaned")
    
    # Save
    doc.save(output_path)
    
    print(f"\n{'='*60}")
    print(f"📄 Input: {input_path}")
    print(f"📄 Output: {output_path}")
    print(f"🧹 Cleaned: {cleaned_count} locations")
    print(f"✅ CLEAN COMPLETE!")
    
    return cleaned_count

if __name__ == "__main__":
    # Clean the V8 output
    cleaned = clean_document(
        "OUTPUT_V8_MERGED_RUNS.docx",
        "OUTPUT_V8_CLEANED.docx"
    )
    
    print(f"\n🎯 Next: Open OUTPUT_V8_CLEANED.docx để kiểm tra!")
