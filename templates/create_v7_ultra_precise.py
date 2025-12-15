#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ULTRA PRECISE V7 - EXACT TEXT MATCHING
=======================================
Match CHÍNH XÁC text trong document để replace
"""

from docx import Document
import re

def replace_in_run_exact(paragraph, old_text, new_text):
    """
    Replace CHÍNH XÁC old_text với new_text trong run
    Preserve format!
    """
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True
    return False

def replace_in_run_pattern(paragraph, pattern, replacement):
    """
    Replace bằng regex pattern
    """
    for run in paragraph.runs:
        if re.search(pattern, run.text):
            run.text = re.sub(pattern, replacement, run.text)
            return True
    return False

def process_document():
    """
    Process với EXACT PATTERNS từ ảnh
    """
    INPUT = "mau-nha-nuoc/Mau-ly-lich-2C-TCTW-98.docx"
    OUTPUT = "mau_2c_V7_ULTRA_PRECISE.docx"
    
    doc = Document(INPUT)
    replacements = 0
    
    # EXACT REPLACEMENTS - Dựa trên text thực tế
    exact_replacements = [
        # P0 - Header
        ("Tỉnh: .......................................", "Tỉnh: {{ tinh }}"),
        
        # P1 - Đơn vị
        ("Đơn vị trực thuộc: ..........................", "Đơn vị trực thuộc: {{ don_vi_truc_thuoc }}"),
        
        # P2 - Đơn vị cơ sở
        ("Đơn vị cơ sở: ................................", "Đơn vị cơ sở: {{ don_vi_co_so }}"),
        ("........................................", " "),
        
        # P3 - Số hiệu
        ("Số hiệu cán bộ, công chức", "Số hiệu: {{ so_hieu }}"),
        
        # P5 - Họ tên + giới tính
        ("1) Họ và tên khai sinh: ……………………………………..", "1) Họ và tên khai sinh: {{ ho_ten }}"),
        ("Nam, nữ: .....................", "Nam, nữ: {{ gioi_tinh }}"),
        
        # P6 - Tên khác
        ("2) Các tên gọi khác:    ..........................................................", "2) Các tên gọi khác: {{ ten_goi_khac }}"),
        
        # P7 - Cấp ủy
        ("3) Cấp ủy hiện tại: .......................................", "3) Cấp ủy hiện tại: {{ cap_uy_hien_tai }}"),
        ("Cấp ủy kiêm: .........................................", "Cấp ủy kiêm: {{ cap_uy_kiem }}"),
        
        # P8 - Chức vụ
        ("Chức vụ (Đảng, đoàn thể, Chính quyền, kể cả chức vụ kiêm nhiệm):       ..................................................", 
         "Chức vụ (Đảng, đoàn thể, Chính quyền, kể cả chức vụ kiêm nhiệm): {{ chuc_vu_full }}"),
        
        # P9 - Phụ cấp
        (".........................................................................   Phụ cấp chức vụ: ...........................",
         " Phụ cấp chức vụ: {{ phu_cap_chuc_vu }}"),
        
        # P10 - Sinh ngày
        ("4) Sinh ngày: .......... tháng .......... năm ...............", 
         "4) Sinh ngày: {{ ngay }} tháng {{ thang }} năm {{ nam }}"),
        ("5) Nơi sinh: ..................................................",
         "5) Nơi sinh: {{ noi_sinh }}"),
        
        # P11 - Quê quán
        ("6) Quê quán (xã, phường): .......................................", "6) Quê quán (xã, phường): {{ que_quan_xa }}"),
        ("(huyện, quận): ........................", "(huyện, quận): {{ que_quan_huyen }}"),
        ("(tỉnh, TP): ...............................", "(tỉnh, TP): {{ que_quan_tinh }}"),
        
        # P12 - Nơi ở + điện thoại
        ("7) Nơi ở hiện nay (Xã, huyện, tỉnh hoặc số nhà, đường phố, TP): ..............................................",
         "7) Nơi ở hiện nay: {{ noi_o_hien_nay }}"),
        ("Đ/thoại: ....................", "Đ/thoại: {{ dien_thoai }}"),
        
        # P13 - Dân tộc + tôn giáo
        ("8) Dân tộc: (Kinh, Tày, Mông, Ê đê...): ..............................", "8) Dân tộc: {{ dan_toc }}"),
        ("9) Tôn giáo: ......................................................", "9) Tôn giáo: {{ ton_giao }}"),
        
        # P14 - Thành phần
        ("10) Thành phần gia đình xuất thân:  .................................................................",
         "10) Thành phần gia đình xuất thân: {{ thanh_phan_xuat_than }}"),
        
        # P16 - Nghề nghiệp
        ("11) Nghề nghiệp bản thân trước khi được tuyển dụng:      .................................................................",
         "11) Nghề nghiệp bản thân trước khi được tuyển dụng: {{ nghe_nghiep_ban_than }}"),
        
        # P18 - Ngày tuyển dụng
        ("12) Ngày được tuyển dụng: ......... / ........... / ..........", 
         "12) Ngày được tuyển dụng: {{ ngay_tuyen_dung }}"),
        ("Vào cơ quan nào, ở đâu: .............................................", 
         "Vào cơ quan: {{ co_quan_tuyen_dung }}"),
        
        # P19 - Vào cơ quan + cách mạng
        ("13) Ngày vào cơ quan hiện đang công tác: ...... / ....... / ......", 
         "13) Ngày vào cơ quan hiện đang công tác: {{ ngay_vao_co_quan }}"),
        ("Ngày tham gia cách mạng: ...... / ....... / ........",
         "Ngày tham gia cách mạng: {{ ngay_tham_gia_cach_mang }}"),
        
        # P20 - Vào Đảng
        ("14) Ngày vào Đảng Cộng sản Việt Nam: ......... / .......... / .......",
         "14) Ngày vào Đảng Cộng sản Việt Nam: {{ ngay_vao_dang }}"),
        ("Ngày chính thức: ........ / .......... / ..............",
         "Ngày chính thức: {{ ngay_chinh_thuc_dang }}"),
        
        # P21 - Tổ chức
        ("15) Ngày tham gia các tổ chức chính trị, xã hội:        ..................................................................",
         "15) Ngày tham gia các tổ chức chính trị, xã hội: {{ ngay_tham_gia_to_chuc }}"),
        
        # P23 - Nhập ngũ
        ("16) Ngày nhập ngũ: ... / ... / ....", "16) Ngày nhập ngũ: {{ ngay_nhap_ngu }}"),
        ("Ngày xuất ngũ: ... / ... / ....", "Ngày xuất ngũ: {{ ngay_xuat_ngu }}"),
        ("Quân hàm, chức vụ cao nhất (năm): ............................", 
         "Quân hàm: {{ quan_ham }}"),
        
        # P24 - Học vấn
        ("17)Trình độ học vấn: Giáo dục phổ thông: ..............", 
         "17)Trình độ học vấn: Giáo dục phổ thông: {{ trinh_do_giao_duc_pho_thong }}"),
        ("Học hàm, học vị cao nhất: .................................................",
         "Học hàm, học vị cao nhất: {{ hoc_ham_hoc_vi }}"),
        
        # P26 - Lý luận + ngoại ngữ
        ("- Lý luận chính trị: ...............................", 
         "- Lý luận chính trị: {{ ly_luan_chinh_tri }}"),
        ("- Ngoại ngữ:  .................................................................",
         "- Ngoại ngữ: {{ ngoai_ngu }}"),
        
        # P28 - Công tác
        ("18) Công tác chính đảng làm:   .................................................................",
         "18) Công tác chính đảng làm: {{ cong_tac_chinh_dang }}"),
        
        # P29 - Bậc lương
        ("19) Ngạch công chức: ...................", "19) Ngạch công chức: {{ nguoi_cong_chuc_vien_chuc }}"),
        ("(mã số: .................)", "(mã số: {{ ma_so }})"),
        ("Bậc lương: ..........", "Bậc lương: {{ bac_luong }}"),
        ("hệ số: ...........", "hệ số: {{ he_so }}"),
        ("từ tháng .... /.......", "từ tháng {{ tu_thang }}"),
        
        # P36 - Sức khỏe
        ("24) Tình trạng sức khỏe:", "24) Tình trạng sức khỏe: {{ tinh_trang_suc_khoe }}"),
        ("Cao: ..... m", "Cao: {{ chieu_cao }} m"),
        ("Cân nặng: ......... (kg)", "Cân nặng: {{ can_nang }} kg"),
        ("Nhóm máu: .......", "Nhóm máu: {{ nhom_mau }}"),
    ]
    
    print(f"📖 Processing {INPUT}")
    print(f"🎯 Total exact replacements: {len(exact_replacements)}")
    
    # Process all paragraphs
    for i, para in enumerate(doc.paragraphs):
        for old, new in exact_replacements:
            if replace_in_run_exact(para, old, new):
                print(f"✅ P{i}: {new[:50]}")
                replacements += 1
    
    # Process tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    for old, new in exact_replacements:
                        if replace_in_run_exact(para, old, new):
                            print(f"✅ T{table_idx}R{row_idx}C{cell_idx}: {new[:50]}")
                            replacements += 1
    
    # Save
    doc.save(OUTPUT)
    print(f"\n{'='*60}")
    print(f"📄 File: {OUTPUT}")
    print(f"🔄 Replacements: {replacements}")
    print(f"✅ V7 ULTRA PRECISE COMPLETE!")

if __name__ == "__main__":
    process_document()
