#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
TEST V9 FINAL TEMPLATE
======================
Render with correct variable mapping
"""

from docxtpl import DocxTemplate
import json
from docx import Document
import re

def flatten_dict(d, parent_key='', sep='_'):
    """Flatten nested dict"""
    items = []
    for k, v in d.items():
        if k.startswith('_'):
            continue
        new_key = f"{parent_key}{sep}{k}" if parent_key else k
        if isinstance(v, dict):
            items.extend(flatten_dict(v, new_key, sep=sep).items())
        else:
            items.append((new_key, v))
    return dict(items)

def main():
    TEMPLATE = "mau_2c_V9_FINAL_TEMPLATE.docx"
    OUTPUT = "OUTPUT_V9_FINAL.docx"
    JSON_FILE = "mau_2c_DATA_RESTRUCTURED.json"
    
    print(f"📖 Loading template: {TEMPLATE}")
    doc = DocxTemplate(TEMPLATE)
    
    print(f"📊 Loading data: {JSON_FILE}")
    with open(JSON_FILE, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # Flatten
    flat_data = flatten_dict(data)
    
    # Keep arrays
    render_data = {**flat_data}
    if 'hoc_tap' in data:
        render_data['hoc_tap'] = data['hoc_tap']
    if 'cong_tac' in data:
        render_data['cong_tac'] = data['cong_tac']
    if 'gia_dinh' in data:
        render_data['gia_dinh'] = data['gia_dinh']
    
    print(f"✅ Loaded {len(render_data)} fields")
    print(f"\n🔍 Key fields check:")
    print(f"   ngach_cong_chuc: {data.get('ngach_cong_chuc')}")
    print(f"   ma_ngach: {data.get('ma_ngach')}")
    print(f"   bac_luong: {data.get('bac_luong')}")
    print(f"   he_so_luong: {data.get('he_so_luong')}")
    print(f"   tu_thang_nam: {data.get('tu_thang_nam')}")
    
    print(f"\n🎨 Rendering...")
    doc.render(render_data)
    
    print(f"💾 Saving: {OUTPUT}")
    doc.save(OUTPUT)
    
    # Clean dots
    print(f"\n🧹 Cleaning dots...")
    output_doc = Document(OUTPUT)
    cleaned = 0
    
    for para in output_doc.paragraphs:
        for run in para.runs:
            if re.search(r'[.…]{3,}', run.text):
                run.text = re.sub(r'[.…]{3,}', '', run.text).strip()
                cleaned += 1
    
    for table in output_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if re.search(r'[.…]{3,}', run.text):
                            run.text = re.sub(r'[.…]{3,}', '', run.text).strip()
                            cleaned += 1
    
    output_doc.save(OUTPUT.replace('.docx', '_CLEANED.docx'))
    
    # Check result
    print(f"\n{'='*60}")
    print(f"✅ Cleaned {cleaned} locations")
    print(f"📄 Output: {OUTPUT.replace('.docx', '_CLEANED.docx')}")
    
    # Verify key fields rendered
    print(f"\n🔍 Verifying paragraph 29 (work info)...")
    final_doc = Document(OUTPUT.replace('.docx', '_CLEANED.docx'))
    p29 = final_doc.paragraphs[29].text if len(final_doc.paragraphs) > 29 else "NOT FOUND"
    print(f"   P29: {p29[:150]}")
    
    if "Chuyên viên" in p29 and "01.003" in p29 and "2.34" in p29 and "10/2022" in p29:
        print(f"\n🎉 SUCCESS! All work fields rendered correctly!")
    else:
        print(f"\n⚠️  Some fields may be missing, check manually")

if __name__ == "__main__":
    main()
