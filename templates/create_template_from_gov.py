"""
Phân tích và tạo template từ mẫu Sơ yếu lý lịch chuẩn nhà nước
Thông tư 06/2023/TT-BNV
"""
from docx import Document
from docx.shared import Pt, RGBColor
import re

# Đọc file gốc
source_file = r"d:\thang\utility-server\templates\mau-nha-nuoc\mau-so-yeu-ly-lich-16852372646782068835917.docx"
output_file = r"d:\thang\utility-server\templates\so_yeu_ly_lich_nha_nuoc.docx"

print("📖 Đang đọc mẫu nhà nước...")
doc = Document(source_file)

print(f"✅ Tìm thấy {len(doc.paragraphs)} đoạn văn và {len(doc.tables)} bảng")

# Thay thế các dấu ... bằng {{variable}}
replacements = {
    # Header info
    r'Cơ quan quản lý cán bộ.*?:.*?\.+': 'Cơ quan quản lý cán bộ, công chức, viên chức: {{co_quan_quan_ly}}',
    r'Cơ quan, đơn vị sử dụng.*?:.*?\.+': 'Cơ quan, đơn vị sử dụng cán bộ, công chức, viên chức: {{don_vi_su_dung}}',
    r'Số hiệu:.*?\.+': 'Số hiệu: {{so_hieu}}',
    r'Mã số định danh:.*?\.+': 'Mã số định danh: {{ma_so_dinh_danh}}',
    
    # Basic info
    r'6\) Dân tộc:.*?\.+': '6) Dân tộc: {{dan_toc}}',
    r'7\) Tôn giáo:.*?\.+': '7) Tôn giáo: {{ton_giao}}',
    r'8\) Số CCCD:.*?\.+ Ngày cấp:.*?/.*/.*? SĐT liên hệ:.*?\.+': '8) Số CCCD: {{cccd}} Ngày cấp: {{cccd_ngay_cap}} SĐT liên hệ: {{sdt}}',
    r'9\) Số BHXH:.*?\.+ Số thẻ BHYT:.*?\.+': '9) Số BHXH: {{so_bhxh}} Số thẻ BHYT: {{so_bhyt}}',
    r'10\) Nơi ở hiện nay:.*?\.+': '10) Nơi ở hiện nay: {{noi_o_hien_nay}}',
    r'11\) Thành phần gia đình xuất thân:.*?\.+': '11) Thành phần gia đình xuất thân: {{thanh_phan_xuat_than}}',
    r'12\) Nghề nghiệp trước khi được tuyển dụng:.*?\.+': '12) Nghề nghiệp trước khi được tuyển dụng: {{nghe_nghiep_truoc}}',
    r'13\) Ngày được tuyển dụng lần đầu:.*?/.*/.*? Cơ quan.*?:.*': '13) Ngày được tuyển dụng lần đầu: {{ngay_tuyen_dung}} Cơ quan, tổ chức, đơn vị tuyển dụng: {{co_quan_tuyen_dung}}',
    r'14\) Ngày vào cơ quan hiện đang công tác:.*?\.+': '14) Ngày vào cơ quan hiện đang công tác: {{ngay_vao_co_quan}}',
    r'15\) Ngày vào Đảng.*?: .*?/.*/.*? Ngày chính thức: .*?/.*/.*': '15) Ngày vào Đảng Cộng sản Việt Nam: {{ngay_vao_dang}} Ngày chính thức: {{ngay_chinh_thuc}}',
}

# Clone document
new_doc = Document(source_file)

# Replace in paragraphs
for i, para in enumerate(new_doc.paragraphs):
    text = para.text
    for pattern, replacement in replacements.items():
        if re.search(pattern, text):
            # Clear runs
            for run in para.runs:
                run.text = ''
            # Add new text
            run = para.add_run(replacement)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            break

# Xử lý bảng - thay ... bằng {{variable}}
for table_idx, table in enumerate(new_doc.tables):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                text = para.text
                if '…' in text or '.....' in text or '......' in text:
                    # Replace dots with mustache variables
                    new_text = re.sub(r'\.{3,}|…+', '{{data}}', text)
                    for run in para.runs:
                        run.text = ''
                    run = para.add_run(new_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)

# Save
new_doc.save(output_file)
print(f"✅ Đã tạo template: {output_file}")
print("📝 Đã thay thế các dấu chấm bằng {{variable}}")
