"""
Script tạo template Word chuyên nghiệp cho Hợp đồng lao động
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_border(section):
    """Thêm viền trang"""
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'text')
    
    for border_name in ('top', 'left', 'bottom', 'right'):
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single')
        border_el.set(qn('w:sz'), '12')
        border_el.set(qn('w:space'), '24')
        border_el.set(qn('w:color'), '2E75B6')
        pgBorders.append(border_el)
    
    sectPr.append(pgBorders)

def set_cell_border(cell, **kwargs):
    """Set border for table cell"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            element = OxmlElement(tag)
            for key in ['sz', 'val', 'color', 'space']:
                if key in edge_data:
                    element.set(qn(f'w:{key}'), str(edge_data[key]))
            tcBorders.append(element)
    
    tcPr.append(tcBorders)

# Tạo document mới
doc = Document()

# Setup margins
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# Thêm viền trang
add_page_border(section)

# ================ HEADER ================
# Tiêu đề trên cùng
header1 = doc.add_paragraph()
header1.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = header1.add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM')
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0, 0, 0)

header2 = doc.add_paragraph()
header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = header2.add_run('Độc lập - Tự do - Hạnh phúc')
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0, 0, 0)

# Đường kẻ dưới header
header3 = doc.add_paragraph()
header3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = header3.add_run('═' * 30)
run.font.size = Pt(12)

# Space
doc.add_paragraph()

# ================ TITLE ================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('HỢP ĐỒNG LAO ĐỘNG')
run.bold = True
run.font.size = Pt(18)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(46, 117, 182)  # Xanh dương đậm

# Số hợp đồng
contract_no = doc.add_paragraph()
contract_no.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = contract_no.add_run('Số: ')
run.italic = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run2 = contract_no.add_run('{{contractNumber}}')
run2.italic = True
run2.bold = True
run2.font.size = Pt(12)
run2.font.name = 'Times New Roman'
run2.font.color.rgb = RGBColor(192, 0, 0)  # Đỏ

doc.add_paragraph()

# ================ PREAMBLE ================
intro = doc.add_paragraph()
intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = intro.add_run('Căn cứ: ')
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Times New Roman'
run2 = intro.add_run('Bộ luật Lao động năm 2019')
run2.font.size = Pt(13)
run2.font.name = 'Times New Roman'

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = date_para.add_run('Hôm nay, ngày ')
run.font.size = Pt(13)
run.font.name = 'Times New Roman'
run2 = date_para.add_run('{{signDate}}')
run2.bold = True
run2.font.size = Pt(13)
run2.font.name = 'Times New Roman'
run2.font.color.rgb = RGBColor(192, 0, 0)
run3 = date_para.add_run(', chúng tôi gồm có:')
run3.font.size = Pt(13)
run3.font.name = 'Times New Roman'

doc.add_paragraph()

# ================ BÊN A ================
heading_a = doc.add_paragraph()
run = heading_a.add_run('BÊN A - NGƯỜI SỬ DỤNG LAO ĐỘNG')
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(46, 117, 182)

# Bảng thông tin Bên A
table_a = doc.add_table(rows=5, cols=2)
table_a.style = 'Light Grid Accent 1'

# Set column widths
table_a.columns[0].width = Cm(5)
table_a.columns[1].width = Cm(11)

# Row 1
table_a.rows[0].cells[0].text = 'Tên công ty:'
table_a.rows[0].cells[1].text = '{{company.name}}'

# Row 2
table_a.rows[1].cells[0].text = 'Địa chỉ:'
table_a.rows[1].cells[1].text = '{{company.address}}'

# Row 3
table_a.rows[2].cells[0].text = 'Điện thoại:'
table_a.rows[2].cells[1].text = '{{company.phone}}'

# Row 4
table_a.rows[3].cells[0].text = 'Mã số thuế:'
table_a.rows[3].cells[1].text = '{{company.taxId}}'

# Row 5
table_a.rows[4].cells[0].text = 'Người đại diện:'
table_a.rows[4].cells[1].text = '{{company.representative}}'

# Format table
for row in table_a.rows:
    for idx, cell in enumerate(row.cells):
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.size = Pt(12)
        if idx == 0:  # First column - bold
            cell.paragraphs[0].runs[0].bold = True
        cell.vertical_alignment = 1  # Center vertically

doc.add_paragraph()

# ================ BÊN B ================
heading_b = doc.add_paragraph()
run = heading_b.add_run('BÊN B - NGƯỜI LAO ĐỘNG')
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(46, 117, 182)

# Bảng thông tin Bên B
table_b = doc.add_table(rows=6, cols=2)
table_b.style = 'Light Grid Accent 1'

table_b.columns[0].width = Cm(5)
table_b.columns[1].width = Cm(11)

table_b.rows[0].cells[0].text = 'Họ và tên:'
table_b.rows[0].cells[1].text = '{{employee.fullName}}'

table_b.rows[1].cells[0].text = 'Ngày sinh:'
table_b.rows[1].cells[1].text = '{{employee.birthDate}}'

table_b.rows[2].cells[0].text = 'CMND/CCCD:'
table_b.rows[2].cells[1].text = '{{employee.idNumber}}'

table_b.rows[3].cells[0].text = 'Địa chỉ:'
table_b.rows[3].cells[1].text = '{{employee.address}}'

table_b.rows[4].cells[0].text = 'Điện thoại:'
table_b.rows[4].cells[1].text = '{{employee.phone}}'

table_b.rows[5].cells[0].text = 'Email:'
table_b.rows[5].cells[1].text = '{{employee.email}}'

for row in table_b.rows:
    for idx, cell in enumerate(row.cells):
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.size = Pt(12)
        if idx == 0:
            cell.paragraphs[0].runs[0].bold = True
        cell.vertical_alignment = 1

doc.add_paragraph()

# ================ ĐIỀU 1 ================
article1 = doc.add_paragraph()
run = article1.add_run('ĐIỀU 1: VỊ TRÍ VÀ CÔNG VIỆC')
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(46, 117, 182)

# Thông tin vị trí
p1 = doc.add_paragraph(style='List Bullet')
run = p1.add_run('Chức danh: ')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run2 = p1.add_run('{{position.title}}')
run2.font.name = 'Times New Roman'
run2.font.size = Pt(12)

p2 = doc.add_paragraph(style='List Bullet')
run = p2.add_run('Bộ phận: ')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run2 = p2.add_run('{{position.department}}')
run2.font.name = 'Times New Roman'
run2.font.size = Pt(12)

p3 = doc.add_paragraph(style='List Bullet')
run = p3.add_run('Ngày bắt đầu: ')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run2 = p3.add_run('{{position.startDate}}')
run2.font.name = 'Times New Roman'
run2.font.size = Pt(12)

p4 = doc.add_paragraph(style='List Bullet')
run = p4.add_run('Loại hợp đồng: ')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run2 = p4.add_run('{{position.contractType}}')
run2.font.name = 'Times New Roman'
run2.font.size = Pt(12)

# Nhiệm vụ
tasks_heading = doc.add_paragraph()
run = tasks_heading.add_run('Nhiệm vụ cụ thể:')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

# Loop tasks
task_loop = doc.add_paragraph()
run = task_loop.add_run('{% for task in tasks %}')
run.font.name = 'Times New Roman'
run.font.size = Pt(1)  # Hidden

task_item = doc.add_paragraph(style='List Number')
run = task_item.add_run('{{task.name}}')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

task_desc = doc.add_paragraph()
task_desc.paragraph_format.left_indent = Cm(1)
run = task_desc.add_run('{{task.description}}')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

task_endloop = doc.add_paragraph()
run = task_endloop.add_run('{% endfor %}')
run.font.name = 'Times New Roman'
run.font.size = Pt(1)

doc.add_paragraph()

# ================ ĐIỀU 2 ================
article2 = doc.add_paragraph()
run = article2.add_run('ĐIỀU 2: THỜI GIAN LÀM VIỆC')
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(46, 117, 182)

p1 = doc.add_paragraph(style='List Bullet')
run = p1.add_run('Làm việc 8 giờ/ngày, 5 ngày/tuần (Thứ 2 - Thứ 6)')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

p2 = doc.add_paragraph(style='List Bullet')
run = p2.add_run('Giờ làm việc: 08:00 - 17:00 (nghỉ trưa 12:00 - 13:00)')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

p3 = doc.add_paragraph(style='List Bullet')
run = p3.add_run('Được nghỉ các ngày lễ, Tết theo quy định Nhà nước')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

doc.add_paragraph()

# ================ ĐIỀU 3 ================
article3 = doc.add_paragraph()
run = article3.add_run('ĐIỀU 3: LƯƠNG VÀ PHÚ CẤP')
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(46, 117, 182)

# Bảng lương
salary_table = doc.add_table(rows=4, cols=2)
salary_table.style = 'Light Grid Accent 1'

salary_table.columns[0].width = Cm(8)
salary_table.columns[1].width = Cm(8)

# Header row
salary_table.rows[0].cells[0].text = 'Khoản'
salary_table.rows[0].cells[1].text = 'Số tiền (VNĐ)'

# Data rows
salary_table.rows[1].cells[0].text = 'Lương cơ bản:'
salary_table.rows[1].cells[1].text = '{{salary.base}}'

salary_table.rows[2].cells[0].text = 'Phụ cấp:'
salary_table.rows[2].cells[1].text = '{{salary.allowance}}'

salary_table.rows[3].cells[0].text = 'TỔNG CỘNG:'
salary_table.rows[3].cells[1].text = '{{salary.total}}'

# Format salary table
for idx, row in enumerate(salary_table.rows):
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                if idx == 0 or idx == 3:  # Header or total
                    run.bold = True
        cell.vertical_alignment = 1
        
# Right align amount column
for row in salary_table.rows:
    row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

p1 = doc.add_paragraph(style='List Bullet')
run = p1.add_run('Lương được trả vào ngày 5 hàng tháng')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

p2 = doc.add_paragraph(style='List Bullet')
run = p2.add_run('Được xét tăng lương định kỳ 6 tháng/lần')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

doc.add_paragraph()

# ================ ĐIỀU 4 ================
article4 = doc.add_paragraph()
run = article4.add_run('ĐIỀU 4: CHẾ ĐỘ PHÚC LỢI')
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(46, 117, 182)

# Benefits loop
benefit_loop = doc.add_paragraph()
run = benefit_loop.add_run('{% for benefit in benefits %}')
run.font.name = 'Times New Roman'
run.font.size = Pt(1)

benefit_item = doc.add_paragraph(style='List Bullet')
run = benefit_item.add_run('{{benefit}}')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

benefit_endloop = doc.add_paragraph()
run = benefit_endloop.add_run('{% endfor %}')
run.font.name = 'Times New Roman'
run.font.size = Pt(1)

doc.add_paragraph()

# ================ ĐIỀU 5 ================
article5 = doc.add_paragraph()
run = article5.add_run('ĐIỀU 5: QUYỀN VÀ NGHĨA VỤ')
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(46, 117, 182)

# Bên A
duty_a = doc.add_paragraph()
run = duty_a.add_run('Bên A có trách nhiệm:')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

p1 = doc.add_paragraph(style='List Bullet')
run = p1.add_run('Tạo điều kiện làm việc và đào tạo cho Bên B')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

p2 = doc.add_paragraph(style='List Bullet')
run = p2.add_run('Trả lương đúng hạn, đầy đủ theo thỏa thuận')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

p3 = doc.add_paragraph(style='List Bullet')
run = p3.add_run('Đóng bảo hiểm xã hội, y tế, thất nghiệp theo quy định')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

# Bên B
duty_b = doc.add_paragraph()
run = duty_b.add_run('Bên B có trách nhiệm:')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

p1 = doc.add_paragraph(style='List Bullet')
run = p1.add_run('Hoàn thành tốt công việc được giao')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

p2 = doc.add_paragraph(style='List Bullet')
run = p2.add_run('Tuân thủ nội quy, quy định của công ty')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

p3 = doc.add_paragraph(style='List Bullet')
run = p3.add_run('Bảo mật thông tin công ty')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

doc.add_paragraph()

# ================ ĐIỀU 6 ================
article6 = doc.add_paragraph()
run = article6.add_run('ĐIỀU 6: ĐIỀU KHOẢN CHUNG')
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(46, 117, 182)

p1 = doc.add_paragraph(style='List Bullet')
run = p1.add_run('Hợp đồng có hiệu lực từ ngày ký')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

p2 = doc.add_paragraph(style='List Bullet')
run = p2.add_run('Mọi sửa đổi, bổ sung phải được hai bên thỏa thuận bằng văn bản')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

p3 = doc.add_paragraph(style='List Bullet')
run = p3.add_run('Hợp đồng được lập thành 02 bản, mỗi bên giữ 01 bản')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

# Add page break before signatures
doc.add_page_break()

# ================ SIGNATURES ================
doc.add_paragraph()
doc.add_paragraph()

signature_table = doc.add_table(rows=1, cols=2)
signature_table.autofit = False
signature_table.allow_autofit = False

# Left cell - Bên A
cell_a = signature_table.rows[0].cells[0]
p = cell_a.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

run = p.add_run('ĐẠI DIỆN BÊN A')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

p.add_run('\n')
run = p.add_run('(Ký, ghi rõ họ tên)')
run.italic = True
run.font.name = 'Times New Roman'
run.font.size = Pt(11)

p.add_run('\n\n\n\n\n')

run = p.add_run('{{company.representative}}')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

# Right cell - Bên B
cell_b = signature_table.rows[0].cells[1]
p = cell_b.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

run = p.add_run('ĐẠI DIỆN BÊN B')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

p.add_run('\n')
run = p.add_run('(Ký, ghi rõ họ tên)')
run.italic = True
run.font.name = 'Times New Roman'
run.font.size = Pt(11)

p.add_run('\n\n\n\n\n')

run = p.add_run('{{employee.fullName}}')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

# Save
output_path = r'd:\thang\utility-server\templates\hop_dong_lao_dong.docx'
doc.save(output_path)
print(f'✅ Template created successfully: {output_path}')
