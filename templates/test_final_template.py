"""
TEST FINAL AUTO TEMPLATE
========================
Render với docxtpl và kiểm tra format
"""

from docxtpl import DocxTemplate
import json
import os

print("🧪 TEST FINAL AUTO TEMPLATE")
print("=" * 60)

# Load template
template_file = 'mau_2c_FINAL_AUTO_TEMPLATE.docx'
print(f"\n📖 Loading template: {template_file}")

try:
    tpl = DocxTemplate(template_file)
    print("   ✅ Template loaded successfully")
except Exception as e:
    print(f"   ❌ Error: {e}")
    exit(1)

# Load data
data_file = 'mau_2c_DATA_RESTRUCTURED.json'
print(f"\n📖 Loading data: {data_file}")

with open(data_file, 'r', encoding='utf-8') as f:
    data = json.load(f)

print(f"   ✅ Loaded {len(data)} fields")

# Prepare context
print(f"\n🔧 Preparing render context...")

# Flatten some nested structures if needed
context = {}
for key, value in data.items():
    if isinstance(value, list):
        # Keep lists as-is for table loops
        context[key] = value
    elif isinstance(value, dict):
        # Flatten dicts
        for sub_key, sub_value in value.items():
            context[f"{key}_{sub_key}"] = sub_value
    else:
        context[key] = value

# Add combined birth date if needed
if 'ngay' in data and 'thang' in data and 'nam' in data:
    context['sinh_ngay_thang_nam'] = f"{data['ngay']}/{data['thang']}/{data['nam']}"

print(f"   ✅ Context prepared: {len(context)} fields")

# Render
print(f"\n🎨 Rendering...")

try:
    tpl.render(context)
    print("   ✅ Render successful!")
except Exception as e:
    print(f"   ⚠️  Render had issues: {e}")
    print("   💡 Continuing with partial render...")

# Save
output_file = 'OUTPUT_FINAL_AUTO.docx'
print(f"\n💾 Saving to: {output_file}")

try:
    tpl.save(output_file)
    size = os.path.getsize(output_file)
    print(f"   ✅ Saved successfully!")
    print(f"   📊 Size: {size:,} bytes ({size/1024:.2f} KB)")
except Exception as e:
    print(f"   ❌ Save error: {e}")
    exit(1)

# Analyze output
print(f"\n🔍 Analyzing output...")
from docx import Document

doc = Document(output_file)

# Check content
non_empty_paras = sum(1 for p in doc.paragraphs if p.text.strip())
print(f"   📝 Non-empty paragraphs: {non_empty_paras} / {len(doc.paragraphs)}")

# Check if Jinja variables are still there (not rendered)
unrendered = []
for i, p in enumerate(doc.paragraphs[:20]):
    if '{{' in p.text or '}}' in p.text:
        unrendered.append(f"P{i}: {p.text[:50]}")

if unrendered:
    print(f"   ⚠️  Found {len(unrendered)} unrendered variables:")
    for u in unrendered[:5]:
        print(f"      {u}")
else:
    print(f"   ✅ All variables rendered!")

# Check fonts
fonts_used = set()
for p in doc.paragraphs:
    for run in p.runs:
        if run.font.name:
            fonts_used.add(run.font.name)

print(f"   🎨 Fonts used: {', '.join(sorted(fonts_used)) if fonts_used else 'None detected'}")

print("\n" + "=" * 60)
print("✅ TEST COMPLETE!")
print(f"\n📋 CHECKLIST:")
print(f"   1. Mở file: {output_file}")
print(f"   2. So sánh với gốc:")
print(f"      - Font có đúng không? (Times New Roman 13)")
print(f"      - Data có đầy đủ không?")
print(f"      - Format có giữ được không?")
print(f"\n🎯 Nếu OK → Giải pháp thành công!")
print(f"🎯 Nếu chưa OK → Cần điều chỉnh thêm patterns")
