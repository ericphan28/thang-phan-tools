#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Create Template V5 với cấu trúc bảng gia đình ĐÚNG
Thay vì 1 loop chung, phải có 4 sections riêng biệt
"""

from docx import Document
from pathlib import Path

# Load V4 template
template_file = Path("mau_2c_template_FINAL_V4.docx")
doc = Document(template_file)

print("=" * 80)
print("🔧 TẠO TEMPLATE V5 - SỬA BẢNG GIA ĐÌNH")
print("=" * 80)

# TABLE 3: Gia đình - RESTRUCTURE
print("\n📋 Bảng 3: Gia đình")
if len(doc.tables) > 2:
    table3 = doc.tables[2]
    if len(table3.rows) > 1:
        row = table3.rows[1]
        
        # Column 0: Keep the structure labels (Bố mẹ, Vợ, Các con, Anh chị em)
        # This column is already correct in the original
        # We DON'T touch it
        
        # Columns 1-3: Split into 4 sections
        # Format cho từng section:
        
        # Section 1: Bố mẹ (first 2 entries)
        # Section 2: Vợ/Chồng (next 1 entry)
        # Section 3: Các con (multiple entries)
        # Section 4: Anh chị em ruột (multiple entries)
        
        col1_template = """{% for member in bo_me %}{{ member.ho_ten }}
{% endfor %}....................
{% for member in vo_chong %}{{ member.ho_ten }}
{% endfor %}


{% for child in cac_con %}{{ child.ho_ten }}
{% endfor %}


{% for sib in anh_chi_em %}{{ sib.ho_ten }}
{% endfor %}"""
        
        col2_template = """{% for member in bo_me %}{{ member.nam_sinh }}
{% endfor %}
{% for member in vo_chong %}{{ member.nam_sinh }}
{% endfor %}


{% for child in cac_con %}{{ child.nam_sinh }}
{% endfor %}


{% for sib in anh_chi_em %}{{ sib.nam_sinh }}
{% endfor %}"""
        
        col3_template = """{% for member in bo_me %}{{ member.thong_tin }}
{% endfor %}
{% for member in vo_chong %}{{ member.thong_tin }}
{% endfor %}


{% for child in cac_con %}{{ child.thong_tin }}
{% endfor %}


{% for sib in anh_chi_em %}{{ sib.thong_tin }}
{% endfor %}"""
        
        row.cells[1].text = col1_template
        row.cells[2].text = col2_template
        row.cells[3].text = col3_template
        
        print("   ✅ Đã sửa: 4 sections (bố mẹ, vợ/chồng, các con, anh chị em)")

# TABLE 4: Gia đình vợ/chồng - RESTRUCTURE
print("\n📋 Bảng 4: Gia đình vợ/chồng")
if len(doc.tables) > 3:
    table4 = doc.tables[3]
    if len(table4.rows) > 1:
        row = table4.rows[1]
        
        # Column 0: Keep structure (Bố mẹ, Anh chị em)
        # Don't touch
        
        # Columns 1-3: Split into 2 sections
        
        col1_template = """{% for member in bo_me_vo_chong %}{{ member.ho_ten }}
{% endfor %}....................


{% for sib in anh_chi_em_vo_chong %}{{ sib.ho_ten }}
{% endfor %}"""
        
        col2_template = """{% for member in bo_me_vo_chong %}{{ member.nam_sinh }}
{% endfor %}


{% for sib in anh_chi_em_vo_chong %}{{ sib.nam_sinh }}
{% endfor %}"""
        
        col3_template = """{% for member in bo_me_vo_chong %}{{ member.thong_tin }}
{% endfor %}


{% for sib in anh_chi_em_vo_chong %}{{ sib.thong_tin }}
{% endfor %}"""
        
        row.cells[1].text = col1_template
        row.cells[2].text = col2_template
        row.cells[3].text = col3_template
        
        print("   ✅ Đã sửa: 2 sections (bố mẹ, anh chị em)")

# Save as V5
output_file = Path("mau_2c_template_FINAL_V5.docx")
doc.save(str(output_file))

file_size = output_file.stat().st_size

print("\n" + "=" * 80)
print("✅ ĐÃ TẠO TEMPLATE V5!")
print("=" * 80)
print(f"📄 File: {output_file}")
print(f"📊 Size: {file_size:,} bytes ({file_size/1024:.2f} KB)")
print()
print("💡 THAY ĐỔI:")
print("   - Bảng 3: 4 sections riêng biệt")
print("     • Bố mẹ (2 người)")
print("     • Vợ/Chồng (1 người)")
print("     • Các con (nhiều người)")
print("     • Anh chị em ruột (nhiều người)")
print()
print("   - Bảng 4: 2 sections riêng biệt")
print("     • Bố mẹ vợ/chồng (2 người)")
print("     • Anh chị em vợ/chồng (nhiều người)")
print()
print("📝 SỬ DỤNG:")
print("   - Template: mau_2c_template_FINAL_V5.docx")
print("   - Data: mau_2c_DATA_RESTRUCTURED.json")
print("=" * 80)
