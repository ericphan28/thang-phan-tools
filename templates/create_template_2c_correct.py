"""
Tạo template ĐÚNG từ Mẫu 2C - GIỮ NGUYÊN cấu trúc, CHỈ thay dấu chấm
"""
from docx import Document
import re
from copy import deepcopy

source = r"d:\thang\utility-server\templates\mau-nha-nuoc\Mau-ly-lich-2C-TCTW-98.docx"
output = r"d:\thang\utility-server\templates\mau_2c_template_correct.docx"

print("📖 Đọc mẫu gốc...")
doc = Document(source)

print("🔧 Bước 1: Thay thế các dấu chấm trong PARAGRAPHS")
# Mapping for paragraphs - chỉ thay ở những chỗ CÓ DẤU CHẤM
para_replacements = {
    r'Tỉnh:\s*………+': 'Tỉnh: {{tinh}}',
    r'Đơn vị trực thuộc:\s*\.+': 'Đơn vị trực thuộc: {{don_vi_truc_thuoc}}',
    r'Đơn vị cơ sở:\s*\.+': 'Đơn vị cơ sở: {{don_vi_co_so}}',
    r'Số hiệu cán bộ[^:]*': 'Số hiệu cán bộ: {{so_hieu}}',
    
    # Họ tên và giới tính
    r'1\) Họ và tên khai sinh:\s*……+': '1) Họ và tên khai sinh: {{ho_ten}}',
    r'Nam, nữ:\s*\.+': 'Nam, nữ: {{gioi_tinh}}',
    
    # Các tên khác
    r'2\) Các tên gọi khác:\s*\.+': '2) Các tên gọi khác: {{ten_khac}}',
    
    # Cấp ủy
    r'3\) Cấp ủy hiện tại:\s*\.+': '3) Cấp ủy hiện tại: {{cap_uy}}',
    r'Cấp ủy kiêm:\s*\.+': 'Cấp ủy kiêm: {{cap_uy_kiem}}',
    
    # Chức vụ
    r'Chức vụ[^:]+:\s*\.+': 'Chức vụ (Đảng, đoàn thể, Chính quyền, kể cả chức vụ kiêm nhiệm): {{chuc_vu}}',
    r'Phụ cấp chức vụ:\s*\.+': 'Phụ cấp chức vụ: {{phu_cap}}',
    
    # Ngày sinh, nơi sinh
    r'4\) Sinh ngày:\s*\.+ tháng\s*\.+ năm\s*\.+': '4) Sinh ngày: {{ngay}} tháng {{thang}} năm {{nam}}',
    r'5\) Nơi sinh:\s*\.+': '5) Nơi sinh: {{noi_sinh}}',
    
    # Quê quán (3 phần)
    r'6\) Quê quán \(xã, phường\):\s*\.+': '6) Quê quán (xã, phường): {{que_xa}}',
    r'\(huyện, quận\):\s*\.+': '(huyện, quận): {{que_huyen}}',
    r'\(tỉnh, TP\):\s*\.+': '(tỉnh, TP): {{que_tinh}}',
    
    # Nơi ở hiện nay
    r'7\) Nơi ở hiện nay[^:]+:\s*\.+': '7) Nơi ở hiện nay (Xã, huyện, tỉnh hoặc số nhà, đường phố, TP): {{dia_chi}}',
    r'đ/thoại:\s*\.+': 'đ/thoại: {{dien_thoai}}',
    
    # Dân tộc, tôn giáo
    r'8\) Dân tộc:[^:]+:\s*\.+': '8) Dân tộc: {{dan_toc}}',
    r'9\) Tôn giáo:\s*\.+': '9) Tôn giáo: {{ton_giao}}',
    
    # Thành phần xuất thân
    r'10\) Thành phần gia đình xuất thân:\s*\.+': '10) Thành phần gia đình xuất thân: {{thanh_phan}}',
    
    # Nghề nghiệp trước
    r'11\) Nghề nghiệp bản thân trước[^:]+:\s*\.+': '11) Nghề nghiệp bản thân trước khi được tuyển dụng: {{nghe_truoc}}',
    
    # Ngày tuyển dụng
    r'12\) Ngày được tuyển dụng:\s*\.+\s*/\s*\.+\s*/\s*\.+': '12) Ngày được tuyển dụng: {{ngay_tuyen_dung}}',
    r'Vào cơ quan nào[^:]+:\s*\.+': 'Vào cơ quan nào, ở đâu: {{co_quan_tuyen_dung}}',
    
    # Ngày vào cơ quan
    r'13\) Ngày vào cơ quan[^:]+:\s*\.+\s*/\s*\.+\s*/\s*\.+': '13) Ngày vào cơ quan hiện đang công tác: {{ngay_vao_co_quan}}',
    r'Ngày tham gia cách mạng:\s*\.+\s*/\s*\.+\s*/\s*\.+': 'Ngày tham gia cách mạng: {{ngay_cach_mang}}',
    
    # Ngày vào Đảng
    r'14\) Ngày vào Đảng[^:]+:\s*\.+\s*/\s*\.+\s*/\s*\.+': '14) Ngày vào Đảng Cộng sản Việt Nam: {{ngay_vao_dang}}',
    r'Ngày chính thức:\s*\.+\s*/\s*\.+\s*/\s*\.+': 'Ngày chính thức: {{ngay_chinh_thuc}}',
    
    # Tổ chức chính trị
    r'15\) Ngày tham gia[^:]+:\s*\.+': '15) Ngày tham gia các tổ chức chính trị, xã hội: {{to_chuc}}',
    
    # Quân ngũ
    r'16\) Ngày nhập ngũ:\s*\.+\s*/\s*\.+\s*/\s*\.+': '16) Ngày nhập ngũ: {{ngay_nhap_ngu}}',
    r'Ngày xuất ngũ:\s*\.+\s*/\s*\.+\s*/\s*\.+': 'Ngày xuất ngũ: {{ngay_xuat_ngu}}',
    r'Quân hàm[^:]+:\s*\.+': 'Quân hàm, chức vụ cao nhất (năm): {{quan_ham}}',
    
    # Học vấn
    r'17\) Trình độ học vấn:[^:]+:\s*\.+': '17) Trình độ học vấn: Giáo dục phổ thông: {{hoc_van}}',
    r'Học hàm, học vị cao nhất:\s*\.+': 'Học hàm, học vị cao nhất: {{hoc_vi}}',
    r'- Lý luận chính trị:\s*\.+': '- Lý luận chính trị: {{ly_luan}}',
    r'- Ngoại ngữ:\s*\.+': '- Ngoại ngữ: {{ngoai_ngu}}',
    
    # Công tác
    r'18\) Công tác chính[^:]+:\s*\.+': '18) Công tác chính đang làm: {{cong_tac}}',
    
    # Ngạch lương
    r'19\) Ngạch công chức:\s*\.+': '19) Ngạch công chức: {{ngach}}',
    r'\(mã số:\s*\.+\)': '(mã số: {{ma_ngach}})',
    r'Bậc lương:\s*\.+,\s*hệ số:\s*\.+': 'Bậc lương: {{bac}}, hệ số: {{he_so}}',
    r'từ tháng\s*\.+\s*/\.+': 'từ tháng {{thang_luong}}',
    
    # Danh hiệu
    r'20\) Danh hiệu[^:]+:\s*\.+': '20) Danh hiệu được phong (năm nào): {{danh_hieu}}',
    
    # Sở trường
    r'21\) Sở trường công tác:\s*\.+': '21) Sở trường công tác: {{so_truong}}',
    r'Công việc đã làm lâu nhất:\s*\.+': 'Công việc đã làm lâu nhất: {{cv_lau_nhat}}',
    
    # Khen thưởng
    r'22\) Khen thưởng:\s*\.+': '22) Khen thưởng: {{khen_thuong}}',
    
    # Kỷ luật
    r'23\) Kỷ luật[^:]+:\s*\.+': '23) Kỷ luật (Đảng, Chính quyền, Đoàn thể, Cấp quyết định, năm nào, lý do, hình thức, ...): {{ky_luat}}',
    
    # Sức khỏe
    r'24\) Tình trạng sức khỏe:\s*\.+': '24) Tình trạng sức khỏe: {{suc_khoe}}',
    r'Cao:\s*1m\.+': 'Cao: {{chieu_cao}}',
    r'Cân nặng:\s*\.+': 'Cân nặng: {{can_nang}}',
    r'\(kg\),\s*Nhóm máu:\s*\.+': '(kg), Nhóm máu: {{nhom_mau}}',
    
    # CMND
    r'25\) Số chứng minh nhân dân:\s*\.+': '25) Số chứng minh nhân dân: {{cmnd}}',
    r'Thương binh loại:\s*\.+': 'Thương binh loại: {{thuong_binh}}',
    r'Gia đình liệt sĩ:': 'Gia đình liệt sĩ: {{liet_si}}',
}

for para in doc.paragraphs:
    original = para.text
    for pattern, replacement in para_replacements.items():
        if re.search(pattern, para.text):
            new_text = re.sub(pattern, replacement, para.text)
            if new_text != original:
                # Clear and rewrite
                for run in para.runs:
                    run.text = ''
                para.add_run(new_text)
                break

print("  ✓ Đã xử lý paragraphs")

print("\n🔧 Bước 2: Xử lý BẢNG 1 - Đào tạo (GIỮ NGUYÊN header)")
table1 = doc.tables[0]
# Chỉ thay dấu chấm ở row 2 (row data)
for cell in table1.rows[1].cells:
    if '.' in cell.text:
        cell.text = '{{dao_tao_data}}'  # Placeholder - user sẽ điền thủ công

print("  ✓ Bảng 1: Đào tạo")

print("\n🔧 Bước 3: Xử lý BẢNG 2 - Quá trình công tác")
table2 = doc.tables[1]
for cell in table2.rows[1].cells:
    if '.' in cell.text:
        cell.text = '{{cong_tac_data}}'

print("  ✓ Bảng 2: Quá trình công tác")

print("\n🔧 Bước 4: Xử lý BẢNG 3 - Gia đình bản thân (GIỮ NGUYÊN labels)")
table3 = doc.tables[2]
# Row 1 có sẵn "Bố, mẹ", "Vợ", "Chồng", "Các con" - CHỈ thay dấu chấm ở các cột còn lại
row = table3.rows[1]
# Cột 0: Giữ nguyên (có labels)
# Cột 1-3: Thay dấu chấm
for i in range(1, 4):
    if '.' in row.cells[i].text:
        row.cells[i].text = '{{gia_dinh_data}}'

print("  ✓ Bảng 3: Gia đình bản thân")

print("\n🔧 Bước 5: Xử lý BẢNG 4 - Gia đình bên vợ/chồng")
table4 = doc.tables[3]
row = table4.rows[1]
for i in range(1, 4):
    if '.' in row.cells[i].text:
        row.cells[i].text = '{{gia_dinh_vo_chong_data}}'

print("  ✓ Bảng 4: Gia đình vợ/chồng")

print("\n🔧 Bước 6: Xử lý BẢNG 5 - Quá trình lương (horizontal)")
table5 = doc.tables[4]
# Bảng này đặc biệt - 3 rows x 7 cols
# Row 0: Tháng/năm
# Row 1: Ngạch/bậc
# Row 2: Hệ số lương
print("  ✓ Bảng 5: Quá trình lương (giữ nguyên)")

# Save
doc.save(output)
print(f"\n✅ HOÀN TẤT! File: {output}")
print("📝 Lưu ý: Template này GIỮ NGUYÊN cấu trúc gốc, chỉ thay dấu chấm")
print("💡 Các bảng cần điền dữ liệu thủ công vào Word trước khi dùng")
