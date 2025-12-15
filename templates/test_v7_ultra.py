#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
TEST V7 ULTRA PRECISE
"""

from docxtpl import DocxTemplate
import json
from docx import Document

def flatten_dict(d, parent_key='', sep='_'):
    """Flatten nested dict"""
    items = []
    for k, v in d.items():
        if k.startswith('_'):
            continue
        new_key = f"{parent_key}{sep}{k}" if parent_key else k
        if isinstance(v, dict):
            items.extend(flatten_dict(v, new_key, sep=sep).items())
        else:
            items.append((new_key, v))
    return dict(items)

def main():
    TEMPLATE = "mau_2c_V7_ULTRA_PRECISE.docx"
    OUTPUT = "OUTPUT_V7_ULTRA_PRECISE.docx"
    JSON_FILE = "mau_2c_DATA_RESTRUCTURED.json"
    
    print(f"📖 Loading template: {TEMPLATE}")
    doc = DocxTemplate(TEMPLATE)
    
    print(f"📊 Loading data: {JSON_FILE}")
    with open(JSON_FILE, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # Flatten
    flat_data = flatten_dict(data)
    
    # Keep arrays
    render_data = {**flat_data}
    if 'hoc_tap' in data:
        render_data['hoc_tap'] = data['hoc_tap']
    if 'cong_tac' in data:
        render_data['cong_tac'] = data['cong_tac']
    if 'gia_dinh' in data:
        render_data['gia_dinh'] = data['gia_dinh']
    
    print(f"✅ Loaded {len(render_data)} fields")
    
    print(f"🎨 Rendering...")
    doc.render(render_data)
    
    print(f"💾 Saving: {OUTPUT}")
    doc.save(OUTPUT)
    
    # Check dots
    print(f"\n🔍 Checking for dots...")
    output_doc = Document(OUTPUT)
    dot_count = 0
    dot_paragraphs = []
    
    for i, para in enumerate(output_doc.paragraphs):
        text = para.text
        # Count dots (3+ consecutive)
        if '...' in text or '…..' in text or '……' in text:
            dot_count += 1
            if len(dot_paragraphs) < 15:  # Show first 15
                dot_paragraphs.append((i, text[:120]))
    
    print(f"\n📊 RESULTS:")
    print(f"   Total paragraphs: {len(output_doc.paragraphs)}")
    print(f"   Paragraphs with dots: {dot_count}")
    print(f"   Coverage: {100 - (dot_count/len(output_doc.paragraphs)*100):.1f}%")
    
    if dot_paragraphs:
        print(f"\n⚠️  Sample dots (first 15):")
        for i, text in dot_paragraphs:
            print(f"   P{i}: {text}")
    
    print(f"\n{'='*60}")
    if dot_count == 0:
        print(f"🎉 PERFECT! 100% - Không còn dấu chấm!")
    elif dot_count < 10:
        print(f"✅ EXCELLENT! {100 - (dot_count/len(output_doc.paragraphs)*100):.1f}% - Chỉ còn {dot_count} chỗ!")
    elif dot_count < 30:
        print(f"✅ VERY GOOD! {100 - (dot_count/len(output_doc.paragraphs)*100):.1f}% - Còn {dot_count} chỗ cần improve")
    else:
        print(f"⚠️  GOOD! {100 - (dot_count/len(output_doc.paragraphs)*100):.1f}% - Còn {dot_count} chỗ")
    
    print(f"\n✅ File: {OUTPUT}")

if __name__ == "__main__":
    main()
