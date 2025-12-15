#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
V9 - FIXED MAPPING
==================
Fix variable names to match JSON keys:
- ngach_cong_chuc (not nguoi_cong_chuc_vien_chuc)
- ma_ngach (not ma_so)
- he_so_luong (not he_so)
- tu_thang_nam (not tu_thang)
"""

from docx import Document
from copy import deepcopy

def merge_runs_in_paragraph(paragraph):
    """Merge all runs thành 1 run duy nhất"""
    if len(paragraph.runs) <= 1:
        return False
    
    full_text = ''.join([run.text for run in paragraph.runs])
    first_run = paragraph.runs[0]
    
    for i in range(len(paragraph.runs) - 1, 0, -1):
        paragraph._element.remove(paragraph.runs[i]._element)
    
    first_run.text = full_text
    return True

def replace_exact_in_paragraph(paragraph, old_text, new_text):
    """Replace EXACT text in paragraph"""
    if old_text in paragraph.text:
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                return True
    return False

def process_document():
    INPUT = "mau-nha-nuoc/Mau-ly-lich-2C-TCTW-98.docx"
    OUTPUT = "mau_2c_V9_FINAL_TEMPLATE.docx"
    
    print(f"📖 Loading: {INPUT}")
    doc = Document(INPUT)
    
    print(f"🔄 Step 1: Merge runs...")
    merged_count = 0
    for para in doc.paragraphs:
        if merge_runs_in_paragraph(para):
            merged_count += 1
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if merge_runs_in_paragraph(para):
                        merged_count += 1
    print(f"   ✅ Merged {merged_count} locations")
    
    print(f"\n🔄 Step 2: Replace with CORRECT variable names...")
    
    replacements = [
        # Header
        ("Tỉnh: ………………….", "Tỉnh: {{ tinh }}"),
        ("Đơn vị trực thuộc: ..........................", "Đơn vị trực thuộc: {{ don_vi_truc_thuoc }}"),
        ("Đơn vị cơ sở: ................................", "Đơn vị cơ sở: {{ don_vi_co_so }}"),
        ("Số hiệu cán bộ, công chức", "Số hiệu: {{ so_hieu }}"),
        
        # Personal info
        ("1) Họ và tên khai sinh: ……………………………………..", "1) Họ và tên khai sinh: {{ ho_ten }}"),
        ("Nam, nữ: .....................", "Nam, nữ: {{ gioi_tinh }}"),
        ("2) Các tên gọi khác:", "2) Các tên gọi khác: {{ ten_goi_khac }}"),
        ("3) Cấp ủy hiện tại: .......................................", "3) Cấp ủy hiện tại: {{ cap_uy_hien_tai }}"),
        ("Cấp ủy kiêm: .........................................", "Cấp ủy kiêm: {{ cap_uy_kiem }}"),
        ("Chức vụ (Đảng, đoàn thể, Chính quyền, kể cả chức vụ kiêm nhiệm):", "Chức vụ: {{ chuc_vu_full }}"),
        ("Phụ cấp chức vụ: ...........................", "Phụ cấp chức vụ: {{ phu_cap_chuc_vu }}"),
        
        # Birth
        ("4) Sinh ngày: ..........", "4) Sinh ngày: {{ ngay }}"),
        ("tháng ..........", "tháng {{ thang }}"),
        ("năm ...............", "năm {{ nam }}"),
        ("5) Nơi sinh: ..................................................", "5) Nơi sinh: {{ noi_sinh }}"),
        
        # Origin
        ("6) Quê quán (xã, phường): .......................................", "6) Quê quán (xã, phường): {{ que_quan_xa }}"),
        ("(huyện, quận): ........................", "(huyện, quận): {{ que_quan_huyen }}"),
        ("(tỉnh, TP): ...............................", "(tỉnh, TP): {{ que_quan_tinh }}"),
        
        # Contact
        ("7) Nơi ở hiện nay (Xã, huyện, tỉnh hoặc số nhà, đường phố, TP): ..............................................", "7) Nơi ở hiện nay: {{ noi_o_hien_nay }}"),
        ("đ/thoại: ....................", "đ/thoại: {{ dien_thoai }}"),
        
        # Background
        ("8) Dân tộc: (Kinh, Tày, Mông, Ê đê...): ..............................", "8) Dân tộc: {{ dan_toc }}"),
        ("9) Tôn giáo: ......................................................", "9) Tôn giáo: {{ ton_giao }}"),
        ("10) Thành phần gia đình xuất thân:", "10) Thành phần gia đình xuất thân: {{ thanh_phan_xuat_than }}"),
        ("11) Nghề nghiệp bản thân trước khi được tuyển dụng:", "11) Nghề nghiệp bản thân: {{ nghe_nghiep_ban_than }}"),
        
        # Career dates
        ("12) Ngày được tuyển dụng:", "12) Ngày được tuyển dụng: {{ ngay_tuyen_dung }}"),
        ("Vào cơ quan nào, ở đâu:", "Vào cơ quan: {{ co_quan_tuyen_dung }}"),
        ("13) Ngày vào cơ quan hiện đang công tác:", "13) Ngày vào cơ quan: {{ ngay_vao_co_quan }}"),
        ("Ngày tham gia cách mạng:", "Ngày tham gia cách mạng: {{ ngay_tham_gia_cach_mang }}"),
        ("14) Ngày vào Đảng Cộng sản Việt Nam:", "14) Ngày vào Đảng: {{ ngay_vao_dang }}"),
        ("Ngày chính thức:", "Ngày chính thức: {{ ngay_chinh_thuc_dang }}"),
        ("15) Ngày tham gia các tổ chức chính trị, xã hội:", "15) Tổ chức: {{ ngay_tham_gia_to_chuc }}"),
        ("16) Ngày nhập ngũ:", "16) Ngày nhập ngũ: {{ ngay_nhap_ngu }}"),
        ("Ngày xuất ngũ:", "Ngày xuất ngũ: {{ ngay_xuat_ngu }}"),
        ("Quân hàm, chức vụ cao nhất (năm):", "Quân hàm: {{ quan_ham }}"),
        
        # Education
        ("17)Trình độ học vấn: Giáo dục phổ thông:", "17) Giáo dục phổ thông: {{ trinh_do_giao_duc_pho_thong }}"),
        ("Học hàm, học vị cao nhất:", "Học hàm, học vị: {{ hoc_ham_hoc_vi }}"),
        ("- Lý luận chính trị:", "- Lý luận chính trị: {{ ly_luan_chinh_tri }}"),
        ("- Ngoại ngữ:", "- Ngoại ngữ: {{ ngoai_ngu }}"),
        
        # Work - FIXED MAPPING!
        ("18) Công tác chính đảng làm:", "18) Công tác: {{ cong_tac_chinh }}"),
        ("19) Ngạch công chức:", "19) Ngạch: {{ ngach_cong_chuc }}"),  # FIXED!
        ("(mã số:", "(mã số: {{ ma_ngach }}"),  # FIXED!
        ("Bậc lương:", "Bậc lương: {{ bac_luong }}"),
        ("hệ số:", "hệ số: {{ he_so_luong }}"),  # FIXED!
        ("từ tháng", "từ tháng {{ tu_thang_nam }}"),  # FIXED!
        
        # More work details
        ("20) Danh hiệu được phong (năm nào):", "20) Danh hiệu: {{ danh_hieu }}"),
        ("21) Sở trường công tác:", "21) Sở trường: {{ so_truong_cong_tac }}"),
        ("Công việc đã làm lâu nhất:", "Công việc lâu nhất: {{ cong_viec_lau_nhat }}"),
        ("22) Khen thưởng:", "22) Khen thưởng: {{ khen_thuong }}"),
        ("23) Kỷ luật (Đảng, Chính quyền, Đoàn thể, Cấp quyết định, năm nào, lý do, hình thức, ):", "23) Kỷ luật: {{ ky_luat }}"),
        
        # Health
        ("24) Tình trạng sức khỏe:", "24) Sức khỏe: {{ suc_khoe }}"),
        ("Cao:", "Cao: {{ chieu_cao }}"),
        ("Cân nặng:", "Cân nặng: {{ can_nang }}"),
        ("Nhóm máu:", "Nhóm máu: {{ nhom_mau }}"),
        
        # ID
        ("25) Số chứng minh nhân dân:", "25) Số CMND: {{ so_cmnd }}"),
        ("Ngày cấp:", "Ngày cấp: {{ ngay_cap }}"),
        ("Nơi cấp:", "Nơi cấp: {{ noi_cap }}"),
        ("Thương binh loại:", "Thương binh: {{ thuong_binh_loai }}"),
        ("Gia đình liệt sĩ:", "Gia đình liệt sĩ: {{ gia_dinh_liet_si }}"),
        
        # History
        ("28) ĐẶC ĐIỂM LỊCH SỬ BẢN THÂN", "28) ĐẶC ĐIỂM LỊCH SỬ BẢN THÂN"),
        ("a) Khai rõ: bị, bị, bị (tự ạy đoàng mấm nâo đến điếp đoàng mấm nâo, ở đẩu), dẫ khai báo cho aỉ, những văn đế", 
         "a) Bị bắt, bị tù: {{ lich_su_bi_bat }}"),
        ("b) Bản thân có làm việc trong chế độ cũ (Cơ quan, đơn vị nào, địa điểm, chức danh, chức vụ, thời gian làm việc ...)", 
         "b) Làm việc chế độ cũ: {{ lam_viec_che_do_cu }}"),
        
        # Foreign relations
        ("29) QUAN HỆ VỚI NƯỚC NGOÀI", "29) QUAN HỆ VỚI NƯỚC NGOÀI"),
        ("- Tham gia hoặc cơ quan hệ với cơ sở tổ chức nào ở nước ngoại (tầm giá đề, địa chỉ dậu, diếp tốn chưng và hoạt đang cụy số)", 
         "{{ quan_he_nuoc_ngoai }}"),
        ("- Có thân nhân (Bả, mấ và đũmg cha, sãm đả nói rõ) ở nước ngoại (tầm giá đề ăí ...)?", 
         "{{ than_nhan_nuoc_ngoai }}"),
        
        # Family
        ("30) VỀ GIA ĐÌNH", "30) VỀ GIA ĐÌNH"),
        ("31) HOÀN CẢNH KINH TẾ GIA ĐÌNH", "31) HOÀN CẢNH KINH TẾ GIA ĐÌNH"),
    ]
    
    replace_count = 0
    for i, para in enumerate(doc.paragraphs):
        for old, new in replacements:
            if replace_exact_in_paragraph(para, old, new):
                print(f"✅ P{i}: {new[:70]}")
                replace_count += 1
    
    # Tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    for old, new in replacements:
                        if replace_exact_in_paragraph(para, old, new):
                            print(f"✅ T{table_idx}R{row_idx}C{cell_idx}: {new[:50]}")
                            replace_count += 1
    
    # Save
    doc.save(OUTPUT)
    
    print(f"\n{'='*60}")
    print(f"📄 Output: {OUTPUT}")
    print(f"🔄 Total replacements: {replace_count}")
    print(f"✅ V9 COMPLETE with FIXED MAPPING!")
    print(f"\n🎯 Key fixes:")
    print(f"   - ngach_cong_chuc (not nguoi_cong_chuc_vien_chuc)")
    print(f"   - ma_ngach (not ma_so)")
    print(f"   - he_so_luong (not he_so)")
    print(f"   - tu_thang_nam (not tu_thang)")

if __name__ == "__main__":
    process_document()
