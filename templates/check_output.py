"""Check what's in OUTPUT_MAILMERGE.docx"""
from docx import Document
import os

doc = Document('OUTPUT_MAILMERGE.docx')

print(f"📄 OUTPUT_MAILMERGE.docx")
print(f"Size: {os.path.getsize('OUTPUT_MAILMERGE.docx'):,} bytes")
print(f"\n📝 First 20 paragraphs with content:")

count = 0
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"P{i}: {p.text[:150]}")
        count += 1
        if count >= 20:
            break

print(f"\n📊 Table 3 (Family) content:")
if len(doc.tables) >= 3:
    table = doc.tables[2]
    print(f"Rows: {len(table.rows)}")
    for i, row in enumerate(table.rows[:5]):
        print(f"Row {i}: {[cell.text[:50] for cell in row.cells]}")
