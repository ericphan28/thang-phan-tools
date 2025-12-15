#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
V10 - ADD JINJA2 LOOPS TO TABLES
=================================
Manually add Jinja2 loops to table templates
"""

from docx import Document

def add_loops_to_tables():
    """
    Add Jinja2 loops to tables in V9 template
    """
    INPUT = "mau_2c_V9_FINAL_TEMPLATE.docx"
    OUTPUT = "mau_2c_V10_WITH_LOOPS.docx"
    
    doc = Document(INPUT)
    
    print(f"📖 Loading: {INPUT}")
    print(f"📊 Found {len(doc.tables)} tables\n")
    
    # TABLE 1: Học tập (26) Đào tạo, bồi dưỡng
    print("🔄 TABLE 1: Adding hoc_tap loop...")
    table1 = doc.tables[0]
    row2 = table1.rows[1]
    
    # Clear dots and add Jinja2 loop
    row2.cells[0].text = "{% for ht in hoc_tap %}{{ ht.ten_truong }}{% if not loop.last %}\n{% endif %}{% endfor %}"
    row2.cells[1].text = "{% for ht in hoc_tap %}{{ ht.nganh_hoc }}{% if not loop.last %}\n{% endif %}{% endfor %}"
    row2.cells[2].text = "{% for ht in hoc_tap %}{{ ht.thoi_gian }}{% if not loop.last %}\n{% endif %}{% endfor %}"
    row2.cells[3].text = "{% for ht in hoc_tap %}{{ ht.hinh_thuc }}{% if not loop.last %}\n{% endif %}{% endfor %}"
    row2.cells[4].text = "{% for ht in hoc_tap %}{{ ht.van_bang }}{% if not loop.last %}\n{% endif %}{% endfor %}"
    print("   ✅ Done")
    
    # TABLE 2: Công tác (27) Tóm tắt quá trình công tác
    print("🔄 TABLE 2: Adding cong_tac loop...")
    table2 = doc.tables[1]
    row2 = table2.rows[1]
    
    row2.cells[0].text = "{% for ct in cong_tac %}{{ ct.tu_thang_nam }}{% if ct.den_thang_nam %} - {{ ct.den_thang_nam }}{% endif %}{% if not loop.last %}\n{% endif %}{% endfor %}"
    row2.cells[1].text = "{% for ct in cong_tac %}{{ ct.chuc_danh }}{% if not loop.last %}\n{% endif %}{% endfor %}"
    print("   ✅ Done")
    
    # TABLE 3: Gia đình bản thân (30a)
    print("🔄 TABLE 3: Adding gia_dinh.ban_than loop...")
    table3 = doc.tables[2]
    row2 = table3.rows[1]
    
    # Family - bố mẹ section
    family_text = """Bố, mẹ
{% for gd in gia_dinh.ban_than %}{% if gd.quan_he in ['Bố', 'Mẹ'] %}{{ gd.quan_he }}: {{ gd.ho_ten }}
{% endif %}{% endfor %}
Vợ
Chồng
{% for gd in gia_dinh.ban_than %}{% if gd.quan_he in ['Vợ', 'Chồng'] %}{{ gd.quan_he }}: {{ gd.ho_ten }}
{% endif %}{% endfor %}

Các con:
{% for gd in gia_dinh.ban_than %}{% if gd.quan_he == 'Con' %}{{ gd.ho_ten }}
{% endif %}{% endfor %}

Anh chị em ruột
{% for gd in gia_dinh.ban_than %}{% if gd.quan_he in ['Anh', 'Chị', 'Em'] %}{{ gd.ho_ten }}
{% endif %}{% endfor %}"""
    
    row2.cells[0].text = "{% for gd in gia_dinh.ban_than %}{{ gd.quan_he }}{% if not loop.last %}\n{% endif %}{% endfor %}"
    row2.cells[1].text = "{% for gd in gia_dinh.ban_than %}{{ gd.ho_ten }}{% if not loop.last %}\n{% endif %}{% endfor %}"
    row2.cells[2].text = "{% for gd in gia_dinh.ban_than %}{{ gd.nam_sinh }}{% if not loop.last %}\n{% endif %}{% endfor %}"
    row2.cells[3].text = "{% for gd in gia_dinh.ban_than %}{{ gd.thong_tin }}{% if not loop.last %}\n{% endif %}{% endfor %}"
    print("   ✅ Done")
    
    # TABLE 4: Gia đình vợ/chồng (30b)
    print("🔄 TABLE 4: Adding gia_dinh.vo_chong loop...")
    table4 = doc.tables[3]
    row2 = table4.rows[1]
    
    row2.cells[0].text = "{% for gd in gia_dinh.vo_chong %}{{ gd.quan_he }}{% if not loop.last %}\n{% endif %}{% endfor %}"
    row2.cells[1].text = "{% for gd in gia_dinh.vo_chong %}{{ gd.ho_ten }}{% if not loop.last %}\n{% endif %}{% endfor %}"
    row2.cells[2].text = "{% for gd in gia_dinh.vo_chong %}{{ gd.nam_sinh }}{% if not loop.last %}\n{% endif %}{% endfor %}"
    row2.cells[3].text = "{% for gd in gia_dinh.vo_chong %}{{ gd.thong_tin }}{% if not loop.last %}\n{% endif %}{% endfor %}"
    print("   ✅ Done")
    
    # TABLE 5: Hoàn cảnh kinh tế (31)
    print("🔄 TABLE 5: Adding he_so_luong_history loop...")
    table5 = doc.tables[4]
    
    # Row 3 needs loop for salary history
    if len(table5.rows) >= 3:
        row3 = table5.rows[2]
        row3.cells[0].text = "{% if he_so_luong_history %}{% for hs in he_so_luong_history %}{{ hs.tu_thang_nam }}{% if not loop.last %}\n{% endif %}{% endfor %}{% endif %}"
        row3.cells[1].text = "{% if he_so_luong_history %}{% for hs in he_so_luong_history %}{{ hs.chuc_danh }}, Bậc {{ hs.bac }}{% if not loop.last %}\n{% endif %}{% endfor %}{% endif %}"
        row3.cells[2].text = "{% if he_so_luong_history %}{% for hs in he_so_luong_history %}{{ hs.he_so }}{% if not loop.last %}\n{% endif %}{% endfor %}{% endif %}"
    print("   ✅ Done")
    
    # Save
    doc.save(OUTPUT)
    
    print(f"\n{'='*60}")
    print(f"📄 Output: {OUTPUT}")
    print(f"✅ V10 COMPLETE with Jinja2 loops in all tables!")
    print(f"\n🎯 Next: python test_v10_with_loops.py")

if __name__ == "__main__":
    add_loops_to_tables()
