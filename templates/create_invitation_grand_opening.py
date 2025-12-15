"""
Tạo thiệp mời KHAI TRƯƠNG chuyên nghiệp
- Màu gold + red sang trọng
- Border đẹp
- Logo position
- Formal layout
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_border(section, color='C41E3A', width='24'):
    """Thêm viền trang màu đỏ sang trọng"""
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'text')
    
    for border_name in ('top', 'left', 'bottom', 'right'):
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'double')  # Double border cho sang
        border_el.set(qn('w:sz'), width)
        border_el.set(qn('w:space'), '24')
        border_el.set(qn('w:color'), color)
        pgBorders.append(border_el)
    
    sectPr.append(pgBorders)

def add_shading(paragraph, color):
    """Thêm background color cho paragraph"""
    pPr = paragraph._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    pPr.append(shd)

# Tạo document
doc = Document()

# Setup margins - narrow để có nhiều space cho design
section = doc.sections[0]
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(2)
section.right_margin = Cm(2)

# Page size A5 (phổ biến cho thiệp mời)
section.page_width = Cm(14.8)
section.page_height = Cm(21)

# Thêm viền đỏ đậm
add_page_border(section, color='C41E3A', width='36')

# ================ LOGO SPACE ================
logo_space = doc.add_paragraph()
logo_space.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = logo_space.add_run('[LOGO CÔNG TY]')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)
run.italic = True
doc.add_paragraph()  # Space

# ================ DECORATIVE TOP ================
decor1 = doc.add_paragraph()
decor1.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = decor1.add_run('✦ ✦ ✦ ✦ ✦')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(255, 215, 0)  # Gold

doc.add_paragraph()

# ================ MAIN TITLE ================
title1 = doc.add_paragraph()
title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_shading(title1, 'C41E3A')  # Red background
run = title1.add_run('TRÂN TRỌNG KÍNH MỜI')
run.bold = True
run.font.size = Pt(18)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(255, 255, 255)  # White text

doc.add_paragraph()

# ================ DECORATIVE LINE ================
line1 = doc.add_paragraph()
line1.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = line1.add_run('═══════════════════════')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(255, 215, 0)  # Gold

doc.add_paragraph()

# ================ GUEST INFO ================
guest_label = doc.add_paragraph()
guest_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = guest_label.add_run('Quý khách:')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

guest_name = doc.add_paragraph()
guest_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = guest_name.add_run('{{guest.name}}')
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(196, 30, 58)  # Red

guest_title = doc.add_paragraph()
guest_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = guest_title.add_run('{{guest.title}}')
run.italic = True
run.font.size = Pt(11)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# ================ EVENT INFO ================
event_intro = doc.add_paragraph()
event_intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = event_intro.add_run('Tham dự buổi lễ khai trương')
run.font.size = Pt(13)
run.font.name = 'Times New Roman'

event_name = doc.add_paragraph()
event_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = event_name.add_run('{{business.name}}')
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(196, 30, 58)  # Red

doc.add_paragraph()

# ================ DECORATIVE LINE 2 ================
line2 = doc.add_paragraph()
line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = line2.add_run('❈ ❈ ❈')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(255, 215, 0)  # Gold

doc.add_paragraph()

# ================ DETAILS BOX ================
# Venue
venue = doc.add_paragraph()
venue.alignment = WD_ALIGN_PARAGRAPH.CENTER
run1 = venue.add_run('🏛️  ')
run1.font.size = Pt(14)
run2 = venue.add_run('Địa điểm: ')
run2.bold = True
run2.font.size = Pt(12)
run2.font.name = 'Times New Roman'
run3 = venue.add_run('{{venue.address}}')
run3.font.size = Pt(12)
run3.font.name = 'Times New Roman'

# Date Time
datetime = doc.add_paragraph()
datetime.alignment = WD_ALIGN_PARAGRAPH.CENTER
run1 = datetime.add_run('📅  ')
run1.font.size = Pt(14)
run2 = datetime.add_run('Thời gian: ')
run2.bold = True
run2.font.size = Pt(12)
run2.font.name = 'Times New Roman'
run3 = datetime.add_run('{{event.datetime}}')
run3.font.size = Pt(12)
run3.font.name = 'Times New Roman'
run3.font.color.rgb = RGBColor(196, 30, 58)

doc.add_paragraph()

# ================ PROGRAM ================
program_title = doc.add_paragraph()
program_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_shading(program_title, 'FFD700')  # Gold background
run = program_title.add_run('🎁  CHƯƠNG TRÌNH')
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(196, 30, 58)

program_list = doc.add_paragraph()
program_list.alignment = WD_ALIGN_PARAGRAPH.CENTER
program_list.paragraph_format.left_indent = Cm(3)
program_list.paragraph_format.right_indent = Cm(3)

# Static program (có thể thay bằng loop)
program_text = """• Lễ cắt băng khai trương
• Tiệc buffet
• Tham quan showroom
• Quà tặng tri ân"""

run = program_list.add_run(program_text)
run.font.size = Pt(11)
run.font.name = 'Times New Roman'

doc.add_paragraph()
doc.add_paragraph()

# ================ RSVP ================
rsvp = doc.add_paragraph()
rsvp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run1 = rsvp.add_run('📞  Liên hệ: ')
run1.font.size = Pt(11)
run1.font.name = 'Times New Roman'
run2 = rsvp.add_run('{{contact.phone}}')
run2.bold = True
run2.font.size = Pt(11)
run2.font.name = 'Times New Roman'
run2.font.color.rgb = RGBColor(196, 30, 58)

rsvp2 = doc.add_paragraph()
rsvp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run1 = rsvp2.add_run('📧  ')
run1.font.size = Pt(11)
run2 = rsvp2.add_run('{{contact.email}}')
run2.font.size = Pt(10)
run2.font.name = 'Times New Roman'
run2.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
doc.add_paragraph()

# ================ FOOTER ================
decor_bottom = doc.add_paragraph()
decor_bottom.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = decor_bottom.add_run('✦ ✦ ✦ ✦ ✦')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(255, 215, 0)  # Gold

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('{{business.slogan}}')
run.italic = True
run.font.size = Pt(10)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(100, 100, 100)

# Save
output_path = r'd:\thang\utility-server\templates\thiep_khai_truong.docx'
doc.save(output_path)
print(f'✅ Thiệp khai trương created: {output_path}')
