#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Validate V5 output - check if data is in correct positions"""

from docx import Document
from pathlib import Path

output_file = Path("OUTPUT_MAU_2C_V5.docx")
doc = Document(output_file)

print("=" * 80)
print("📋 KIỂM TRA BẢNG 3: GIA ĐÌNH (V5)")
print("=" * 80)

t3 = doc.tables[2]
print(f"Rows: {len(t3.rows)}, Cols: {len(t3.rows[0].cells)}")

print("\n📝 ROW 0 (Header):")
for i in range(4):
    print(f"  Col {i}: {t3.rows[0].cells[i].text[:50]}")

print("\n📝 ROW 1 (Data):")
for i in range(4):
    col_text = t3.rows[1].cells[i].text
    lines = [line.strip() for line in col_text.split('\n') if line.strip()]
    
    print(f"\n  Col {i}:")
    if i == 0:
        # Column 0: Structure labels
        print("    (Cấu trúc nhãn - không đếm)")
        for j, line in enumerate(lines[:10], 1):
            if line and line != "....................":
                print(f"      {j}. {line}")
    else:
        # Columns 1-3: Data
        print(f"    Total lines: {len(lines)}")
        
        # Show first 10 lines
        for j, line in enumerate(lines[:12], 1):
            if line and line != "....................":
                print(f"      {j}. {line[:80]}")

print("\n" + "=" * 80)
print("📋 KIỂM TRA BẢNG 4: GIA ĐÌNH VỢ/CHỒNG (V5)")
print("=" * 80)

t4 = doc.tables[3]
print(f"Rows: {len(t4.rows)}, Cols: {len(t4.rows[0].cells)}")

print("\n📝 ROW 1 (Data):")
for i in range(4):
    col_text = t4.rows[1].cells[i].text
    lines = [line.strip() for line in col_text.split('\n') if line.strip()]
    
    print(f"\n  Col {i}:")
    if i == 0:
        # Column 0: Structure labels
        print("    (Cấu trúc nhãn)")
        for j, line in enumerate(lines[:10], 1):
            if line and line != "....................":
                print(f"      {j}. {line}")
    else:
        # Columns 1-3: Data
        print(f"    Total lines: {len(lines)}")
        
        # Show all lines
        for j, line in enumerate(lines, 1):
            if line and line != "....................":
                print(f"      {j}. {line[:80]}")

print("\n" + "=" * 80)
print("✅ VALIDATION SUMMARY")
print("=" * 80)

# Check if structure looks correct
t3_col1_lines = [l.strip() for l in t3.rows[1].cells[1].text.split('\n') if l.strip() and l.strip() != '....................']
t4_col1_lines = [l.strip() for l in t4.rows[1].cells[1].text.split('\n') if l.strip() and l.strip() != '....................']

print(f"\n📊 Bảng 3 - Tổng {len(t3_col1_lines)} tên người")
print(f"   - Expected: 7 (2 bố mẹ + 1 vợ + 2 con + 2 anh chị em)")
print(f"   - Status: {'✅ PASS' if len(t3_col1_lines) == 7 else '❌ FAIL'}")

print(f"\n📊 Bảng 4 - Tổng {len(t4_col1_lines)} tên người")
print(f"   - Expected: 4 (2 bố mẹ vợ + 2 anh chị em vợ)")
print(f"   - Status: {'✅ PASS' if len(t4_col1_lines) == 4 else '❌ FAIL'}")

if len(t3_col1_lines) == 7 and len(t4_col1_lines) == 4:
    print("\n🎉 TẤT CẢ ĐÚNG!")
else:
    print("\n⚠️ CẦN KIỂM TRA LẠI!")

print("=" * 80)
