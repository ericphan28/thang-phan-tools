#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
TẠO TEMPLATE HOÀN CHỈNH V6 - 100% FIELDS
=========================================
Replace TẤT CẢ các field có trong JSON (100+ fields)
Để output KHÔNG còn dấu chấm nữa!
"""

from docx import Document
import re
import json

def load_json_data(file_path):
    """Load JSON để lấy ALL field names"""
    with open(file_path, 'r', encoding='utf-8') as f:
        return json.load(f)

def get_all_fields(data, prefix=''):
    """Recursive extract ALL field names từ nested JSON"""
    fields = []
    for key, value in data.items():
        if key.startswith('_'):  # Skip comments
            continue
        
        full_key = f"{prefix}.{key}" if prefix else key
        
        if isinstance(value, dict):
            fields.extend(get_all_fields(value, full_key))
        elif isinstance(value, list) and len(value) > 0 and isinstance(value[0], dict):
            # Array of objects - add array name
            fields.append(full_key)
            # Add first item fields as template
            for item_key in value[0].keys():
                if not item_key.startswith('_'):
                    fields.append(f"{full_key}[].{item_key}")
        else:
            fields.append(full_key)
    
    return fields

def simple_replace_in_run(paragraph, pattern, jinja_var):
    """Replace trong run - preserve format"""
    for run in paragraph.runs:
        if re.search(pattern, run.text, re.IGNORECASE):
            run.text = re.sub(pattern, f'{{{{ {jinja_var} }}}}', run.text, flags=re.IGNORECASE)
            return True
    return False

def process_document(doc_path, output_path, field_list):
    """
    Process document với FULL field list
    """
    doc = Document(doc_path)
    replacements = 0
    
    # MAPPING: pattern -> jinja_var
    # Tự động tạo từ field_list
    field_mapping = {}
    
    for field in field_list:
        # Clean field name for Jinja2
        jinja_name = field.replace('.', '_').replace('[]', '')
        
        # Create regex pattern
        # Ví dụ: "ho_ten" -> match "Họ và tên: ......"
        patterns = [
            # Pattern 1: Exact field name with dots
            (rf'{re.escape(field)}[\s:\.]+', jinja_name),
            # Pattern 2: Vietnamese label
            (rf'(?i)(họ và tên|tên khai sinh)[\s:]+\.+', 'ho_ten'),
            (rf'(?i)(nam, nữ)[\s:]+\.+', 'gioi_tinh'),
            (rf'(?i)sinh ngày[\s:]+\.+', 'ngay_sinh'),
            (rf'(?i)(nơi sinh|phường)[\s:]+\.+', 'noi_sinh'),
            (rf'(?i)(quê quán|huyện|tỉnh)[\s:]+\.+', 'que_quan'),
            (rf'(?i)(dân tộc|tộc)[\s:]+\.+', 'dan_toc'),
            (rf'(?i)(tôn giáo|giáo)[\s:]+\.+', 'ton_giao'),
            (rf'(?i)(điện thoại|phone)[\s:]+\.+', 'dien_thoai'),
            (rf'(?i)(email|mail)[\s:]+\.+', 'email'),
            # Pattern 3: Section numbers
            (r'1\)[\s]+\.{3,}', 'ho_ten'),
            (r'2\)[\s]+\.{3,}', 'ten_goi_khac'),
            (r'3\)[\s]+\.{3,}', 'cap_uy_hien_tai'),
            (r'4\)[\s]+\.{3,}', 'ngay_sinh'),
            (r'5\)[\s]+\.{3,}', 'noi_sinh'),
            (r'6\)[\s]+\.{3,}', 'que_quan'),
            (r'7\)[\s]+\.{3,}', 'noi_o_hien_nay'),
            (r'8\)[\s]+\.{3,}', 'dan_toc'),
            (r'9\)[\s]+\.{3,}', 'ton_giao'),
            (r'10\)[\s]+\.{3,}', 'thanh_phan_xuat_than'),
            (r'11\)[\s]+\.{3,}', 'nghe_nghiep_ban_than'),
            (r'12\)[\s]+\.{3,}', 'ngay_tuyen_dung'),
            (r'13\)[\s]+\.{3,}', 'ngay_vao_co_quan'),
            (r'14\)[\s]+\.{3,}', 'ngay_vao_dang'),
            (r'15\)[\s]+\.{3,}', 'ngay_tham_gia_to_chuc'),
            (r'16\)[\s]+\.{3,}', 'ngay_nhap_ngu'),
            (r'17\)[\s]+\.{3,}', 'trinh_do_hoc_van'),
        ]
        
        for pattern, var in patterns:
            if var == jinja_name or var in field:
                field_mapping[pattern] = var
    
    # Add comprehensive Vietnamese label patterns
    comprehensive_patterns = {
        # Header fields
        r'(?i)(tỉnh|thành phố)[\s:]+\.{2,}': 'tinh',
        r'(?i)(đơn vị trực thuộc)[\s:]+\.{2,}': 'don_vi_truc_thuoc',
        r'(?i)(đơn vị cơ sở)[\s:]+\.{2,}': 'don_vi_co_so',
        r'(?i)(số hiệu.*bộ.*công chức)[\s:]+\.{2,}': 'so_hieu',
        
        # Personal info
        r'(?i)(họ và tên khai sinh)[\s:]+\.{2,}': 'ho_ten',
        r'(?i)(nam.*nữ)[\s:]+\.{2,}': 'gioi_tinh',
        r'(?i)(các tên gọi khác)[\s:]+\.{2,}': 'ten_goi_khac',
        r'(?i)(cấp ủy hiện tại)[\s:]+\.{2,}': 'cap_uy_hien_tai',
        r'(?i)(cấp ủy kiêm)[\s:]+\.{2,}': 'cap_uy_kiem',
        r'(?i)(chức vụ.*chánh quyền)[\s:]+\.{2,}': 'chuc_vu_full',
        r'(?i)(phụ cấp chức vụ)[\s:]+\.{2,}': 'phu_cap_chuc_vu',
        
        # Birth info
        r'(?i)sinh ngày[\s:]+\.+\s+tháng': 'ngay',
        r'(?i)tháng[\s:]+\.+\s+năm': 'thang',
        r'(?i)năm[\s:]+\.+\s+(?:5\)|\n)': 'nam',
        r'(?i)nơi sinh[\s:]*\.{2,}': 'noi_sinh',
        
        # Origin
        r'(?i)quê quán.*xã[\s:]+\.{2,}': 'que_quan_xa',
        r'(?i)quê quán.*huyện[\s:]+\.{2,}': 'que_quan_huyen',
        r'(?i)quê quán.*tỉnh[\s:]+\.{2,}': 'que_quan_tinh',
        
        # Contact
        r'(?i)nơi ở hiện nay[\s:]+\.{2,}': 'noi_o_hien_nay',
        r'(?i)điện thoại[\s:]+\.{2,}': 'dien_thoai',
        r'(?i)email[\s:]+\.{2,}': 'email',
        
        # Background
        r'(?i)dân tộc[\s:]+\.{2,}': 'dan_toc',
        r'(?i)tôn giáo[\s:]+\.{2,}': 'ton_giao',
        r'(?i)thành phần.*xuất thân[\s:]+\.{2,}': 'thanh_phan_xuat_than',
        r'(?i)nghề nghiệp.*trước.*tuyển dụng[\s:]+\.{2,}': 'nghe_nghiep_ban_than',
        
        # Career dates
        r'(?i)ngày.*tuyển dụng[\s:]+\.{2,}': 'ngay_tuyen_dung',
        r'(?i)vào cơ quan[\s:]+\.{2,}': 'co_quan_tuyen_dung',
        r'(?i)ngày vào cơ quan[\s:]+\.{2,}': 'ngay_vao_co_quan',
        r'(?i)ngày tham gia.*cách mạng[\s:]+\.{2,}': 'ngay_tham_gia_cach_mang',
        r'(?i)ngày vào.*Đảng[\s:]+\.{2,}': 'ngay_vao_dang',
        r'(?i)ngày chính thức[\s:]+\.{2,}': 'ngay_chinh_thuc_dang',
        r'(?i)ngày tham gia.*tổ chức[\s:]+\.{2,}': 'ngay_tham_gia_to_chuc',
        r'(?i)ngày nhập ngũ[\s:]+\.{2,}': 'ngay_nhap_ngu',
        r'(?i)ngày xuất ngũ[\s:]+\.{2,}': 'ngay_xuat_ngu',
        r'(?i)quân hàm[\s:]+\.{2,}': 'quan_ham',
        
        # Education
        r'(?i)trình độ.*giáo dục phổ thông[\s:]+\.{2,}': 'trinh_do_giao_duc_pho_thong',
        r'(?i)học hàm.*học vị[\s:]+\.{2,}': 'hoc_ham_hoc_vi',
        r'(?i)lý luận chính trị[\s:]+\.{2,}': 'ly_luan_chinh_tri',
        r'(?i)ngoại ngữ[\s:]+\.{2,}': 'ngoai_ngu',
        r'(?i)quản lý nhà nước[\s:]+\.{2,}': 'quan_ly_nha_nuoc',
        r'(?i)tin học[\s:]+\.{2,}': 'tin_hoc',
        
        # Work
        r'(?i)công tác chính đảng[\s:]+\.{2,}': 'cong_tac_chinh_dang',
        r'(?i)người công chức[\s:]+\.{2,}': 'nguoi_cong_chuc_vien_chuc',
        r'(?i)bậc lương[\s:]+\.{2,}': 'bac_luong',
        r'(?i)hệ số[\s:]+\.{2,}': 'he_so',
        r'(?i)từ tháng[\s:]+\.{2,}': 'tu_thang',
        r'(?i)sơ lược.*công tác[\s:]+\.{2,}': 'so_luoc_cong_tac',
        
        # Awards
        r'(?i)khen thưởng[\s:]+\.{2,}': 'khen_thuong',
        r'(?i)kỷ luật[\s:]+\.{2,}': 'ky_luat',
        r'(?i)tình trạng sức khỏe[\s:]+\.{2,}': 'tinh_trang_suc_khoe',
        r'(?i)cao[\s:]+\.{2,}\s*cm': 'chieu_cao',
        r'(?i)cân nặng[\s:]+\.{2,}\s*kg': 'can_nang',
        r'(?i)nhóm máu[\s:]+\.{2,}': 'nhom_mau',
        
        # Family - section 29
        r'(?i)đặc điểm lịch sử.*bản thân[\s:]+\.{2,}': 'dac_diem_lich_su_ban_than',
        r'(?i)đặc điểm.*gia đình[\s:]+\.{2,}': 'dac_diem_lich_su_gia_dinh',
        
        # Foreign relations - section 30
        r'(?i)quan hệ.*nước ngoài[\s:]+\.{2,}': 'quan_he_voi_nuoc_ngoai',
        r'(?i)thân nhân.*đảng viên[\s:]+\.{2,}': 'than_nhan_dang_vien',
        r'(?i)có thân nhân.*nước ngoài[\s:]+\.{2,}': 'co_than_nhan_o_nuoc_ngoai',
        
        # Family details - section 31
        r'(?i)tháng năm[\s:]+\.{2,}': 'thang_nam',
        r'(?i)nơi sinh[\s:]+\.{2,}': 'nguoi_than_noi_sinh',
        r'(?i)hệ số lương[\s:]+\.{2,}': 'nguoi_than_he_so_luong',
        r'(?i)nghề nghiệp.*chức danh[\s:]+\.{2,}': 'nguoi_than_nghe_nghiep',
        r'(?i)nơi công tác[\s:]+\.{2,}': 'nguoi_than_noi_cong_tac',
        
        # Family member type labels
        r'(?i)bố mẹ.*anh chị em ruột[\s:]+\.{2,}': 'family_label_1',
        r'(?i)vợ.*chồng[\s:]+\.{2,}': 'family_label_2',
        
        # Economic status
        r'(?i)nhà ở[\s:]+\.{2,}': 'nha_o',
        r'(?i)đặc cấp[\s:]+\.{2,}': 'dac_cap',
        r'(?i)tổng diện tích[\s:]+\.{2,}\s*m2': 'tong_dien_tich',
        r'(?i)nhà tư nhân[\s:]+\.{2,}': 'nha_tu_nhan',
        r'(?i)đất được cấp[\s:]+\.{2,}': 'dat_duoc_cap',
        r'(?i)đất mua[\s:]+\.{2,}': 'dat_mua',
        
        # Signature fields
        r'(?i)nơi khai[\s:]+\.{2,}': 'noi_khai',
        r'(?i)tổ xã.*hoặc nơi làm việc[\s:]+\.{2,}': 'to_xa_lam_viec',
        r'(?i)xác nhận của cơ quan[\s:]+\.{2,}': 'xac_nhan_co_quan',
    }
    
    # Merge all patterns
    all_patterns = {**field_mapping, **comprehensive_patterns}
    
    print(f"📖 Loaded: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
    print(f"🎯 Total patterns: {len(all_patterns)}")
    
    # Process paragraphs
    for i, para in enumerate(doc.paragraphs):
        for pattern, var in all_patterns.items():
            if simple_replace_in_run(para, pattern, var):
                print(f"✅ P{i}: {var}")
                replacements += 1
    
    # Process tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    for pattern, var in all_patterns.items():
                        if simple_replace_in_run(para, pattern, var):
                            print(f"✅ T{table_idx}R{row_idx}C{cell_idx}: {var}")
                            replacements += 1
    
    # Save
    doc.save(output_path)
    print(f"\n📄 File: {output_path}")
    print(f"🔄 Replacements: {replacements}")
    print(f"✅ Template V6 HOÀN CHỈNH!")
    
    return replacements

if __name__ == "__main__":
    INPUT_FILE = "mau-nha-nuoc/Mau-ly-lich-2C-TCTW-98.docx"
    OUTPUT_FILE = "mau_2c_V6_COMPLETE_TEMPLATE.docx"
    JSON_FILE = "mau_2c_DATA_RESTRUCTURED.json"
    
    # Load JSON to get field list
    print("📊 Loading JSON data...")
    data = load_json_data(JSON_FILE)
    fields = get_all_fields(data)
    print(f"🎯 Found {len(fields)} fields in JSON")
    
    # Process document
    print(f"\n🔄 Processing {INPUT_FILE}...")
    replacements = process_document(INPUT_FILE, OUTPUT_FILE, fields)
    
    print(f"\n{'='*60}")
    print(f"✅ HOÀN THÀNH!")
    print(f"📊 Tổng cộng: {replacements} replacements")
    print(f"📄 Template: {OUTPUT_FILE}")
    print(f"🎯 Test với: python test_v6_complete.py")
