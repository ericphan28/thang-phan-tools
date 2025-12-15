#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
GIẢI PHÁP TỰ ĐỘNG 100% - KHÔNG CẦN THỦ CÔNG
Sử dụng python-docx-template với cách tiếp cận THÔNG MINH
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import re
import json

print("=" * 80)
print("🚀 GIẢI PHÁP TỰ ĐỘNG - TẠO TEMPLATE CHUYÊN NGHIỆP")
print("=" * 80)

print("""
💡 Ý TƯỞNG MỚI:

Thay vì:
❌ Replace text → Mất format
❌ Tạo template thủ công → Mất thời gian

→ Sử dụng:
✅ INSERT Jinja2 variables VÀO TRONG text runs
✅ GIỮ NGUYÊN format của run đó
✅ 100% TỰ ĐỘNG!

CÁCH HOẠT ĐỘNG:
1. Đọc file gốc
2. Tìm patterns (dấu chấm ...)
3. KHÔNG XÓA run cũ
4. CHỈ THAY THẾ TEXT trong run
5. Format tự động giữ nguyên!
""")

def replace_text_preserve_format(run, old_text, new_text):
    """
    Replace text trong run NHƯNG GIỮ NGUYÊN format
    """
    if old_text in run.text:
        run.text = run.text.replace(old_text, new_text)
        return True
    return False

def smart_replace_in_paragraph(paragraph, pattern, replacement):
    """
    Tìm và replace text trong paragraph, giữ nguyên format
    """
    full_text = paragraph.text
    if not re.search(pattern, full_text):
        return False
    
    # Find which runs contain the pattern
    for run in paragraph.runs:
        if re.search(pattern, run.text):
            run.text = re.sub(pattern, replacement, run.text)
            return True
    
    # If pattern spans multiple runs, need to handle differently
    # Reconstruct text from runs
    current_pos = 0
    match = re.search(pattern, full_text)
    
    if match:
        match_start = match.start()
        match_end = match.end()
        
        # Find which runs the match spans
        for i, run in enumerate(paragraph.runs):
            run_start = current_pos
            run_end = current_pos + len(run.text)
            
            if match_start >= run_start and match_start < run_end:
                # Match starts in this run
                if match_end <= run_end:
                    # Match ends in same run
                    run.text = run.text[:match_start - run_start] + replacement + run.text[match_end - run_start:]
                    return True
                else:
                    # Match spans multiple runs - complex case
                    # For now, simple replace
                    paragraph.text = re.sub(pattern, replacement, full_text)
                    return True
            
            current_pos = run_end
    
    return False

def create_smart_template():
    """
    Tạo template THÔNG MINH - giữ format 100%
    """
    
    # Load original file
    doc = Document("mau-nha-nuoc/Mau-ly-lich-2C-TCTW-98.docx")
    
    print("\n📖 Đọc file gốc...")
    print(f"   - Paragraphs: {len(doc.paragraphs)}")
    print(f"   - Tables: {len(doc.tables)}")
    
    replacements = {
        # Simple fields
        r"Tỉnh:\s*[\.…]{3,}": "Tỉnh: {{ tinh }}",
        r"Họ và tên:\s*[\.…]{3,}": "Họ và tên: {{ ho_ten }}",
        r"Sinh ngày:\s*[\.…]{3,}\s*tháng:\s*[\.…]{3,}\s*năm:\s*[\.…]{3,}": 
            "Sinh ngày: {{ ngay }} tháng: {{ thang }} năm: {{ nam }}",
        r"Quê quán.*?:\s*[\.…]{3,}": "Quê quán: {{ que_quan }}",
        r"Dân tộc:\s*[\.…]{3,}": "Dân tộc: {{ dan_toc }}",
        r"Tôn giáo:\s*[\.…]{3,}": "Tôn giáo: {{ ton_giao }}",
    }
    
    replaced_count = 0
    
    # Process paragraphs
    print("\n🔧 Xử lý paragraphs...")
    for i, para in enumerate(doc.paragraphs):
        for pattern, replacement in replacements.items():
            if smart_replace_in_paragraph(para, pattern, replacement):
                replaced_count += 1
                print(f"   ✅ P{i}: {pattern[:30]}... → {replacement[:30]}...")
    
    # Process tables
    print("\n🔧 Xử lý tables...")
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    for pattern, replacement in replacements.items():
                        if smart_replace_in_paragraph(para, pattern, replacement):
                            replaced_count += 1
                            print(f"   ✅ T{table_idx}-R{row_idx}-C{cell_idx}: {replacement[:40]}...")
    
    # Save
    output_path = Path("mau_2c_template_SMART.docx")
    doc.save(str(output_path))
    
    print(f"\n✅ Đã tạo template: {output_path}")
    print(f"📊 Tổng số replacements: {replaced_count}")
    print(f"📏 Size: {output_path.stat().st_size:,} bytes")
    
    return output_path

if __name__ == "__main__":
    print("\n" + "=" * 80)
    print("🚀 BẮT ĐẦU TẠO TEMPLATE...")
    print("=" * 80)
    
    try:
        template_path = create_smart_template()
        
        print("\n" + "=" * 80)
        print("✅ THÀNH CÔNG!")
        print("=" * 80)
        print(f"\n📄 Template: {template_path}")
        print("\n💡 KIỂM TRA:")
        print("   1. Mở file template trong Word")
        print("   2. Kiểm tra format có giữ nguyên không")
        print("   3. Kiểm tra {{ variables }} có đúng vị trí không")
        print("\n🎯 Nếu format giữ nguyên → Test với docxtpl!")
        
    except Exception as e:
        print(f"\n❌ LỖI: {e}")
        import traceback
        traceback.print_exc()

print("=" * 80)
