"""
Phân tích chi tiết Mẫu 2C-TCTW-98 để tạo template logic
"""
from docx import Document
import json

doc = Document(r"d:\thang\utility-server\templates\mau-nha-nuoc\Mau-ly-lich-2C-TCTW-98.docx")

print("=" * 80)
print("PHÂN TÍCH MẪU 2C-TCTW-98 - SƠ YẾU LÝ LỊCH CÁN BỘ")
print("=" * 80)

print(f"\n📊 Tổng quan:")
print(f"   - Số đoạn văn: {len(doc.paragraphs)}")
print(f"   - Số bảng: {len(doc.tables)}")

print("\n" + "=" * 80)
print("📋 PHÂN TÍCH CÁC BẢNG")
print("=" * 80)

for i, table in enumerate(doc.tables):
    print(f"\n{'='*60}")
    print(f"BẢNG {i+1}: {len(table.rows)} hàng x {len(table.columns)} cột")
    print(f"{'='*60}")
    
    # Print first 3 rows to understand structure
    for row_idx in range(min(3, len(table.rows))):
        print(f"\nHàng {row_idx + 1}:")
        for col_idx, cell in enumerate(table.rows[row_idx].cells):
            text = cell.text.strip()
            if text:
                print(f"  Cột {col_idx + 1}: {text[:100]}")

print("\n" + "=" * 80)
print("📝 CẤU TRÚC THÔNG TIN (26-31 mục)")
print("=" * 80)

# Extract structure from paragraphs
sections = {}
current_section = None

for para in doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue
    
    # Check if it's a numbered section (1), 2), 3), etc.)
    if ')' in text and any(text.startswith(f"{i})") for i in range(1, 32)):
        section_num = text.split(')')[0].strip()
        section_content = text.split(')', 1)[1].strip() if ')' in text else ''
        sections[section_num] = section_content[:150]

print("\n📌 CÁC MỤC CHÍNH:")
for num in sorted(sections.keys(), key=lambda x: int(x) if x.isdigit() else 0):
    print(f"   {num}) {sections[num]}")

print("\n" + "=" * 80)
print("🎯 KẾT LUẬN: CẤU TRÚC MẪU 2C")
print("=" * 80)
print("""
Mẫu 2C bao gồm:
- PHẦN I: Thông tin cơ bản (26 mục)
  + Mục 1-5: Định danh (Họ tên, giới tính, tên khác, cấp ủy, chức vụ)
  + Mục 6-10: Thông tin cá nhân (Sinh, quê, địa chỉ, dân tộc, tôn giáo)
  + Mục 11-15: Lý lịch nghề nghiệp (Tuyển dụng, vào Đảng, vào Đoàn...)
  + Mục 16-20: Học vấn và ngạch lương
  + Mục 21-25: Thành tích và sức khỏe
  
- PHẦN II: Các bảng chi tiết
  + Bảng 1 (Mục 26): Đào tạo, bồi dưỡng
  + Bảng 2 (Mục 27): Quá trình công tác
  + Bảng 3 (Mục 28): Đặc điểm lịch sử bản thân
  + Bảng 4 (Mục 29): Quan hệ với nước ngoài
  + Bảng 5 (Mục 30): Quan hệ gia đình
  
- PHẦN III: Hoàn cảnh kinh tế (Mục 31)
""")
