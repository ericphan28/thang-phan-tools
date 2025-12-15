"""
Tạo thiệp mời SINH NHẬT vui nhộn
- Màu pastel bright
- Fun decorations
- Playful layout
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_border(section):
    """Viền trang màu vui nhộn"""
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'text')
    
    # Rainbow border effect - dùng art border
    for border_name in ('top', 'left', 'bottom', 'right'):
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'triple')  # Triple line
        border_el.set(qn('w:sz'), '24')
        border_el.set(qn('w:space'), '24')
        border_el.set(qn('w:color'), 'FF69B4')  # Hot pink
        pgBorders.append(border_el)
    
    sectPr.append(pgBorders)

def add_shading(paragraph, color):
    """Background color"""
    pPr = paragraph._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    pPr.append(shd)

# Create doc
doc = Document()

# A5 size
section = doc.sections[0]
section.page_width = Cm(14.8)
section.page_height = Cm(21)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(2)
section.right_margin = Cm(2)

# Pink border
add_page_border(section)

# Set background color (light yellow)
# Note: Word watermark would be better, but this is code-only

# ================ TOP DECORATIONS ================
decor_top = doc.add_paragraph()
decor_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = decor_top.add_run('🎈 🎉 🎂 🎁 🎊')
run.font.size = Pt(20)

doc.add_paragraph()

# ================ MAIN TITLE ================
title1 = doc.add_paragraph()
title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title1.add_run("YOU'RE INVITED!")
run.bold = True
run.font.size = Pt(24)
run.font.name = 'Comic Sans MS'  # Fun font
run.font.color.rgb = RGBColor(255, 105, 180)  # Hot pink

doc.add_paragraph()

# ================ SUBTITLE ================
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_shading(subtitle, 'FFD700')  # Gold background
run = subtitle.add_run('🎂  BIRTHDAY CELEBRATION  🎂')
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Comic Sans MS'
run.font.color.rgb = RGBColor(255, 255, 255)

doc.add_paragraph()
doc.add_paragraph()

# ================ FOR WHO ================
for_label = doc.add_paragraph()
for_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = for_label.add_run('for')
run.font.size = Pt(14)
run.font.name = 'Comic Sans MS'
run.italic = True
run.font.color.rgb = RGBColor(100, 100, 100)

celebrant = doc.add_paragraph()
celebrant.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = celebrant.add_run('{{celebrant.name}}')
run.bold = True
run.font.size = Pt(22)
run.font.name = 'Comic Sans MS'
run.font.color.rgb = RGBColor(255, 105, 180)

age = doc.add_paragraph()
age.alignment = WD_ALIGN_PARAGRAPH.CENTER
run1 = age.add_run('Turning ')
run1.font.size = Pt(16)
run1.font.name = 'Comic Sans MS'
run2 = age.add_run('{{celebrant.age}}')
run2.bold = True
run2.font.size = Pt(20)
run2.font.name = 'Comic Sans MS'
run2.font.color.rgb = RGBColor(255, 215, 0)  # Gold
run3 = age.add_run('!')
run3.font.size = Pt(16)
run3.font.name = 'Comic Sans MS'

doc.add_paragraph()
doc.add_paragraph()

# ================ BALLOONS DECORATION ================
balloons = doc.add_paragraph()
balloons.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = balloons.add_run('🎈 🎈 🎈 🎈 🎈')
run.font.size = Pt(18)

doc.add_paragraph()

# ================ EVENT DETAILS ================
# Date
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run1 = date_p.add_run('📅  ')
run1.font.size = Pt(14)
run2 = date_p.add_run('Date: ')
run2.bold = True
run2.font.size = Pt(13)
run2.font.name = 'Arial'
run3 = date_p.add_run('{{event.date}}')
run3.font.size = Pt(13)
run3.font.name = 'Arial'
run3.font.color.rgb = RGBColor(255, 105, 180)

# Time
time_p = doc.add_paragraph()
time_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run1 = time_p.add_run('🕐  ')
run1.font.size = Pt(14)
run2 = time_p.add_run('Time: ')
run2.bold = True
run2.font.size = Pt(13)
run2.font.name = 'Arial'
run3 = time_p.add_run('{{event.time}}')
run3.font.size = Pt(13)
run3.font.name = 'Arial'
run3.font.color.rgb = RGBColor(255, 105, 180)

# Venue
venue_p = doc.add_paragraph()
venue_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run1 = venue_p.add_run('📍  ')
run1.font.size = Pt(14)
run2 = venue_p.add_run('Venue: ')
run2.bold = True
run2.font.size = Pt(13)
run2.font.name = 'Arial'
run3 = venue_p.add_run('{{event.venue}}')
run3.font.size = Pt(13)
run3.font.name = 'Arial'
run3.font.color.rgb = RGBColor(255, 105, 180)

doc.add_paragraph()
doc.add_paragraph()

# ================ MESSAGE ================
message = doc.add_paragraph()
message.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_shading(message, 'FFF5EE')  # Seashell background
run = message.add_run('Please join us to make this day special!')
run.font.size = Pt(12)
run.font.name = 'Comic Sans MS'
run.italic = True
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# ================ RSVP ================
rsvp_title = doc.add_paragraph()
rsvp_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = rsvp_title.add_run('RSVP')
run.bold = True
run.font.size = Pt(13)
run.font.name = 'Arial'
run.font.color.rgb = RGBColor(255, 105, 180)

rsvp_contact = doc.add_paragraph()
rsvp_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
run1 = rsvp_contact.add_run('📞  ')
run1.font.size = Pt(12)
run2 = rsvp_contact.add_run('{{contact.phone}}')
run2.bold = True
run2.font.size = Pt(12)
run2.font.name = 'Arial'

rsvp_email = doc.add_paragraph()
rsvp_email.alignment = WD_ALIGN_PARAGRAPH.CENTER
run1 = rsvp_email.add_run('📧  ')
run1.font.size = Pt(12)
run2 = rsvp_email.add_run('{{contact.email}}')
run2.font.size = Pt(11)
run2.font.name = 'Arial'
run2.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
doc.add_paragraph()

# ================ BOTTOM DECORATIONS ================
decor_bottom = doc.add_paragraph()
decor_bottom.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = decor_bottom.add_run('🎊 🎁 🎂 🎉 🎈')
run.font.size = Pt(20)

# Save
output_path = r'd:\thang\utility-server\templates\thiep_sinh_nhat.docx'
doc.save(output_path)
print(f'✅ Thiệp sinh nhật created: {output_path}')
