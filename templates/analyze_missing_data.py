"""
Script phân tích dữ liệu thiếu trong OUTPUT_MAU_2C_DOCXTPL.docx
"""

from docx import Document
import re
from pathlib import Path

def analyze_missing_data():
    """
    Phân tích và thống kê những trường nào còn thiếu dữ liệu
    """
    
    print("🔍 PHÂN TÍCH DỮ LIỆU THIẾU TRONG OUTPUT")
    print("="*80)
    
    doc_path = Path("OUTPUT_MAU_2C_DOCXTPL.docx")
    if not doc_path.exists():
        print(f"❌ Không tìm thấy {doc_path}")
        return
    
    doc = Document(doc_path)
    
    # Pattern để tìm trường thiếu dữ liệu
    patterns = {
        'dots_3': r'\.{3,}',  # 3+ dots
        'dots_unicode': r'…{2,}',  # 2+ unicode dots
        'mixed': r'[\.…]{3,}',  # Mixed dots
        'colon_dots': r':\s*[\.…]{3,}',  # ": ..."
        'parentheses_dots': r'\([\.…]{3,}\)',  # "(...)"
    }
    
    missing_fields = []
    field_count = 0
    
    print("\n📋 PHÂN TÍCH PARAGRAPHS:")
    print("-"*80)
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        # Tìm các pattern thiếu dữ liệu
        has_missing = False
        missing_types = []
        
        for pattern_name, pattern in patterns.items():
            matches = re.findall(pattern, text)
            if matches:
                has_missing = True
                missing_types.append(f"{pattern_name}({len(matches)})")
        
        if has_missing:
            field_count += 1
            # Tìm tên field (text trước dấu ":")
            field_name = text.split(':')[0] if ':' in text else text[:50]
            
            missing_fields.append({
                'index': i + 1,
                'field_name': field_name,
                'text': text[:100] + ('...' if len(text) > 100 else ''),
                'missing_types': ', '.join(missing_types)
            })
            
            print(f"\n❌ Para {i+1}: {field_name}")
            print(f"   Text: {text[:80]}...")
            print(f"   Thiếu: {', '.join(missing_types)}")
    
    # Phân tích tables
    print("\n\n📊 PHÂN TÍCH TABLES:")
    print("-"*80)
    
    table_missing = []
    
    for table_idx, table in enumerate(doc.tables):
        print(f"\n📋 Bảng {table_idx + 1}: {len(table.rows)}x{len(table.columns)}")
        
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                
                # Check if cell has missing data
                has_missing = False
                for pattern in patterns.values():
                    if re.search(pattern, text):
                        has_missing = True
                        break
                
                if has_missing and text:
                    table_missing.append({
                        'table': table_idx + 1,
                        'row': row_idx + 1,
                        'col': col_idx + 1,
                        'text': text[:50]
                    })
                    print(f"   ❌ Cell [{row_idx+1},{col_idx+1}]: {text[:50]}...")
    
    # Summary
    print("\n\n" + "="*80)
    print("📊 TỔNG KẾT:")
    print("="*80)
    print(f"\n✅ Tổng paragraphs có dữ liệu: {len([p for p in doc.paragraphs if p.text.strip()])}")
    print(f"❌ Paragraphs thiếu dữ liệu: {field_count}")
    print(f"❌ Table cells thiếu dữ liệu: {len(table_missing)}")
    
    # Group by section
    print("\n\n📋 DANH SÁCH FIELDS THIẾU DỮ LIỆU:")
    print("-"*80)
    
    # Extract field names
    field_names = []
    for item in missing_fields:
        text = item['field_name']
        
        # Extract meaningful field name
        if ')' in text:
            # "1) Họ và tên: ..." → "Họ và tên"
            field_name = text.split(')')[-1].strip().split(':')[0].strip()
        else:
            field_name = text.split(':')[0].strip()
        
        if field_name and len(field_name) > 3:
            field_names.append(field_name)
    
    # Remove duplicates and print
    unique_fields = []
    seen = set()
    for field in field_names:
        if field not in seen:
            unique_fields.append(field)
            seen.add(field)
    
    print(f"\n🔢 Tổng cộng: {len(unique_fields)} fields thiếu dữ liệu\n")
    
    for idx, field in enumerate(unique_fields, 1):
        print(f"{idx:2d}. {field}")
    
    # Suggest missing fields for JSON
    print("\n\n💡 GỢI Ý THÊM VÀO JSON:")
    print("-"*80)
    print("\nCác trường cần thêm vào mau_2c_DATA_FULL.json:")
    print()
    
    suggested_json_fields = []
    
    # Map Vietnamese field names to JSON keys
    field_mapping = {
        'Đơn vị trực thuộc': 'don_vi_truc_thuoc',
        'Đơn vị cơ sở': 'don_vi_co_so',
        'Số hiệu': 'so_hieu',
        'Nam, nữ': 'gioi_tinh',
        'Các tên gọi khác': 'ten_goi_khac',
        'Cấp ủy hiện tại': 'cap_uy_hien_tai',
        'Cấp ủy kiêm': 'cap_uy_kiem',
        'Chức vụ': 'chuc_vu_full',
        'Phụ cấp chức vụ': 'phu_cap_chuc_vu',
        'Nơi sinh': 'noi_sinh',
        'Quê quán': 'que_quan',
        'Nơi ở hiện nay': 'noi_o_hien_nay',
        'đ/thoại': 'dien_thoai',
        'Dân tộc': 'dan_toc',
        'Tôn giáo': 'ton_giao',
        'Thành phần gia đình xuất thân': 'thanh_phan_xuat_than',
        'Nghề nghiệp bản thân': 'nghe_nghiep_ban_than',
        'Ngày được tuyển dụng': 'ngay_tuyen_dung',
        'Ngày vào cơ quan hiện đang công tác': 'ngay_vao_co_quan',
        'Ngày tham gia cách mạng': 'ngay_tham_gia_cach_mang',
        'Ngày chính thức': 'ngay_chinh_thuc_dang',
        'Quân hàm': 'quan_ham',
        'Trình độ học vấn': 'trinh_do_hoc_van',
        'Học hàm, học vị': 'hoc_ham_hoc_vi',
        'Công tác chính đang làm': 'cong_tac_chinh',
        'Ngạch công chức': 'ngach_cong_chuc',
        'Bậc lương': 'bac_luong',
        'Hệ số': 'he_so_luong',
        'Danh hiệu được phong': 'danh_hieu',
        'Sở trường công tác': 'so_truong',
        'Công việc đã làm lâu nhất': 'cong_viec_lau_nhat',
        'Khen thưởng': 'khen_thuong',
        'Kỷ luật': 'ky_luat',
        'Tình trạng sức khỏe': 'suc_khoe',
        'Cao': 'chieu_cao',
        'Cân nặng': 'can_nang',
        'Nhóm máu': 'nhom_mau',
        'Số chứng minh nhân dân': 'so_cmnd',
        'Thương binh loại': 'thuong_binh',
        'Gia đình liệt sĩ': 'gia_dinh_liet_si',
    }
    
    for field in unique_fields:
        if field in field_mapping:
            json_key = field_mapping[field]
            suggested_json_fields.append(f'  "{json_key}": "",  # {field}')
    
    for suggestion in suggested_json_fields[:20]:  # First 20
        print(suggestion)
    
    if len(suggested_json_fields) > 20:
        print(f"\n... và {len(suggested_json_fields) - 20} fields khác")
    
    print("\n\n" + "="*80)
    print("✅ PHÂN TÍCH HOÀN TẤT!")
    print("="*80)

if __name__ == "__main__":
    try:
        analyze_missing_data()
    except Exception as e:
        print(f"\n❌ LỖI: {e}")
        import traceback
        traceback.print_exc()
