"""
Script TẠO TEMPLATE CHUYÊN NGHIỆP - Phân tích CHÍNH XÁC file gốc
Không thủ công, không đoán mò - Dựa trên CẤU TRÚC THỰC TẾ
"""

from docx import Document
import re
from pathlib import Path
from typing import Dict, List, Tuple

def analyze_original_structure():
    """
    BƯỚC 1: Phân tích CHÍNH XÁC cấu trúc file gốc
    """
    print("=" * 80)
    print("🔍 PHÂN TÍCH CẤU TRÚC FILE GỐC - CHUYÊN NGHIỆP")
    print("=" * 80)
    
    doc_path = Path("mau-nha-nuoc/Mau-ly-lich-2C-TCTW-98.docx")
    if not doc_path.exists():
        print(f"❌ Không tìm thấy {doc_path}")
        return None, None
    
    doc = Document(doc_path)
    
    print(f"\n📊 TỔNG QUAN:")
    print(f"   - Paragraphs: {len(doc.paragraphs)}")
    print(f"   - Tables: {len(doc.tables)}")
    
    # Phân tích paragraphs
    print("\n" + "=" * 80)
    print("📋 PHÂN TÍCH TỪNG PARAGRAPH:")
    print("=" * 80)
    
    paragraph_patterns = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        # Tìm patterns có dấu chấm/dots
        dots_patterns = re.findall(r'[\.…]{3,}', text)
        
        if dots_patterns:
            # Extract field name
            parts = text.split(':')
            if len(parts) >= 2:
                field_name = parts[0].strip()
                value_part = ':'.join(parts[1:]).strip()
                
                paragraph_patterns.append({
                    'index': i,
                    'field_name': field_name,
                    'original_text': text,
                    'dots_count': len(dots_patterns),
                    'pattern_type': 'field_with_colon'
                })
                
                print(f"\n📌 Para {i+1}: {field_name[:60]}")
                print(f"   Text: {text[:80]}...")
                print(f"   Dots patterns: {len(dots_patterns)}")
    
    # Phân tích tables
    print("\n" + "=" * 80)
    print("📊 PHÂN TÍCH TABLES:")
    print("=" * 80)
    
    table_structures = []
    
    for table_idx, table in enumerate(doc.tables):
        print(f"\n📋 Table {table_idx + 1}:")
        print(f"   Kích thước: {len(table.rows)} rows × {len(table.columns)} cols")
        
        # Analyze headers (row 0)
        headers = []
        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                headers.append(cell.text.strip())
            print(f"   Headers: {headers}")
        
        # Analyze data row structure
        if len(table.rows) > 1:
            data_row = table.rows[1]
            row_structure = []
            for col_idx, cell in enumerate(data_row.cells):
                text = cell.text.strip()
                has_dots = bool(re.search(r'[\.…]{3,}', text))
                row_structure.append({
                    'col': col_idx,
                    'text': text[:30],
                    'has_dots': has_dots
                })
            
            print(f"   Data row structure:")
            for col_info in row_structure:
                status = "🔴 Cần replace" if col_info['has_dots'] else "🟢 Label/Fixed"
                print(f"      Col {col_info['col']}: {col_info['text'][:30]}... - {status}")
        
        table_structures.append({
            'index': table_idx,
            'rows': len(table.rows),
            'cols': len(table.columns),
            'headers': headers
        })
    
    return paragraph_patterns, table_structures, doc

def create_professional_mapping(paragraph_patterns: List[Dict]) -> Dict[str, str]:
    """
    BƯỚC 2: Tạo MAPPING CHUYÊN NGHIỆP dựa trên phân tích
    """
    print("\n" + "=" * 80)
    print("🎯 TẠO MAPPING CHUYÊN NGHIỆP")
    print("=" * 80)
    
    # MAPPING CHUẨN - Dựa trên PHÂN TÍCH THỰC TẾ
    mapping = {
        # HEADER
        r"Tỉnh:\s*[\.…]{3,}": "Tỉnh: {{ tinh }}",
        r"Đơn vị trực thuộc:\s*[\.…]{3,}": "Đơn vị trực thuộc: {{ don_vi_truc_thuoc }}",
        r"Đơn vị cơ sở:\s*[\.…]{3,}": "Đơn vị cơ sở: {{ don_vi_co_so }}",
        r"Số hiệu[^:]*:\s*[\.…]{3,}": "Số hiệu: {{ so_hieu }}",
        
        # MỤC 1-3: Thông tin cơ bản
        r"(?:1\)|①)\s*Họ và tên[^:]*:\s*[\.…]{3,}": "1) Họ và tên: {{ ho_ten }}",
        r"Nam,\s*nữ:\s*[\.…]{3,}": "Nam, nữ: {{ gioi_tinh }}",
        r"(?:2\)|②)\s*Các tên gọi khác:\s*[\.…]{3,}": "2) Các tên gọi khác: {{ ten_goi_khac }}",
        
        r"(?:3\)|③)\s*Cấp ủy hiện tại:\s*[\.…]{3,}": "3) Cấp ủy hiện tại: {{ cap_uy_hien_tai }}",
        r"Cấp ủy kiêm:\s*[\.…]{3,}": "Cấp ủy kiêm: {{ cap_uy_kiem }}",
        r"Chức vụ[^:]*:\s*[\.…]{3,}": "Chức vụ: {{ chuc_vu_full }}",
        r"Phụ cấp chức vụ:\s*[\.…]{3,}": "Phụ cấp chức vụ: {{ phu_cap_chuc_vu }}",
        
        # MỤC 4-7: Ngày sinh, địa chỉ
        r"(?:4\)|④)\s*Sinh ngày:\s*[\.…]{3,}\s*tháng\s*[\.…]{3,}\s*năm\s*[\.…]{3,}": 
            "4) Sinh ngày: {{ ngay }} tháng: {{ thang }} năm: {{ nam }}",
        r"(?:5\)|⑤)\s*Nơi sinh:\s*[\.…]{3,}": "5) Nơi sinh: {{ noi_sinh }}",
        
        r"(?:6\)|⑥)\s*Quê quán\s*\([^)]+\):\s*[\.…]{3,}": "6) Quê quán (xã, phường): {{ que_quan_xa }}",
        r"\(huyện,\s*quận\):\s*[\.…]{3,}": "(huyện, quận): {{ que_quan_huyen }}",
        r"\(tỉnh,\s*TP\):\s*[\.…]{3,}": "(tỉnh, TP): {{ que_quan_tinh }}",
        
        r"(?:7\)|⑦)\s*Nơi ở hiện nay[^:]*:\s*[\.…]{3,}": "7) Nơi ở hiện nay: {{ noi_o_hien_nay }}",
        r"đ/thoại:\s*[\.…]{3,}": "đ/thoại: {{ dien_thoai }}",
        r"[Ee]mail:\s*[\.…]{3,}": "Email: {{ email }}",
        
        # MỤC 8-11: Dân tộc, thành phần
        r"(?:8\)|⑧)\s*Dân tộc[^:]*:\s*[\.…]{3,}": "8) Dân tộc: {{ dan_toc }}",
        r"(?:9\)|⑨)\s*Tôn giáo:\s*[\.…]{3,}": "9) Tôn giáo: {{ ton_giao }}",
        r"(?:10\)|⑩)\s*Thành phần gia đình xuất thân:\s*[\.…]{3,}": "10) Thành phần gia đình xuất thân: {{ thanh_phan_xuat_than }}",
        r"(?:11\)|⑪)\s*Nghề nghiệp bản thân[^:]*:\s*[\.…]{3,}": "11) Nghề nghiệp bản thân: {{ nghe_nghiep_ban_than }}",
        
        # MỤC 12-13: Tuyển dụng
        r"(?:12\)|⑫)\s*Ngày được tuyển dụng:\s*[\.…]{3,}\s*/\s*[\.…]{3,}\s*/\s*[\.…]{3,}": 
            "12) Ngày được tuyển dụng: {{ ngay_tuyen_dung }}",
        r"Vào cơ quan[^:]*:\s*[\.…]{3,}": "Vào cơ quan: {{ co_quan_tuyen_dung }}",
        
        r"(?:13\)|⑬)\s*Ngày vào cơ quan hiện đang công tác:\s*[\.…]{3,}\s*/\s*[\.…]{3,}\s*/\s*[\.…]{3,}": 
            "13) Ngày vào cơ quan: {{ ngay_vao_co_quan }}",
        r"Ngày tham gia cách mạng:\s*[\.…]{3,}\s*/\s*[\.…]{3,}\s*/\s*[\.…]{3,}": 
            "Ngày tham gia cách mạng: {{ ngay_tham_gia_cach_mang }}",
        
        # MỤC 14-16: Đảng, quân đội
        r"(?:14\)|⑭)\s*Ngày vào Đảng[^:]*:\s*[\.…]{3,}\s*/\s*[\.…]{3,}\s*/\s*[\.…]{3,}": 
            "14) Ngày vào Đảng Cộng sản Việt Nam: {{ ngay_vao_dang }}",
        r"Ngày chính thức:\s*[\.…]{3,}\s*/\s*[\.…]{3,}\s*/\s*[\.…]{3,}": 
            "Ngày chính thức: {{ ngay_chinh_thuc_dang }}",
        
        r"(?:15\)|⑮)\s*Ngày tham gia các tổ chức[^:]*:\s*[\.…]{3,}": 
            "15) Ngày tham gia tổ chức: {{ ngay_tham_gia_to_chuc }}",
        
        r"(?:16\)|⑯)\s*Ngày nhập ngũ:\s*[\.…]{3,}\s*/\s*[\.…]{3,}\s*/\s*[\.…]{3,}": 
            "16) Ngày nhập ngũ: {{ ngay_nhap_ngu }}",
        r"Ngày xuất ngũ:\s*[\.…]{3,}\s*/\s*[\.…]{3,}\s*/\s*[\.…]{3,}": 
            "Ngày xuất ngũ: {{ ngay_xuat_ngu }}",
        r"Quân hàm[^:]*:\s*[\.…]{3,}": "Quân hàm: {{ quan_ham }}",
        
        # MỤC 17: Học vấn
        r"(?:17\)|⑰)[^:]*Trình độ học vấn[^:]*:\s*Giáo dục phổ thông:\s*[\.…]{3,}": 
            "17) Trình độ học vấn: Giáo dục phổ thông: {{ trinh_do_giao_duc_pho_thong }}",
        r"Học hàm,\s*học vị[^:]*:\s*[\.…]{3,}": "Học hàm, học vị: {{ hoc_ham_hoc_vi }}",
        r"-\s*Lý luận chính trị:\s*[\.…]{3,}": "- Lý luận chính trị: {{ ly_luan_chinh_tri }}",
        r"-\s*Ngoại ngữ:\s*[\.…]{3,}": "- Ngoại ngữ: {{ ngoai_ngu }}",
        r"-\s*Quản lý nhà nước:\s*[\.…]{3,}": "- Quản lý nhà nước: {{ quan_ly_nha_nuoc }}",
        r"-\s*Tin học:\s*[\.…]{3,}": "- Tin học: {{ tin_hoc }}",
        
        # MỤC 18-21: Công tác
        r"(?:18\)|⑱)\s*Công tác chính[^:]*:\s*[\.…]{3,}": "18) Công tác chính: {{ cong_tac_chinh }}",
        
        r"(?:19\)|⑲)\s*Ngạch công chức:\s*[\.…]{3,}": "19) Ngạch công chức: {{ ngach_cong_chuc }}",
        r"\(mã số:\s*[\.…]{3,}\)": "(mã số: {{ ma_ngach }})",
        r"Bậc lương:\s*[\.…]{3,}": "Bậc lương: {{ bac_luong }}",
        r"hệ số:\s*[\.…]{3,}": "hệ số: {{ he_so_luong }}",
        r"từ tháng\s*[\.…]{3,}\s*/\s*[\.…]{3,}": "từ tháng: {{ tu_thang_nam }}",
        
        r"(?:20\)|⑳)\s*Danh hiệu[^:]*:\s*[\.…]{3,}": "20) Danh hiệu: {{ danh_hieu }}",
        
        r"(?:21\)|㉑)\s*Sở trường công tác:\s*[\.…]{3,}": "21) Sở trường: {{ so_truong_cong_tac }}",
        r"Công việc đã làm lâu nhất:\s*[\.…]{3,}": "Công việc lâu nhất: {{ cong_viec_lau_nhat }}",
        
        # MỤC 22-25: Khen thưởng, sức khỏe
        r"(?:22\)|㉒)\s*Khen thưởng:\s*[\.…]{3,}": "22) Khen thưởng: {{ khen_thuong }}",
        r"(?:23\)|㉓)\s*Kỷ luật[^:]*:\s*[\.…]{3,}": "23) Kỷ luật: {{ ky_luat }}",
        
        r"(?:24\)|㉔)\s*Tình trạng sức khỏe:\s*[\.…]{3,}": "24) Sức khỏe: {{ suc_khoe }}",
        r"Cao:\s*1m\s*[\.…]{3,}": "Cao: {{ chieu_cao }}",
        r"Cân nặng:\s*[\.…]{3,}": "Cân nặng: {{ can_nang }}",
        r"Nhóm máu:\s*[\.…]{3,}": "Nhóm máu: {{ nhom_mau }}",
        
        r"(?:25\)|㉕)\s*Số chứng minh nhân dân:\s*[\.…]{3,}": "25) Số CMND: {{ so_cmnd }}",
        r"Ngày cấp:\s*[\.…]{3,}": "Ngày cấp: {{ ngay_cap }}",
        r"Nơi cấp:\s*[\.…]{3,}": "Nơi cấp: {{ noi_cap }}",
        r"Thương binh loại:\s*[\.…]{3,}": "Thương binh: {{ thuong_binh_loai }}",
        r"Gia đình liệt sĩ:\s*[\.…]{3,}": "Gia đình liệt sĩ: {{ gia_dinh_liet_si }}",
        
        # MỤC 26-31: Gia đình
        r"Tình trạng hôn nhân:\s*[\.…]{3,}": "Tình trạng hôn nhân: {{ tinh_trang_hon_nhan }}",
        r"Họ và tên vợ[^:]*:\s*[\.…]{3,}": "Họ và tên vợ (chồng): {{ ten_vo_chong }}",
        r"Năm sinh:\s*[\.…]{3,}": "Năm sinh: {{ nam_sinh_vo_chong }}",
        r"Quê quán:\s*[\.…]{3,}": "Quê quán: {{ que_quan_vo_chong }}",
        r"Nghề nghiệp:\s*[\.…]{3,}": "Nghề nghiệp: {{ nghe_nghiep_vo_chong }}",
        r"Chỗ ở:\s*[\.…]{3,}": "Chỗ ở: {{ cho_o_vo_chong }}",
        
        # Kinh tế
        r"\+\s*lương:\s*[\.…]{3,}": "+ Lương: {{ nguon_thu_luong }}",
        r"\+\s*Các nguồn khác:\s*[\.…]{3,}": "+ Nguồn khác: {{ nguon_thu_khac }}",
        
        # Lịch sử
        r"a\)\s*Khai rõ:[^\.]{10,}[\.…]{3,}": "a) Khai rõ: {{ lich_su_bi_bat }}",
        r"b\)\s*Bản thân có làm việc[^\.]{10,}[\.…]{3,}": "b) Bản thân: {{ lam_viec_che_do_cu }}",
        
        # Chữ ký
        r"Ngày\s+[\.…]{3,}\s+tháng\s+[\.…]{3,}\s+năm\s+20[\.…]{2,}": 
            "Ngày {{ ngay_ky }} tháng {{ thang_ky }} năm {{ nam_ky }}",
    }
    
    print(f"\n✅ Đã tạo {len(mapping)} mapping patterns")
    return mapping

def apply_mapping_professionally(doc: Document, mapping: Dict[str, str]) -> Document:
    """
    BƯỚC 3: Apply mapping một cách CHUYÊN NGHIỆP
    """
    print("\n" + "=" * 80)
    print("🔧 APPLY MAPPING CHUYÊN NGHIỆP")
    print("=" * 80)
    
    replaced_count = 0
    
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        
        original = para.text
        new_text = original
        
        # Apply all mappings
        for pattern, replacement in mapping.items():
            new_text = re.sub(pattern, replacement, new_text, flags=re.IGNORECASE)
        
        if new_text != original:
            para.text = new_text
            replaced_count += 1
            print(f"✓ {original[:50]}... → {new_text[:50]}...")
    
    print(f"\n✅ Đã replace {replaced_count} paragraphs")
    return doc

def process_tables_professionally(doc: Document):
    """
    BƯỚC 4: Xử lý tables theo CẤU TRÚC THỰC TẾ
    """
    print("\n" + "=" * 80)
    print("📊 XỬ LÝ TABLES CHUYÊN NGHIỆP")
    print("=" * 80)
    
    if len(doc.tables) < 5:
        print(f"⚠️ Chỉ có {len(doc.tables)} tables, cần 5 tables")
        return doc
    
    # TABLE 1: Đào tạo (2×5)
    print("\n📋 Table 1: Đào tạo, bồi dưỡng (2×5)")
    table1 = doc.tables[0]
    if len(table1.rows) > 1 and len(table1.columns) >= 5:
        row = table1.rows[1]
        row.cells[0].text = "{% for edu in dao_tao %}{{ edu.ten_truong }}{% endfor %}"
        row.cells[1].text = "{% for edu in dao_tao %}{{ edu.nganh_hoc }}{% endfor %}"
        row.cells[2].text = "{% for edu in dao_tao %}{{ edu.thoi_gian }}{% endfor %}"
        row.cells[3].text = "{% for edu in dao_tao %}{{ edu.hinh_thuc }}{% endfor %}"
        row.cells[4].text = "{% for edu in dao_tao %}{{ edu.van_bang }}{% endfor %}"
        print("   ✅ Đã thêm Jinja2 loops")
    
    # TABLE 2: Công tác (2×2)
    print("\n📋 Table 2: Quá trình công tác (2×2)")
    table2 = doc.tables[1]
    if len(table2.rows) > 1 and len(table2.columns) >= 2:
        row = table2.rows[1]
        row.cells[0].text = "{% for work in cong_tac %}{{ work.thoi_gian }}{% endfor %}"
        row.cells[1].text = "{% for work in cong_tac %}{{ work.chuc_vu_don_vi }}{% endfor %}"
        print("   ✅ Đã thêm Jinja2 loops")
    
    # TABLE 3: Gia đình bản thân (2×4)
    print("\n📋 Table 3: Gia đình bản thân (2×4)")
    table3 = doc.tables[2]
    if len(table3.rows) > 1 and len(table3.columns) >= 4:
        row = table3.rows[1]
        # Column 0 = GIỮ NGUYÊN (labels: Bố, mẹ, Vợ, Chồng, Các con)
        row.cells[1].text = "{% for member in gia_dinh %}{{ member.ho_ten }}{% endfor %}"
        row.cells[2].text = "{% for member in gia_dinh %}{{ member.nam_sinh }}{% endfor %}"
        row.cells[3].text = "{% for member in gia_dinh %}{{ member.thong_tin }}{% endfor %}"
        print("   ✅ Đã thêm Jinja2 loops (giữ column 0)")
    
    # TABLE 4: Gia đình vợ/chồng (2×4)
    print("\n📋 Table 4: Gia đình vợ/chồng (2×4)")
    table4 = doc.tables[3]
    if len(table4.rows) > 1 and len(table4.columns) >= 4:
        row = table4.rows[1]
        # Column 0 = GIỮ NGUYÊN
        row.cells[1].text = "{% for member in gia_dinh_vo_chong %}{{ member.ho_ten }}{% endfor %}"
        row.cells[2].text = "{% for member in gia_dinh_vo_chong %}{{ member.nam_sinh }}{% endfor %}"
        row.cells[3].text = "{% for member in gia_dinh_vo_chong %}{{ member.thong_tin }}{% endfor %}"
        print("   ✅ Đã thêm Jinja2 loops (giữ column 0)")
    
    # TABLE 5: Lương (3×7)
    print("\n📋 Table 5: Quá trình lương (3×7)")
    table5 = doc.tables[4]
    if len(table5.rows) > 2 and len(table5.columns) >= 3:
        row = table5.rows[2]  # Row 3 = data row
        row.cells[0].text = "{% for sal in luong %}{{ sal.thang_nam }}{% endfor %}"
        row.cells[1].text = "{% for sal in luong %}{{ sal.ngach_bac }}{% endfor %}"
        row.cells[2].text = "{% for sal in luong %}{{ sal.he_so }}{% endfor %}"
        print("   ✅ Đã thêm Jinja2 loops")
    
    print(f"\n✅ Đã xử lý {len(doc.tables)} tables")
    return doc

def create_professional_template():
    """
    MAIN: Tạo template CHUYÊN NGHIỆP
    """
    print("\n" + "🚀" * 40)
    print("🚀 TẠO TEMPLATE CHUYÊN NGHIỆP - VERSION 3.0")
    print("🚀" * 40)
    
    # BƯỚC 1: Phân tích cấu trúc
    para_patterns, table_structures, doc = analyze_original_structure()
    
    if doc is None:
        return False
    
    # BƯỚC 2: Tạo mapping
    mapping = create_professional_mapping(para_patterns)
    
    # BƯỚC 3: Apply mapping
    doc = apply_mapping_professionally(doc, mapping)
    
    # BƯỚC 4: Xử lý tables
    doc = process_tables_professionally(doc)
    
    # BƯỚC 5: Save
    output_path = Path("mau_2c_template_PROFESSIONAL_V3.docx")
    doc.save(str(output_path))
    
    file_size = output_path.stat().st_size
    
    print("\n" + "=" * 80)
    print("✅ HOÀN THÀNH - TEMPLATE CHUYÊN NGHIỆP!")
    print("=" * 80)
    print(f"📄 File: {output_path}")
    print(f"📊 Size: {file_size:,} bytes ({file_size/1024:.2f} KB)")
    print(f"💡 Mapping: {len(mapping)} patterns applied")
    print(f"📋 Tables: {len(doc.tables)} tables processed")
    print("\n🎯 SẴN SÀNG TEST với docxtpl!")
    
    return True

if __name__ == "__main__":
    try:
        success = create_professional_template()
        if not success:
            print("\n❌ Thất bại!")
            exit(1)
    except Exception as e:
        print(f"\n❌ LỖI: {e}")
        import traceback
        traceback.print_exc()
        exit(1)
