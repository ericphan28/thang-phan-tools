"""
Script cải thiện template để xử lý XUỐNG DÒNG trong bảng đúng cách
"""

from docx import Document
from pathlib import Path

def improve_table_templates():
    """
    Cải thiện template để mỗi item trong loop XUỐNG DÒNG
    """
    
    print("="*80)
    print("🔧 CẢI THIỆN TEMPLATE - XUỐNG DÒNG TRONG BẢNG")
    print("="*80)
    
    template_path = Path("mau_2c_template_PROFESSIONAL_V3.docx")
    doc = Document(template_path)
    
    # TABLE 1: Đào tạo - Mỗi entry 1 dòng
    print("\n📋 Bảng 1: Đào tạo")
    if len(doc.tables) > 0:
        table1 = doc.tables[0]
        if len(table1.rows) > 1:
            row = table1.rows[1]
            # Thêm \n giữa các items
            row.cells[0].text = "{% for edu in dao_tao %}{{ edu.ten_truong }}\n{% endfor %}"
            row.cells[1].text = "{% for edu in dao_tao %}{{ edu.nganh_hoc }}\n{% endfor %}"
            row.cells[2].text = "{% for edu in dao_tao %}{{ edu.thoi_gian }}\n{% endfor %}"
            row.cells[3].text = "{% for edu in dao_tao %}{{ edu.hinh_thuc }}\n{% endfor %}"
            row.cells[4].text = "{% for edu in dao_tao %}{{ edu.van_bang }}\n{% endfor %}"
            print("   ✅ Đã thêm \\n để xuống dòng")
    
    # TABLE 2: Công tác - Mỗi entry 1 dòng
    print("\n📋 Bảng 2: Công tác")
    if len(doc.tables) > 1:
        table2 = doc.tables[1]
        if len(table2.rows) > 1:
            row = table2.rows[1]
            row.cells[0].text = "{% for work in cong_tac %}{{ work.thoi_gian }}\n{% endfor %}"
            row.cells[1].text = "{% for work in cong_tac %}{{ work.chuc_vu_don_vi }}\n\n{% endfor %}"
            print("   ✅ Đã thêm \\n để xuống dòng")
    
    # TABLE 3: Gia đình - Mỗi entry 1 dòng
    print("\n📋 Bảng 3: Gia đình bản thân")
    if len(doc.tables) > 2:
        table3 = doc.tables[2]
        if len(table3.rows) > 1:
            row = table3.rows[1]
            # Column 0 = GIỮ NGUYÊN labels
            row.cells[1].text = "{% for member in gia_dinh %}{{ member.ho_ten }}\n{% endfor %}"
            row.cells[2].text = "{% for member in gia_dinh %}{{ member.nam_sinh }}\n{% endfor %}"
            row.cells[3].text = "{% for member in gia_dinh %}{{ member.thong_tin }}\n{% endfor %}"
            print("   ✅ Đã thêm \\n để xuống dòng")
    
    # TABLE 4: Gia đình vợ/chồng - Mỗi entry 1 dòng
    print("\n📋 Bảng 4: Gia đình vợ/chồng")
    if len(doc.tables) > 3:
        table4 = doc.tables[3]
        if len(table4.rows) > 1:
            row = table4.rows[1]
            # Column 0 = GIỮ NGUYÊN
            row.cells[1].text = "{% for member in gia_dinh_vo_chong %}{{ member.ho_ten }}\n{% endfor %}"
            row.cells[2].text = "{% for member in gia_dinh_vo_chong %}{{ member.nam_sinh }}\n{% endfor %}"
            row.cells[3].text = "{% for member in gia_dinh_vo_chong %}{{ member.thong_tin }}\n{% endfor %}"
            print("   ✅ Đã thêm \\n để xuống dòng")
    
    # TABLE 5: Lương - Mỗi entry 1 dòng
    print("\n📋 Bảng 5: Quá trình lương")
    if len(doc.tables) > 4:
        table5 = doc.tables[4]
        if len(table5.rows) > 2:
            row = table5.rows[2]
            row.cells[0].text = "{% for sal in luong %}{{ sal.thang_nam }}\n{% endfor %}"
            row.cells[1].text = "{% for sal in luong %}{{ sal.ngach_bac }}\n{% endfor %}"
            row.cells[2].text = "{% for sal in luong %}{{ sal.he_so }}\n{% endfor %}"
            print("   ✅ Đã thêm \\n để xuống dòng")
    
    # Save
    output_path = Path("mau_2c_template_FINAL_V4.docx")
    doc.save(str(output_path))
    
    file_size = output_path.stat().st_size
    
    print("\n" + "="*80)
    print("✅ ĐÃ CẢI THIỆN TEMPLATE!")
    print("="*80)
    print(f"📄 File: {output_path}")
    print(f"📊 Size: {file_size:,} bytes ({file_size/1024:.2f} KB)")
    print("\n💡 Thay đổi:")
    print("   - Thêm \\n sau mỗi {{ variable }} trong loop")
    print("   - Mỗi entry sẽ xuống 1 dòng riêng")
    print("\n🎯 Test ngay: python test_docxtpl.py")

if __name__ == "__main__":
    improve_table_templates()
