#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Validate all table data in output"""

from docx import Document
from pathlib import Path

output_file = Path("OUTPUT_MAU_2C_DOCXTPL.docx")

if not output_file.exists():
    print(f"❌ File not found: {output_file}")
    exit(1)

doc = Document(output_file)

print("=" * 80)
print("📋 KIỂM TRA TẤT CẢ CÁC BẢNG")
print("=" * 80)

# Table 1: Education (Đào tạo)
print("\n📋 BẢNG 1: ĐÀO TẠO")
print("-" * 80)
t1 = doc.tables[0]
for i in range(5):
    print(f"\nCột {i+1}:")
    print(t1.rows[1].cells[i].text[:200])  # First 200 chars

# Table 2: Work History (Công tác)
print("\n\n📋 BẢNG 2: QUÁ TRÌNH CÔNG TÁC")
print("-" * 80)
t2 = doc.tables[1]
for i in range(2):
    print(f"\nCột {i+1}:")
    print(t2.rows[1].cells[i].text[:300])

# Table 3: Family (Gia đình)
print("\n\n📋 BẢNG 3: GIA ĐÌNH")
print("-" * 80)
t3 = doc.tables[2]
for i in range(4):
    print(f"\nCột {i+1}:")
    lines = t3.rows[1].cells[i].text.split('\n')
    for j, line in enumerate(lines[:10], 1):  # First 10 lines
        if line.strip():
            print(f"  {j}. {line.strip()}")

# Table 4: Spouse's Family (Gia đình vợ/chồng)
print("\n\n📋 BẢNG 4: GIA ĐÌNH VỢ/CHỒNG")
print("-" * 80)
t4 = doc.tables[3]
for i in range(4):
    print(f"\nCột {i+1}:")
    lines = t4.rows[1].cells[i].text.split('\n')
    for j, line in enumerate(lines[:10], 1):
        if line.strip():
            print(f"  {j}. {line.strip()}")

# Table 5: Salary (Lương)
print("\n\n📋 BẢNG 5: LƯƠNG")
print("-" * 80)
t5 = doc.tables[4]
print(f"Rows: {len(t5.rows)}, Cols: {len(t5.rows[0].cells)}")
if len(t5.rows) > 1:
    for i in range(min(7, len(t5.rows[1].cells))):
        print(f"\nCột {i+1}:")
        print(t5.rows[1].cells[i].text[:150])

print("\n" + "=" * 80)
print("✅ HOÀN TẤT KIỂM TRA")
print("=" * 80)
