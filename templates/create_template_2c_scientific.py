"""
Tạo template KHOA HỌC từ Mẫu 2C-TCTW-98
Thay thế TẤT CẢ dấu chấm (...) bằng {{variables}} có logic
"""
from docx import Document
from docx.shared import Pt
from copy import deepcopy
import re

# Đọc file gốc
source = r"d:\thang\utility-server\templates\mau-nha-nuoc\Mau-ly-lich-2C-TCTW-98.docx"
output = r"d:\thang\utility-server\templates\so_yeu_ly_lich_2c_template.docx"

print("📖 Đọc mẫu 2C-TCTW-98...")
doc = Document(source)

print(f"✅ Tìm thấy {len(doc.paragraphs)} đoạn, {len(doc.tables)} bảng\n")

# =============================================================================
# BƯỚC 1: Mapping variables logic cho 31 mục
# =============================================================================
print("🔧 BƯỚC 1: Thay thế các mục chính (1-31)")

replacements = {
    # Header
    r'Tỉnh: ………………….': 'Tỉnh: {{tinh}}',
    r'Đơn vị trực thuộc: \.+': 'Đơn vị trực thuộc: {{don_vi_truc_thuoc}}',
    r'Đơn vị cơ sở: \.+': 'Đơn vị cơ sở: {{don_vi_co_so}}',
    r'Số hiệu cán bộ, công chức': 'Số hiệu cán bộ, công chức: {{so_hieu_can_bo}}',
    
    # Mục 1-5
    r'1\) Họ và tên khai sinh: ……+': '1) Họ và tên khai sinh: {{ho_ten}}',
    r'Nam, nữ: \.+': 'Nam, nữ: {{gioi_tinh}}',
    r'2\) Các tên gọi khác:\s+\.+': '2) Các tên gọi khác: {{ten_goi_khac}}',
    r'3\) Cấp ủy hiện tại: \.+': '3) Cấp ủy hiện tại: {{cap_uy_hien_tai}}',
    r'Cấp ủy kiêm: \.+': 'Cấp ủy kiêm: {{cap_uy_kiem}}',
    r'Chức vụ \(Đảng, đoàn thể, Chính quyền[^:]+: \.+': 'Chức vụ (Đảng, đoàn thể, Chính quyền, kể cả chức vụ kiêm nhiệm): {{chuc_vu}}',
    r'Phụ cấp chức vụ: \.+': 'Phụ cấp chức vụ: {{phu_cap_chuc_vu}}',
    
    # Mục 4-10
    r'4\) Sinh ngày: \.+ tháng \.+ năm \.+': '4) Sinh ngày: {{ngay_sinh}} tháng {{thang_sinh}} năm {{nam_sinh}}',
    r'5\) Nơi sinh: \.+': '5) Nơi sinh: {{noi_sinh}}',
    r'6\) Quê quán \(xã, phường\): \.+\(huyện, quận\): \.+\(tỉnh, TP\): \.+': 
        '6) Quê quán (xã, phường): {{que_quan_xa}} (huyện, quận): {{que_quan_huyen}} (tỉnh, TP): {{que_quan_tinh}}',
    r'7\) Nơi ở hiện nay[^:]+: \.+': '7) Nơi ở hiện nay (Xã, huyện, tỉnh hoặc số nhà, đường phố, TP): {{noi_o_hien_nay}}',
    r'đ/thoại: \.+': 'đ/thoại: {{dien_thoai}}',
    r'8\) Dân tộc:[^:]+: \.+': '8) Dân tộc: {{dan_toc}}',
    r'9\) Tôn giáo: \.+': '9) Tôn giáo: {{ton_giao}}',
    r'10\) Thành phần gia đình xuất thân:\s+\.+': '10) Thành phần gia đình xuất thân: {{thanh_phan_xuat_than}}',
    
    # Mục 11-20
    r'11\) Nghề nghiệp bản thân trước khi được tuyển dụng:\s+\.+': 
        '11) Nghề nghiệp bản thân trước khi được tuyển dụng: {{nghe_nghiep_truoc}}',
    r'12\) Ngày được tuyển dụng: \.+ / \.+ / \.+': '12) Ngày được tuyển dụng: {{ngay_tuyen_dung}}',
    r'Vào cơ quan nào, ở d[âấ]u: \.+': 'Vào cơ quan nào, ở đâu: {{co_quan_tuyen_dung}}',
    r'13\) Ngày vào cơ quan hiện đang công tác: \.+ / \.+ / \.+': 
        '13) Ngày vào cơ quan hiện đang công tác: {{ngay_vao_co_quan}}',
    r'Ngày tham gia cách mạng: \.+ / \.+ / \.+': 'Ngày tham gia cách mạng: {{ngay_tham_gia_cach_mang}}',
    r'14\) Ngày vào Đảng Cộng sản Việt Nam: \.+ / \.+ / \.+': 
        '14) Ngày vào Đảng Cộng sản Việt Nam: {{ngay_vao_dang}}',
    r'Ngày chính thức: \.+ / \.+ / \.+': 'Ngày chính thức: {{ngay_chinh_thuc_dang}}',
    r'15\) Ngày tham gia các tổ chức chính trị, xã hội:\s+\.+': 
        '15) Ngày tham gia các tổ chức chính trị, xã hội: {{to_chuc_chinh_tri_xa_hoi}}',
    r'16\) Ngày nhập ngũ: \.+ / \.+ / \.+': '16) Ngày nhập ngũ: {{ngay_nhap_ngu}}',
    r'Ngày xuất ngũ: \.+ / \.+ / \.+': 'Ngày xuất ngũ: {{ngay_xuat_ngu}}',
    r'Quân hàm, chức vụ cao nhất \(năm\): \.+': 'Quân hàm, chức vụ cao nhất (năm): {{quan_ham_chuc_vu}}',
    r'17\) Trình độ học vấn: Giáo dục phổ thông: \.+': '17) Trình độ học vấn: Giáo dục phổ thông: {{hoc_van_pho_thong}}',
    r'Học hàm, học vị cao nhất: \.+': 'Học hàm, học vị cao nhất: {{hoc_ham_hoc_vi}}',
    r'- Lý luận chính trị: \.+': '- Lý luận chính trị: {{ly_luan_chinh_tri}}',
    r'- Ngoại ngữ: \.+': '- Ngoại ngữ: {{ngoai_ngu}}',
    r'18\) Công tác chính đang làm:\s+\.+': '18) Công tác chính đang làm: {{cong_tac_chinh}}',
    r'19\) Ngạch công chức: \.+\(mã số: \.+\)': '19) Ngạch công chức: {{ngach_cong_chuc}} (mã số: {{ma_ngach}})',
    r'Bậc lương: \.+, hệ số: \.+ từ tháng \.+ /\.+': 
        'Bậc lương: {{bac_luong}}, hệ số: {{he_so_luong}} từ tháng {{thang_huong_luong}}',
    r'20\) Danh hiệu được phong \(năm nào\):\s+\.+': '20) Danh hiệu được phong (năm nào): {{danh_hieu}}',
    
    # Mục 21-25
    r'21\) Sở trường công tác: \.+': '21) Sở trường công tác: {{so_truong}}',
    r'Công việc đã làm lâu nhất: \.+': 'Công việc đã làm lâu nhất: {{cong_viec_lau_nhat}}',
    r'22\) Khen thưởng:\s+\.+': '22) Khen thưởng: {{khen_thuong}}',
    r'23\) Kỷ luật[^:]+:\s+\.+': '23) Kỷ luật (Đảng, Chính quyền, Đoàn thể, Cấp quyết định, năm nào, lý do, hình thức, ...): {{ky_luat}}',
    r'24\) Tình trạng sức khỏe: \.+': '24) Tình trạng sức khỏe: {{suc_khoe}}',
    r'Cao: 1m\.+, Cân nặng: \.+ \(kg\), Nhóm máu: \.+': 
        'Cao: {{chieu_cao}}, Cân nặng: {{can_nang}} (kg), Nhóm máu: {{nhom_mau}}',
    r'25\) Số chứng minh nhân dân: \.+': '25) Số chứng minh nhân dân: {{so_cmnd}}',
    r'Thương binh loại: \.+': 'Thương binh loại: {{thuong_binh_loai}}',
    r'Gia đình liệt sĩ:': 'Gia đình liệt sĩ: {{gia_dinh_liet_si}}',
}

# Apply replacements to paragraphs
for para in doc.paragraphs:
    for pattern, replacement in replacements.items():
        if re.search(pattern, para.text):
            # Clear all runs
            for run in para.runs[:]:
                run.text = ''
            # Add replacement text
            new_run = para.add_run(re.sub(pattern, replacement, para.text))
            new_run.font.name = 'Times New Roman'
            new_run.font.size = Pt(13)

print("  ✓ Đã thay thế thông tin cơ bản")

# =============================================================================
# BƯỚC 2: Xử lý 5 BẢNG
# =============================================================================
print("\n🔧 BƯỚC 2: Xử lý 5 bảng chi tiết")

# BẢNG 1: Đào tạo, bồi dưỡng (26)
print("  📋 Bảng 1: Đào tạo")
table1 = doc.tables[0]
# Add template row with mustache loop
row = table1.rows[1]
for i, cell in enumerate(row.cells):
    cell.text = ''
    if i == 0:
        cell.text = '{{#dao_tao}}{{ten_truong}}{{/dao_tao}}'
    elif i == 1:
        cell.text = '{{#dao_tao}}{{nganh_hoc}}{{/dao_tao}}'
    elif i == 2:
        cell.text = '{{#dao_tao}}{{thoi_gian}}{{/dao_tao}}'
    elif i == 3:
        cell.text = '{{#dao_tao}}{{hinh_thuc}}{{/dao_tao}}'
    elif i == 4:
        cell.text = '{{#dao_tao}}{{van_bang}}{{/dao_tao}}'

# BẢNG 2: Quá trình công tác (27)
print("  📋 Bảng 2: Quá trình công tác")
table2 = doc.tables[1]
row = table2.rows[1]
row.cells[0].text = '{{#qua_trinh_cong_tac}}{{thoi_gian}}{{/qua_trinh_cong_tac}}'
row.cells[1].text = '{{#qua_trinh_cong_tac}}{{chuc_danh_don_vi}}{{/qua_trinh_cong_tac}}'

# BẢNG 3: Quan hệ gia đình - Bản thân (30a)
print("  📋 Bảng 3: Gia đình bản thân")
table3 = doc.tables[2]
row = table3.rows[1]
for i, cell in enumerate(row.cells):
    cell.text = ''
    if i == 0:
        cell.text = '{{#gia_dinh_ban_than}}{{quan_he}}{{/gia_dinh_ban_than}}'
    elif i == 1:
        cell.text = '{{#gia_dinh_ban_than}}{{ho_ten}}{{/gia_dinh_ban_than}}'
    elif i == 2:
        cell.text = '{{#gia_dinh_ban_than}}{{nam_sinh}}{{/gia_dinh_ban_than}}'
    elif i == 3:
        cell.text = '{{#gia_dinh_ban_than}}{{thong_tin}}{{/gia_dinh_ban_than}}'

# BẢNG 4: Quan hệ gia đình - Vợ/Chồng (30b)
print("  📋 Bảng 4: Gia đình bên vợ/chồng")
table4 = doc.tables[3]
row = table4.rows[1]
for i, cell in enumerate(row.cells):
    cell.text = ''
    if i == 0:
        cell.text = '{{#gia_dinh_vo_chong}}{{quan_he}}{{/gia_dinh_vo_chong}}'
    elif i == 1:
        cell.text = '{{#gia_dinh_vo_chong}}{{ho_ten}}{{/gia_dinh_vo_chong}}'
    elif i == 2:
        cell.text = '{{#gia_dinh_vo_chong}}{{nam_sinh}}{{/gia_dinh_vo_chong}}'
    elif i == 3:
        cell.text = '{{#gia_dinh_vo_chong}}{{thong_tin}}{{/gia_dinh_vo_chong}}'

# BẢNG 5: Quá trình lương (31)
print("  📋 Bảng 5: Quá trình lương")
table5 = doc.tables[4]
# This is a horizontal timeline table - handle differently
# Replace column headers with variables
for col_idx in range(1, min(7, len(table5.columns))):
    cell = table5.cell(0, col_idx)
    cell.text = f'{{{{#qua_trinh_luong}}}}{{{{thang_nam_{col_idx}}}}}{{{{/qua_trinh_luong}}}}'

print("  ✓ Đã thêm loop syntax cho tất cả bảng")

# =============================================================================
# BƯỚC 3: Phần kinh tế (31)
# =============================================================================
print("\n🔧 BƯỚC 3: Phần hoàn cảnh kinh tế")
# Find and replace in later paragraphs
for para in doc.paragraphs[60:]:  # Rough area where section 31 is
    text = para.text
    if 'lương' in text.lower() and '........' in text:
        text = re.sub(r'\+ lương: \.+', '+ lương: {{luong_hang_nam}}', text)
        text = re.sub(r'\+ Các nguồn khác:\s+\.+', '+ Các nguồn khác: {{nguon_khac}}', text)
        para.text = text
    if 'Nhà ở:' in text:
        text = re.sub(r'Được cấp, được thuê, loại nhà: \.+', 
                     'Được cấp, được thuê, loại nhà: {{nha_duoc_cap}}', text)
        text = re.sub(r'tổng diện tích sử dụng: \.+ m2', 
                     'tổng diện tích sử dụng: {{dien_tich_nha_cap}} m2', text)
        para.text = text
    if 'tự mua, tự xây' in text:
        text = re.sub(r'Nhà tự mua, tự xây, loại nhà: \.+', 
                     'Nhà tự mua, tự xây, loại nhà: {{nha_tu_mua}}', text)
        text = re.sub(r'tổng diện tích sử dụng: \.+ m2', 
                     'tổng diện tích sử dụng: {{dien_tich_nha_mua}} m2', text)
        para.text = text
    if 'Đất ở:' in text:
        text = re.sub(r'\+ Đất được cấp: \.+ m2', '+ Đất được cấp: {{dat_duoc_cap}} m2', text)
        text = re.sub(r'\+ Đất tự mua: \.+ m2', '+ Đất tự mua: {{dat_tu_mua}} m2', text)
        para.text = text
    if 'Đất sản xuất' in text:
        text = re.sub(r'Đất sản xuất, kinh doanh:[^\.]+\.+', 
                     'Đất sản xuất, kinh doanh: {{dat_san_xuat}}', text)
        para.text = text

print("  ✓ Đã thay thế thông tin kinh tế")

# Save
doc.save(output)
print(f"\n✅ HOÀN TẤT! Đã tạo template: {output}")
print(f"📊 Template có {len(doc.paragraphs)} đoạn, {len(doc.tables)} bảng với {{variables}} đầy đủ")
