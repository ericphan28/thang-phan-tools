"""
Test Extended Document API Features
Kiểm tra các tính năng nâng cao: watermark, password, batch operations, etc.
"""

import requests
import json
from pathlib import Path
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from docx import Document
from openpyxl import Workbook
from PIL import Image

BASE_URL = "http://localhost:8000/api/v1/documents"

# Colors for output
GREEN = "\033[92m"
RED = "\033[91m"
YELLOW = "\033[93m"
BLUE = "\033[94m"
RESET = "\033[0m"

def print_test(name, status, message=""):
    """Print test result with colors"""
    color = GREEN if status == "✅" else RED if status == "❌" else YELLOW
    print(f"{color}{status} {name}{RESET} {message}")

def print_category(name):
    """Print category header"""
    print(f"\n{BLUE}{'='*60}")
    print(f"📦 {name}")
    print(f"{'='*60}{RESET}\n")

def create_test_pdf(filename="test.pdf", pages=3):
    """Tạo file PDF test"""
    pdf_path = Path("uploads") / filename
    pdf_path.parent.mkdir(exist_ok=True)
    
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    for i in range(pages):
        c.drawString(100, 750, f"Test PDF - Page {i+1}")
        c.drawString(100, 730, "Tiếng Việt: Đây là nội dung test")
        c.showPage()
    c.save()
    
    with open(pdf_path, 'wb') as f:
        f.write(buffer.getvalue())
    
    return pdf_path

def create_protected_pdf(filename="protected.pdf", password="test123"):
    """Tạo PDF có password"""
    import pypdf
    
    pdf_path = create_test_pdf("temp_for_protect.pdf")
    protected_path = Path("uploads") / filename
    
    reader = pypdf.PdfReader(pdf_path)
    writer = pypdf.PdfWriter()
    
    for page in reader.pages:
        writer.add_page(page)
    
    writer.encrypt(password)
    
    with open(protected_path, 'wb') as f:
        writer.write(f)
    
    pdf_path.unlink()  # Delete temp file
    return protected_path

def create_test_word(filename="test.docx"):
    """Tạo file Word test"""
    doc_path = Path("uploads") / filename
    doc_path.parent.mkdir(exist_ok=True)
    
    doc = Document()
    doc.add_heading("Test Document", 0)
    doc.add_paragraph("This is test content")
    doc.add_paragraph("Tiếng Việt: Đây là văn bản test")
    doc.save(doc_path)
    
    return doc_path

def create_test_excel(filename="test.xlsx"):
    """Tạo file Excel test"""
    excel_path = Path("uploads") / filename
    excel_path.parent.mkdir(exist_ok=True)
    
    wb = Workbook()
    ws = wb.active
    ws['A1'] = "Test Data"
    ws['A2'] = "Row 1"
    ws['B2'] = "Value 1"
    wb.save(excel_path)
    
    return excel_path

# ==================== PDF Operations ====================

def test_watermark_text():
    """Test thêm watermark text vào PDF"""
    try:
        pdf_path = create_test_pdf("test_watermark.pdf")
        
        with open(pdf_path, 'rb') as f:
            response = requests.post(
                f"{BASE_URL}/pdf/watermark-text",
                files={"file": ("test.pdf", f, "application/pdf")},
                data={
                    "watermark_text": "CONFIDENTIAL",
                    "opacity": "0.3",
                    "position": "center"
                }
            )
        
        if response.status_code == 200:
            print_test("Watermark Text", "✅", f"({len(response.content)} bytes)")
            return True
        else:
            print_test("Watermark Text", "❌", f"Status {response.status_code}: {response.text[:100]}")
            return False
    except Exception as e:
        print_test("Watermark Text", "❌", str(e))
        return False

def test_protect_pdf():
    """Test bảo vệ PDF bằng password"""
    try:
        pdf_path = create_test_pdf("test_protect.pdf")
        
        with open(pdf_path, 'rb') as f:
            response = requests.post(
                f"{BASE_URL}/pdf/protect",
                files={"file": ("test.pdf", f, "application/pdf")},
                data={"password": "test123"}
            )
        
        if response.status_code == 200:
            print_test("Protect PDF", "✅", f"({len(response.content)} bytes)")
            return True
        else:
            print_test("Protect PDF", "❌", f"Status {response.status_code}: {response.text[:100]}")
            return False
    except Exception as e:
        print_test("Protect PDF", "❌", str(e))
        return False

def test_unlock_pdf():
    """Test mở khóa PDF có password"""
    try:
        pdf_path = create_protected_pdf("test_locked.pdf", "test123")
        
        with open(pdf_path, 'rb') as f:
            response = requests.post(
                f"{BASE_URL}/pdf/unlock",
                files={"file": ("locked.pdf", f, "application/pdf")},
                data={"password": "test123"}
            )
        
        if response.status_code == 200:
            print_test("Unlock PDF", "✅", f"({len(response.content)} bytes)")
            return True
        else:
            print_test("Unlock PDF", "❌", f"Status {response.status_code}: {response.text[:100]}")
            return False
    except Exception as e:
        print_test("Unlock PDF", "❌", str(e))
        return False

def test_add_page_numbers():
    """Test thêm số trang vào PDF"""
    try:
        pdf_path = create_test_pdf("test_page_numbers.pdf")
        
        with open(pdf_path, 'rb') as f:
            response = requests.post(
                f"{BASE_URL}/pdf/add-page-numbers",
                files={"file": ("test.pdf", f, "application/pdf")},
                data={
                    "position": "bottom-center",
                    "format": "Page {page} of {total}"
                }
            )
        
        if response.status_code == 200:
            print_test("Add Page Numbers", "✅", f"({len(response.content)} bytes)")
            return True
        else:
            print_test("Add Page Numbers", "❌", f"Status {response.status_code}: {response.text[:100]}")
            return False
    except Exception as e:
        print_test("Add Page Numbers", "❌", str(e))
        return False

def test_pdf_info():
    """Test lấy thông tin PDF"""
    try:
        pdf_path = create_test_pdf("test_info.pdf")
        
        with open(pdf_path, 'rb') as f:
            response = requests.post(
                f"{BASE_URL}/info/pdf",
                files={"file": ("test.pdf", f, "application/pdf")}
            )
        
        if response.status_code == 200:
            data = response.json()
            pages = data.get('pages', 0)
            print_test("PDF Info", "✅", f"({pages} pages)")
            return True
        else:
            print_test("PDF Info", "❌", f"Status {response.status_code}: {response.text[:100]}")
            return False
    except Exception as e:
        print_test("PDF Info", "❌", str(e))
        return False

def test_word_info():
    """Test lấy thông tin Word"""
    try:
        doc_path = create_test_word("test_word_info.docx")
        
        with open(doc_path, 'rb') as f:
            response = requests.post(
                f"{BASE_URL}/info/word",
                files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            )
        
        if response.status_code == 200:
            data = response.json()
            paragraphs = data.get('paragraphs', 0)
            print_test("Word Info", "✅", f"({paragraphs} paragraphs)")
            return True
        else:
            print_test("Word Info", "❌", f"Status {response.status_code}: {response.text[:100]}")
            return False
    except Exception as e:
        print_test("Word Info", "❌", str(e))
        return False

# ==================== Batch Operations ====================

def test_batch_word_to_pdf():
    """Test batch Word → PDF"""
    try:
        doc1 = create_test_word("batch1.docx")
        doc2 = create_test_word("batch2.docx")
        
        with open(doc1, 'rb') as f1, open(doc2, 'rb') as f2:
            response = requests.post(
                f"{BASE_URL}/batch/word-to-pdf",
                files=[
                    ("files", ("batch1.docx", f1, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")),
                    ("files", ("batch2.docx", f2, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))
                ]
            )
        
        if response.status_code == 200:
            print_test("Batch Word → PDF", "✅", f"({len(response.content)} bytes ZIP)")
            return True
        else:
            print_test("Batch Word → PDF", "❌", f"Status {response.status_code}: {response.text[:100]}")
            return False
    except Exception as e:
        print_test("Batch Word → PDF", "❌", str(e))
        return False

def test_batch_pdf_to_word():
    """Test batch PDF → Word"""
    try:
        pdf1 = create_test_pdf("batch_pdf1.pdf")
        pdf2 = create_test_pdf("batch_pdf2.pdf")
        
        with open(pdf1, 'rb') as f1, open(pdf2, 'rb') as f2:
            response = requests.post(
                f"{BASE_URL}/batch/pdf-to-word",
                files=[
                    ("files", ("batch1.pdf", f1, "application/pdf")),
                    ("files", ("batch2.pdf", f2, "application/pdf"))
                ]
            )
        
        if response.status_code == 200:
            print_test("Batch PDF → Word", "✅", f"({len(response.content)} bytes ZIP)")
            return True
        else:
            print_test("Batch PDF → Word", "❌", f"Status {response.status_code}: {response.text[:100]}")
            return False
    except Exception as e:
        print_test("Batch PDF → Word", "❌", str(e))
        return False

def test_batch_compress_pdf():
    """Test batch nén PDF"""
    try:
        pdf1 = create_test_pdf("compress1.pdf")
        pdf2 = create_test_pdf("compress2.pdf")
        
        with open(pdf1, 'rb') as f1, open(pdf2, 'rb') as f2:
            response = requests.post(
                f"{BASE_URL}/batch/compress-pdf",
                files=[
                    ("files", ("compress1.pdf", f1, "application/pdf")),
                    ("files", ("compress2.pdf", f2, "application/pdf"))
                ],
                data={"quality": "medium"}
            )
        
        if response.status_code == 200:
            print_test("Batch Compress PDF", "✅", f"({len(response.content)} bytes ZIP)")
            return True
        else:
            print_test("Batch Compress PDF", "❌", f"Status {response.status_code}: {response.text[:100]}")
            return False
    except Exception as e:
        print_test("Batch Compress PDF", "❌", str(e))
        return False

def test_merge_word_to_pdf():
    """Test gộp nhiều Word thành 1 PDF"""
    try:
        doc1 = create_test_word("merge1.docx")
        doc2 = create_test_word("merge2.docx")
        
        with open(doc1, 'rb') as f1, open(doc2, 'rb') as f2:
            response = requests.post(
                f"{BASE_URL}/batch/merge-word-to-pdf",
                files=[
                    ("files", ("merge1.docx", f1, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")),
                    ("files", ("merge2.docx", f2, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))
                ]
            )
        
        if response.status_code == 200:
            print_test("Merge Word → PDF", "✅", f"({len(response.content)} bytes)")
            return True
        else:
            print_test("Merge Word → PDF", "❌", f"Status {response.status_code}: {response.text[:100]}")
            return False
    except Exception as e:
        print_test("Merge Word → PDF", "❌", str(e))
        return False

# ==================== Advanced PDF Operations ====================

def test_pdf_extract_content():
    """Test trích xuất nội dung PDF (images, fonts)"""
    try:
        pdf_path = create_test_pdf("test_extract_content.pdf")
        
        with open(pdf_path, 'rb') as f:
            response = requests.post(
                f"{BASE_URL}/pdf/extract-content",
                files={"file": ("test.pdf", f, "application/pdf")}
            )
        
        if response.status_code == 200:
            print_test("Extract PDF Content", "✅", f"({len(response.content)} bytes ZIP)")
            return True
        else:
            print_test("Extract PDF Content", "❌", f"Status {response.status_code}: {response.text[:100]}")
            return False
    except Exception as e:
        print_test("Extract PDF Content", "❌", str(e))
        return False

def test_html_to_pdf():
    """Test HTML → PDF"""
    try:
        html_content = """
        <html>
        <head><title>Test HTML</title></head>
        <body>
            <h1>Test HTML to PDF</h1>
            <p>This is a test paragraph.</p>
            <p>Tiếng Việt: Đây là nội dung HTML</p>
        </body>
        </html>
        """
        
        response = requests.post(
            f"{BASE_URL}/convert/html-to-pdf",
            data={"html_content": html_content}
        )
        
        if response.status_code == 200:
            print_test("HTML → PDF", "✅", f"({len(response.content)} bytes)")
            return True
        else:
            print_test("HTML → PDF", "❌", f"Status {response.status_code}: {response.text[:100]}")
            return False
    except Exception as e:
        print_test("HTML → PDF", "❌", str(e))
        return False

# ==================== Main Test Runner ====================

def main():
    """Chạy tất cả extended tests"""
    print("\n" + "="*60)
    print("🧪 TESTING EXTENDED DOCUMENT API FEATURES")
    print("="*60)
    
    # PDF Operations
    print_category("PDF Operations")
    pdf_ops = [
        ("Watermark Text", test_watermark_text),
        ("Protect PDF (Password)", test_protect_pdf),
        ("Unlock PDF", test_unlock_pdf),
        ("Add Page Numbers", test_add_page_numbers),
        ("PDF Info", test_pdf_info),
        ("Word Info", test_word_info),
        ("Extract PDF Content", test_pdf_extract_content),
        ("HTML → PDF", test_html_to_pdf),
    ]
    
    # Batch Operations
    print_category("Batch Operations")
    batch_ops = [
        ("Batch Word → PDF", test_batch_word_to_pdf),
        ("Batch PDF → Word", test_batch_pdf_to_word),
        ("Batch Compress PDF", test_batch_compress_pdf),
        ("Merge Word → PDF", test_merge_word_to_pdf),
    ]
    
    all_tests = pdf_ops + batch_ops
    results = []
    
    for name, test_func in all_tests:
        try:
            result = test_func()
            results.append((name, result))
        except Exception as e:
            print_test(name, "❌", f"Exception: {e}")
            results.append((name, False))
    
    # Summary
    print("\n" + "="*60)
    print("📊 EXTENDED TEST SUMMARY")
    print("="*60)
    
    passed = sum(1 for _, result in results if result)
    failed = len(results) - passed
    
    print(f"\n{GREEN}✅ Passed: {passed}{RESET}")
    print(f"{RED}❌ Failed: {failed}{RESET}")
    print(f"Total: {len(results)}")
    print(f"Success Rate: {passed/len(results)*100:.1f}%")
    
    if failed > 0:
        print(f"\n{RED}Failed tests:{RESET}")
        for name, result in results:
            if not result:
                print(f"  - {name}")
    
    print("\n" + "="*60 + "\n")
    
    return passed, failed

if __name__ == "__main__":
    passed, failed = main()
    exit(0 if failed == 0 else 1)
