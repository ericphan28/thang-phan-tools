"""
Test Gemini API for PDF to Word conversion
"""

import asyncio
import os
import sys
from pathlib import Path
from docx import Document
from dotenv import load_dotenv

# Load environment variables from backend/.env
load_dotenv(dotenv_path=Path("backend/.env"))

# Add backend to path
sys.path.append(str(Path(__file__).parent / "backend"))
from app.services.document_service import DocumentService

async def test_gemini_pdf_to_word():
    """Test Gemini PDF to Word conversion"""
    
    # Create test PDF with Vietnamese text
    import fitz  # PyMuPDF
    
    # Create a simple PDF with Vietnamese text
    test_pdf = Path("test_vietnamese.pdf")
    doc = fitz.open()
    page = doc.new_page()
    
    vietnamese_text = """
    Xin chào! Đây là tài liệu tiếng Việt.
    
    Nội dung bao gồm:
    - Văn bản tiếng Việt có dấu
    - Các ký tự đặc biệt: áéíóúàèìòùâêîôûăđĩõũ
    - Số liệu: 123,456.78
    - Ngày tháng: 15/12/2024
    
    Đây là đoạn văn bản để test Gemini API.
    """
    
    page.insert_text((50, 100), vietnamese_text, fontsize=12)
    doc.save(str(test_pdf))
    doc.close()
    
    print(f"✅ Created test PDF: {test_pdf}")
    
    # Test DocumentService
    output_dir = Path("./test_outputs")
    output_dir.mkdir(exist_ok=True)
    
    service = DocumentService(output_dir=output_dir)
    
    if not service.gemini_model:
        print("❌ Gemini API not configured!")
        return
    
    print("✅ Gemini API configured successfully")
    
    try:
        # Convert PDF to Word
        print("\n🔄 Converting PDF to Word with Gemini...")
        result_path = await service.pdf_to_word_with_gemini(test_pdf)
        
        print(f"✅ Conversion completed: {result_path}")
        
        # Check if file exists and has content
        if result_path.exists():
            file_size = result_path.stat().st_size
            print(f"📊 File size: {file_size} bytes")
            
            if file_size > 1000:  # More than 1KB
                print("✅ File size looks good!")
                
                # Try to read the Word document
                try:
                    doc = Document(str(result_path))
                    paragraph_count = len(doc.paragraphs)
                    text_content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                    
                    print(f"📄 Paragraphs: {paragraph_count}")
                    print(f"📝 Content preview (first 200 chars):")
                    print(f"   {text_content[:200]}...")
                    
                    if text_content.strip():
                        print("✅ Word document contains text!")
                    else:
                        print("❌ Word document is empty!")
                        
                except Exception as e:
                    print(f"❌ Error reading Word document: {e}")
            else:
                print("⚠️ File size is too small, likely empty")
        else:
            print("❌ Output file was not created")
            
    except Exception as e:
        print(f"❌ Error during conversion: {e}")
        import traceback
        traceback.print_exc()
    
    # Cleanup
    try:
        test_pdf.unlink()
        print(f"\n🧹 Cleaned up test file: {test_pdf}")
    except:
        pass

if __name__ == "__main__":
    asyncio.run(test_gemini_pdf_to_word())