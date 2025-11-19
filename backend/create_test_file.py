from docx import Document
from pathlib import Path

# Create test Word document
doc = Document()
doc.add_heading('Test Document Conversion', 0)

doc.add_paragraph('Đây là văn bản tiếng Việt có dấu.')
doc.add_paragraph('This is English text.')

doc.add_heading('Features to test:', level=1)
doc.add_paragraph('✓ Text extraction', style='List Bullet')
doc.add_paragraph('✓ PDF conversion', style='List Bullet')
doc.add_paragraph('✓ Word info retrieval', style='List Bullet')

# Save
output_dir = Path('D:/thang/utility-server/backend/uploads/test_files')
output_dir.mkdir(parents=True, exist_ok=True)

output_path = output_dir / 'test_document.docx'
doc.save(output_path)

print(f"✅ Created test file: {output_path}")
print(f"📊 File size: {output_path.stat().st_size} bytes")
