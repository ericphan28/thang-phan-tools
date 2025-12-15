#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Service for Mẫu 2C-TCTW-98 (Sơ yếu lý lịch)
"""

from docxtpl import DocxTemplate
from docx import Document
import os
import re
from datetime import datetime
from pathlib import Path

class Mau2CService:
    """Service to generate Mẫu 2C-TCTW-98 documents"""
    
    def __init__(self):
        self.template_path = Path(__file__).parent.parent / "templates" / "mau_2c_V10_WITH_LOOPS.docx"
        
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template not found: {self.template_path}")
    
    def flatten_dict(self, d, parent_key='', sep='_'):
        """Flatten nested dictionary"""
        items = []
        for k, v in d.items():
            if k.startswith('_'):
                continue
            new_key = f"{parent_key}{sep}{k}" if parent_key else k
            if isinstance(v, dict):
                items.extend(self.flatten_dict(v, new_key, sep=sep).items())
            else:
                items.append((new_key, v))
        return dict(items)
    
    def map_arrays_for_template(self, data):
        """
        Map JSON arrays to template variable names
        
        JSON -> Template mapping:
        - dao_tao -> hoc_tap
        - cong_tac -> cong_tac (with key mapping)
        - gia_dinh -> gia_dinh.ban_than
        - gia_dinh_vo_chong -> gia_dinh.vo_chong
        - luong -> he_so_luong_history
        """
        render_data = {}
        
        # Map dao_tao -> hoc_tap
        if 'dao_tao' in data:
            render_data['hoc_tap'] = data['dao_tao']
        
        # Map cong_tac with key transformation
        if 'cong_tac' in data:
            render_data['cong_tac'] = [
                {
                    'tu_thang_nam': ct.get('thoi_gian', ''),
                    'den_thang_nam': ct.get('den_thang_nam', ''),
                    'chuc_danh': ct.get('chuc_vu_don_vi', '')
                }
                for ct in data['cong_tac']
            ]
        
        # Map gia_dinh arrays
        if 'gia_dinh' in data or 'gia_dinh_vo_chong' in data:
            render_data['gia_dinh'] = {
                'ban_than': data.get('gia_dinh', []),
                'vo_chong': data.get('gia_dinh_vo_chong', [])
            }
        
        # Map luong -> he_so_luong_history
        if 'luong' in data:
            render_data['he_so_luong_history'] = [
                {
                    'tu_thang_nam': l.get('thang_nam', ''),
                    'chuc_danh': 'Chuyên viên',
                    'bac': l.get('ngach_bac', '').split('Bậc ')[-1] if 'Bậc' in l.get('ngach_bac', '') else '',
                    'he_so': l.get('he_so', '')
                }
                for l in data['luong']
            ]
        
        return render_data
    
    def clean_dots(self, doc):
        """Remove sequences of 3+ dots from document"""
        cleaned_count = 0
        
        # Clean paragraphs
        for para in doc.paragraphs:
            for run in para.runs:
                if re.search(r'[.…]{3,}', run.text):
                    run.text = re.sub(r'[.…]{3,}', '', run.text).strip()
                    cleaned_count += 1
        
        # Clean tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if re.search(r'[.…]{3,}', run.text):
                                run.text = re.sub(r'[.…]{3,}', '', run.text).strip()
                                cleaned_count += 1
        
        return cleaned_count
    
    def generate(self, data: dict, output_dir: str = None) -> str:
        """
        Generate Mẫu 2C document from data
        
        Args:
            data: Dictionary containing all required fields
            output_dir: Directory to save output file (default: temp directory)
        
        Returns:
            Path to generated file
        """
        try:
            # Load template
            doc_template = DocxTemplate(str(self.template_path))
            
            # Flatten simple fields
            flat_data = self.flatten_dict(data)
            
            # Prepare render data
            render_data = {**flat_data}
            
            # Map arrays with correct keys
            array_mappings = self.map_arrays_for_template(data)
            render_data.update(array_mappings)
            
            # Render template
            doc_template.render(render_data)
            
            # Save to temporary file
            if output_dir is None:
                output_dir = Path(__file__).parent.parent.parent / "output"
            else:
                output_dir = Path(output_dir)
            
            output_dir.mkdir(parents=True, exist_ok=True)
            
            # Generate filename with timestamp (ASCII-safe)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            # Use timestamp only for filename to avoid encoding issues
            filename = f"Mau_2C_{timestamp}.docx"
            temp_path = output_dir / filename
            
            doc_template.save(str(temp_path))
            
            # Clean dots
            doc = Document(str(temp_path))
            cleaned = self.clean_dots(doc)
            
            # Save cleaned version
            final_path = output_dir / filename.replace('.docx', '_cleaned.docx')
            doc.save(str(final_path))
            
            # Remove temp file
            temp_path.unlink()
            
            return str(final_path)
            
        except Exception as e:
            raise Exception(f"Failed to generate Mẫu 2C: {str(e)}")
