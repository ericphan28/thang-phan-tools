"""
Service kiểm tra thể thức văn bản hành chính
"""
import json
import logging
from pathlib import Path
from typing import Dict, Optional
from sqlalchemy.orm import Session

from app.services.gemini_service import GeminiService
from app.services.prompts.vb_checker_prompts import PROMPT_KIEM_TRA_THE_THUC, PROMPT_KIEM_TRA_CO_BAN
from app.models.van_ban_models import VanBanHanhChinh, MucDoKhan, TrangThaiVanBan

logger = logging.getLogger(__name__)


class VBCheckerService:
    """
    Service kiểm tra thể thức văn bản hành chính theo Nghị định 30/2020
    """
    
    def __init__(self, db: Session, user_id: int):
        self.db = db
        self.user_id = user_id
        self.gemini = GeminiService(db, user_id)
    
    async def check_the_thuc(
        self, 
        file_path: str, 
        file_name: str,
        chi_tiet_cao: bool = False,
        luu_database: bool = True
    ) -> Dict:
        """
        Kiểm tra thể thức văn bản
        
        Args:
            file_path: Đường dẫn file văn bản
            file_name: Tên file gốc
            chi_tiet_cao: Kiểm tra chi tiết (dùng prompt phức tạp)
            luu_database: Có lưu kết quả vào DB không
            
        Returns:
            Dict chứa kết quả kiểm tra
        """
        logger.info(f"===== BẮT ĐẦU KIỂM TRA VĂN BẢN: {file_name} =====")
        logger.info(f"File path: {file_path}, Chi tiết cao: {chi_tiet_cao}, Lưu DB: {luu_database}")
        
        file_path_obj = Path(file_path)
        file_ext = file_path_obj.suffix.lower()
        
        try:
            # CHIẾN LƯỢC TỐI ƯU:
            # - TXT: Extract local (giữ 100% format, free)
            # - DOCX/DOC/PDF/Image: Gemini Vision (giữ layout 2 cột, header/footer)
            
            if file_ext == '.txt':
                # ===== LUỒNG 1: TXT FILES (simple text) =====
                logger.info(f"BƯỚC 1: Reading TXT file (local)...")
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                logger.info(f"✅ Text read: {len(text)} characters")
                
                if not text or len(text) < 100:
                    return {
                        "success": False,
                        "message": "File text quá ngắn hoặc rỗng."
                    }
                
                # BƯỚC 2: Gọi Gemini với TEXT
                logger.info(f"BƯỚC 2: Analyzing text with Gemini...")
                prompt = PROMPT_KIEM_TRA_THE_THUC if chi_tiet_cao else PROMPT_KIEM_TRA_THE_THUC
                prompt_filled = prompt.format(noi_dung_van_ban=text[:8000])
                
                response = self.gemini.generate_content(
                    prompt=prompt_filled,
                    model="gemini-2.0-flash-exp",
                    operation="check_van_ban_text"
                )
                text_for_db = text
                
            else:
                # ===== LUỒNG 2: DOCX/PDF/IMAGE =====
                # Gemini Vision: OCR + phân tích trong 1 call, GIỮ NGUYÊN LAYOUT
                logger.info(f"BƯỚC 1+2 (tối ưu): Gemini Vision OCR + analysis for {file_name}...")
                logger.info(f"Lý do: File {file_ext} có layout phức tạp (2 cột, header/footer)")
                
                prompt_template = PROMPT_KIEM_TRA_THE_THUC if chi_tiet_cao else PROMPT_KIEM_TRA_THE_THUC
                # Bỏ placeholder vì Gemini tự OCR
                prompt_direct = prompt_template.replace(
                    "Văn bản cần kiểm tra:\n{noi_dung_van_ban}",
                    "Hãy OCR toàn bộ nội dung văn bản từ file đính kèm (bao gồm cả header, footer, 2 cột) và kiểm tra thể thức."
                )
                
                if file_ext == '.pdf':
                    response = self.gemini.generate_content_with_pdf(
                        prompt=prompt_direct,
                        pdf_path=file_path,
                        model="gemini-2.5-flash",
                        operation="check_van_ban_pdf"
                    )
                elif file_ext in ['.docx', '.doc']:
                    # DOCX: Cố gắng convert → PDF (giữ layout), nếu fail → extract text
                    logger.info(f"Processing DOCX file...")
                    
                    try:
                        # Option 1: Convert DOCX → PDF (preserve layout, cần Gotenberg/LibreOffice)
                        logger.info(f"Trying to convert DOCX to PDF for layout preservation...")
                        from app.services.document_service import DocumentService
                        doc_service = DocumentService()
                        
                        pdf_path = await doc_service.word_to_pdf(
                            input_file=Path(file_path)
                        )
                        logger.info(f"✅ Converted to PDF: {pdf_path}")
                        
                        # Gemini analyze PDF
                        response = self.gemini.generate_content_with_pdf(
                            prompt=prompt_direct,
                            pdf_path=str(pdf_path),
                            model="gemini-2.5-flash",
                            operation="check_van_ban_docx_as_pdf"
                        )
                        text_for_db = "(Analyzed via PDF - full layout preserved)"
                        
                        # Cleanup
                        try:
                            pdf_path.unlink()
                        except:
                            pass
                    
                    except Exception as convert_error:
                        # Option 2: Fallback - Extract text trực tiếp (mất layout 2 cột)
                        logger.warning(f"PDF conversion failed: {str(convert_error)}")
                        logger.info(f"Fallback: Extracting text with python-docx (layout may be lost)")
                        
                        try:
                            from docx import Document
                            doc = Document(file_path)
                            text = "\n".join([p.text for p in doc.paragraphs])
                            logger.info(f"✅ Text extracted: {len(text)} chars (WARNING: 2-column layout may be merged)")
                            
                            if not text or len(text) < 100:
                                return {
                                    "success": False,
                                    "message": "File DOCX quá ngắn hoặc rỗng."
                                }
                            
                            # Gemini phân tích text (không có layout)
                            prompt_filled = prompt_template.format(noi_dung_van_ban=text[:8000])
                            response = self.gemini.generate_content(
                                prompt=prompt_filled,
                                model="gemini-2.0-flash-exp",
                                operation="check_van_ban_docx_text"
                            )
                            text_for_db = text
                            
                        except Exception as docx_error:
                            logger.error(f"python-docx also failed: {str(docx_error)}")
                            return {
                                "success": False,
                                "message": f"Không thể đọc file Word. Vui lòng upload file PDF hoặc kiểm tra file có hợp lệ không."
                            }
                else:  # .jpg, .jpeg, .png
                    response = self.gemini.generate_content_with_image(
                        prompt=prompt_direct,
                        image_path=file_path,
                        model="gemini-2.5-flash",
                        operation="check_van_ban_image"
                    )
                
                # Text placeholder cho DB
                text_for_db = "(Analyzed via Gemini Vision - preserves original layout)"
                logger.info(f"✅ Gemini Vision processed file (1 API call, layout preserved)")
            
            logger.info(f"✅ Gemini response received, length: {len(response.text) if response.text else 0}")
            if response.text:
                preview = response.text[:300].replace('\n', ' ')
                logger.debug(f"Response preview: {preview}...")
            
            # BƯỚC 3: Parse JSON response
            logger.info(f"BƯỚC 3: Parsing JSON response...")
            result = self._parse_gemini_response(response.text)
            
            if not result["success"]:
                return result
            
            # BƯỚC 4: Lưu vào database (nếu cần)
            logger.info(f"BƯỚC 4: Saving to database..." if luu_database else "BƯỚC 4: Skip saving (luu_database=False)")
            if luu_database:
                van_ban_id = await self._save_to_database(
                    file_path=file_path,
                    file_name=file_name,
                    text_content=text_for_db,
                    result=result
                )
                result["van_ban_id"] = van_ban_id
            
            logger.info(f"Document check completed. Score: {result['tong_diem']}/100")
            return result
            
        except Exception as e:
            logger.error(f"Error checking document format: {str(e)}", exc_info=True)
            return {
                "success": False,
                "message": f"Lỗi khi kiểm tra văn bản: {str(e)}"
            }
    
    async def _extract_text(self, file_path: str) -> str:
        """
        Trích xuất text từ file (PDF, DOCX, hoặc image)
        """
        file_path_obj = Path(file_path)
        file_ext = file_path_obj.suffix.lower()
        
        try:
            if file_ext in ['.pdf', '.jpg', '.jpeg', '.png']:
                # Dùng OCR service có sẵn
                from app.services.ocr_service import OCRService
                ocr = OCRService(self.db, self.user_id)
                
                # OCR với Gemini (sử dụng vision)
                result = await ocr.process_with_gemini(
                    file_path=file_path,
                    prompt="Trích xuất toàn bộ nội dung văn bản. Giữ nguyên định dạng và cấu trúc."
                )
                return result.get("text", "")
                
            elif file_ext in ['.docx', '.doc']:
                # Đọc DOCX - với fallback nếu file corrupt
                try:
                    from docx import Document
                    doc = Document(file_path)
                    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    return text
                except Exception as docx_error:
                    logger.warning(f"python-docx failed, fallback to OCR: {str(docx_error)}")
                    # Fallback: Convert doc -> OCR (xử lý file corrupt/old format)
                    from app.services.ocr_service import OCRService
                    ocr = OCRService(self.db, self.user_id)
                    result = await ocr.process_with_gemini(
                        file_path=file_path,
                        prompt="Trích xuất toàn bộ nội dung văn bản. Giữ nguyên định dạng và cấu trúc."
                    )
                    return result.get("text", "")
                
            elif file_ext == '.txt':
                # Đọc text thuần
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                raise ValueError(f"Định dạng file không được hỗ trợ: {file_ext}")
                
        except ValueError:
            # Re-raise format errors
            raise
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise ValueError(f"Không thể đọc file. Vui lòng kiểm tra file có hợp lệ không.")
    
    def _parse_gemini_response(self, response: str) -> Dict:
        """
        Parse response từ Gemini API
        """
        try:
            # Loại bỏ markdown code block nếu có
            response_clean = response.strip()
            if response_clean.startswith("```json"):
                response_clean = response_clean[7:]
            if response_clean.startswith("```"):
                response_clean = response_clean[3:]
            if response_clean.endswith("```"):
                response_clean = response_clean[:-3]
            
            response_clean = response_clean.strip()
            
            # Parse JSON
            data = json.loads(response_clean)
            
            # Validate required fields
            required_fields = ["tong_diem", "loai_van_ban", "vi_pham", "dat_yeu_cau"]
            for field in required_fields:
                if field not in data:
                    raise ValueError(f"Missing required field: {field}")
            
            return {
                "success": True,
                "tong_diem": data["tong_diem"],
                "loai_van_ban": data["loai_van_ban"],
                "vi_pham": data["vi_pham"],
                "dat_yeu_cau": data["dat_yeu_cau"],
                "message": "Kiểm tra thành công"
            }
            
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse Gemini response as JSON: {str(e)}")
            logger.error(f"Response content: {response[:500]}")
            return {
                "success": False,
                "message": "AI trả về dữ liệu không hợp lệ. Vui lòng thử lại."
            }
        except Exception as e:
            logger.error(f"Error parsing Gemini response: {str(e)}")
            return {
                "success": False,
                "message": f"Lỗi xử lý kết quả: {str(e)}"
            }
    
    async def _save_to_database(
        self, 
        file_path: str, 
        file_name: str,
        text_content: str,
        result: Dict
    ) -> int:
        """
        Lưu kết quả kiểm tra vào database
        """
        try:
            file_ext = Path(file_path).suffix.lower()
            file_type = "PDF" if file_ext == ".pdf" else "DOCX" if file_ext in [".docx", ".doc"] else "IMAGE"
            
            van_ban = VanBanHanhChinh(
                user_id=self.user_id,
                file_url=file_path,  # TODO: Upload lên S3/cloud storage
                file_name=file_name,
                file_type=file_type,
                loai_van_ban=result.get("loai_van_ban", "KHAC"),
                noi_dung_full=text_content[:5000],  # Giới hạn lưu 5000 ký tự
                diem_the_thuc=result["tong_diem"],
                vi_pham_the_thuc=result["vi_pham"],
                dat_yeu_cau=result["dat_yeu_cau"],
                trang_thai=TrangThaiVanBan.CHO_XU_LY
            )
            
            self.db.add(van_ban)
            self.db.commit()
            self.db.refresh(van_ban)
            
            logger.info(f"Saved document check result to database. ID: {van_ban.id}")
            return van_ban.id
            
        except Exception as e:
            logger.error(f"Error saving to database: {str(e)}")
            self.db.rollback()
            raise
