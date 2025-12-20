# -*- coding: utf-8 -*-
"""
AI Report Generator Service
Generates beautiful Word reports from text using AI + python-docx styling
"""

import os
import json
import time
from pathlib import Path
from typing import Dict, Any, List
from fastapi import HTTPException
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import logging

logger = logging.getLogger(__name__)


class AIReportGenerator:
    """Generate beautiful Word reports using AI reasoning + python-docx styling"""
    
    def __init__(self, output_dir: Path = None):
        self.output_dir = output_dir or Path("temp_outputs")
        self.output_dir.mkdir(exist_ok=True)
        
    async def generate_comparison_report(
        self, 
        text_input: str, 
        title: str = None,
        language: str = "vi"
    ) -> Path:
        """
        Generate comparison report (like Microsoft Graph API vs REST API example)
        
        Args:
            text_input: Raw text to analyze
            title: Report title (auto-generated if not provided)
            language: vi or en
            
        Returns:
            Path to generated .docx file
        """
        start_time = time.time()
        
        try:
            # Step 1: Analyze text with Gemini/Claude
            logger.info("Analyzing text with AI...")
            analysis = await self._analyze_text_for_comparison(text_input, language)
            
            # Step 2: Generate Word document
            logger.info("Creating Word document...")
            output_path = self.output_dir / f"report_{int(time.time())}.docx"
            self._create_comparison_report_docx(analysis, output_path)
            
            processing_time = time.time() - start_time
            logger.info(f"Report generated in {processing_time:.2f}s")
            
            return output_path
            
        except Exception as e:
            logger.error(f"Report generation failed: {e}")
            raise HTTPException(500, f"Report generation failed: {str(e)}")
    
    async def _analyze_text_for_comparison(self, text: str, language: str) -> Dict[str, Any]:
        """
        Use Gemini to analyze text and structure into comparison format
        
        Returns JSON structure:
        {
            "title": "...",
            "introduction": "...",
            "comparison_table": {
                "headers": ["Feature", "Option A", "Option B"],
                "rows": [[...], [...]]
            },
            "sections": [
                {"title": "...", "content": "...", "bullets": [...]},
                ...
            ],
            "conclusion": "..."
        }
        """
        try:
            import google.generativeai as genai
            
            # Get API key
            api_key = None
            try:
                from app.services.ai_usage_service import get_api_key
                api_key = get_api_key("gemini")
            except:
                api_key = os.getenv("GOOGLE_API_KEY")
            
            if not api_key:
                raise HTTPException(500, "Gemini API key not configured")
                
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel("gemini-2.0-flash-exp")
            
            # Prompt for comparison analysis
            lang_desc = "Vietnamese" if language == "vi" else "English"
            prompt = f"""Analyze the following text and structure it into a comparison report format.

INPUT TEXT:
{text}

OUTPUT INSTRUCTIONS:
Return ONLY a valid JSON object (no markdown, no code blocks) with this exact structure:

{{
    "title": "Main title of the report",
    "introduction": "Brief introduction paragraph",
    "comparison_table": {{
        "headers": ["Feature/Aspect", "Option/Method 1", "Option/Method 2"],
        "rows": [
            ["Feature 1", "Description for option 1", "Description for option 2"],
            ["Feature 2", "Description for option 1", "Description for option 2"]
        ]
    }},
    "sections": [
        {{
            "title": "Section Title",
            "content": "Paragraph explaining this aspect",
            "bullets": ["Key point 1", "Key point 2", "Key point 3"]
        }}
    ],
    "conclusion": "Summary and recommendations"
}}

RULES:
- Output language: {lang_desc}
- Extract key comparison points from the text
- Create 3-5 sections covering different aspects
- Each section should have 3-5 bullet points
- Make the comparison clear and structured
- Preserve technical terms in original language

JSON:"""

            response = model.generate_content(prompt)
            result_text = response.text.strip()
            
            # Remove markdown code blocks if present
            if result_text.startswith('```'):
                lines = result_text.split('\n')
                result_text = '\n'.join(lines[1:-1])  # Remove first and last line
            
            # Parse JSON
            analysis = json.loads(result_text)
            
            # Calculate usage
            usage = response.usage_metadata
            total_tokens = usage.prompt_token_count + usage.candidates_token_count
            cost = (usage.prompt_token_count / 1_000_000 * 0.075) + (usage.candidates_token_count / 1_000_000 * 0.30)
            
            analysis['ai_usage'] = {
                "engine": "gemini",
                "model": "gemini-2.0-flash-exp",
                "total_tokens": total_tokens,
                "cost_usd": round(cost, 6)
            }
            
            return analysis
            
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse AI response as JSON: {e}")
            logger.error(f"AI response: {result_text[:500]}")
            raise HTTPException(500, "AI returned invalid JSON format")
        except Exception as e:
            logger.error(f"Text analysis failed: {e}")
            raise HTTPException(500, f"Text analysis failed: {str(e)}")
    
    def _create_comparison_report_docx(self, analysis: Dict[str, Any], output_path: Path):
        """Create beautiful Word document from analysis data"""
        try:
            doc = Document()
            
            # Set narrow margins
            for section in doc.sections:
                section.top_margin = Inches(0.75)
                section.bottom_margin = Inches(0.75)
                section.left_margin = Inches(1.0)
                section.right_margin = Inches(1.0)
            
            # Title (Blue, centered, large)
            title_para = doc.add_heading(analysis['title'], level=1)
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in title_para.runs:
                run.font.size = Pt(20)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 112, 192)  # Blue
            
            # Add spacing
            self._add_spacing(title_para, after=12)
            
            # Introduction
            if analysis.get('introduction'):
                intro_para = doc.add_paragraph(analysis['introduction'])
                intro_para.paragraph_format.space_after = Pt(12)
            
            # Comparison Table
            if analysis.get('comparison_table'):
                self._add_comparison_table(doc, analysis['comparison_table'])
                # Add spacing after table
                doc.add_paragraph()
            
            # Sections
            for section_data in analysis.get('sections', []):
                self._add_section(doc, section_data)
            
            # Conclusion
            if analysis.get('conclusion'):
                # Conclusion heading
                conclusion_heading = doc.add_heading('Kết luận', level=2)
                for run in conclusion_heading.runs:
                    run.font.size = Pt(14)
                    run.font.color.rgb = RGBColor(0, 112, 192)
                
                # Conclusion text
                conclusion_para = doc.add_paragraph(analysis['conclusion'])
                conclusion_para.paragraph_format.space_after = Pt(12)
            
            # Save document
            doc.save(str(output_path))
            logger.info(f"Word document saved: {output_path}")
            
        except Exception as e:
            logger.error(f"Word document creation failed: {e}")
            raise HTTPException(500, f"Document creation failed: {str(e)}")
    
    def _add_comparison_table(self, doc: Document, table_data: Dict[str, Any]):
        """Add styled comparison table"""
        headers = table_data.get('headers', [])
        rows = table_data.get('rows', [])
        
        if not headers or not rows:
            return
        
        # Create table
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # Header row with blue background
        header_cells = table.rows[0].cells
        for i, header_text in enumerate(headers):
            cell = header_cells[i]
            cell.text = header_text
            
            # Style header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Blue background
            self._set_cell_background(cell, RGBColor(0, 112, 192))
        
        # Data rows
        for row_idx, row_data in enumerate(rows, start=1):
            cells = table.rows[row_idx].cells
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < len(headers):
                    cells[col_idx].text = str(cell_text)
                    # Style data cells
                    for paragraph in cells[col_idx].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
    
    def _add_section(self, doc: Document, section_data: Dict[str, Any]):
        """Add section with heading, content, and bullet points"""
        # Section heading (blue, bold)
        heading = doc.add_heading(section_data.get('title', ''), level=2)
        for run in heading.runs:
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 112, 192)
        self._add_spacing(heading, after=6)
        
        # Section content
        if section_data.get('content'):
            content_para = doc.add_paragraph(section_data['content'])
            content_para.paragraph_format.space_after = Pt(6)
        
        # Bullet points
        for bullet in section_data.get('bullets', []):
            p = doc.add_paragraph(bullet, style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
        
        # Add spacing after section
        doc.add_paragraph()
    
    def _set_cell_background(self, cell, color: RGBColor):
        """Set cell background color"""
        try:
            cell_xml_element = cell._element
            cell_properties = cell_xml_element.get_or_add_tcPr()
            shade_obj = OxmlElement('w:shd')
            shade_obj.set(qn('w:fill'), f"{color.rgb:06X}")
            cell_properties.append(shade_obj)
        except Exception as e:
            logger.warning(f"Failed to set cell background: {e}")
    
    def _add_spacing(self, paragraph, before: int = 0, after: int = 0):
        """Add spacing before/after paragraph (in points)"""
        if before:
            paragraph.paragraph_format.space_before = Pt(before)
        if after:
            paragraph.paragraph_format.space_after = Pt(after)
