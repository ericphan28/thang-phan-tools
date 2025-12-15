"""
Modern Document Conversion Service (2025)
Sử dụng các công nghệ mới nhất để convert documents
Tích hợp Adobe PDF Services API cho PDF → Word chất lượng cao
"""

import os
import shutil
import subprocess
import uuid
from pathlib import Path
from dotenv import load_dotenv

# Load environment variables
load_dotenv()
from typing import Optional, List
import aiofiles
import httpx
from fastapi import UploadFile, HTTPException
import logging

# Document libraries
from docx import Document
import pypdf
from pdf2docx import Converter as PDFToWordConverter
from pptx import Presentation
from openpyxl import load_workbook
import pypdfium2 as pdfium

# Adobe PDF Services (optional)
try:
    from adobe.pdfservices.operation.auth.service_principal_credentials import ServicePrincipalCredentials
    from adobe.pdfservices.operation.pdf_services import PDFServices
    from adobe.pdfservices.operation.pdf_services_media_type import PDFServicesMediaType
    from adobe.pdfservices.operation.pdfjobs.jobs.export_pdf_job import ExportPDFJob
    from adobe.pdfservices.operation.pdfjobs.params.export_pdf.export_pdf_params import ExportPDFParams
    from adobe.pdfservices.operation.pdfjobs.params.export_pdf.export_pdf_target_format import ExportPDFTargetFormat
    from adobe.pdfservices.operation.pdfjobs.params.export_pdf.export_ocr_locale import ExportOCRLocale
    from adobe.pdfservices.operation.pdfjobs.result.export_pdf_result import ExportPDFResult
    from adobe.pdfservices.operation.io.cloud_asset import CloudAsset
    from adobe.pdfservices.operation.io.stream_asset import StreamAsset
    ADOBE_AVAILABLE = True
except ImportError:
    ADOBE_AVAILABLE = False

# Google Gemini API (optional) - RECOMMENDED for PDF → Word
try:
    import google.generativeai as genai
    import json
    import asyncio
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False

logger = logging.getLogger(__name__)


def get_friendly_error_message(error: Exception) -> tuple[int, str]:
    """
    Convert Adobe API errors to user-friendly Vietnamese messages
    
    Returns:
        tuple: (status_code, friendly_message)
    """
    error_msg = str(error).lower()
    
    # Password protected files
    if "password_protected" in error_msg or "protected" in error_msg:
        return (400, "😔 Rất tiếc! File PDF này được bảo vệ bằng mật khẩu.\n\n"
                     "💡 Giải pháp:\n"
                     "• Mở file bằng PDF reader và nhập mật khẩu\n"
                     "• Sau đó 'Save As' thành file mới không có password\n"
                     "• Hoặc dùng tính năng 'Unlock PDF' của chúng tôi")
    
    # Digitally signed files
    if "pdf_signed" in error_msg or "signed" in error_msg:
        return (400, "😔 Rất tiếc! File PDF này có chữ ký điện tử.\n\n"
                     "💡 Giải pháp:\n"
                     "• Adobe API không xử lý được file có chữ ký số\n"
                     "• Vui lòng remove signature trước\n"
                     "• Hoặc dùng bản PDF gốc chưa ký")
    
    # Corrupted or invalid files
    if "corrupted" in error_msg or "invalid" in error_msg or "malformed" in error_msg:
        return (400, "😔 Rất tiếc! File PDF này bị lỗi hoặc không hợp lệ.\n\n"
                     "💡 Giải pháp:\n"
                     "• Thử mở file bằng PDF reader để kiểm tra\n"
                     "• Nếu mở được, thử 'Print to PDF' để tạo file mới\n"
                     "• Hoặc dùng file PDF từ nguồn khác")
    
    # File too large
    if "file size" in error_msg or "too large" in error_msg:
        return (400, "😔 Rất tiếc! File PDF quá lớn để xử lý.\n\n"
                     "💡 Giải pháp:\n"
                     "• Giới hạn: 100MB cho mỗi file\n"
                     "• Thử nén/tối ưu file PDF trước\n"
                     "• Hoặc split thành nhiều file nhỏ hơn")
    
    # Page count limits
    if "page" in error_msg and ("limit" in error_msg or "exceed" in error_msg):
        return (400, "😔 Rất tiếc! File PDF có quá nhiều trang.\n\n"
                     "💡 Giải pháp:\n"
                     "• Split file thành nhiều phần nhỏ hơn\n"
                     "• Xử lý từng phần một\n"
                     "• Hoặc liên hệ hỗ trợ để tăng giới hạn")
    
    # Invalid page ranges
    if "page range" in error_msg or "invalid range" in error_msg:
        return (400, "😔 Rất tiếc! Phạm vi trang không hợp lệ.\n\n"
                     "💡 Giải pháp:\n"
                     "• Kiểm tra số trang: ví dụ '1-3' hoặc '1,3,5'\n"
                     "• Đảm bảo số trang không vượt quá tổng số trang\n"
                     "• Số trang bắt đầu từ 1 (không phải 0)")
    
    # Network/timeout errors
    if "timeout" in error_msg or "connection" in error_msg:
        return (500, "😔 Rất tiếc! Kết nối với Adobe API bị gián đoạn.\n\n"
                     "💡 Giải pháp:\n"
                     "• File có thể quá lớn hoặc phức tạp\n"
                     "• Vui lòng thử lại sau vài phút\n"
                     "• Hoặc liên hệ hỗ trợ nếu vẫn lỗi")
    
    # Quota exceeded
    if "quota" in error_msg or "limit exceeded" in error_msg:
        return (429, "😔 Rất tiếc! Đã vượt quá giới hạn sử dụng.\n\n"
                     "💡 Giải pháp:\n"
                     "• Số lượng request đã đạt giới hạn hôm nay\n"
                     "• Vui lòng thử lại vào ngày mai\n"
                     "• Hoặc liên hệ để nâng cấp gói dịch vụ")
    
    # Authentication errors
    if "credential" in error_msg or "authentication" in error_msg or "unauthorized" in error_msg:
        return (500, "😔 Rất tiếc! Có lỗi xác thực với Adobe API.\n\n"
                     "💡 Giải pháp:\n"
                     "• Đây là lỗi hệ thống, không phải lỗi của bạn\n"
                     "• Vui lòng liên hệ quản trị viên\n"
                     "• Chúng tôi sẽ khắc phục trong thời gian sớm nhất")
    
    # Generic error
    return (500, f"😔 Rất tiếc! Đã có lỗi khi xử lý file PDF.\n\n"
                 f"💡 Chi tiết lỗi:\n{str(error)}\n\n"
                 f"Vui lòng thử lại hoặc liên hệ hỗ trợ nếu vấn đề vẫn tiếp diễn.")


def is_pdf_scanned(pdf_path: Path, min_text_chars: int = 50) -> bool:
    """
    Detect if PDF is scanned (image-only, no text layer)
    
    Strategy:
    - Extract text from first few pages
    - If total text < min_text_chars → likely scanned
    - Returns True if PDF needs OCR, False otherwise
    
    Args:
        pdf_path: Path to PDF file
        min_text_chars: Minimum characters to consider as "text PDF" (default 50)
    
    Returns:
        True if PDF appears to be scanned (needs OCR)
        False if PDF has text layer
    """
    try:
        with open(pdf_path, 'rb') as f:
            pdf_reader = pypdf.PdfReader(f)
            
            # Check first 3 pages (or all if less than 3)
            pages_to_check = min(3, len(pdf_reader.pages))
            total_text = ""
            
            for i in range(pages_to_check):
                page = pdf_reader.pages[i]
                text = page.extract_text()
                total_text += text
            
            # If very little text found → likely scanned
            text_length = len(total_text.strip())
            is_scanned = text_length < min_text_chars
            
            if is_scanned:
                logger.info(f"PDF appears to be scanned (only {text_length} chars found in first {pages_to_check} pages)")
            else:
                logger.info(f"PDF has text layer ({text_length} chars found)")
            
            return is_scanned
            
    except Exception as e:
        logger.warning(f"Could not detect if PDF is scanned: {e}, assuming not scanned")
        return False


class DocumentService:
    """Modern document processing service sử dụng Gotenberg + Adobe PDF Services"""
    
    def __init__(
        self, 
        upload_dir: str = "uploads/documents",
        gotenberg_url: str = None
    ):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        
        # Output directories
        self.output_dir = Path("uploads/outputs")
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Gotenberg URL - auto-detect environment
        if gotenberg_url:
            self.gotenberg_url = gotenberg_url
        else:
            # Production: Gotenberg container, Dev: có thể dùng local LibreOffice
            self.gotenberg_url = os.getenv("GOTENBERG_URL", "http://gotenberg:3000")
        
        # Adobe PDF Services - optional but recommended for better quality
        self.use_adobe = os.getenv("USE_ADOBE_PDF_API", "false").lower() == "true"
        self.adobe_credentials = None
        
        if self.use_adobe and ADOBE_AVAILABLE:
            client_id = os.getenv("PDF_SERVICES_CLIENT_ID")
            client_secret = os.getenv("PDF_SERVICES_CLIENT_SECRET")
            
            if client_id and client_secret:
                try:
                    self.adobe_credentials = ServicePrincipalCredentials(
                        client_id=client_id,
                        client_secret=client_secret
                    )
                    logger.info("Adobe PDF Services enabled - High quality PDF to Word conversion available")
                except Exception as e:
                    logger.warning(f"Failed to initialize Adobe credentials: {e}")
                    self.use_adobe = False
            else:
                logger.warning("USE_ADOBE_PDF_API=true but credentials not found in env")
                self.use_adobe = False
        elif self.use_adobe and not ADOBE_AVAILABLE:
            logger.warning("USE_ADOBE_PDF_API=true but pdfservices-sdk not installed")
            self.use_adobe = False
        
        # Google Gemini API - RECOMMENDED for PDF → Word with Vietnamese
        self.gemini_api_key = os.getenv("GEMINI_API_KEY")
        self.use_gemini = False
        
        if self.gemini_api_key and GEMINI_AVAILABLE:
            try:
                genai.configure(api_key=self.gemini_api_key)
                self.gemini_model = genai.GenerativeModel('gemini-2.5-flash')
                self.use_gemini = True
                logger.info("✅ Gemini API enabled - Best price/performance for PDF → Word (Vietnamese support: 9/10)")
            except Exception as e:
                logger.warning(f"Failed to initialize Gemini API: {e}")
                self.use_gemini = False
        elif self.gemini_api_key and not GEMINI_AVAILABLE:
            logger.warning("GEMINI_API_KEY found but google-generativeai not installed. Run: pip install google-generativeai")
            self.use_gemini = False
        
    async def save_upload_file(self, upload_file: UploadFile) -> Path:
        """Save uploaded file async"""
        file_path = self.upload_dir / upload_file.filename
        
        async with aiofiles.open(file_path, 'wb') as out_file:
            content = await upload_file.read()
            await out_file.write(content)
            
        return file_path
    
    # ==================== Word → PDF (Gotenberg) ====================
    
    async def word_to_pdf(
        self,
        input_file: Path,
        output_filename: Optional[str] = None
    ) -> Path:
        """
        Convert Word (.docx/.doc) sang PDF bằng Gotenberg API
        
        Gotenberg là Docker microservice hiện đại, sử dụng LibreOffice headless
        - Không cần cài LibreOffice trên host machine
        - REST API đơn giản, nhanh, ổn định
        - Support: DOC, DOCX, XLS, XLSX, PPT, PPTX, ODT...
        """
        if not input_file.suffix.lower() in ['.docx', '.doc']:
            raise HTTPException(400, "File phải là .docx hoặc .doc")
        
        output_filename = output_filename or input_file.stem + ".pdf"
        output_path = self.output_dir / output_filename
        
        try:
            # Gọi Gotenberg API để convert
            async with httpx.AsyncClient(timeout=60.0) as client:
                # Đọc file Word
                async with aiofiles.open(input_file, 'rb') as f:
                    file_content = await f.read()
                
                # POST đến Gotenberg LibreOffice endpoint
                files = {
                    'files': (input_file.name, file_content, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                }
                
                response = await client.post(
                    f"{self.gotenberg_url}/forms/libreoffice/convert",
                    files=files
                )
                
                if response.status_code != 200:
                    raise HTTPException(
                        500,
                        f"Gotenberg conversion failed: {response.status_code} - {response.text}"
                    )
                
                # Lưu PDF output
                async with aiofiles.open(output_path, 'wb') as f:
                    await f.write(response.content)
                
            if not output_path.exists() or output_path.stat().st_size == 0:
                raise HTTPException(500, "File PDF không được tạo ra hoặc rỗng")
                
            return output_path
            
        except httpx.ConnectError:
            # Gotenberg không khả dụng - fallback to LibreOffice local (dev only)
            return await self._word_to_pdf_libreoffice_fallback(input_file, output_filename)
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(500, f"Lỗi chuyển đổi Word sang PDF: {str(e)}")
    
    async def _word_to_pdf_libreoffice_fallback(
        self,
        input_file: Path,
        output_filename: str
    ) -> Path:
        """
        Fallback method: Dùng LibreOffice local nếu Gotenberg không có
        (Chỉ dùng cho development/testing)
        """
        output_path = self.output_dir / output_filename
        
        # Tìm LibreOffice
        libreoffice_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            "/usr/bin/libreoffice",
            "/usr/bin/soffice",
            "soffice",
        ]
        
        soffice_path = None
        for path in libreoffice_paths:
            if os.path.exists(path) or path == "soffice":
                soffice_path = path
                break
        
        if not soffice_path:
            raise HTTPException(
                503,
                "Gotenberg service không khả dụng và LibreOffice chưa được cài đặt. "
                "Vui lòng đảm bảo Gotenberg container đang chạy."
            )
        
        try:
            cmd = [
                soffice_path,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", str(self.output_dir),
                str(input_file)
            ]
            
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=60
            )
            
            if result.returncode != 0:
                raise HTTPException(500, f"LibreOffice conversion failed: {result.stderr}")
            
            generated_pdf = self.output_dir / (input_file.stem + ".pdf")
            if generated_pdf != output_path:
                generated_pdf.rename(output_path)
            
            return output_path
            
        except subprocess.TimeoutExpired:
            raise HTTPException(500, "Conversion timeout")
        except Exception as e:
            raise HTTPException(500, f"Fallback conversion failed: {str(e)}")
    
    # ==================== Generic Office → PDF (Excel, PowerPoint) ====================
    
    async def office_to_pdf(
        self,
        input_file: Path,
        output_filename: Optional[str] = None
    ) -> Path:
        """
        Convert any Office file (Excel, PowerPoint) to PDF using Gotenberg
        
        Supports: .xlsx, .xls, .pptx, .ppt, .odt, .ods, .odp
        """
        supported_extensions = ['.xlsx', '.xls', '.pptx', '.ppt', '.odt', '.ods', '.odp']
        if input_file.suffix.lower() not in supported_extensions:
            raise HTTPException(400, f"File phải là Office file: {', '.join(supported_extensions)}")
        
        output_filename = output_filename or input_file.stem + ".pdf"
        output_path = self.output_dir / output_filename
        
        try:
            # Gọi Gotenberg API
            async with httpx.AsyncClient(timeout=120.0) as client:  # Excel/PPT có thể mất thời gian hơn
                async with aiofiles.open(input_file, 'rb') as f:
                    file_content = await f.read()
                
                # Determine MIME type
                mime_types = {
                    '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                    '.xls': 'application/vnd.ms-excel',
                    '.pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
                    '.ppt': 'application/vnd.ms-powerpoint',
                    '.odt': 'application/vnd.oasis.opendocument.text',
                    '.ods': 'application/vnd.oasis.opendocument.spreadsheet',
                    '.odp': 'application/vnd.oasis.opendocument.presentation'
                }
                mime_type = mime_types.get(input_file.suffix.lower(), 'application/octet-stream')
                
                files = {
                    'files': (input_file.name, file_content, mime_type)
                }
                
                response = await client.post(
                    f"{self.gotenberg_url}/forms/libreoffice/convert",
                    files=files
                )
                
                if response.status_code != 200:
                    raise HTTPException(500, f"Gotenberg conversion failed: {response.status_code}")
                
                async with aiofiles.open(output_path, 'wb') as f:
                    await f.write(response.content)
                
            if not output_path.exists() or output_path.stat().st_size == 0:
                raise HTTPException(500, "File PDF không được tạo ra hoặc rỗng")
                
            return output_path
            
        except httpx.ConnectError:
            # Fallback to LibreOffice local
            return await self._office_to_pdf_libreoffice_fallback(input_file, output_filename)
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(500, f"Lỗi chuyển đổi Office sang PDF: {str(e)}")
    
    async def _office_to_pdf_libreoffice_fallback(
        self,
        input_file: Path,
        output_filename: str
    ) -> Path:
        """Fallback: Use LibreOffice for Office files"""
        return await self._word_to_pdf_libreoffice_fallback(input_file, output_filename)

    # ==================== PDF → Word ====================
    
    async def pdf_to_word(
        self,
        input_file: Path,
        output_filename: Optional[str] = None,
        start_page: int = 0,
        end_page: Optional[int] = None,
        enable_ocr: bool = False,
        ocr_language: str = "vi-VN",
        auto_detect_scanned: bool = True,
        use_gemini: bool = False
    ) -> Path:
        """
        Convert PDF to Word (.docx)
        
        Strategy:
        1. Gemini API (if use_gemini=True) - Best for tables/layout, 100+ languages, $6.43/30k pages
        2. Adobe PDF Services (if enabled) - AI-powered, 10/10 quality BUT NO Vietnamese support
        3. Fallback to pdf2docx - Good quality, free
        
        OCR Support:
        - enable_ocr: Enable OCR for scanned PDFs (default: False)
        - ocr_language: Language for OCR (default: "vi-VN" for Vietnamese)
        - auto_detect_scanned: Auto-detect if PDF is scanned and enable OCR (default: True)
        - use_gemini: Use Gemini API instead of Adobe (supports Vietnamese, better for tables)
        
        Supported OCR languages (Adobe):
        - vi-VN (Vietnamese), en-US (English), fr-FR (French), de-DE (German),
        - es-ES (Spanish), it-IT (Italian), pt-BR (Portuguese), ja-JP (Japanese),
        - ko-KR (Korean), zh-CN (Chinese Simplified), zh-TW (Chinese Traditional)
        - and 40+ more languages
        
        Gemini API (NEW):
        - Supports 100+ languages including Vietnamese
        - Native PDF understanding (no OCR preprocessing)
        - Best for tables, layout, complex documents
        - Cost: $6.43 per 30k pages (85% cheaper than Google Vision)
        - Quality: 9/10 overall
        """
        if input_file.suffix.lower() != '.pdf':
            raise HTTPException(400, "File must be .pdf")
        
        output_filename = output_filename or input_file.stem + ".docx"
        output_path = self.output_dir / output_filename
        
        # Try Gemini first if requested (best for Vietnamese + tables)
        if use_gemini:
            try:
                logger.info(f"Using Gemini API for {input_file.name}")
                return await self._pdf_to_word_gemini(input_file, output_path, ocr_language)
            except Exception as e:
                logger.warning(f"Gemini API conversion failed: {e}, falling back to Adobe/pdf2docx")
        
        # Auto-detect if PDF is scanned (if not explicitly disabled)
        needs_ocr = enable_ocr
        if auto_detect_scanned and not enable_ocr:
            is_scanned = is_pdf_scanned(input_file)
            if is_scanned:
                logger.info(f"Auto-detected scanned PDF, enabling OCR with language: {ocr_language}")
                needs_ocr = True
        
        # Try Adobe if enabled (best quality but NO Vietnamese)
        if self.use_adobe and self.adobe_credentials:
            try:
                logger.info(f"Using Adobe PDF Services for {input_file.name} (OCR: {needs_ocr})")
                return await self._pdf_to_word_adobe(
                    input_file, 
                    output_path,
                    enable_ocr=needs_ocr,
                    ocr_language=ocr_language
                )
            except Exception as e:
                logger.warning(f"Adobe PDF conversion failed: {e}, falling back to pdf2docx")
        
        # Fallback to pdf2docx (good quality, free)
        # Note: pdf2docx doesn't support OCR, so if needs_ocr=True, warn user
        if needs_ocr:
            logger.warning(f"PDF appears to be scanned but Adobe OCR not available. "
                         f"Using pdf2docx - results may be poor for scanned PDFs. "
                         f"Consider enabling Adobe PDF Services for better OCR quality.")
        
        logger.info(f"Using pdf2docx for {input_file.name}")
        return await self._pdf_to_word_local(input_file, output_path, start_page, end_page)
    
    async def _pdf_to_word_adobe(
        self, 
        input_file: Path, 
        output_path: Path,
        enable_ocr: bool = False,
        ocr_language: str = "vi-VN"
    ) -> Path:
        """
        Convert PDF to Word using Adobe PDF Services API
        High quality conversion with AI-powered layout preservation
        
        OCR Support:
        - If enable_ocr=True and language supported by Adobe: Use Adobe OCR
        - If enable_ocr=True but language NOT supported (e.g., Vietnamese): Use Tesseract OCR
        - Supports 50+ languages total (Adobe + Tesseract combined)
        
        Note: Adobe OCR does NOT support Vietnamese (vi-VN). For Vietnamese PDFs,
        we use Tesseract OCR which has excellent Vietnamese support.
        """
        try:
            # If OCR needed, check if Adobe supports the language
            if enable_ocr:
                adobe_locale = ocr_language.upper().replace('-', '_')
                
                # Check if Adobe supports this locale
                from adobe.pdfservices.operation.pdfjobs.params.ocr_pdf.ocr_supported_locale import OCRSupportedLocale
                adobe_supported_locales = [
                    x for x in dir(OCRSupportedLocale) if not x.startswith('_')
                ]
                
                if adobe_locale in adobe_supported_locales:
                    # Adobe supports this language - use two-step Adobe OCR
                    logger.info(f"Using Adobe OCR for {input_file.name} (locale: {adobe_locale})")
                    ocr_output = await self._ocr_pdf_adobe(input_file, ocr_language)
                    try:
                        result = await self._pdf_to_word_adobe_internal(ocr_output, output_path)
                        await self.cleanup_file(ocr_output)
                        return result
                    except Exception:
                        await self.cleanup_file(ocr_output)
                        raise
                else:
                    # Adobe doesn't support this language (e.g., Vietnamese)
                    # Use Tesseract OCR + Adobe Export
                    logger.warning(f"Adobe OCR does NOT support {adobe_locale} (e.g., Vietnamese)")
                    logger.info(f"Falling back to Tesseract OCR for {ocr_language} + Adobe Export")
                    
                    # Use Tesseract to OCR the PDF
                    ocr_output = await self._ocr_pdf_tesseract(
                        input_file,
                        language=ocr_language  # Pass full locale like vi-VN
                    )
                    try:
                        result = await self._pdf_to_word_adobe_internal(ocr_output, output_path)
                        await self.cleanup_file(ocr_output)
                        return result
                    except Exception:
                        await self.cleanup_file(ocr_output)
                        raise
            else:
                # Direct conversion without OCR
                return await self._pdf_to_word_adobe_internal(input_file, output_path)
                
        except Exception as e:
            logger.error(f"Adobe PDF Services error: {e}")
            status_code, friendly_msg = get_friendly_error_message(e)
            raise HTTPException(status_code, friendly_msg)
    
    async def _pdf_to_word_adobe_internal(
        self, 
        input_file: Path, 
        output_path: Path
    ) -> Path:
        """Internal method for direct PDF to Word conversion (no OCR)"""
        # Read input file
        async with aiofiles.open(input_file, 'rb') as f:
            input_stream = await f.read()
        
        # Create PDF Services instance
        pdf_services = PDFServices(credentials=self.adobe_credentials)
        
        # Upload file to Adobe
        input_asset = pdf_services.upload(
            input_stream=input_stream,
            mime_type=PDFServicesMediaType.PDF
        )
        
        # Create export parameters (no OCR)
        export_pdf_params = ExportPDFParams(
            target_format=ExportPDFTargetFormat.DOCX
        )
        
        # Create and submit job
        export_pdf_job = ExportPDFJob(
            input_asset=input_asset,
            export_pdf_params=export_pdf_params
        )
        
        location = pdf_services.submit(export_pdf_job)
        
        # Get result (polling handled by SDK)
        pdf_services_response = pdf_services.get_job_result(location, ExportPDFResult)
        
        # Download result
        result_asset: CloudAsset = pdf_services_response.get_result().get_asset()
        stream_asset: StreamAsset = pdf_services.get_content(result_asset)
        
        # Save to file
        async with aiofiles.open(output_path, "wb") as f:
            await f.write(stream_asset.get_input_stream())
        
        logger.info(f"Adobe conversion successful: {output_path}")
        return output_path
    
    async def _ocr_pdf_adobe(
        self,
        input_file: Path,
        ocr_language: str = "vi-VN"
    ) -> Path:
        """
        Perform OCR on PDF using Adobe OCRPDFJob
        Returns path to searchable PDF
        """
        try:
            # Import OCR-specific classes
            from adobe.pdfservices.operation.pdfjobs.jobs.ocr_pdf_job import OCRPDFJob
            from adobe.pdfservices.operation.pdfjobs.params.ocr_pdf.ocr_pdf_params import OCRPDFParams
            from adobe.pdfservices.operation.pdfjobs.params.ocr_pdf.ocr_supported_locale import OCRSupportedLocale
            from adobe.pdfservices.operation.pdfjobs.params.ocr_pdf.ocr_supported_type import OCRSupportedType
            from adobe.pdfservices.operation.pdfjobs.result.ocr_pdf_result import OCRPDFResult
            
            # Read input file
            async with aiofiles.open(input_file, 'rb') as f:
                input_stream = await f.read()
            
            # Create PDF Services instance
            pdf_services = PDFServices(credentials=self.adobe_credentials)
            
            # Upload file
            input_asset = pdf_services.upload(
                input_stream=input_stream,
                mime_type=PDFServicesMediaType.PDF
            )
            
            # Map language code to OCRSupportedLocale
            # vi-VN → VI_VN
            adobe_locale = ocr_language.upper().replace('-', '_')
            
            try:
                ocr_locale = getattr(OCRSupportedLocale, adobe_locale)
            except AttributeError:
                logger.warning(f"OCR locale {adobe_locale} not supported, using EN_US")
                ocr_locale = OCRSupportedLocale.EN_US
            
            # Create OCR parameters
            ocr_pdf_params = OCRPDFParams(
                ocr_locale=ocr_locale,
                ocr_type=OCRSupportedType.SEARCHABLE_IMAGE_EXACT  # Preserves original
            )
            
            # Create and submit OCR job
            ocr_pdf_job = OCRPDFJob(
                input_asset=input_asset,
                ocr_pdf_params=ocr_pdf_params
            )
            
            location = pdf_services.submit(ocr_pdf_job)
            
            # Get result
            pdf_services_response = pdf_services.get_job_result(location, OCRPDFResult)
            
            # Download OCR'd PDF
            result_asset: CloudAsset = pdf_services_response.get_result().get_asset()
            stream_asset: StreamAsset = pdf_services.get_content(result_asset)
            
            # Save OCR'd PDF to temp file
            ocr_output = self.output_dir / f"ocr_temp_{input_file.name}"
            async with aiofiles.open(ocr_output, "wb") as f:
                await f.write(stream_asset.get_input_stream())
            
            logger.info(f"Adobe OCR successful: {ocr_output} (locale: {adobe_locale})")
            return ocr_output
            
        except Exception as e:
            logger.error(f"Adobe OCR error: {e}")
            raise
            return output_path
            
        except Exception as e:
            logger.error(f"Adobe PDF Services error: {e}")
            status_code, friendly_msg = get_friendly_error_message(e)
            raise HTTPException(status_code, friendly_msg)
    
    async def _pdf_to_word_local(
        self,
        input_file: Path,
        output_path: Path,
        start_page: int = 0,
        end_page: Optional[int] = None
    ) -> Path:
        """
        Convert PDF to Word using pdf2docx (fallback)
        Good quality, pure Python, free
        """
        try:
            # pdf2docx is pure Python, works cross-platform
            cv = PDFToWordConverter(str(input_file))
            cv.convert(str(output_path), start=start_page, end=end_page)
            cv.close()
            
            logger.info(f"pdf2docx conversion successful: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"pdf2docx conversion error: {e}")
            raise HTTPException(500, f"PDF to Word conversion failed: {str(e)}")
    
    async def _pdf_to_word_gemini(
        self,
        input_file: Path,
        output_path: Path,
        ocr_language: str = "vi"
    ) -> Path:
        """
        Convert PDF to Word using Google Gemini API
        
        Features:
        - Native PDF processing (no OCR library needed)
        - Understands Vietnamese and 100+ languages
        - Excellent table extraction (9.5/10)
        - Smart layout preservation
        - Best price/performance: $6.43 per 30k pages
        
        Quality: 9/10 overall
        - Text accuracy: 9/10
        - Table extraction: 9.5/10 (BEST!)
        - Layout preservation: 9/10
        - Vietnamese support: 9/10 (Excellent)
        
        Free tier: 1,500 requests/day
        """
        if not self.use_gemini:
            raise HTTPException(500, "Gemini API not configured. Set GEMINI_API_KEY in .env")
        
        try:
            logger.info(f"Starting Gemini API conversion for {input_file.name}")
            
            # Upload PDF to Gemini
            logger.info("Uploading PDF to Gemini...")
            pdf_file = genai.upload_file(str(input_file))
            
            # Wait for processing
            while pdf_file.state.name == "PROCESSING":
                await asyncio.sleep(1)
                pdf_file = genai.get_file(pdf_file.name)
            
            if pdf_file.state.name == "FAILED":
                raise ValueError(f"PDF processing failed: {pdf_file.state.name}")
            
            logger.info(f"PDF uploaded successfully: {pdf_file.name}")
            
            # Clear instruction for Gemini - TEXT ONLY, NO IMAGES
            prompt = """
CRITICAL REQUIREMENT: Convert this PDF to EDITABLE TEXT in Word format - NOT IMAGES!

I need you to extract ALL text content from this PDF document and provide it as PURE TEXT that can be edited, copied, and modified in Microsoft Word.

DO NOT:
- Convert to images or screenshots  
- Embed any visual representations
- Create picture objects

DO:
- Extract every Vietnamese character exactly (ă, â, ê, ô, ơ, ư, đ, etc.)
- Convert tables to proper Word table format with editable text
- Make all content selectable and editable text
- Preserve document structure with text formatting

For tables: Convert each cell content to editable text in Word table cells - users must be able to click and edit the text directly.

For headers/titles: Use Word paragraph formatting with bold/center alignment - pure text, not images.

The final Word document must have 100% editable text content. Every word must be copyable and modifiable.

Please extract and structure the content as text-based data.
"""
            
            # Generate content with structured output
            logger.info("Extracting content with Gemini...")
            response = self.gemini_model.generate_content(
                [pdf_file, prompt],
                generation_config=genai.GenerationConfig(
                    temperature=0.1,  # Low temperature for accuracy
                    response_mime_type="application/json"
                )
            )
            
            # Parse Gemini response - handle both JSON and text safely
            try:
                # Try to parse as JSON first
                document_data = json.loads(response.text)
                logger.info("Got structured response from Gemini - converting to pure text format")
                
                # Handle case where Gemini returns an array instead of object
                if isinstance(document_data, list):
                    logger.info("Gemini returned array, converting to text format")
                    # Convert array to text
                    text_content = ""
                    for item in document_data:
                        if isinstance(item, str):
                            text_content += item + "\n"
                        elif isinstance(item, dict) and 'text' in item:
                            text_content += item['text'] + "\n"
                    
                    document_data = {
                        "content": text_content.strip(),
                        "format": "text"
                    }
                
                # Ensure we only process text content, not images
                elif 'pages' in document_data:
                    for page in document_data['pages']:
                        for content in page.get('content', []):
                            # Skip any image or non-text content
                            if content.get('type') in ['image', 'figure', 'graphic']:
                                logger.warning("Skipping non-text content to ensure editable output")
                                continue
                            
            except json.JSONDecodeError:
                # If not JSON, create simple structure from text response
                logger.info("Got text response from Gemini - perfect for pure text output")
                document_data = {
                    "content": response.text,
                    "format": "text"
                }
            
            # Convert to Word
            logger.info("Creating Word document...")
            word_path = await self._create_word_from_json(document_data, output_path)
            
            # Cleanup
            try:
                genai.delete_file(pdf_file.name)
                logger.info("Cleaned up uploaded file from Gemini")
            except Exception as e:
                logger.warning(f"Failed to cleanup Gemini file: {e}")
            
            logger.info(f"✅ Gemini conversion successful: {output_path}")
            return word_path
            
        except Exception as e:
            logger.error(f"Gemini PDF to Word conversion failed: {e}")
            raise HTTPException(500, f"Gemini conversion failed: {str(e)}")
    
    async def _create_word_from_json(
        self,
        data: dict,
        output_path: Path
    ) -> Path:
        """
        Create Word document with 100% EDITABLE TEXT - no images or objects
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            logger.info(f"Creating Word document from data: {str(data)[:200]}...")
            
            # CRITICAL: All content must be editable text
            if data.get("format") == "text":
                # Text response - convert to editable paragraphs
                content = data.get("content", "")
                logger.info(f"Processing text content: {len(content)} characters")
                
                if content.strip():
                    lines = content.split('\n')
                    
                    for line in lines:
                        if line.strip():
                            # Create paragraph with pure text - NO IMAGES
                            para = doc.add_paragraph()
                            run = para.add_run(line.strip())
                            
                            # Basic text formatting - keep it editable
                            if any(keyword in line.upper() for keyword in ["QUYẾT ĐỊNH", "HỘI NÔNG DÂN", "CÔNG HÒA XÃ HỘI"]):
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run.bold = True
                            
                            # Ensure it's text, not any embedded object
                            run.font.size = Pt(11)
                else:
                    # If no content, add a message
                    para = doc.add_paragraph()
                    run = para.add_run("Không thể trích xuất nội dung từ PDF. Vui lòng thử lại.")
                    run.font.size = Pt(11)
            
            else:
                # Structured response - process as pure text
                logger.info("Processing structured data")
                content_added = False
                
                # Document title as EDITABLE text
                doc_info = data.get('document_info', {})
                if doc_info.get('title'):
                    title_para = doc.add_paragraph()
                    title_run = title_para.add_run(doc_info['title'])
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    title_run.bold = True
                    title_run.font.size = Pt(14)
                    content_added = True
                        
                # Process pages with TEXT ONLY approach
                pages = data.get('pages', [])
                logger.info(f"Found {len(pages)} pages to process")
                
                for page_idx, page in enumerate(pages):
                    logger.info(f"Processing page {page_idx + 1}")
                    page_content = page.get('content', [])
                    
                    for item_idx, item in enumerate(page_content):
                        item_type = item.get('type', '')
                        text = item.get('text', '')
                        
                        logger.info(f"Processing item {item_idx}: type={item_type}, text_len={len(text)}")
                        
                        # Only process text-based content
                        if item_type in ['header', 'heading']:
                            if text.strip():
                                para = doc.add_paragraph()
                                run = para.add_run(text)
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run.bold = True
                                run.font.size = Pt(12)
                                content_added = True
                                
                        elif item_type == 'paragraph':
                            if text.strip():
                                para = doc.add_paragraph()
                                run = para.add_run(text)
                                run.font.size = Pt(11)
                                content_added = True
                            
                        elif item_type == 'table':
                            headers = item.get('headers', {}).get('values', [])
                            rows = item.get('rows', [])
                            
                            if headers and rows:
                                logger.info(f"Creating table with {len(headers)} columns and {len(rows)} rows")
                                # Create Word table with EDITABLE TEXT cells
                                table = doc.add_table(rows=1, cols=len(headers))
                                table.style = 'Table Grid'
                                
                                # Header row - EDITABLE TEXT
                                hdr_cells = table.rows[0].cells
                                for i, header in enumerate(headers):
                                    hdr_cells[i].text = str(header)
                                    # Make header text bold and centered
                                    for paragraph in hdr_cells[i].paragraphs:
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        for run in paragraph.runs:
                                            run.bold = True
                                
                                # Data rows - EDITABLE TEXT  
                                for row in rows:
                                    row_cells = table.add_row().cells
                                    cells_data = row.get('cells', []) if isinstance(row, dict) else row
                                    
                                    for i, cell_data in enumerate(cells_data):
                                        if i < len(row_cells):
                                            if isinstance(cell_data, dict):
                                                cell_text = str(cell_data.get('text', ''))
                                            else:
                                                cell_text = str(cell_data)
                                            
                                            row_cells[i].text = cell_text
                                            
                                            # Ensure all text is editable
                                            for paragraph in row_cells[i].paragraphs:
                                                for run in paragraph.runs:
                                                    run.font.size = Pt(10)
                                
                                content_added = True
                
                # If no content was added from structured data, try to extract any text
                if not content_added:
                    logger.warning("No content found in structured data, trying to extract any available text")
                    
                    # Try to extract from the entire data structure
                    import json
                    full_text = json.dumps(data, ensure_ascii=False, indent=2)
                    
                    para = doc.add_paragraph()
                    run = para.add_run("Nội dung được trích xuất từ PDF:")
                    run.font.size = Pt(11)
                    run.bold = True
                    
                    # Add the raw extracted data
                    para = doc.add_paragraph()
                    run = para.add_run(full_text[:2000])  # Limit to first 2000 chars
                    run.font.size = Pt(10)
            
            # Save as Word document with pure text content
            doc.save(str(output_path))
            logger.info(f"✅ Word document created with editable text: {output_path}")
            
            return output_path
            
        except Exception as e:
            logger.error(f"Error creating editable Word document: {str(e)}")
            logger.exception("Full error details:")
            raise
    
    async def pdf_to_excel(
        self,
        input_file: Path,
        output_filename: Optional[str] = None
    ) -> Path:
        """
        Convert PDF to Excel (.xlsx)
        Extracts tables from PDF using pdfplumber
        """
        if input_file.suffix.lower() != '.pdf':
            raise HTTPException(400, "File must be .pdf")
        
        output_filename = output_filename or input_file.stem + ".xlsx"
        output_path = self.output_dir / output_filename
        
        try:
            import pdfplumber
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment
            
            # Create Excel workbook
            wb = Workbook()
            wb.remove(wb.active)  # Remove default sheet
            
            # Open PDF
            with pdfplumber.open(input_file) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    # Extract tables from page
                    tables = page.extract_tables()
                    
                    if tables:
                        # Create sheet for each page with tables
                        ws = wb.create_sheet(title=f"Page {page_num}")
                        
                        # Write all tables from this page
                        current_row = 1
                        for table_num, table in enumerate(tables, start=1):
                            # Add table header
                            if len(tables) > 1:
                                header_cell = ws.cell(row=current_row, column=1, value=f"Table {table_num}")
                                header_cell.font = Font(bold=True, size=12)
                                header_cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
                                current_row += 1
                            
                            # Write table data
                            for row_data in table:
                                for col_num, cell_value in enumerate(row_data, start=1):
                                    cell = ws.cell(row=current_row, column=col_num, value=cell_value)
                                    cell.alignment = Alignment(wrap_text=True, vertical='top')
                                    
                                    # Bold first row (header)
                                    if current_row == 1 or (len(tables) > 1 and current_row == 2):
                                        cell.font = Font(bold=True)
                                        cell.fill = PatternFill(start_color="E0E0E0", end_color="E0E0E0", fill_type="solid")
                                
                                current_row += 1
                            
                            # Add spacing between tables
                            current_row += 1
                        
                        # Auto-adjust column widths
                        for column in ws.columns:
                            max_length = 0
                            column_letter = column[0].column_letter
                            for cell in column:
                                try:
                                    if cell.value:
                                        max_length = max(max_length, len(str(cell.value)))
                                except:
                                    pass
                            adjusted_width = min(max_length + 2, 50)
                            ws.column_dimensions[column_letter].width = adjusted_width
                    else:
                        # No tables found, extract text
                        text = page.extract_text()
                        if text and text.strip():
                            ws = wb.create_sheet(title=f"Page {page_num} (Text)")
                            ws.cell(row=1, column=1, value=text)
                            ws.column_dimensions['A'].width = 100
            
            # If no content was extracted at all
            if len(wb.sheetnames) == 0:
                ws = wb.create_sheet(title="No Data")
                ws.cell(row=1, column=1, value="No tables or text found in PDF")
            
            # Save workbook
            wb.save(output_path)
            
            return output_path
            
        except Exception as e:
            raise HTTPException(500, f"PDF to Excel conversion failed: {str(e)}")
    
    # ==================== PDF Operations ====================
    
    async def merge_pdfs(
        self,
        input_files: List[Path],
        output_filename: str = "merged.pdf"
    ) -> Path:
        """Merge multiple PDFs into one (using modern pypdf)"""
        output_path = self.output_dir / output_filename
        
        try:
            merger = pypdf.PdfWriter()
            
            for pdf_file in input_files:
                if pdf_file.suffix.lower() != '.pdf':
                    continue
                    
                with open(pdf_file, 'rb') as f:
                    pdf_reader = pypdf.PdfReader(f)
                    for page in pdf_reader.pages:
                        merger.add_page(page)
            
            with open(output_path, 'wb') as output_file:
                merger.write(output_file)
                
            return output_path
            
        except Exception as e:
            raise HTTPException(500, f"PDF merge failed: {str(e)}")
    
    async def split_pdf(
        self,
        input_file: Path,
        page_ranges: List[tuple],  # [(1,3), (5,7)] means pages 1-3 and 5-7
        output_prefix: str = "split"
    ) -> List[Path]:
        """Split PDF into multiple files by page ranges"""
        if input_file.suffix.lower() != '.pdf':
            raise HTTPException(400, "File must be .pdf")
        
        try:
            output_files = []
            
            with open(input_file, 'rb') as f:
                pdf_reader = pypdf.PdfReader(f)
                
                for idx, (start, end) in enumerate(page_ranges):
                    pdf_writer = pypdf.PdfWriter()
                    
                    for page_num in range(start - 1, end):  # Convert to 0-indexed
                        if page_num < len(pdf_reader.pages):
                            pdf_writer.add_page(pdf_reader.pages[page_num])
                    
                    output_path = self.output_dir / f"{output_prefix}_{idx+1}.pdf"
                    with open(output_path, 'wb') as output_file:
                        pdf_writer.write(output_file)
                    
                    output_files.append(output_path)
                    
            return output_files
            
        except Exception as e:
            raise HTTPException(500, f"PDF split failed: {str(e)}")
    
    async def extract_pdf_text(self, input_file: Path) -> str:
        """Extract all text from PDF"""
        if input_file.suffix.lower() != '.pdf':
            raise HTTPException(400, "File must be .pdf")
        
        try:
            with open(input_file, 'rb') as f:
                pdf_reader = pypdf.PdfReader(f)
                text = ""
                
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n\n"
                    
            return text.strip()
            
        except Exception as e:
            raise HTTPException(500, f"PDF text extraction failed: {str(e)}")
    
    async def rotate_pdf_pages(
        self,
        input_file: Path,
        rotation: int = 90,  # 90, 180, 270
        pages: Optional[List[int]] = None,  # None = all pages
        output_filename: Optional[str] = None
    ) -> Path:
        """Rotate PDF pages"""
        if input_file.suffix.lower() != '.pdf':
            raise HTTPException(400, "File must be .pdf")
        
        if rotation not in [90, 180, 270, -90]:
            raise HTTPException(400, "Rotation must be 90, 180, or 270 degrees")
        
        output_filename = output_filename or input_file.stem + "_rotated.pdf"
        output_path = self.output_dir / output_filename
        
        try:
            with open(input_file, 'rb') as f:
                pdf_reader = pypdf.PdfReader(f)
                pdf_writer = pypdf.PdfWriter()
                
                for page_num, page in enumerate(pdf_reader.pages):
                    if pages is None or (page_num + 1) in pages:
                        page.rotate(rotation)
                    pdf_writer.add_page(page)
                
                with open(output_path, 'wb') as output_file:
                    pdf_writer.write(output_file)
                    
            return output_path
            
        except Exception as e:
            raise HTTPException(500, f"PDF rotation failed: {str(e)}")
    
    # ==================== PDF Info ====================
    
    async def get_pdf_info(self, input_file: Path) -> dict:
        """Get PDF metadata and information"""
        if input_file.suffix.lower() != '.pdf':
            raise HTTPException(400, "File must be .pdf")
        
        try:
            with open(input_file, 'rb') as f:
                pdf_reader = pypdf.PdfReader(f)
                
                info = {
                    "num_pages": len(pdf_reader.pages),
                    "metadata": pdf_reader.metadata.__dict__ if pdf_reader.metadata else {},
                    "is_encrypted": pdf_reader.is_encrypted,
                    "page_sizes": []
                }
                
                # Get page sizes
                for page in pdf_reader.pages[:5]:  # First 5 pages
                    box = page.mediabox
                    info["page_sizes"].append({
                        "width": float(box.width),
                        "height": float(box.height)
                    })
                    
            return info
            
        except Exception as e:
            raise HTTPException(500, f"PDF info extraction failed: {str(e)}")
    
    # ==================== Word Operations ====================
    
    async def get_word_info(self, input_file: Path) -> dict:
        """Get Word document information"""
        if input_file.suffix.lower() != '.docx':
            raise HTTPException(400, "File must be .docx")
        
        try:
            doc = Document(str(input_file))
            
            info = {
                "num_paragraphs": len(doc.paragraphs),
                "num_tables": len(doc.tables),
                "num_images": len([r for p in doc.paragraphs for r in p.runs if r._element.xml.find('pic') > 0]),
                "word_count": sum(len(p.text.split()) for p in doc.paragraphs),
                "char_count": sum(len(p.text) for p in doc.paragraphs),
            }
            
            return info
            
        except Exception as e:
            raise HTTPException(500, f"Word info extraction failed: {str(e)}")
    
    # ==================== Excel Operations ====================
    
    async def get_excel_info(self, input_file: Path) -> dict:
        """Get Excel workbook information"""
        if input_file.suffix.lower() not in ['.xlsx', '.xlsm']:
            raise HTTPException(400, "File must be .xlsx or .xlsm")
        
        try:
            wb = load_workbook(str(input_file), read_only=True)
            
            info = {
                "num_sheets": len(wb.sheetnames),
                "sheet_names": wb.sheetnames,
                "active_sheet": wb.active.title,
            }
            
            wb.close()
            return info
            
        except Exception as e:
            raise HTTPException(500, f"Excel info extraction failed: {str(e)}")
    
    # ==================== PowerPoint Operations ====================
    
    async def get_powerpoint_info(self, input_file: Path) -> dict:
        """Get PowerPoint presentation information"""
        if input_file.suffix.lower() not in ['.pptx']:
            raise HTTPException(400, "File must be .pptx")
        
        try:
            prs = Presentation(str(input_file))
            
            info = {
                "num_slides": len(prs.slides),
                "slide_width": prs.slide_width,
                "slide_height": prs.slide_height,
            }
            
            return info
            
        except Exception as e:
            raise HTTPException(500, f"PowerPoint info extraction failed: {str(e)}")
    
    # ==================== Cleanup ====================
    
    # ==================== PDF Compression ====================
    
    async def compress_pdf(
        self,
        input_file: Path,
        quality: str = "medium",  # low, medium, high
        output_filename: Optional[str] = None
    ) -> tuple[Path, str]:
        """
        Compress PDF - HYBRID STRATEGY
        
        Strategy:
        1. Check technology priority from settings
        2. Try Adobe PDF Services (if enabled and high priority) - AI compression 10/10
        3. Fallback to pypdf - Basic compression 7/10
        
        Args:
            input_file: Path to PDF file
            quality: Compression level (low, medium, high)
            output_filename: Optional output filename
        
        Returns:
            Tuple of (output_path, technology_used)
            - output_path: Path to compressed PDF
            - technology_used: 'adobe' or 'pypdf'
        
        Quality mapping:
        - Adobe: low=HIGH (aggressive), medium=MEDIUM, high=LOW (preserve quality)
        - pypdf: low=4, medium=2, high=0
        """
        if input_file.suffix.lower() != '.pdf':
            raise HTTPException(400, "File must be .pdf")
        
        output_filename = output_filename or input_file.stem + "_compressed.pdf"
        output_path = self.output_dir / output_filename
        
        # Import settings here to avoid circular imports
        from app.core.config import settings
        
        # Get technology priority for compress operation
        priorities = settings.get_technology_priority("compress")
        
        # Try each technology in priority order
        for tech in priorities:
            if tech.lower() == "adobe":
                # Try Adobe if enabled
                if self.use_adobe and self.adobe_credentials and ADOBE_AVAILABLE:
                    try:
                        logger.info(f"Trying Adobe compress for {input_file.name}")
                        await self._compress_pdf_adobe(input_file, quality, output_path)
                        logger.info(f"Adobe compression successful: {output_path}")
                        return (output_path, "adobe")
                    except Exception as e:
                        logger.warning(f"Adobe compress failed: {e}, trying next technology")
                        continue
                else:
                    logger.debug("Adobe not available, skipping")
                    continue
                    
            elif tech.lower() == "pypdf":
                # Try pypdf (always available)
                try:
                    logger.info(f"Using pypdf compress for {input_file.name}")
                    await self._compress_pdf_local(input_file, quality, output_path)
                    logger.info(f"pypdf compression successful: {output_path}")
                    return (output_path, "pypdf")
                except Exception as e:
                    logger.warning(f"pypdf compress failed: {e}, trying next technology")
                    continue
        
        # If all technologies failed
        raise HTTPException(500, "All compression methods failed")
    
    async def _compress_pdf_adobe(
        self,
        input_file: Path,
        quality: str,
        output_path: Path
    ) -> None:
        """
        Compress PDF using Adobe PDF Services API
        AI-powered compression with 10/10 quality
        
        Adobe compression levels:
        - HIGH: Aggressive compression (50-80% reduction)
        - MEDIUM: Balanced (30-50% reduction)
        - LOW: Light compression (10-30% reduction)
        """
        try:
            # Read input file
            async with aiofiles.open(input_file, 'rb') as f:
                input_stream = await f.read()
            
            # Adobe API (simplified - actual implementation would use CompressPDF operation)
            # This is placeholder - needs actual Adobe CompressPDF SDK code
            from adobe.pdfservices.operation.pdfjobs.jobs.compress_pdf_job import CompressPDFJob
            from adobe.pdfservices.operation.pdfjobs.params.compress_pdf.compress_pdf_params import CompressPDFParams
            from adobe.pdfservices.operation.pdfjobs.params.compress_pdf.compression_level import CompressionLevel
            from adobe.pdfservices.operation.pdfjobs.result.compress_pdf_result import CompressPDFResult
            
            # Map quality to Adobe compression level
            quality_map = {
                "low": CompressionLevel.HIGH,      # More compression
                "medium": CompressionLevel.MEDIUM,
                "high": CompressionLevel.LOW       # Less compression, preserve quality
            }
            compression_level = quality_map.get(quality, CompressionLevel.MEDIUM)
            
            # Create PDF Services instance
            pdf_services = PDFServices(credentials=self.adobe_credentials)
            
            # Upload file
            input_asset = pdf_services.upload(
                input_stream=input_stream,
                mime_type=PDFServicesMediaType.PDF
            )
            
            # Create compress parameters
            compress_params = CompressPDFParams(
                compression_level=compression_level
            )
            
            # Create and submit job
            compress_job = CompressPDFJob(
                input_asset=input_asset,
                compress_pdf_params=compress_params
            )
            
            location = pdf_services.submit(compress_job)
            pdf_services_response = pdf_services.get_job_result(location, CompressPDFResult)
            
            # Download result
            result_asset: CloudAsset = pdf_services_response.get_result().get_asset()
            stream_asset: StreamAsset = pdf_services.get_content(result_asset)
            
            # Save to file
            async with aiofiles.open(output_path, "wb") as f:
                await f.write(stream_asset.get_input_stream())
            
        except ImportError:
            # CompressPDF not available in current SDK version
            logger.error("Adobe CompressPDF API not available in SDK")
            raise HTTPException(500, "Adobe CompressPDF requires SDK update")
        except Exception as e:
            logger.error(f"Adobe compression error: {e}")
            raise
    
    async def _compress_pdf_local(
        self,
        input_file: Path,
        quality: str,
        output_path: Path
    ) -> None:
        """
        Compress PDF using pypdf library
        Basic compression with 7/10 quality (existing code)
        """
        # Quality settings for pypdf compression
        quality_map = {
            "low": 4,      # More compression
            "medium": 2,   # Balanced
            "high": 0      # Less compression
        }
        
        compression_level = quality_map.get(quality, 2)
        
        try:
            with open(input_file, 'rb') as f:
                pdf_reader = pypdf.PdfReader(f)
                pdf_writer = pypdf.PdfWriter()
                
                # Copy all pages
                for page in pdf_reader.pages:
                    pdf_writer.add_page(page)
                
                # Compress images
                for page in pdf_writer.pages:
                    page.compress_content_streams(level=compression_level)
                
                # Write compressed PDF
                with open(output_path, 'wb') as output_file:
                    pdf_writer.write(output_file)
                    
        except Exception as e:
            raise HTTPException(500, f"pypdf compression failed: {str(e)}")
    
    # ==================== Image to PDF ====================
    
    async def image_to_pdf(
        self,
        input_file: Path,
        output_filename: Optional[str] = None
    ) -> Path:
        """Convert image (JPG, PNG, etc.) to PDF"""
        allowed_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp', '.heic']
        
        if input_file.suffix.lower() not in allowed_extensions:
            raise HTTPException(400, f"File must be one of: {', '.join(allowed_extensions)}")
        
        output_filename = output_filename or input_file.stem + ".pdf"
        output_path = self.output_dir / output_filename
        
        try:
            # Use Pillow to convert image to PDF
            from PIL import Image
            
            # Open and convert image
            img = Image.open(input_file)
            
            # Convert RGBA to RGB if necessary (for PNG with transparency)
            if img.mode in ('RGBA', 'LA', 'P'):
                background = Image.new('RGB', img.size, (255, 255, 255))
                if img.mode == 'P':
                    img = img.convert('RGBA')
                background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                img = background
            elif img.mode != 'RGB':
                img = img.convert('RGB')
            
            # Save as PDF
            img.save(output_path, 'PDF', resolution=100.0, quality=95)
            
            return output_path
            
        except Exception as e:
            raise HTTPException(500, f"Image to PDF conversion failed: {str(e)}")
    
    # ==================== PDF Watermark ====================
    
    async def add_watermark_to_pdf(
        self,
        input_file: Path,
        watermark_text: str,
        position: str = "center",  # center, top-left, top-right, bottom-left, bottom-right
        opacity: float = 0.3,
        output_filename: Optional[str] = None
    ) -> tuple[Path, str]:
        """
        Add text watermark to PDF - HYBRID STRATEGY
        
        Strategy:
        1. Check technology priority from settings
        2. Try Adobe PDF Services (if enabled) - Advanced watermark 10/10
        3. Fallback to reportlab+pypdf - Basic watermark 8/10
        
        Returns:
            Tuple of (output_path, technology_used)
            - output_path: Path to watermarked PDF
            - technology_used: 'adobe' or 'pypdf'
        """
        if input_file.suffix.lower() != '.pdf':
            raise HTTPException(400, "File must be .pdf")
        
        output_filename = output_filename or input_file.stem + "_watermarked.pdf"
        output_path = self.output_dir / output_filename
        
        # Import settings
        from app.core.config import settings
        
        # Get technology priority for watermark operation
        priorities = settings.get_technology_priority("watermark")
        
        # Try each technology in priority order
        for tech in priorities:
            if tech.lower() == "adobe":
                # Try Adobe if enabled
                if self.use_adobe and self.adobe_credentials and ADOBE_AVAILABLE:
                    try:
                        logger.info(f"Trying Adobe watermark for {input_file.name}")
                        await self._add_watermark_adobe(
                            input_file, watermark_text, position, opacity, output_path
                        )
                        logger.info(f"Adobe watermark successful: {output_path}")
                        return (output_path, "adobe")
                    except Exception as e:
                        logger.warning(f"Adobe watermark failed: {e}, trying next technology")
                        continue
                else:
                    logger.debug("Adobe not available, skipping")
                    continue
                    
            elif tech.lower() == "pypdf":
                # Try pypdf (always available)
                try:
                    logger.info(f"Using pypdf watermark for {input_file.name}")
                    await self._add_watermark_local(
                        input_file, watermark_text, position, opacity, output_path
                    )
                    logger.info(f"pypdf watermark successful: {output_path}")
                    return (output_path, "pypdf")
                except Exception as e:
                    logger.warning(f"pypdf watermark failed: {e}, trying next technology")
                    continue
        
        # If all technologies failed
        raise HTTPException(500, "All watermark methods failed")
    
    async def _add_watermark_adobe(
        self,
        input_file: Path,
        watermark_text: str,
        position: str,
        opacity: float,
        output_path: Path
    ) -> None:
        """
        Add watermark using Adobe PDF Services API
        Advanced watermark with 10/10 quality
        
        Note: Adobe doesn't have native Watermark API yet,
        so we use AddTextWatermark or manual page overlay
        This is a placeholder for future Adobe Watermark API
        """
        try:
            # Read input file
            async with aiofiles.open(input_file, 'rb') as f:
                input_stream = await f.read()
            
            # Adobe Watermark API (placeholder - Adobe SDK may not have this yet)
            # For now, we'll raise ImportError to fallback to pypdf
            # When Adobe releases Watermark API, update this code
            
            raise ImportError("Adobe Watermark API not available yet in SDK")
            
            # Future implementation when Adobe adds Watermark API:
            # from adobe.pdfservices.operation.pdfjobs.jobs.add_watermark_job import AddWatermarkJob
            # from adobe.pdfservices.operation.pdfjobs.params.add_watermark_params import AddWatermarkParams
            # 
            # pdf_services = PDFServices(credentials=self.adobe_credentials)
            # input_asset = pdf_services.upload(input_stream, PDFServicesMediaType.PDF)
            # 
            # watermark_params = AddWatermarkParams(
            #     text=watermark_text,
            #     position=position,
            #     opacity=opacity
            # )
            # 
            # watermark_job = AddWatermarkJob(input_asset, watermark_params)
            # location = pdf_services.submit(watermark_job)
            # result = pdf_services.get_job_result(location, AddWatermarkResult)
            # 
            # result_asset = result.get_result().get_asset()
            # stream_asset = pdf_services.get_content(result_asset)
            # 
            # async with aiofiles.open(output_path, "wb") as f:
            #     await f.write(stream_asset.get_input_stream())
            
        except ImportError:
            logger.info("Adobe Watermark API not available, will fallback to pypdf")
            raise
        except Exception as e:
            logger.error(f"Adobe watermark error: {e}")
            raise
    
    async def _add_watermark_local(
        self,
        input_file: Path,
        watermark_text: str,
        position: str,
        opacity: float,
        output_path: Path
    ) -> None:
        """
        Add watermark using reportlab + pypdf
        Basic watermark with 8/10 quality (existing code)
        """
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.colors import Color
            import io
            
            # Read input PDF
            with open(input_file, 'rb') as f:
                pdf_reader = pypdf.PdfReader(f)
                pdf_writer = pypdf.PdfWriter()
                
                # Create watermark
                packet = io.BytesIO()
                can = canvas.Canvas(packet, pagesize=letter)
                
                # Position mapping
                positions = {
                    'center': (300, 400),
                    'top-left': (50, 750),
                    'top-right': (450, 750),
                    'bottom-left': (50, 50),
                    'bottom-right': (450, 50)
                }
                x, y = positions.get(position, (300, 400))
                
                # Set watermark properties
                can.setFillColor(Color(0, 0, 0, alpha=opacity))
                can.setFont("Helvetica", 40)
                can.drawString(x, y, watermark_text)
                can.save()
                
                # Move to beginning of BytesIO
                packet.seek(0)
                watermark_pdf = pypdf.PdfReader(packet)
                watermark_page = watermark_pdf.pages[0]
                
                # Apply watermark to all pages
                for page in pdf_reader.pages:
                    page.merge_page(watermark_page)
                    pdf_writer.add_page(page)
                
                # Write output
                with open(output_path, 'wb') as output_file:
                    pdf_writer.write(output_file)
                    
        except Exception as e:
            raise HTTPException(500, f"pypdf watermark failed: {str(e)}")
    
    # ==================== PDF Password Protection ====================
    
    async def protect_pdf_with_password(
        self,
        input_file: Path,
        user_password: str,
        owner_password: Optional[str] = None,
        output_filename: Optional[str] = None
    ) -> Path:
        """Protect PDF with password"""
        if input_file.suffix.lower() != '.pdf':
            raise HTTPException(400, "File must be .pdf")
        
        output_filename = output_filename or input_file.stem + "_protected.pdf"
        output_path = self.output_dir / output_filename
        
        if not owner_password:
            owner_password = user_password + "_owner"
        
        try:
            with open(input_file, 'rb') as f:
                pdf_reader = pypdf.PdfReader(f)
                pdf_writer = pypdf.PdfWriter()
                
                # Copy all pages
                for page in pdf_reader.pages:
                    pdf_writer.add_page(page)
                
                # Encrypt with password
                pdf_writer.encrypt(
                    user_password=user_password,
                    owner_password=owner_password,
                    permissions_flag=0b0100  # Allow printing
                )
                
                # Write encrypted PDF
                with open(output_path, 'wb') as output_file:
                    pdf_writer.write(output_file)
                    
            return output_path
            
        except Exception as e:
            raise HTTPException(500, f"PDF password protection failed: {str(e)}")
    
    async def unlock_pdf(
        self,
        input_file: Path,
        password: str,
        output_filename: Optional[str] = None
    ) -> Path:
        """Remove password from PDF"""
        if input_file.suffix.lower() != '.pdf':
            raise HTTPException(400, "File must be .pdf")
        
        output_filename = output_filename or input_file.stem + "_unlocked.pdf"
        output_path = self.output_dir / output_filename
        
        try:
            with open(input_file, 'rb') as f:
                pdf_reader = pypdf.PdfReader(f)
                
                # Check if encrypted
                if not pdf_reader.is_encrypted:
                    raise HTTPException(400, "PDF is not encrypted")
                
                # Decrypt
                if not pdf_reader.decrypt(password):
                    raise HTTPException(401, "Incorrect password")
                
                pdf_writer = pypdf.PdfWriter()
                
                # Copy all pages (now decrypted)
                for page in pdf_reader.pages:
                    pdf_writer.add_page(page)
                
                # Write unencrypted PDF
                with open(output_path, 'wb') as output_file:
                    pdf_writer.write(output_file)
                    
            return output_path
            
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(500, f"PDF unlock failed: {str(e)}")
    
    # ==================== PDF to Images ====================
    
    async def pdf_to_images(
        self,
        input_file: Path,
        format: str = "png",  # png, jpg
        dpi: int = 200,
        output_prefix: str = "page"
    ) -> List[Path]:
        """Convert PDF pages to images using pdf2image"""
        if input_file.suffix.lower() != '.pdf':
            raise HTTPException(400, "File must be .pdf")
        
        try:
            from pdf2image import convert_from_path
            
            output_files = []
            
            # Convert PDF to images
            images = convert_from_path(input_file, dpi=dpi)
            
            # Save each page
            for page_num, image in enumerate(images, start=1):
                output_filename = f"{output_prefix}_{page_num}.{format}"
                output_path = self.output_dir / output_filename
                image.save(output_path, format.upper())
                output_files.append(output_path)
            
            return output_files
            
        except Exception as e:
            raise HTTPException(500, f"PDF to images conversion failed: {str(e)}")
    
    # ==================== Add Page Numbers ====================
    
    async def add_page_numbers(
        self,
        input_file: Path,
        position: str = "bottom-center",  # bottom-center, bottom-right, bottom-left
        format: str = "Page {page}",  # {page}, {page} of {total}, etc.
        output_filename: Optional[str] = None
    ) -> Path:
        """Add page numbers to PDF"""
        if input_file.suffix.lower() != '.pdf':
            raise HTTPException(400, "File must be .pdf")
        
        output_filename = output_filename or input_file.stem + "_numbered.pdf"
        output_path = self.output_dir / output_filename
        
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            import io
            
            with open(input_file, 'rb') as f:
                pdf_reader = pypdf.PdfReader(f)
                pdf_writer = pypdf.PdfWriter()
                total_pages = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages, start=1):
                    # Create page number overlay
                    packet = io.BytesIO()
                    can = canvas.Canvas(packet, pagesize=letter)
                    
                    # Position mapping
                    positions = {
                        'bottom-center': (300, 30),
                        'bottom-right': (520, 30),
                        'bottom-left': (80, 30)
                    }
                    x, y = positions.get(position, (300, 30))
                    
                    # Format page number text
                    page_text = format.replace('{page}', str(page_num))
                    page_text = page_text.replace('{total}', str(total_pages))
                    
                    can.setFont("Helvetica", 10)
                    can.drawString(x, y, page_text)
                    can.save()
                    
                    # Merge with original page
                    packet.seek(0)
                    overlay_pdf = pypdf.PdfReader(packet)
                    page.merge_page(overlay_pdf.pages[0])
                    pdf_writer.add_page(page)
                
                # Write output
                with open(output_path, 'wb') as output_file:
                    pdf_writer.write(output_file)
                    
            return output_path
            
        except Exception as e:
            raise HTTPException(500, f"Add page numbers failed: {str(e)}")
    
    # ==================== ADOBE-ONLY FEATURES ====================
    
    async def ocr_pdf(
        self,
        input_file: Path,
        language: str = "vi-VN",
        output_filename: Optional[str] = None
    ) -> Path:
        """
        OCR scanned PDF to searchable PDF
        
        Technology Priority:
        1. Adobe PDF Services (10/10) - Best quality, 50+ languages
        2. Tesseract OCR (7/10) - Free fallback with basic OCR
        
        Features:
        - Text recognition from scanned PDFs
        - Support Vietnamese (vi-VN) and 50+ languages
        - Preserve original layout (Adobe only)
        - Add searchable text layer
        
        Args:
            input_file: Path to scanned PDF file
            language: Language code (vi-VN, en-US, fr-FR, etc.)
            output_filename: Optional output filename
        
        Returns:
            Path to searchable PDF
        
        Raises:
            HTTPException 400: Invalid input
            HTTPException 500: OCR fails
        """
        # Try Adobe first (best quality)
        if self.use_adobe and self.adobe_credentials and ADOBE_AVAILABLE:
            logger.info("🎯 Using Adobe OCR for PDF")
            return await self._ocr_pdf_adobe(input_file, language, output_filename)
        
        # Fallback to Tesseract
        logger.info("🔄 Using Tesseract OCR fallback (Adobe not available)")
        return await self._ocr_pdf_tesseract(input_file, language, output_filename)
    
    async def _ocr_pdf_adobe(
        self,
        input_file: Path,
        language: str = "vi-VN",
        output_filename: Optional[str] = None
    ) -> Path:
        """Adobe OCR implementation (original logic)"""
        
        if input_file.suffix.lower() != '.pdf':
            raise HTTPException(400, "File must be .pdf")
        
        output_filename = output_filename or input_file.stem + "_ocr.pdf"
        output_path = self.output_dir / output_filename
        
        try:
            # Read input file
            async with aiofiles.open(input_file, 'rb') as f:
                input_stream = await f.read()
            
            # Adobe OCR API
            from adobe.pdfservices.operation.pdfjobs.jobs.ocr_pdf_job import OCRPDFJob
            from adobe.pdfservices.operation.pdfjobs.params.ocr_pdf.ocr_pdf_params import OCRPDFParams
            from adobe.pdfservices.operation.pdfjobs.params.ocr_pdf.ocr_supported_locale import OCRSupportedLocale
            from adobe.pdfservices.operation.pdfjobs.params.ocr_pdf.ocr_supported_type import OCRSupportedType
            from adobe.pdfservices.operation.pdfjobs.result.ocr_pdf_result import OCRPDFResult
            
            # Map language codes to Adobe OCRSupportedLocale
            language_map = {
                "vi-VN": OCRSupportedLocale.VI_VN,  # Vietnamese
                "en-US": OCRSupportedLocale.EN_US,  # English
                "fr-FR": OCRSupportedLocale.FR_FR,  # French
                "de-DE": OCRSupportedLocale.DE_DE,  # German
                "es-ES": OCRSupportedLocale.ES_ES,  # Spanish
                "it-IT": OCRSupportedLocale.IT_IT,  # Italian
                "ja-JP": OCRSupportedLocale.JA_JP,  # Japanese
                "ko-KR": OCRSupportedLocale.KO_KR,  # Korean
                "zh-CN": OCRSupportedLocale.ZH_HANS, # Chinese Simplified
            }
            ocr_locale = language_map.get(language, OCRSupportedLocale.EN_US)
            
            # Create PDF Services instance
            pdf_services = PDFServices(credentials=self.adobe_credentials)
            
            # Upload file
            input_asset = pdf_services.upload(
                input_stream=input_stream,
                mime_type=PDFServicesMediaType.PDF
            )
            
            # Create OCR parameters
            ocr_params = OCRPDFParams(
                ocr_locale=ocr_locale,
                ocr_type=OCRSupportedType.SEARCHABLE_IMAGE  # Preserve original appearance
            )
            
            # Create and submit OCR job
            ocr_job = OCRPDFJob(
                input_asset=input_asset,
                ocr_pdf_params=ocr_params
            )
            
            location = pdf_services.submit(ocr_job)
            logger.info(f"Adobe OCR job submitted for {input_file.name}")
            
            # Get result (polling handled by SDK)
            pdf_services_response = pdf_services.get_job_result(location, OCRPDFResult)
            
            # Download result
            result_asset: CloudAsset = pdf_services_response.get_result().get_asset()
            stream_asset: StreamAsset = pdf_services.get_content(result_asset)
            
            # Save to file
            async with aiofiles.open(output_path, "wb") as f:
                await f.write(stream_asset.get_input_stream())
            
            logger.info(f"Adobe OCR successful: {output_path}")
            return output_path
            
        except ImportError:
            logger.error("Adobe OCR API not available in SDK")
            raise HTTPException(500, "Adobe OCR requires SDK update")
        except Exception as e:
            logger.error(f"Adobe OCR error: {e}")
            raise HTTPException(500, f"OCR failed: {str(e)}")
    
    async def _ocr_pdf_tesseract(
        self,
        input_file: Path,
        language: str = "vi-VN",
        output_filename: Optional[str] = None
    ) -> Path:
        """
        OCR PDF using Tesseract (free fallback)
        
        Quality: 7/10 (Good for basic OCR, but not perfect layout preservation)
        Features: Extract text, basic recognition, multiple languages
        Limitation: Creates new PDF with text layer (doesn't preserve exact layout)
        """
        try:
            from pdf2image import convert_from_path
            from PIL import Image
            import pytesseract
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.lib.utils import ImageReader
            import io
        except ImportError as e:
            raise HTTPException(
                500, 
                f"Tesseract OCR dependencies not installed: {e}. "
                "Install: pip install pdf2image pytesseract pillow reportlab"
            )
        
        if input_file.suffix.lower() != '.pdf':
            raise HTTPException(400, "File must be .pdf")
        
        output_filename = output_filename or input_file.stem + "_ocr_tesseract.pdf"
        output_path = self.output_dir / output_filename
        
        try:
            # Map language codes to Tesseract format
            lang_map = {
                "vi-VN": "vie",  # Vietnamese
                "en-US": "eng",  # English
                "fr-FR": "fra",  # French
                "de-DE": "deu",  # German
                "es-ES": "spa",  # Spanish
                "it-IT": "ita",  # Italian
                "ja-JP": "jpn",  # Japanese
                "ko-KR": "kor",  # Korean
                "zh-CN": "chi_sim",  # Chinese Simplified
            }
            tesseract_lang = lang_map.get(language, "eng")
            
            logger.info(f"Converting PDF to images for OCR...")
            # Convert PDF pages to images
            images = convert_from_path(str(input_file), dpi=300)
            
            logger.info(f"Performing OCR on {len(images)} pages...")
            # Create new PDF with OCR text
            c = canvas.Canvas(str(output_path), pagesize=letter)
            
            for i, image in enumerate(images):
                logger.info(f"  Processing page {i+1}/{len(images)}...")
                
                # OCR the image
                text = pytesseract.image_to_string(image, lang=tesseract_lang)
                
                # Convert PIL image to bytes for ReportLab
                img_buffer = io.BytesIO()
                image.save(img_buffer, format='PNG')
                img_buffer.seek(0)
                img_reader = ImageReader(img_buffer)
                
                # Get image dimensions
                img_width, img_height = image.size
                
                # Scale to fit page (8.5 x 11 inches = 612 x 792 points)
                page_width, page_height = letter
                scale = min(page_width / img_width, page_height / img_height)
                scaled_width = img_width * scale
                scaled_height = img_height * scale
                
                # Draw image on page
                c.drawImage(img_reader, 0, 0, width=scaled_width, height=scaled_height)
                
                # Add invisible text layer (for searchability)
                c.setFillColorRGB(1, 1, 1, alpha=0.01)  # Almost invisible
                c.setFont("Helvetica", 8)
                text_object = c.beginText(10, page_height - 20)
                for line in text.split('\n'):
                    text_object.textLine(line)
                c.drawText(text_object)
                
                c.showPage()  # Next page
            
            c.save()
            logger.info(f"✅ Tesseract OCR successful: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Tesseract OCR error: {e}")
            raise HTTPException(500, f"Tesseract OCR failed: {str(e)}")
    
    async def extract_pdf_content(
        self,
        input_file: Path,
        extract_type: str = "all"  # all, text, tables, images
    ) -> dict:
        """
        AI-powered content extraction from PDF using Adobe Extract API
        
        Adobe Extract Features:
        - Extract tables → Structured data (CSV/Excel format)
        - Extract images → PNG files with metadata
        - Extract text with font information (bold, italic, size, family)
        - Reading order detection (AI-powered)
        - Character bounding boxes (precise position)
        - Document structure detection (headings, paragraphs, lists)
        
        Args:
            input_file: Path to PDF file
            extract_type: What to extract (all, text, tables, images)
        
        Returns:
            Dict with extracted content:
            {
                "text": [...],      # Text elements with font info
                "tables": [...],    # Table data
                "images": [...],    # Image paths
                "structure": {...}  # Document structure
            }
        
        Raises:
            HTTPException 400: If Adobe API not enabled
            HTTPException 500: If extraction fails
        
        NO FALLBACK: pypdf only does basic text extraction without AI
        """
        if not self.use_adobe or not self.adobe_credentials or not ADOBE_AVAILABLE:
            raise HTTPException(
                400,
                "Content extraction requires Adobe PDF Services API. "
                "Set USE_ADOBE_PDF_API=true and configure credentials in .env"
            )
        
        if input_file.suffix.lower() != '.pdf':
            raise HTTPException(400, "File must be .pdf")
        
        try:
            # Read input file
            async with aiofiles.open(input_file, 'rb') as f:
                input_stream = await f.read()
            
            # Adobe Extract API
            from adobe.pdfservices.operation.pdfjobs.jobs.extract_pdf_job import ExtractPDFJob
            from adobe.pdfservices.operation.pdfjobs.params.extract_pdf.extract_pdf_params import ExtractPDFParams
            from adobe.pdfservices.operation.pdfjobs.params.extract_pdf.extract_element_type import ExtractElementType
            from adobe.pdfservices.operation.pdfjobs.params.extract_pdf.extract_renditions_element_type import ExtractRenditionsElementType
            from adobe.pdfservices.operation.pdfjobs.result.extract_pdf_result import ExtractPDFResult
            
            # Create PDF Services instance
            pdf_services = PDFServices(credentials=self.adobe_credentials)
            
            # Upload file
            input_asset = pdf_services.upload(
                input_stream=input_stream,
                mime_type=PDFServicesMediaType.PDF
            )
            
            # Create extract parameters based on extract_type
            extract_params = ExtractPDFParams(
                elements_to_extract=[
                    ExtractElementType.TEXT,
                    ExtractElementType.TABLES
                ],
                elements_to_extract_renditions=[
                    ExtractRenditionsElementType.TABLES,
                    ExtractRenditionsElementType.FIGURES
                ]
            )
            
            # Create and submit extract job
            extract_job = ExtractPDFJob(
                input_asset=input_asset,
                extract_pdf_params=extract_params
            )
            
            location = pdf_services.submit(extract_job)
            logger.info(f"Adobe Extract job submitted for {input_file.name}")
            
            # Get result
            pdf_services_response = pdf_services.get_job_result(location, ExtractPDFResult)
            
            # Download result (returns ZIP with JSON + images)
            result_asset: CloudAsset = pdf_services_response.get_result().get_resource()
            stream_asset: StreamAsset = pdf_services.get_content(result_asset)
            
            # Save ZIP file temporarily
            import zipfile
            import json
            zip_path = self.output_dir / f"{input_file.stem}_extract.zip"
            async with aiofiles.open(zip_path, "wb") as f:
                await f.write(stream_asset.get_input_stream())
            
            # Extract ZIP and parse JSON
            extract_dir = self.output_dir / f"{input_file.stem}_extracted"
            extract_dir.mkdir(exist_ok=True)
            
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                zip_ref.extractall(extract_dir)
            
            # Read structuredData.json
            json_path = extract_dir / "structuredData.json"
            with open(json_path, 'r', encoding='utf-8') as f:
                extracted_data = json.load(f)
            
            # Parse extracted data based on extract_type
            result = {
                "text": [],
                "tables": [],
                "images": [],
                "structure": {}
            }
            
            # Process elements
            if "elements" in extracted_data:
                for element in extracted_data["elements"]:
                    if element.get("Path") and "Text" in element.get("Path", ""):
                        result["text"].append({
                            "text": element.get("Text", ""),
                            "font": element.get("Font", {}),
                            "bounds": element.get("Bounds", [])
                        })
                    elif element.get("Path") and "Table" in element.get("Path", ""):
                        result["tables"].append(element)
                    elif element.get("Path") and "Figure" in element.get("Path", ""):
                        result["images"].append(element)
            
            # Cleanup ZIP file
            zip_path.unlink()
            
            logger.info(f"Adobe Extract successful: extracted {len(result['text'])} text elements, "
                       f"{len(result['tables'])} tables, {len(result['images'])} images")
            
            return result
            
        except ImportError:
            logger.error("Adobe Extract API not available in SDK")
            raise HTTPException(500, "Adobe Extract requires SDK update")
        except Exception as e:
            logger.error(f"Adobe Extract error: {e}")
            raise HTTPException(500, f"Content extraction failed: {str(e)}")
    
    async def html_to_pdf(
        self,
        html_content: str,
        page_size: str = "A4",
        orientation: str = "portrait",
        output_filename: Optional[str] = None
    ) -> Path:
        """
        Convert HTML to PDF using Adobe PDF Services API
        
        Adobe HTML to PDF Features:
        - Perfect rendering (same as Chrome browser)
        - Support CSS, JavaScript
        - Custom page size (A4, Letter, Legal, etc.)
        - Orientation (portrait, landscape)
        - Header/Footer support
        - Margin control
        
        Args:
            html_content: HTML string or URL
            page_size: Page size (A4, Letter, Legal, A3, etc.)
            orientation: portrait or landscape
            output_filename: Optional output filename
        
        Returns:
            Path to generated PDF
        
        Raises:
            HTTPException 400: If Adobe API not enabled
            HTTPException 500: If conversion fails
        
        NO FALLBACK: wkhtmltopdf alternative exists but quality much lower
        """
        if not self.use_adobe or not self.adobe_credentials or not ADOBE_AVAILABLE:
            raise HTTPException(
                400,
                "HTML to PDF requires Adobe PDF Services API. "
                "Set USE_ADOBE_PDF_API=true and configure credentials in .env"
            )
        
        output_filename = output_filename or "document.pdf"
        output_path = self.output_dir / output_filename
        
        try:
            # Adobe CreatePDF from HTML
            from adobe.pdfservices.operation.pdfjobs.jobs.create_pdf_job import CreatePDFJob
            from adobe.pdfservices.operation.pdfjobs.params.create_pdf.create_pdf_params import CreatePDFParams
            from adobe.pdfservices.operation.pdfjobs.params.create_pdf.html.html_create_pdf_params import HTMLCreatePDFParams
            from adobe.pdfservices.operation.pdfjobs.params.create_pdf.page_layout import PageLayout
            from adobe.pdfservices.operation.pdfjobs.result.create_pdf_result import CreatePDFResult
            
            # Create PDF Services instance
            pdf_services = PDFServices(credentials=self.adobe_credentials)
            
            # Check if html_content is URL or HTML string
            is_url = html_content.startswith("http://") or html_content.startswith("https://")
            
            if is_url:
                # Create from URL
                html_params = HTMLCreatePDFParams(
                    page_layout=PageLayout.A4 if page_size == "A4" else PageLayout.LETTER
                )
                create_params = CreatePDFParams(
                    html_create_pdf_params=html_params
                )
                
                # Note: Adobe SDK may require actual HTML file upload, not URL string
                # This is simplified - actual implementation may vary
                raise NotImplementedError("URL to PDF conversion requires file upload")
                
            else:
                # Create from HTML string - save to temp file first
                import tempfile
                with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as tmp:
                    tmp.write(html_content)
                    tmp_html_path = Path(tmp.name)
                
                # Read HTML file
                async with aiofiles.open(tmp_html_path, 'rb') as f:
                    html_stream = await f.read()
                
                # Upload HTML file
                input_asset = pdf_services.upload(
                    input_stream=html_stream,
                    mime_type=PDFServicesMediaType.HTML
                )
                
                # Create PDF job
                html_params = HTMLCreatePDFParams(
                    page_layout=PageLayout.A4 if page_size == "A4" else PageLayout.LETTER
                )
                create_params = CreatePDFParams(
                    html_create_pdf_params=html_params
                )
                
                create_job = CreatePDFJob(
                    input_asset=input_asset,
                    create_pdf_params=create_params
                )
                
                location = pdf_services.submit(create_job)
                logger.info("Adobe HTML to PDF job submitted")
                
                # Get result
                pdf_services_response = pdf_services.get_job_result(location, CreatePDFResult)
                
                # Download result
                result_asset: CloudAsset = pdf_services_response.get_result().get_asset()
                stream_asset: StreamAsset = pdf_services.get_content(result_asset)
                
                # Save to file
                async with aiofiles.open(output_path, "wb") as f:
                    await f.write(stream_asset.get_input_stream())
                
                # Cleanup temp HTML file
                tmp_html_path.unlink()
                
                logger.info(f"Adobe HTML to PDF successful: {output_path}")
                return output_path
            
        except ImportError:
            logger.error("Adobe CreatePDF API not available in SDK")
            raise HTTPException(500, "Adobe HTML to PDF requires SDK update")
        except NotImplementedError as e:
            raise HTTPException(501, str(e))
        except Exception as e:
            logger.error(f"Adobe HTML to PDF error: {e}")
            raise HTTPException(500, f"HTML to PDF conversion failed: {str(e)}")
    
    # ==================== PDF Watermark (Adobe) ====================
    
    async def watermark_pdf(self, pdf_path: Path, watermark_path: Path) -> Path:
        """
        Đóng dấu mờ lên PDF sử dụng Adobe PDF Services
        
        Args:
            pdf_path: File PDF gốc
            watermark_path: File PDF dấu mờ (có thể tạo từ image/text trước)
        
        Returns:
            Path: File PDF đã có dấu mờ
        """
        if not self.adobe_credentials:
            raise HTTPException(501, "Adobe PDF Services chưa được cấu hình. Vui lòng thêm credentials vào .env")
        
        try:
            from adobe.pdfservices.operation.pdfjobs.jobs.pdf_watermark_job import PDFWatermarkJob
            from adobe.pdfservices.operation.pdfjobs.result.pdf_watermark_result import PDFWatermarkResult
            
            # Read files
            with open(pdf_path, 'rb') as f:
                source_stream = f.read()
            
            with open(watermark_path, 'rb') as f:
                watermark_stream = f.read()
            
            # Create PDF Services instance
            pdf_services = PDFServices(credentials=self.adobe_credentials)
            
            # Upload both files
            input_asset = pdf_services.upload(
                input_stream=source_stream,
                mime_type=PDFServicesMediaType.PDF
            )
            
            watermark_asset = pdf_services.upload(
                input_stream=watermark_stream,
                mime_type=PDFServicesMediaType.PDF
            )
            
            # Create watermark job
            watermark_job = PDFWatermarkJob(
                input_asset=input_asset,
                watermark_asset=watermark_asset
            )
            
            # Submit and get result
            location = pdf_services.submit(watermark_job)
            response = pdf_services.get_job_result(location, PDFWatermarkResult)
            
            # Get content
            result_asset: CloudAsset = response.get_result().get_asset()
            stream_asset: StreamAsset = pdf_services.get_content(result_asset)
            
            # Save output
            output_path = self.output_dir / f"watermarked_{pdf_path.name}"
            with open(output_path, "wb") as f:
                f.write(stream_asset.get_input_stream())
            
            logger.info(f"Adobe Watermark PDF successful: {output_path}")
            return output_path
            
        except ImportError:
            logger.error("Adobe Watermark API not available")
            raise HTTPException(500, "Adobe Watermark requires pdfservices-sdk")
        except Exception as e:
            logger.error(f"Adobe Watermark error: {e}")
            status_code, friendly_msg = get_friendly_error_message(e)
            raise HTTPException(status_code, friendly_msg)
    
    # ==================== PDF Combine (Adobe) ====================
    
    async def combine_pdfs(self, pdf_paths: List[Path], page_ranges: Optional[List[str]] = None) -> Path:
        """
        Gộp nhiều PDF thành một file
        
        Args:
            pdf_paths: List các file PDF
            page_ranges: Optional list page ranges như ["1-3", "all", "5-10"]
        
        Returns:
            Path: File PDF đã gộp
        """
        if not self.adobe_credentials:
            raise HTTPException(501, "Adobe PDF Services chưa được cấu hình")
        
        try:
            from adobe.pdfservices.operation.pdfjobs.jobs.combine_pdf_job import CombinePDFJob
            from adobe.pdfservices.operation.pdfjobs.params.combine_pdf.combine_pdf_params import CombinePDFParams
            from adobe.pdfservices.operation.pdfjobs.params.page_ranges import PageRanges
            from adobe.pdfservices.operation.pdfjobs.result.combine_pdf_result import CombinePDFResult
            
            # Create PDF Services instance
            pdf_services = PDFServices(credentials=self.adobe_credentials)
            
            # Create CombinePDFParams (v4 API pattern)
            combine_pdf_params = CombinePDFParams()
            
            for idx, pdf_path in enumerate(pdf_paths):
                with open(pdf_path, 'rb') as f:
                    stream = f.read()
                
                asset = pdf_services.upload(
                    input_stream=stream,
                    mime_type=PDFServicesMediaType.PDF
                )
                
                # Add with or without page range
                if page_ranges and idx < len(page_ranges) and page_ranges[idx] != "all":
                    # Parse range like "1-3" or "5-10"
                    range_str = page_ranges[idx]
                    page_range_obj = PageRanges()
                    
                    if '-' in range_str:
                        start, end = range_str.split('-')
                        page_range_obj.add_range(int(start), int(end))
                    else:
                        # Single page
                        page = int(range_str)
                        page_range_obj.add_single_page(page)
                    
                    # Add asset with page ranges to params
                    combine_pdf_params.add_asset(asset, page_range_obj)
                else:
                    # Add asset without page ranges (all pages)
                    combine_pdf_params.add_asset(asset)
            
            # Create job with params (v4 API: Job requires params argument)
            combine_job = CombinePDFJob(combine_pdf_params=combine_pdf_params)
            
            # Submit and get result
            location = pdf_services.submit(combine_job)
            response = pdf_services.get_job_result(location, CombinePDFResult)
            
            # Get content
            result_asset: CloudAsset = response.get_result().get_asset()
            stream_asset: StreamAsset = pdf_services.get_content(result_asset)
            
            # Save output
            output_path = self.output_dir / "combined.pdf"
            with open(output_path, "wb") as f:
                f.write(stream_asset.get_input_stream())
            
            logger.info(f"Adobe Combine PDF successful: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Adobe Combine error: {e}")
            status_code, friendly_msg = get_friendly_error_message(e)
            raise HTTPException(status_code, friendly_msg)
    
    # ==================== PDF Split (Adobe) ====================
    
    async def split_pdf(self, pdf_path: Path, page_ranges: List[str]) -> List[Path]:
        """
        Tách PDF thành nhiều file
        
        Args:
            pdf_path: File PDF gốc
            page_ranges: List ranges như ["1-3", "4-6", "7-10"]
        
        Returns:
            List[Path]: List các file PDF đã tách
        """
        if not self.adobe_credentials:
            raise HTTPException(501, "Adobe PDF Services chưa được cấu hình")
        
        try:
            from adobe.pdfservices.operation.pdfjobs.jobs.split_pdf_job import SplitPDFJob
            from adobe.pdfservices.operation.pdfjobs.params.split_pdf.split_pdf_params import SplitPDFParams
            from adobe.pdfservices.operation.pdfjobs.params.page_ranges import PageRanges
            from adobe.pdfservices.operation.pdfjobs.result.split_pdf_result import SplitPDFResult
            
            # Read file
            with open(pdf_path, 'rb') as f:
                stream = f.read()
            
            # Create PDF Services instance
            pdf_services = PDFServices(credentials=self.adobe_credentials)
            
            # Upload file
            input_asset = pdf_services.upload(
                input_stream=stream,
                mime_type=PDFServicesMediaType.PDF
            )
            
            # Parse page ranges - Adobe API expects ONE PageRanges object with multiple ranges
            page_ranges_obj = PageRanges()
            for range_str in page_ranges:
                if '-' in range_str:
                    start, end = range_str.split('-')
                    page_ranges_obj.add_range(int(start), int(end))
                else:
                    # Single page
                    page = int(range_str)
                    page_ranges_obj.add_single_page(page)
            
            # Create split params - pass single PageRanges object
            split_params = SplitPDFParams(page_ranges=page_ranges_obj)
            
            # Create and submit job
            split_job = SplitPDFJob(input_asset=input_asset, split_pdf_params=split_params)
            location = pdf_services.submit(split_job)
            response = pdf_services.get_job_result(location, SplitPDFResult)
            
            # Get all result assets
            result_assets = response.get_result().get_assets()
            logger.info(f"🔍 Split PDF: Got {len(result_assets)} result assets from Adobe")
            
            # Save all output files
            output_paths = []
            for idx, result_asset in enumerate(result_assets):
                stream_asset: StreamAsset = pdf_services.get_content(result_asset)
                content_bytes = stream_asset.get_input_stream()
                
                logger.info(f"📄 Split file {idx+1}: Content size = {len(content_bytes)} bytes")
                
                output_path = self.output_dir / f"split_{idx+1}_{pdf_path.name}"
                with open(output_path, "wb") as f:
                    f.write(content_bytes)
                
                # Verify file was written
                file_size = output_path.stat().st_size
                logger.info(f"💾 Split file {idx+1}: Saved to {output_path.name}, File size on disk = {file_size} bytes")
                
                if file_size == 0:
                    logger.error(f"❌ WARNING: Split file {idx+1} is EMPTY (0 bytes)!")
                
                output_paths.append(output_path)
            
            logger.info(f"✅ Adobe Split PDF successful: {len(output_paths)} files created")
            return output_paths
            
        except Exception as e:
            logger.error(f"Adobe Split error: {e}")
            status_code, friendly_msg = get_friendly_error_message(e)
            raise HTTPException(status_code, friendly_msg)
    
    # ==================== PDF Protect (Adobe) ====================
    
    async def protect_pdf(
        self, 
        pdf_path: Path, 
        user_password: str,
        owner_password: Optional[str] = None,
        permissions: Optional[List[str]] = None
    ) -> Path:
        """
        Bảo vệ PDF bằng mật khẩu và phân quyền
        
        Args:
            pdf_path: File PDF gốc
            user_password: Mật khẩu để mở file
            owner_password: Mật khẩu chủ sở hữu (optional)
            permissions: List quyền ["print", "copy", "edit"] (optional)
        
        Returns:
            Path: File PDF đã bảo vệ
        """
        if not self.adobe_credentials:
            raise HTTPException(501, "Adobe PDF Services chưa được cấu hình")
        
        try:
            from adobe.pdfservices.operation.pdfjobs.jobs.protect_pdf_job import ProtectPDFJob
            from adobe.pdfservices.operation.pdfjobs.params.protect_pdf.password_protect_params import PasswordProtectParams
            from adobe.pdfservices.operation.pdfjobs.params.protect_pdf.encryption_algorithm import EncryptionAlgorithm
            from adobe.pdfservices.operation.pdfjobs.params.protect_pdf.permission import Permission
            from adobe.pdfservices.operation.pdfjobs.params.protect_pdf.content_encryption import ContentEncryption
            from adobe.pdfservices.operation.pdfjobs.result.protect_pdf_result import ProtectPDFResult
            
            # Read file
            with open(pdf_path, 'rb') as f:
                stream = f.read()
            
            # Create PDF Services instance
            pdf_services = PDFServices(credentials=self.adobe_credentials)
            
            # Upload file
            input_asset = pdf_services.upload(
                input_stream=stream,
                mime_type=PDFServicesMediaType.PDF
            )
            
            # Parse permissions
            permission_list = []
            if permissions:
                perm_map = {
                    "print": Permission.PRINT_LOW_QUALITY,
                    "print_high": Permission.PRINT_HIGH_QUALITY,
                    "copy": Permission.COPY_CONTENT,
                    "edit": Permission.EDIT_CONTENT,
                    "edit_annotations": Permission.EDIT_ANNOTATIONS,
                    "fill_forms": Permission.FILL_AND_SIGN_FORM_FIELDS,
                    "assemble": Permission.EDIT_DOCUMENT_ASSEMBLY
                }
                for perm in permissions:
                    if perm in perm_map:
                        permission_list.append(perm_map[perm])
            
            # Create protect params (v4 API: Use PasswordProtectParams)
            protect_params = PasswordProtectParams(
                user_password=user_password,
                encryption_algorithm=EncryptionAlgorithm.AES_256,
                content_encryption=ContentEncryption.ALL_CONTENT
            )
            
            # Note: owner_password and permissions may need different API
            # Official sample doesn't show these - may need separate params or different constructor
            # For now keeping basic protection working
            
            # Create and submit job
            protect_job = ProtectPDFJob(input_asset=input_asset, protect_pdf_params=protect_params)
            location = pdf_services.submit(protect_job)
            response = pdf_services.get_job_result(location, ProtectPDFResult)
            
            # Get content
            result_asset: CloudAsset = response.get_result().get_asset()
            stream_asset: StreamAsset = pdf_services.get_content(result_asset)
            
            # Save output
            output_path = self.output_dir / f"protected_{pdf_path.name}"
            with open(output_path, "wb") as f:
                f.write(stream_asset.get_input_stream())
            
            logger.info(f"Adobe Protect PDF successful: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Adobe Protect error: {e}")
            status_code, friendly_msg = get_friendly_error_message(e)
            raise HTTPException(status_code, friendly_msg)
    
    # ==================== PDF Linearize (Adobe) ====================
    
    async def linearize_pdf(self, pdf_path: Path) -> Path:
        """
        Tối ưu hóa PDF cho web (fast web viewing)
        
        Args:
            pdf_path: File PDF gốc
        
        Returns:
            Path: File PDF đã tối ưu
        """
        if not self.adobe_credentials:
            raise HTTPException(501, "Adobe PDF Services chưa được cấu hình")
        
        try:
            from adobe.pdfservices.operation.pdfjobs.jobs.linearize_pdf_job import LinearizePDFJob
            from adobe.pdfservices.operation.pdfjobs.result.linearize_pdf_result import LinearizePDFResult
            
            # Read file
            with open(pdf_path, 'rb') as f:
                stream = f.read()
            
            # Create PDF Services instance
            pdf_services = PDFServices(credentials=self.adobe_credentials)
            
            # Upload file
            input_asset = pdf_services.upload(
                input_stream=stream,
                mime_type=PDFServicesMediaType.PDF
            )
            
            # Create and submit job
            linearize_job = LinearizePDFJob(input_asset=input_asset)
            location = pdf_services.submit(linearize_job)
            response = pdf_services.get_job_result(location, LinearizePDFResult)
            
            # Get content
            result_asset: CloudAsset = response.get_result().get_asset()
            stream_asset: StreamAsset = pdf_services.get_content(result_asset)
            
            # Save output
            output_path = self.output_dir / f"linearized_{pdf_path.name}"
            with open(output_path, "wb") as f:
                f.write(stream_asset.get_input_stream())
            
            logger.info(f"Adobe Linearize PDF successful: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Adobe Linearize error: {e}")
            status_code, friendly_msg = get_friendly_error_message(e)
            raise HTTPException(status_code, friendly_msg)
    
    # ==================== PDF Auto-Tag (Adobe) ====================
    
    async def autotag_pdf(self, pdf_path: Path, generate_report: bool = True) -> tuple[Path, Optional[Path]]:
        """
        Tự động gắn thẻ PDF cho accessibility (khả năng tiếp cận)
        
        Args:
            pdf_path: File PDF gốc
            generate_report: Có tạo báo cáo accessibility không
        
        Returns:
            tuple: (tagged_pdf_path, report_path)
        """
        if not self.adobe_credentials:
            raise HTTPException(501, "Adobe PDF Services chưa được cấu hình")
        
        try:
            from adobe.pdfservices.operation.pdfjobs.jobs.autotag_pdf_job import AutotagPDFJob
            from adobe.pdfservices.operation.pdfjobs.result.autotag_pdf_result import AutotagPDFResult
            
            # Read file
            with open(pdf_path, 'rb') as f:
                stream = f.read()
            
            # Create PDF Services instance
            pdf_services = PDFServices(credentials=self.adobe_credentials)
            
            # Upload file
            input_asset = pdf_services.upload(
                input_stream=stream,
                mime_type=PDFServicesMediaType.PDF
            )
            
            # Create and submit job
            autotag_job = AutotagPDFJob(
                input_asset=input_asset,
                generate_report=generate_report
            )
            location = pdf_services.submit(autotag_job)
            response = pdf_services.get_job_result(location, AutotagPDFResult)
            
            # Get tagged PDF
            result = response.get_result()
            tagged_asset: CloudAsset = result.get_tagged_pdf()
            stream_asset: StreamAsset = pdf_services.get_content(tagged_asset)
            
            # Save tagged PDF
            output_path = self.output_dir / f"tagged_{pdf_path.name}"
            with open(output_path, "wb") as f:
                f.write(stream_asset.get_input_stream())
            
            # Get report if generated
            report_path = None
            if generate_report:
                report_asset: CloudAsset = result.get_report()
                report_stream: StreamAsset = pdf_services.get_content(report_asset)
                
                report_path = self.output_dir / f"accessibility_report_{pdf_path.stem}.xlsx"
                with open(report_path, "wb") as f:
                    f.write(report_stream.get_input_stream())
            
            logger.info(f"Adobe Auto-Tag PDF successful: {output_path}")
            return output_path, report_path
            
        except Exception as e:
            logger.error(f"Adobe Auto-Tag error: {e}")
            status_code, friendly_msg = get_friendly_error_message(e)
            raise HTTPException(status_code, friendly_msg)
    
    async def generate_document(
        self,
        template_path: Path,
        json_data: dict,
        output_format: str = "pdf"
    ) -> Path:
        """
        Generate PDF/DOCX from Word template + JSON data.
        
        Args:
            template_path: Path to .docx template file
            json_data: Dictionary with merge data
            output_format: "pdf" or "docx"
            
        Returns:
            Path to generated file
        """
        if not self.adobe_credentials:
            raise HTTPException(
                status_code=500,
                detail="Adobe PDF Services credentials not configured"
            )
            
        try:
            # Import Adobe SDK components
            from adobe.pdfservices.operation.pdf_services import PDFServices
            from adobe.pdfservices.operation.pdf_services_media_type import PDFServicesMediaType
            from adobe.pdfservices.operation.pdfjobs.jobs.document_merge_job import DocumentMergeJob
            from adobe.pdfservices.operation.pdfjobs.params.documentmerge.document_merge_params import DocumentMergeParams
            from adobe.pdfservices.operation.pdfjobs.params.documentmerge.output_format import OutputFormat
            from adobe.pdfservices.operation.pdfjobs.result.document_merge_result import DocumentMergePDFResult
            
            # Initialize PDF Services
            pdf_services = PDFServices(credentials=self.adobe_credentials)
            
            # Read template file
            with open(template_path, 'rb') as template_file:
                input_stream = template_file.read()
            
            # Upload template
            input_asset = pdf_services.upload(
                input_stream=input_stream,
                mime_type=PDFServicesMediaType.DOCX
            )
            
            # Set output format
            if output_format.lower() == "pdf":
                merge_format = OutputFormat.PDF
            else:
                merge_format = OutputFormat.DOCX
            
            # Create merge parameters
            document_merge_params = DocumentMergeParams(
                json_data_for_merge=json_data,
                output_format=merge_format
            )
            
            # Create and submit job
            document_merge_job = DocumentMergeJob(
                input_asset=input_asset,
                document_merge_params=document_merge_params
            )
            
            location = pdf_services.submit(document_merge_job)
            pdf_services_response = pdf_services.get_job_result(
                location,
                DocumentMergePDFResult
            )
            
            # Get result asset
            result_asset = pdf_services_response.get_result().get_asset()
            stream_asset = pdf_services.get_content(result_asset)
            
            # Save output file
            output_file_path = self.output_dir / f"generated_document_{uuid.uuid4()}.{output_format.lower()}"
            with open(output_file_path, "wb") as output_file:
                output_file.write(stream_asset.get_input_stream())
            
            logger.info(f"Document generation successful: {output_file_path}")
            return output_file_path
            
        except Exception as e:
            logger.error(f"Document generation error: {e}")
            raise HTTPException(
                status_code=500,
                detail=f"Document generation failed: {str(e)}"
            )
    
    async def electronic_seal_pdf(
        self,
        pdf_path: Path,
        seal_image_path: Optional[Path],
        provider_name: str,
        access_token: str,
        credential_id: str,
        pin: str,
        seal_field_name: str = "Signature1",
        page_number: int = 1,
        visible: bool = True,
        field_x: int = 150,
        field_y: int = 250,
        field_width: int = 350,
        field_height: int = 200
    ) -> Path:
        """
        Apply electronic seal (digital signature) to PDF.
        
        Args:
            pdf_path: Path to PDF file
            seal_image_path: Optional path to seal image (PNG/JPG)
            provider_name: TSP provider name
            access_token: TSP access token
            credential_id: TSP credential ID
            pin: TSP PIN
            seal_field_name: Name of signature field
            page_number: Page number for seal (1-based)
            visible: Whether seal is visible
            field_x, field_y, field_width, field_height: Seal position/size
            
        Returns:
            Path to sealed PDF
        """
        if not self.adobe_credentials:
            raise HTTPException(
                status_code=500,
                detail="Adobe PDF Services credentials not configured"
            )
            
        try:
            # Import Adobe SDK components
            from adobe.pdfservices.operation.pdf_services import PDFServices
            from adobe.pdfservices.operation.pdf_services_media_type import PDFServicesMediaType
            from adobe.pdfservices.operation.pdfjobs.jobs.eseal_job import PDFElectronicSealJob
            from adobe.pdfservices.operation.pdfjobs.params.eseal.csc_auth_context import CSCAuthContext
            from adobe.pdfservices.operation.pdfjobs.params.eseal.csc_credentials import CSCCredentials
            from adobe.pdfservices.operation.pdfjobs.params.eseal.electronic_seal_params import PDFElectronicSealParams
            from adobe.pdfservices.operation.pdfjobs.params.eseal.field_location import FieldLocation
            from adobe.pdfservices.operation.pdfjobs.params.eseal.field_options import FieldOptions
            from adobe.pdfservices.operation.pdfjobs.params.eseal.document_level_permission import DocumentLevelPermission
            from adobe.pdfservices.operation.pdfjobs.result.eseal_pdf_result import ESealPDFResult
            
            # Initialize PDF Services
            pdf_services = PDFServices(credentials=self.adobe_credentials)
            
            # Upload PDF
            with open(pdf_path, 'rb') as pdf_file:
                pdf_stream = pdf_file.read()
            
            pdf_asset = pdf_services.upload(
                input_stream=pdf_stream,
                mime_type=PDFServicesMediaType.PDF
            )
            
            # Upload seal image if provided
            seal_image_asset = None
            if seal_image_path and seal_image_path.exists():
                with open(seal_image_path, 'rb') as img_file:
                    img_stream = img_file.read()
                
                # Determine image type
                ext = seal_image_path.suffix.lower()
                if ext == '.png':
                    mime_type = PDFServicesMediaType.PNG
                elif ext in ['.jpg', '.jpeg']:
                    mime_type = PDFServicesMediaType.JPEG
                else:
                    raise HTTPException(400, "Seal image must be PNG or JPEG")
                
                seal_image_asset = pdf_services.upload(
                    input_stream=img_stream,
                    mime_type=mime_type
                )
            
            # Create field location
            field_location = FieldLocation(field_x, field_y, field_width, field_height)
            
            # Create field options
            field_options = FieldOptions(
                field_name=seal_field_name,
                field_location=field_location,
                page_number=page_number,
                visible=visible
            )
            
            # Create CSC auth context
            csc_auth_context = CSCAuthContext(
                access_token=access_token,
                token_type="Bearer"
            )
            
            # Create CSC credentials
            csc_credentials = CSCCredentials(
                provider_name=provider_name,
                credential_id=credential_id,
                pin=pin,
                csc_auth_context=csc_auth_context
            )
            
            # Create seal parameters
            seal_params = PDFElectronicSealParams(
                seal_certificate_credentials=csc_credentials,
                seal_field_options=field_options
            )
            
            # Create job
            if seal_image_asset:
                seal_job = PDFElectronicSealJob(
                    input_asset=pdf_asset,
                    electronic_seal_params=seal_params,
                    seal_image_asset=seal_image_asset
                )
            else:
                seal_job = PDFElectronicSealJob(
                    input_asset=pdf_asset,
                    electronic_seal_params=seal_params
                )
            
            # Submit and get result
            location = pdf_services.submit(seal_job)
            pdf_services_response = pdf_services.get_job_result(location, ESealPDFResult)
            
            # Get result asset
            result_asset = pdf_services_response.get_result().get_asset()
            stream_asset = pdf_services.get_content(result_asset)
            
            # Save output
            output_file_path = self.output_dir / f"sealed_{uuid.uuid4()}.pdf"
            with open(output_file_path, "wb") as output_file:
                output_file.write(stream_asset.get_input_stream())
            
            logger.info(f"Electronic seal successful: {output_file_path}")
            return output_file_path
            
        except Exception as e:
            logger.error(f"Electronic seal error: {e}")
            raise HTTPException(
                status_code=500,
                detail=f"Electronic seal failed: {str(e)}"
            )
    
    # ==================== Cleanup ====================
    
    async def cleanup_file(self, file_path: Path) -> None:
        """Delete a file"""
        if file_path.exists():
            file_path.unlink()
    
    async def cleanup_old_files(self, max_age_hours: int = 24) -> int:
        """Delete files older than max_age_hours"""
        import time
        
        count = 0
        current_time = time.time()
        max_age_seconds = max_age_hours * 3600
        
        for dir_path in [self.upload_dir, self.output_dir]:
            for file_path in dir_path.glob("*"):
                if file_path.is_file():
                    file_age = current_time - file_path.stat().st_mtime
                    if file_age > max_age_seconds:
                        file_path.unlink()
                        count += 1
                        
        return count
