#!/usr/bin/env python3
"""
Auto Test Script - Word/Excel to PDF Conversion
Test trên server 165.99.59.47 với admin credentials
"""

import requests
import json
from pathlib import Path
import time

# Config
BASE_URL = "http://165.99.59.47"
USERNAME = "admin"
PASSWORD = "admin123"

class TestRunner:
    def __init__(self):
        self.base_url = BASE_URL
        self.token = None
        self.session = requests.Session()
        
    def login(self):
        """Login và lấy token"""
        print("🔐 Đang login...")
        url = f"{self.base_url}/api/auth/login"
        
        # JSON format
        data = {
            "username": USERNAME,
            "password": PASSWORD
        }
        
        try:
            response = self.session.post(url, json=data)
            
            print(f"   Status: {response.status_code}")
            
            if response.status_code == 200:
                result = response.json()
                print(f"   Response keys: {result.keys()}")
                
                # Token is in result['token']['access_token']
                if 'token' in result and 'access_token' in result['token']:
                    self.token = result['token']['access_token']
                elif 'access_token' in result:
                    self.token = result['access_token']
                else:
                    print(f"❌ No access_token in response: {result}")
                    return False
                    
                self.session.headers.update({
                    "Authorization": f"Bearer {self.token}"
                })
                print(f"✅ Login thành công!")
                print(f"   User: {result.get('user', {}).get('username', 'Unknown')}")
                print(f"   Token: {self.token[:30]}...")
                return True
            else:
                print(f"❌ Login thất bại: {response.status_code}")
                print(f"   Response: {response.text[:300]}")
                return False
        except Exception as e:
            print(f"❌ Lỗi login: {str(e)}")
            import traceback
            traceback.print_exc()
            return False
    
    def test_health(self):
        """Test health endpoint"""
        print("\n🏥 Testing health endpoint...")
        try:
            response = self.session.get(f"{self.base_url}/health")
            if response.status_code == 200:
                print(f"✅ Health check: {response.json()}")
                return True
            else:
                print(f"❌ Health check failed: {response.status_code}")
                return False
        except Exception as e:
            print(f"❌ Health check error: {str(e)}")
            return False
    
    def test_gotenberg_health(self):
        """Test Gotenberg service"""
        print("\n🔧 Testing Gotenberg service...")
        try:
            response = self.session.get(f"{self.base_url}:3000/health")
            if response.status_code == 200:
                result = response.json()
                print(f"✅ Gotenberg status: {result['status']}")
                print(f"   - LibreOffice: {result['details']['libreoffice']['status']}")
                print(f"   - Chromium: {result['details']['chromium']['status']}")
                return True
            else:
                print(f"❌ Gotenberg not healthy: {response.status_code}")
                return False
        except Exception as e:
            print(f"❌ Gotenberg error: {str(e)}")
            return False
    
    def create_test_docx(self):
        """Tạo file Word test đơn giản"""
        from docx import Document
        
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('Đây là file Word test để convert sang PDF.')
        doc.add_paragraph('Test với tiếng Việt có dấu: àáảãạăằắẳẵặâầấẩẫậ')
        doc.add_heading('Section 1', level=1)
        doc.add_paragraph('Nội dung section 1 với **bold text**.')
        
        test_file = Path("test_word.docx")
        doc.save(test_file)
        print(f"✅ Created test Word file: {test_file}")
        return test_file
    
    def create_test_xlsx(self):
        """Tạo file Excel test đơn giản"""
        from openpyxl import Workbook
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Test Sheet"
        
        # Headers
        ws['A1'] = 'Tên'
        ws['B1'] = 'Tuổi'
        ws['C1'] = 'Địa chỉ'
        
        # Data
        ws['A2'] = 'Nguyễn Văn A'
        ws['B2'] = 25
        ws['C2'] = 'Hà Nội'
        
        ws['A3'] = 'Trần Thị B'
        ws['B3'] = 30
        ws['C3'] = 'TP HCM'
        
        test_file = Path("test_excel.xlsx")
        wb.save(test_file)
        print(f"✅ Created test Excel file: {test_file}")
        return test_file
    
    def test_word_to_pdf(self):
        """Test Word to PDF conversion"""
        print("\n📄 Testing Word to PDF conversion...")
        
        try:
            # Create test file
            test_file = self.create_test_docx()
            
            # Upload and convert
            url = f"{self.base_url}/api/documents/convert/word-to-pdf"
            with open(test_file, 'rb') as f:
                files = {'file': (test_file.name, f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
                response = self.session.post(url, files=files, timeout=30)
            
            if response.status_code == 200:
                # Save output PDF
                output_file = test_file.with_suffix('.pdf')
                with open(output_file, 'wb') as f:
                    f.write(response.content)
                
                print(f"✅ Word to PDF conversion SUCCESS!")
                print(f"   Input: {test_file} ({test_file.stat().st_size} bytes)")
                print(f"   Output: {output_file} ({output_file.stat().st_size} bytes)")
                
                # Cleanup
                test_file.unlink()
                return True, output_file
            else:
                print(f"❌ Word to PDF FAILED: {response.status_code}")
                print(f"   Response: {response.text[:200]}")
                test_file.unlink()
                return False, None
                
        except Exception as e:
            print(f"❌ Word to PDF error: {str(e)}")
            return False, None
    
    def test_excel_to_pdf(self):
        """Test Excel to PDF conversion"""
        print("\n📊 Testing Excel to PDF conversion...")
        
        try:
            # Create test file
            test_file = self.create_test_xlsx()
            
            # Upload and convert
            url = f"{self.base_url}/api/documents/convert/excel-to-pdf"
            with open(test_file, 'rb') as f:
                files = {'file': (test_file.name, f, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')}
                response = self.session.post(url, files=files, timeout=30)
            
            if response.status_code == 200:
                # Save output PDF
                output_file = test_file.with_suffix('.pdf')
                with open(output_file, 'wb') as f:
                    f.write(response.content)
                
                print(f"✅ Excel to PDF conversion SUCCESS!")
                print(f"   Input: {test_file} ({test_file.stat().st_size} bytes)")
                print(f"   Output: {output_file} ({output_file.stat().st_size} bytes)")
                
                # Cleanup
                test_file.unlink()
                return True, output_file
            else:
                print(f"❌ Excel to PDF FAILED: {response.status_code}")
                print(f"   Response: {response.text[:200]}")
                test_file.unlink()
                return False, None
                
        except Exception as e:
            print(f"❌ Excel to PDF error: {str(e)}")
            return False, None
    
    def run_all_tests(self):
        """Chạy tất cả tests"""
        print("=" * 60)
        print("🚀 STARTING AUTO TEST - WORD/EXCEL TO PDF")
        print("=" * 60)
        print(f"Server: {self.base_url}")
        print(f"Username: {USERNAME}")
        print()
        
        results = {}
        
        # Test 1: Health check
        results['health'] = self.test_health()
        
        # Test 2: Login
        if not self.login():
            print("\n❌ Cannot continue without login!")
            return results
        
        # Test 3: Gotenberg health
        results['gotenberg'] = self.test_gotenberg_health()
        
        # Test 4: Word to PDF
        success, pdf_file = self.test_word_to_pdf()
        results['word_to_pdf'] = success
        if pdf_file:
            pdf_file.unlink()  # Cleanup
        
        # Test 5: Excel to PDF
        success, pdf_file = self.test_excel_to_pdf()
        results['excel_to_pdf'] = success
        if pdf_file:
            pdf_file.unlink()  # Cleanup
        
        # Summary
        print("\n" + "=" * 60)
        print("📊 TEST SUMMARY")
        print("=" * 60)
        for test_name, result in results.items():
            status = "✅ PASS" if result else "❌ FAIL"
            print(f"{test_name:20} : {status}")
        
        total = len(results)
        passed = sum(1 for r in results.values() if r)
        print(f"\nTotal: {passed}/{total} tests passed ({passed/total*100:.1f}%)")
        
        if passed == total:
            print("\n🎉 ALL TESTS PASSED! System is working perfectly!")
        else:
            print("\n⚠️ Some tests failed. Check logs above for details.")
        
        return results


if __name__ == "__main__":
    # Install required packages if needed
    try:
        import docx
        import openpyxl
    except ImportError:
        print("📦 Installing required packages...")
        import subprocess
        subprocess.run(["pip", "install", "python-docx", "openpyxl"], check=True)
        print("✅ Packages installed!")
        print()
    
    # Run tests
    tester = TestRunner()
    results = tester.run_all_tests()
    
    # Exit code
    exit(0 if all(results.values()) else 1)
